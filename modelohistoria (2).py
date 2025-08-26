#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import sys
import time
import threading

# Delays used for the introductory typing and progress bar animations.
TYPEWRITER_DELAY = 0.02
PROGRESS_DELAY = 0.02

def typewriter(text):
    """Print text with a short delay between characters."""
    for char in text:
        sys.stdout.write(char)
        sys.stdout.flush()
        time.sleep(TYPEWRITER_DELAY)
    sys.stdout.write("\n")

def show_intro():
    """Display a short welcome message with a progress bar."""
    print("\033[96m", end="")
    typewriter("🩺  Bienvenido a HealthIA - Tu asistente médico inteligente 🧠")
    print("\033[0m", end="")
    for i in range(0, 101, 5):
        bar = '█' * (i // 5) + '-' * ((100 - i) // 5)
        sys.stdout.write(f"\r[{bar}] {i}%")
        sys.stdout.flush()
        time.sleep(PROGRESS_DELAY)
    print("\n")

import random
import string

import string
import shutil
import re
import json
from os.path import expanduser
from datetime import datetime, date, timedelta

from PyQt5 import QtWidgets, QtGui
from PyPDF2 import PdfReader  # Añadido nuevo
from PyQt5.QtGui import QPixmap, QIcon, QPainter, QPen
from functools import partial
try:
    import qtawesome as qta
    _orig_qta_icon = qta.icon

    def _safe_qta_icon(*names, **kwargs):
        prefixes = ("fa5s.", "fa5r.", "fa5.", "fa.")
        for name in names:
            try:
                return _orig_qta_icon(name, **kwargs)
            except Exception:
                base = name.split(".", 1)[-1]
                for p in prefixes:
                    alt = p + base
                    if alt == name:
                        continue
                    try:
                        return _orig_qta_icon(alt, **kwargs)
                    except Exception:
                        continue
        return QIcon()

    qta.icon = _safe_qta_icon
except Exception:  # pragma: no cover - optional dependency
    qta = None
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QLineEdit, QTextEdit,
    QLabel, QMessageBox, QTableWidget, QTableWidgetItem, QDateEdit, QDateTimeEdit, QComboBox, QCheckBox,
    QScrollArea, QFormLayout, QDialog, QFileDialog, QGroupBox, QGridLayout, QTabWidget, QListWidget,
    QListWidgetItem, QShortcut, QSizePolicy, QCompleter, QAction, QInputDialog, QTableView,
    QHeaderView, QSpinBox, QStyle, QToolButton, QGraphicsDropShadowEffect, QFrame, QRadioButton, QProgressBar,
    QPlainTextEdit, QMenu, QAbstractItemView
)
from PyQt5.QtGui import QFont, QPalette, QColor, QKeySequence, QTextCursor, QTextDocument, QFontMetrics, QDesktopServices
from PyQt5.QtPrintSupport import QPrinter, QPrintDialog
from PyQt5.QtCore import (
    Qt,
    QDate,
    QDateTime,
    QSize,
    QCoreApplication,
    QBuffer,
    QByteArray,
    QIODevice,
    QEvent,
    QObject,
)
from PyQt5.QtCore import pyqtSignal, QUrl, QSortFilterProxyModel


from sqlalchemy import (
    create_engine, Column, Integer, String, Date, DateTime, Time, ForeignKey, Text, Boolean, select, func, text, inspect, Float, or_
)
from sqlalchemy.orm import (
    sessionmaker,
    relationship,
    declarative_base,
    Session,
    scoped_session,
    joinedload,
)
from sqlalchemy.exc import OperationalError, SQLAlchemyError


from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import hashlib
import sys
import os
import tempfile
import json
import openai
from validation_module import (
    DiagnosticoValidator,
    ValidationResult,
    configurar_tabla,
    insertar_resultado_en_tabla,
    grafico_pastel_3d,
)
import base64
from bmi_utils import GrowthTables, classify_nutritional_status
from consulta_cedula import consultar_cedula_todo
import estilos

# Lazy pandas loader to avoid heavy import at startup
_pd = None

def _get_pandas():
    global _pd
    if _pd is None:
        import pandas as pd
        _pd = pd
    return _pd

def read_excel(*args, **kwargs):
    pd = _get_pandas()
    return pd.read_excel(*args, **kwargs)

def read_csv(*args, **kwargs):
    pd = _get_pandas()
    return pd.read_csv(*args, **kwargs)


import secrets
from PyQt5.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout, 
                           QPushButton, QLineEdit, QLabel, QStackedWidget, 
                           QCheckBox, QMessageBox, QComboBox)
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QFont, QIcon
from PyQt5.QtCore import QThread  # Importa QThread

from PyQt5.QtCore import Qt, QThread, pyqtSignal, QTimer  # Importa QTimer aquí
from PyQt5.QtGui import QMovie, QColor
from PyQt5.QtWidgets import QApplication, QWidget, QLabel, QVBoxLayout, QMainWindow
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import matplotlib.dates as mdates
import matplotlib.pyplot as plt


print("\r\033[92mSistema listo! ✅\033[0m\n")
detener_progreso = True
os.system('cls')

import ctypes
time.sleep(2)  # Esperar 2 segundos antes de minimizar
ctypes.windll.user32.ShowWindow(ctypes.windll.kernel32.GetConsoleWindow(), 6) 



from sqlalchemy import (
    create_engine, Column, Integer, String, Date, DateTime, Time, ForeignKey, Text, Boolean, select, func, text, inspect, Float, or_
)
from sqlalchemy.orm import sessionmaker, relationship, declarative_base, Session, scoped_session
from sqlalchemy.exc import OperationalError, SQLAlchemyError
import re
CIE10_REGEX = re.compile(
    r'(?:\(|-|:)\s*([A-TV-Z][0-9]{2}(?:\.[0-9A-TV-Z]{1,4})?)\)?',
    re.I
)
Z_BLACK = {"Z00", "Z72", "Z83", "Z86", "Z87", "Z91"}
TRANSLATIONS = {
    'es': {
        'search_label': 'Buscar paciente:',
        'search_patient': 'Buscar Paciente',
        'new_patient': 'Nuevo Paciente',
        'edit_patient': 'Editar Paciente',
        'delete_patient': 'Eliminar Paciente',
        'new_history': 'Historia Clínica',
        'delete_history': 'Eliminar Historia',
        'generate_daily': 'Generar Informe Diario',
        'export_pdf': 'Exportar el Registro Diario de Consulta a PDF',
        'clean_duplicates': 'Eliminar Duplicados',
        'generate_history': 'Generar Historia Clinica',
        'save_history': 'Guardar Historia',
        'export_word': 'Exportar a Word',
        'send_ai': 'Enviar a IA',
        'generate_recipe': 'Generar Receta',
        'generate_lab': 'Generar Laboratorio',
        'auto_adjust': 'Autoajustar Tamaño',
        'load_file': 'Cargar Archivo',
        'analyze_image': 'Analizar Imagen',
        'explain_patient': 'Explicar al Paciente',
        'ask_history': 'Preguntar Historial',
        'fixed_med': 'Medicacion Fija',
        'appointments': 'Citas',
        'billing': 'Facturación',
        'print_invoice': 'Imprimir',
        'data_analysis': 'Análisis de Datos',
        'ai_study': 'Estudio Clínico IA',
        'clinic': 'Consultorio',
        'manage_db': 'Gestionar Base de Datos',
        'manage_insurance': 'Gestionar ARS',
        'manage_services': 'Gestionar Servicios',
        'forms': 'Formularios',
        'medical_certificate': 'Certificado Médico',
        'phq9': 'Cuestionario PHQ-9',
        'gad7': 'Cuestionario GAD-7',
        'mmse': 'Mini-Mental',
        'denver': 'Tamizaje Denver II',
        'today_appointments': 'Citas de Hoy',
        'insurance_stats': 'Pacientes por ARS',
        'housing_stats': 'Estadísticas de Vivienda',
        'stats_dashboard': 'Dashboard de Estadísticas',
        'stats_patients': 'Pacientes',
        'stats_histories': 'Historias',
        'stats_appointments': 'Citas',
        'stats_pending': 'Pendientes',
        'reminders': 'Recordatorios',
        'consents': 'Consentimientos',
        'advanced_search': 'Búsqueda Avanzada',
        'vital_signs': 'Signos Vitales',
        'family_social': 'Salud Familiar y Social',
        'lab_results': 'Resultados de Laboratorio',
        'height': 'Altura',
        'bmi': 'IMC',
        'period_history': 'Informe por Periodo',
        'apply_filters': 'Aplicar Filtros',
        'all': 'Todos',
        'gender': 'Sexo',
        'diagnosis_filter': 'Diagnóstico contiene',
        'total_patients': 'Total de Pacientes',
        'gender_distribution': 'Distribución por Sexo',
        'age_distribution': 'Distribución por Edad',
        'male': 'Masculino',
        'female': 'Femenino',
        'patient': 'Paciente',
        'generate': 'Generar',
        'healthy_patient': 'Paciente sano',
        'load_diag_history': 'Cargar diagnóstico',
        'days_off': 'Días de incapacidad',
        'validate_ai_file': 'Validar Diagnóstico IA (Archivo)',
        'export_results': 'Exportar Resultados'
    },
    'en': {
        'search_label': 'Search patient:',
        'search_patient': 'Search Patient',
        'new_patient': 'New Patient',
        'edit_patient': 'Edit Patient',
        'delete_patient': 'Delete Patient',
        'new_history': 'Medical History',
        'delete_history': 'Delete History',
        'generate_daily': 'Generate Daily Report',
        'export_pdf': 'Export Daily Consultation Record to PDF',
        'clean_duplicates': 'Delete Duplicates',
        'generate_history': 'Generate Medical History',
        'save_history': 'Save History',
        'export_word': 'Export to Word',
        'send_ai': 'Send to AI',
        'generate_recipe': 'Generate Prescription',
        'generate_lab': 'Generate Lab',
        'auto_adjust': 'Auto Adjust Size',
        'load_file': 'Load File',
        'analyze_image': 'Analyze Image',
        'explain_patient': 'Explain to Patient',
        'ask_history': 'Query History',
        'fixed_med': 'Fixed Medication',
        'appointments': 'Appointments',
        'billing': 'Billing',
        'print_invoice': 'Print',
        'data_analysis': 'Data Analysis',
        'ai_study': 'Clinical AI Study',
        'clinic': 'Clinic',
        'manage_db': 'Manage Database',
        'manage_insurance': 'Manage Insurance',
        'manage_services': 'Manage Services',
        'forms': 'Forms',
        'medical_certificate': 'Medical Certificate',
        'phq9': 'PHQ-9 Questionnaire',
        'gad7': 'GAD-7 Questionnaire',
        'mmse': 'Mini-Mental',
        'denver': 'Denver II Screening',
        'today_appointments': "Today's Appointments",
        'insurance_stats': 'Patients by Insurance',
        'housing_stats': 'Housing Statistics',
        'stats_dashboard': 'Statistics Dashboard',
        'stats_patients': 'Patients',
        'stats_histories': 'Histories',
        'stats_appointments': 'Appointments',
        'stats_pending': 'Pending',
        'reminders': 'Reminders',
        'consents': 'Consents',
        'advanced_search': 'Advanced Search',
        'vital_signs': 'Vital Signs',
        'family_social': 'Family & Social Health',
        'lab_results': 'Lab Results',
        'height': 'Height',
        'bmi': 'BMI',
        'period_history': 'History Report by Period',
        'apply_filters': 'Apply Filters',
        'all': 'All',
        'gender': 'Gender',
        'diagnosis_filter': 'Diagnosis contains',
        'total_patients': 'Total Patients',
        'gender_distribution': 'Gender Distribution',
        'age_distribution': 'Age Distribution',
        'male': 'Male',
        'female': 'Female',
        'patient': 'Patient',
        'generate': 'Generate',
        'healthy_patient': 'Healthy patient',
        'load_diag_history': 'Load diagnosis',
        'days_off': 'Days off',
        'validate_ai_file': 'Validate AI Diagnosis (File)',
        'export_results': 'Export Results'
    }
}


class LoadingScreen(QWidget):
    stop_signal = pyqtSignal()  # Signal to stop the animation

    def __init__(self):
        super().__init__()
        self.setWindowFlags(Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint | Qt.Tool)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.movie = QMovie("")
        self.label_animation = QLabel(self)
        self.label_animation.setMovie(self.movie)
        self.label_text = QLabel("", self)
        self.label_text.setAlignment(Qt.AlignCenter)
        self.label_text.setStyleSheet("color: #2980b9; font-size: 20px; font-weight: bold;")
       
        layout = QVBoxLayout()
        layout.addWidget(self.label_animation)
        layout.addWidget(self.label_text)
        layout.setContentsMargins(0, 0, 0, 0)
        self.setLayout(layout)
        
        #self.movie.start()
        self.resize(400, 300)
        self.setWindowTitle("Cargando...")
        
        self.color_value = 0
        self.timer = QTimer()
        self.timer.timeout.connect(self.animate_text_color)
        self.timer.start(60)
        
        self.stop_signal.connect(self.stop_animation)

    def animate_text_color(self):
        self.color_value = (self.color_value + 1) % 100
        blue_intensity = 50 + int(205 * (self.color_value / 100))
        color = QColor(0, blue_intensity, 255)
        self.label_text.setStyleSheet(f"color: {color.name()}; font-size: 20px; font-weight: bold;")

    def stop_animation(self):
        self.timer.stop()
        self.movie.stop()
        self.create_ready_file()  # Crear el archivo de indicador antes de cerrar la ventana de carga
        self.close()

    def create_ready_file(self):
        # Ruta del directorio donde está el ejecutable
        exe_dir = os.path.dirname(sys.argv[0])
        ready_file_path = os.path.join(exe_dir, "app_ready.txt")
        with open(ready_file_path, "w") as f:
            f.write("App is ready")

class InitializationThread(QThread):
    loading_complete = pyqtSignal()

    def run(self):
     
       
        self.loading_complete.emit()
class LoadingThread(QThread):
    def __init__(self, loading_screen):
        super().__init__()
        self.loading_screen = loading_screen

    def run(self):
        self.loading_screen.show()


class HistoryGeneratorThread(QThread):
    """Generate medical history suggestions in a background thread."""
    chunk_received = pyqtSignal(str)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, messages, model):
        super().__init__()
        self.messages = messages
        self.model = model

    def run(self):
        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=self.messages,
                temperature=0.6,
                stream=True,
            )
            text = ""
            for chunk in response:
                delta = chunk["choices"][0].get("delta", {}).get("content", "")
                if delta:
                    text += delta
                    self.chunk_received.emit(delta)
            # Clean basic markdown markers
            text = re.sub(r"#{1,6}\s", "", text)
            text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
            text = re.sub(r"-{3,}", "", text)
            self.finished.emit(text)
        except Exception as e:
            self.error.emit(str(e))


class AIStreamThread(QThread):
    """Generic thread to stream OpenAI responses."""
    chunk_received = pyqtSignal(str)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, messages, model):
        super().__init__()
        self.messages = messages
        self.model = model

    def run(self):
        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=self.messages,
                temperature=0.6,
                stream=True,
            )
            text = ""
            for chunk in response:
                delta = chunk["choices"][0].get("delta", {}).get("content", "")
                if delta:
                    text += delta
                    self.chunk_received.emit(delta)
            self.finished.emit(text)
        except Exception as e:
            self.error.emit(str(e))


class ValidationThread(QThread):
    """Process validation cases in the background."""
    progress = pyqtSignal(int)
    result_ready = pyqtSignal(ValidationResult, dict)
    stats_ready = pyqtSignal(dict)
    error = pyqtSignal(str)

    def __init__(self, validator: DiagnosticoValidator, ruta: str, autosave: str | None = None):
        super().__init__()
        self.validator = validator
        self.ruta = ruta
        self.autosave = autosave

    def run(self):
        try:
            if self.ruta.endswith(".xlsx"):
                df = read_excel(self.ruta)
            else:
                df = read_csv(self.ruta)
        except Exception as e:
            self.error.emit(str(e))
            return

        try:
            for res, metrics in self.validator.procesar_dataframe(df):
                self.result_ready.emit(res, metrics)
                self.progress.emit(metrics.get("casos_validados", 0))
                if self.autosave:
                    try:
                        clean = self.autosave.replace(".xlsx", "_clean.xlsx")
                        self.validator.exportar_excel(self.autosave, clean)
                    except Exception:
                        pass
            stats = self.validator.generar_reporte_final()
            self.stats_ready.emit(stats)
        except Exception as e:
            self.error.emit(str(e))
        

STYLE_SHEET = '''
QWidget {
    background-color: #2c3e50;
    color: #ecf0f1;
    font-family: "Segoe UI", Arial;
}

QLabel {
    color: #ecf0f1;
    font-size: 32px;
}

QLabel#titleLabel {
    font-size: 36px;
    font-weight: bold;
    color: #ecf0f1;
}

QLineEdit {
    padding: 10px;
    border: 1px solid #34495e;
    border-radius: 6px;
    background-color: #34495e;
    color: #ecf0f1;
    font-size: 28px;
}

QLineEdit:focus {
    border-color: #ecf0f1;
}

QPushButton {
    padding: 10px;
    border-radius: 6px;
    font-size: 28px;
    font-weight: bold;
    color: #ecf0f1;
    background-color: #3498db;
    border: none;
}

QPushButton:hover {
    background-color: #2980b9;
}

QPushButton#secondaryButton {
    background-color: white;
    color: #ecf0f1;
    border: 2px solid #3498db;
}

QPushButton#secondaryButton:hover {
    background-color: rgba(52, 152, 219, 0.1);
}

QComboBox {
    padding: 10px;
    border: 1px solid #34495e;
    border-radius: 6px;
    background-color: #34495e;
    color: #ecf0f1;
    font-size: 28px;
}

QTableWidget {
    background-color: #34495e;
    border: 1px solid #34495e;
    border-radius: 6px;
    gridline-color: #2c3e50;
}

QTableWidget::item {
    padding: 8px;
}

QTableWidget::item:selected {
    background-color: #3498db;
    color: #ecf0f1;
}

QHeaderView::section {
    background-color: #2980b9;
    padding: 8px;
    border: none;
    font-weight: bold;
    color: #ecf0f1;
}
'''

class AdaptiveStylesheet(QWidget):

    def __init__(self, theme="dark"):
        super().__init__()
        self.setWindowTitle("Aplicación Dinámica")
        self.theme = theme

        # Obtener resolución de pantalla y ajustar el tamaño
        width, _ = self.get_screen_resolution()
        if theme == "light":
            stylesheet = self.generate_light_stylesheet(width)
        elif theme == "gpt":
            stylesheet = self.generate_gpt_stylesheet(width)
        else:
            stylesheet = self.generate_stylesheet(width)
        self.apply_stylesheet(self, stylesheet)

    @staticmethod
    def apply_stylesheet(widget, stylesheet):
        widget.setStyleSheet(stylesheet)

    @staticmethod
    def get_screen_resolution():
        app = QApplication.instance()
        screen = app.primaryScreen()
        geometry = screen.geometry()
        return geometry.width(), geometry.height()

    @staticmethod
    def adjust_size(size, width):
        """Return the original size without scaling."""
        return size

    def generate_stylesheet(self, width):
        # Ajusta los tamaños en base a la resolución
        font_size_label = self.adjust_size(32, width)
        font_size_title = self.adjust_size(32, width)
        font_size_input = self.adjust_size(28, width)
        font_size_button = self.adjust_size(28, width)
        padding_size = self.adjust_size(12, width)

        # Definición del stylesheet con tamaños ajustables
        return f'''
        QWidget {{
            background-color: #2c3e50;
            font-family: "Segoe UI", Arial;
        }}
        
        QLabel {{
            color: #ecf0f1;
            font-size: {font_size_label}px;
        }}
        
        QLabel#titleLabel {{
            font-size: {font_size_title}px;
            font-weight: bold;
            color: #ecf0f1;
        }}
        
        QLineEdit {{
            padding: {padding_size}px;
            border: 1px solid #34495e;
            border-radius: 8px;
            background-color: #34495e;
            font-size: {font_size_input}px;
        }}
        
        QLineEdit:focus {{
            border-color: #ecf0f1;
        }}
        
        QPushButton {{
            padding: {padding_size}px;
            border-radius: 8px;
            font-size: {font_size_button}px;
            font-weight: bold;
            color: #ecf0f1;
            background-color: #3498db;
            border: none;
        }}
        
        QPushButton:hover {{
            background-color: #2980b9;
        }}
        
        QPushButton#secondaryButton {{
            background-color: white;
            color: #ecf0f1;
            border: 2px solid #3498db;
        }}
        
        QPushButton#secondaryButton:hover {{
            background-color: rgba(52, 152, 219, 0.1);
        }}
        
        QComboBox {{
            padding: {padding_size}px;
            border: 1px solid #34495e;
            border-radius: 8px;
            background-color: #34495e;
            font-size: {font_size_input}px;
        }}
        
        QTableWidget {{
            background-color: #34495e;
            border: 1px solid #34495e;
            border-radius: 8px;
            gridline-color: #2c3e50;
        }}
        
        QTableWidget::item {{
            padding: {padding_size - 4}px;
        }}
        
        QTableWidget::item:selected {{
            background-color: #3498db;
            color: #ecf0f1;
        }}
        
        QHeaderView::section {{
            background-color: #2980b9;
            padding: {padding_size - 4}px;
            border: none;
            font-weight: bold;
        }}
        '''
    def generate_light_stylesheet(self, width):
        font_size_label = self.adjust_size(32, width)
        font_size_title = self.adjust_size(32, width)
        font_size_input = self.adjust_size(28, width)
        font_size_button = self.adjust_size(28, width)
        padding_size = self.adjust_size(12, width)

        return f'''
        QWidget {{
            background-color: #f0f0f0;
            font-family: "Segoe UI", Arial;
        }}

        QLabel {{
            color: #2c3e50;
            font-size: {font_size_label}px;
        }}

        QLabel#titleLabel {{
            font-size: {font_size_title}px;
            font-weight: bold;
            color: #2c3e50;
        }}

        QLineEdit {{
            padding: {padding_size}px;
            border: 1px solid #bdc3c7;
            border-radius: 8px;
            background-color: #ffffff;
            font-size: {font_size_input}px;
            color: #2c3e50;
        }}

        QLineEdit:focus {{
            border-color: #2980b9;
        }}

        QPushButton {{
            padding: {padding_size}px;
            border-radius: 8px;
            font-size: {font_size_button}px;
            font-weight: bold;
            color: #ffffff;
            background-color: #3498db;
            border: none;
        }}

        QPushButton:hover {{
            background-color: #2980b9;
        }}

        QPushButton#secondaryButton {{
            background-color: white;
            color: #3498db;
            border: 2px solid #3498db;
        }}

        QPushButton#secondaryButton:hover {{
            background-color: rgba(52, 152, 219, 0.1);
        }}

        QComboBox {{
            padding: {padding_size}px;
            border: 1px solid #bdc3c7;
            border-radius: 8px;
            background-color: #ffffff;
            font-size: {font_size_input}px;
            color: #2c3e50;
        }}

        QTableWidget {{
            background-color: #ffffff;
            border: 1px solid #bdc3c7;
            border-radius: 8px;
            gridline-color: #bdc3c7;
        }}

        QTableWidget::item {{
            padding: {padding_size - 4}px;
        }}

        QTableWidget::item:selected {{
            background-color: #bdc3c7;
            color: #2c3e50;
        }}

        QHeaderView::section {{
            background-color: #bdc3c7;
            padding: {padding_size - 4}px;
            border: none;
            font-weight: bold;
            color: #2c3e50;
        }}

        QScrollBar:vertical {{
            background: transparent;
            width: 12px;
            margin: 0px;
            border-radius: 6px;
        }}
        QScrollBar::handle:vertical {{
            background: #c1c1c1;
            min-height: 20px;
            border-radius: 6px;
        }}
        QScrollBar::handle:vertical:hover {{
            background: #a6a6a6;
        }}
        QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
            height: 0px;
        }}
        QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {{
            background: none;
        }}
        QScrollBar:horizontal {{
            background: transparent;
            height: 12px;
            margin: 0px;
            border-radius: 6px;
        }}
        QScrollBar::handle:horizontal {{
            background: #c1c1c1;
            min-width: 20px;
            border-radius: 6px;
        }}
        QScrollBar::handle:horizontal:hover {{
            background: #a6a6a6;
        }}
        QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{
            width: 0px;
        }}
        QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal {{
            background: none;
        }}
        QCalendarWidget QToolButton {{
            color: black;
            font-weight: bold;
        }}
        QCalendarWidget QSpinBox {{
            color: black;
        }}
        QCalendarWidget QWidget {{
            background-color: #f9f9f9;
            border-radius: 4px;
        }}
        '''

    def generate_gpt_stylesheet(self, width):
        font_size_label = self.adjust_size(32, width)
        font_size_title = self.adjust_size(36, width)
        font_size_input = self.adjust_size(28, width)
        font_size_button = self.adjust_size(28, width)
        padding_size = self.adjust_size(12, width)
        accent = "#10a37f"

        return f'''
        QWidget {{
            background-color: #f7f7f8;
            font-family: "Segoe UI", Arial;
        }}

        QLabel {{
            color: #202123;
            font-size: {font_size_label}px;
        }}

        QLabel#titleLabel {{
            font-size: {font_size_title}px;
            font-weight: bold;
            color: #202123;
        }}

        QLineEdit, QTextEdit, QDateEdit, QComboBox {{
            padding: {padding_size}px;
            border: 1px solid #cdd0d5;
            border-radius: 8px;
            background-color: #ffffff;
            font-size: {font_size_input}px;
            color: #202123;
        }}

        QLineEdit:focus, QTextEdit:focus, QDateEdit:focus, QComboBox:focus {{
            border-color: {accent};
        }}

        QPushButton {{
            padding: {padding_size}px {padding_size * 2}px;
            border-radius: 20px;
            font-size: {font_size_button}px;
            font-weight: bold;
            color: #ffffff;
            background-color: {accent};
            border: none;
        }}

        QPushButton:hover {{
            background-color: #0d8e6d;
        }}

        QPushButton:pressed {{
            background-color: #0b735a;
        }}

        QPushButton#secondaryButton {{
            background-color: white;
            color: {accent};
            border: 2px solid {accent};
        }}

        QPushButton#secondaryButton:hover {{
            background-color: rgba(16, 163, 127, 0.1);
        }}

        QTableWidget {{
            background-color: #ffffff;
            border: 1px solid #cdd0d5;
            border-radius: 8px;
            gridline-color: #cdd0d5;
        }}

        QTableWidget::item {{
            padding: {padding_size - 4}px;
        }}

        QTableWidget::item:selected {{
            background-color: {accent};
            color: #ffffff;
        }}

        QHeaderView::section {{
            background-color: {accent};
            padding: {padding_size - 4}px;
            border: none;
            font-weight: bold;
            color: #ffffff;
        }}

        QScrollBar:vertical {{
            background: transparent;
            width: 12px;
            margin: 0px;
            border-radius: 6px;
        }}
        QScrollBar::handle:vertical {{
            background: #c1c1c1;
            min-height: 20px;
            border-radius: 6px;
        }}
        QScrollBar::handle:vertical:hover {{
            background: #a6a6a6;
        }}
        QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
            height: 0px;
        }}
        QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {{
            background: none;
        }}
        QScrollBar:horizontal {{
            background: transparent;
            height: 12px;
            margin: 0px;
            border-radius: 6px;
        }}
        QScrollBar::handle:horizontal {{
            background: #c1c1c1;
            min-width: 20px;
            border-radius: 6px;
        }}
        QScrollBar::handle:horizontal:hover {{
            background: #a6a6a6;
        }}
        QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{
            width: 0px;
        }}
        QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal {{
            background: none;
        }}
        '''

def months_between(birth, today=None):
    """Return whole months between two dates. Never negative."""
    if not birth:
        return 0
    if today is None:
        today = date.today()
    if birth > today:
        return 0
    months = (today.year - birth.year) * 12 + (today.month - birth.month)
    if today.day < birth.day:
        months -= 1
    return max(months, 0)

def format_patient_age(patient):
    """Return age string with months if birth date is known."""
    if not patient:
        return ""
    if patient.fecha_nacimiento:
        m = months_between(patient.fecha_nacimiento)
        y = m // 12
        s = f"{y}"
        s += f" ({m} meses)"
        return s
    return str(patient.edad or "")




class UserManagementDialog(QDialog):
    def __init__(self, users, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.user-circle')
        self.users = users
        self.setWindowTitle("Gestión de Usuarios")
        #self.setMinimumSize(600, 400)
        self.initUI()
        self.setStyleSheet("")  # Elimina estilos anteriores
        self.setStyleSheet(STYLE_SHEET)     
        self.setStyleSheet(STYLE_SHEET)  # Aplica el estilo adaptativo
              # Establecer las banderas de la ventana para mostrar los botones
        self.setWindowFlags(self.windowFlags() | Qt.Window | Qt.WindowTitleHint | Qt.WindowSystemMenuHint)

        # Mostrar la ventana maximizada
        self.showMaximized()
    

    def initUI(self):
        layout = QVBoxLayout()
        
        # Título
        title = QLabel("Gestión de Usuarios")
        title.setObjectName("titleLabel")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)
        
        # Tabla de usuarios
        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["Usuario", "Pregunta de Recuperación", "PIN", "Rol", "Acciones"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        StyleHelper.style_table(self.table)
        layout.addWidget(self.table)
        
        self.update_table()
        
        # Botones
        button_layout = QHBoxLayout()
        close_btn = QPushButton("Cerrar")
        close_btn.setObjectName("secondaryButton")
        close_btn.setIcon(self.style().standardIcon(QStyle.SP_DialogCloseButton))
        StyleHelper.style_button(close_btn)
        close_btn.clicked.connect(self.close)
        button_layout.addWidget(close_btn)
        layout.addLayout(button_layout)
        
        self.setLayout(layout)

    def update_table(self):
        self.table.setRowCount(0)
        for row, (username, data) in enumerate(self.users.items()):
            self.table.insertRow(row)
            self.table.setItem(row, 0, QTableWidgetItem(username))
            self.table.setItem(row, 1, QTableWidgetItem(data['recovery_question']))
            self.table.setItem(row, 2, QTableWidgetItem("****"))
            self.table.setItem(row, 3, QTableWidgetItem(data.get('role', 'doctor')))

            delete_btn = QPushButton("Eliminar")
            delete_btn.setIcon(self.style().standardIcon(QStyle.SP_TrashIcon))
            StyleHelper.style_button(delete_btn)
            delete_btn.clicked.connect(lambda checked, u=username: self.delete_user(u))
            self.table.setCellWidget(row, 4, delete_btn)

    def delete_user(self, username):
        reply = QMessageBox.question(self, 'Confirmar', 
                                   f'¿Está seguro de eliminar al usuario {username}?',
                                   QMessageBox.Yes | QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            del self.users[username]
            self.parent().save_users()
            self.update_table()




class ModernAuthWindow(QWidget):
    login_successful = pyqtSignal(str)
    
    def __init__(self):
        super().__init__()
        self.users_file = 'users.json'
        self.failed_attempts = 0
        self.lock_until = None
        self.recovery_questions = [
            "¿Cuál es el nombre de tu primera mascota?",
            "¿En qué ciudad naciste?",
            "¿Cuál es el nombre de tu mejor amigo de la infancia?",
            "¿Cuál fue tu primer coche?",
            "¿Cuál es el segundo nombre de tu madre?"
        ]
        self.load_users()
        self.initUI()
        
       # Crear instancia de AdaptiveStylesheet y aplicar el estilo por defecto claro
        adaptive_stylesheet = AdaptiveStylesheet(theme="light")
        width, _ = adaptive_stylesheet.get_screen_resolution()
        stylesheet = adaptive_stylesheet.generate_light_stylesheet(width)
        adaptive_stylesheet.apply_stylesheet(self, stylesheet)
        
        
        StyleHelper.set_window_icon(self, 'fa5s.user-circle')
        
       # Modificamos esta parte para manejar el caso sin usuarios
        if not self.users:
            self.create_ready_file1()


            self.stacked_widget.setCurrentIndex(1)  # Mostrar registro
            QMessageBox.information(self, 'Bienvenido', 
                'No hay usuarios registrados. Por favor, cree una cuenta de administrador.')
        else:
            self.stacked_widget.setCurrentIndex(0)  # Mostrar login
        
        # Mostrar la ventana
        self.show()

    def create_ready_file1(self):
        # Ruta del directorio donde está el ejecutable
        exe_dir = os.path.dirname(sys.argv[0])
        ready_file_path = os.path.join(exe_dir, "app_ready.txt")
        with open(ready_file_path, "w") as f:
            f.write("App is ready")
   
    def auto_login_if_not_required(self):
    
        return not bool(self.users)
    def initUI(self):
            # Obtener resolución de pantalla y establecer tamaño mínimo
        screen = QApplication.primaryScreen()
        screen_geometry = screen.availableGeometry()  # Definimos screen_geometry aquí
        screen_width, screen_height = screen_geometry.width(), screen_geometry.height()
        self.setWindowTitle('Sistema de Autenticación')
        # Obtener resolución de pantalla y establecer tamaño mínimo
        width, _ = AdaptiveStylesheet.get_screen_resolution()
        if width >= 3800:  # Resolución cercana a 4K
             self.setMinimumSize(800, 800)
        elif 1367 <= width <= 3799:  # Resolución específica de 1366x768
             self.setMinimumSize(800, 800)
        elif 1360 <= width <= 1366:  # Resolución específica de 1366x768  
             self.resize(300, 400)
        elif 700 <= width <= 1300: 
             self.resize(320, 400)
        elif width == 1280:  # Cualquier resolución de ancho 1280 (1280x600 a 1280x1024)             
             self.resize(350, 400)
        else:
              self.resize(320, 400)
        # Centrar la ventana en la pantalla
        window_geometry = self.frameGeometry()
        window_geometry.moveCenter(screen_geometry.center())
        self.move(window_geometry.topLeft())        
        self.stacked_widget = QStackedWidget()
        
        self.login_widget = self.create_login_widget()
        self.register_widget = self.create_register_widget()
        self.recovery_widget = self.create_recovery_widget()
        
        self.stacked_widget.addWidget(self.login_widget)
        self.stacked_widget.addWidget(self.register_widget)
        self.stacked_widget.addWidget(self.recovery_widget)
        
        layout = QVBoxLayout()
        layout.addWidget(self.stacked_widget)
        self.setLayout(layout)
        
    def load_users(self):
        if os.path.exists(self.users_file):
            try:
                with open(self.users_file, 'r') as f:
                    self.users = json.load(f)
            except (json.JSONDecodeError, FileNotFoundError):
                self.users = {}
        else:
            self.users = {}  

    def create_login_widget(self):
            widget = QWidget()
            layout = QVBoxLayout()
            layout.setSpacing(10)
            layout.setAlignment(Qt.AlignCenter)
    
            # Obtener la resolución de la pantalla
            width, height = QApplication.primaryScreen().size().width(), QApplication.primaryScreen().size().height()
            
         # Configuración del logo, títulos, inputs y botones según la resolución
            if width >= 3800:  # Resolución cercana a 4K
                logo_size = 100
                font_size_title1 = 42
                font_size_title = 40
                input_height = 60
                input_width = 500  # Definimos el ancho del campo de entrada
                button_height = 60
                button_font_size = 28
                input_font_size = 28
            elif 1367 <= width <= 3799:  # Resolución específica de 1366x768
                logo_size = 100
                font_size_title1 = 42
                font_size_title = 40
                input_height = 60
                input_width = 500  # Definimos el ancho del campo de entrada
                button_height = 60
                button_font_size = 28
                input_font_size = 28
            elif 1360 <= width <= 1366:  # Resolución específica de 1366x768
                logo_size = 50
                font_size_title1 = 16
                font_size_title = 16
                input_height = 30
                input_width = 340
                input_font_size = 10
                button_height = 50
                button_font_size = 10
            elif 700 <= width <= 1300:
                logo_size = 50
                font_size_title1 = 14
                font_size_title = 14
                input_height = 40
                input_width = 300
                button_height = 45
                button_font_size = 9
                input_font_size = 12
                self.setMinimumSize(150, 100)
            elif width == 1280:  # Cualquier resolución de ancho 1280 (1280x600 a 1280x1024) 
                logo_size = 50
                font_size_title1 = 14
                font_size_title = 14
                input_height = 35
                input_font_size = 14  
                input_width = 300
                button_height = 45
                button_font_size = 10
               
            else:  # Otras resoluciones más pequeñas
                logo_size = 50
                font_size_title1 = 15
                font_size_title = 14
                input_height = 40
                input_width = 200
                button_height = 40
                button_font_size = 10
                input_font_size = 14  
               
            
            # Agregar el logo al login
            logo_label = QLabel()
            logo_icon = QIcon.fromTheme('medical-symbol')
            if logo_icon.isNull():
                logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'icon2.ico')
                pixmap = QPixmap(logo_path)
            else:
                pixmap = logo_icon.pixmap(logo_size, logo_size)
            if not pixmap.isNull():
                logo_label.setPixmap(
                    pixmap.scaled(
                        logo_size,
                        logo_size,
                        Qt.KeepAspectRatio,
                        Qt.SmoothTransformation,
                    )
                )
            logo_label.setAlignment(Qt.AlignCenter)
            layout.addWidget(logo_label)
            
            # Título principal (nombre de la app)
            title1 = QLabel('HealthIA')
            title1.setAlignment(Qt.AlignCenter)
            title1.setStyleSheet(f"font-size: {font_size_title1}px; font-weight: bold;")

            # Título secundario (descripción de la acción actual)
            title = QLabel('Iniciar Sesión')
            title.setAlignment(Qt.AlignCenter)
            title.setStyleSheet(f"font-size: {font_size_title}px;")

            if qta:
                header_layout = QHBoxLayout()
                header_layout.setAlignment(Qt.AlignCenter)
                icon_label = QLabel()
                icon_label.setPixmap(qta.icon('fa5s.stethoscope', color='#17A589').pixmap(int(font_size_title1*1.5), int(font_size_title1*1.5)))
                header_layout.addWidget(icon_label)
                header_layout.addWidget(title1)
                header_widget = QWidget()
                header_widget.setLayout(header_layout)
                layout.addWidget(header_widget)
            else:
                layout.addWidget(title1)

            layout.addWidget(title)
    
            # Campos de entrada ajustados
            self.login_username = QLineEdit()
            self.login_username.setPlaceholderText('Usuario')
            self.login_username.setFixedHeight(input_height)
            self.login_username.setFixedWidth(input_width)
            if qta:
                self.login_username.addAction(qta.icon('fa5s.user', color='gray'), QLineEdit.LeadingPosition)
            StyleHelper.style_input(self.login_username)
            layout.addWidget(self.login_username)
            
            self.login_password = QLineEdit()
            self.login_password.setPlaceholderText('Contraseña')
            self.login_password.setEchoMode(QLineEdit.Password)
            self.login_password.setFixedHeight(input_height)
            self.login_password.setFixedWidth(input_width)
            if qta:
                self.login_password.addAction(qta.icon('fa5s.lock', color='gray'), QLineEdit.LeadingPosition)
            StyleHelper.style_input(self.login_password)
            layout.addWidget(self.login_password)

            self.remember_check = QCheckBox('Recordar usuario y contraseña')
            self.remember_check.setFixedWidth(input_width)
            if qta:
                # Font Awesome 5 icons for checkbox states
                self.remember_check.setIcon(qta.icon('fa5s.square', color='gray'))
                def _update_icon(state):
                    icon_name = 'fa5s.check-square' if state == Qt.Checked else 'fa5s.square'
                    self.remember_check.setIcon(qta.icon(icon_name, color='gray'))
                self.remember_check.stateChanged.connect(_update_icon)
            layout.addWidget(self.remember_check)
            if os.path.exists(REMEMBER_FILE):
                try:
                    with open(REMEMBER_FILE, 'r') as f:
                        data = json.load(f)
                        self.login_username.setText(data.get('username', ''))
                        self.login_password.setText(data.get('password', ''))
                        self.remember_check.setChecked(True)
                except Exception:
                    pass
            # Botón de inicio de sesión ajustado
            login_btn = QPushButton('Iniciar Sesión')
            StyleHelper.style_button(login_btn)
            login_btn.clicked.connect(self.login)
            layout.addWidget(login_btn)
            
            # Conectar el campo de contraseña al atajo de Enter
            self.login_password.returnPressed.connect(self.login)
    
            # Botón de registro
            self.register_btn = QPushButton('Registrarse')
            StyleHelper.style_button(self.register_btn)
            self.register_btn.clicked.connect(lambda: self.stacked_widget.setCurrentIndex(1))
            self.register_btn.hide()  # Ocultamos el botón por defecto
            layout.addWidget(self.register_btn)
    
            # Botón de recuperación de contraseña
            forgot_btn = QPushButton('¿Olvidaste tu contraseña?')
            StyleHelper.style_button(forgot_btn)
            forgot_btn.clicked.connect(lambda: self.stacked_widget.setCurrentIndex(2))
            layout.addWidget(forgot_btn)
            
            widget.setLayout(layout)
            return widget
        
    def keyPressEvent(self, event):
       
        if event.key() == Qt.Key_Return or event.key() == Qt.Key_Enter:
            self.login()  # Llamamos a la función login si se presiona Enter
            
    def create_register_widget(self):
        # Obtener la resolución de la pantalla
        width, height = QApplication.primaryScreen().size().width(), QApplication.primaryScreen().size().height()
        self.create_ready_file1()    
        if width >= 3800:  # Resolución cercana a 4K
            logo_size = 100
            font_size_title1 = 50
            font_size_title = 50
            input_height = 80
            button_height = 80
            button_font_size = 28
            input_font_size = 28
            input_font_size_desp = 24  # Aumentado para mayor legibilidad
        elif 1360 <= width <= 1366:  # Resolución específica de 1366x768
            logo_size = 50
            font_size_title1 = 14
            font_size_title = 14
            input_height = 40
            input_font_size = 12
            button_height = 40
            button_font_size = 14
            
            input_font_size_desp = 14  # Ajustado para uniformidad
        elif 700 <= width <= 1300:
            logo_size = 50
            font_size_title1 = 14
            font_size_title = 14
            input_height = 60
            button_height = 35
            button_font_size = 10
            input_font_size = 14
            input_font_size_desp = 12  # Ajustado para resoluciones menores
        elif width == 1280:
            logo_size = 50
            font_size_title1 = 14
            font_size_title = 14
            input_height = 50
            button_height = 45
            button_font_size = 10
            input_font_size = 14
            input_font_size_desp = 12
        else:
            logo_size = 40
            font_size_title1 = 15
            font_size_title = 14
            input_height = 60
            button_height = 40
            button_font_size = 12
            input_font_size = 14
            input_font_size_desp = 12  # Ajustado para uniformidad
        # Crear un widget contenedor
        widget = QWidget()
        layout = QVBoxLayout()
        layout.setSpacing(20)
        layout.setAlignment(Qt.AlignCenter)
        # Ruta del directorio donde está el ejecutable
        exe_dir = os.path.dirname(sys.argv[0])
        ready_file_path = os.path.join(exe_dir, "app_ready.txt")
        with open(ready_file_path, "w") as f:
            f.write("App is ready")            
        # Crear un widget contenedor con ancho máximo
        content_widget = QWidget()
        content_widget.setMaximumWidth(400)
        content_layout = QVBoxLayout()
        content_widget.setLayout(content_layout)
        
        # Título
        title = QLabel('Registro de Usuario')
        title.setObjectName("titleLabel")
        title.setAlignment(Qt.AlignCenter)
        content_layout.addWidget(title)
        
        # Campos de entrada
        self.reg_username = QLineEdit()
        self.reg_username.setPlaceholderText('Usuario')
        self.reg_username.setFixedHeight(input_height)
        self.reg_username.setStyleSheet(f"font-size: {input_font_size}px; padding: 5px;")  # Tamaño de fuente del input
        content_layout.addWidget(self.reg_username)
        
        self.reg_password = QLineEdit()
        self.reg_password.setPlaceholderText('Contraseña')
        self.reg_password.setEchoMode(QLineEdit.Password)
        self.reg_password.setFixedHeight(input_height)
        self.reg_password.setStyleSheet(f"font-size: {input_font_size}px; padding: 5px;")  # Tamaño de fuente del input
        content_layout.addWidget(self.reg_password)
        
        self.reg_confirm_password = QLineEdit()
        self.reg_confirm_password.setPlaceholderText('Confirmar Contraseña')
        self.reg_confirm_password.setEchoMode(QLineEdit.Password)
        self.reg_confirm_password.setFixedHeight(input_height)
        self.reg_confirm_password.setStyleSheet(f"font-size: {input_font_size}px; padding: 5px;")  # Tamaño de fuente del input
        content_layout.addWidget(self.reg_confirm_password)
        
        self.recovery_question = QComboBox()
        self.recovery_question.addItems(self.recovery_questions)
        self.recovery_question.setStyleSheet(f"font-size: {input_font_size_desp}px; padding: 5px;")  # Tamaño de fuente del input
        content_layout.addWidget(self.recovery_question)
        
        self.recovery_answer = QLineEdit()
        self.recovery_answer.setPlaceholderText('Respuesta de recuperación')
        self.recovery_answer.setFixedHeight(input_height)
        self.recovery_answer.setStyleSheet(f"font-size: {input_font_size_desp}px; padding: 5px;")  # Tamaño de fuente del input
        content_layout.addWidget(self.recovery_answer)
    
        self.pin_code = QLineEdit()
        self.pin_code.setPlaceholderText('PIN de 4 dígitos')
        self.pin_code.setEchoMode(QLineEdit.Password)
        self.pin_code.setFixedHeight(input_height)
        self.pin_code.setStyleSheet(f"font-size: {input_font_size}px; padding: 5px;")  # Tamaño de fuente del input
        content_layout.addWidget(self.pin_code)

        self.role_combo = QComboBox()
        self.role_combo.addItems(["doctor", "recepcionista", "admin"])
        self.role_combo.setFixedHeight(input_height)
        self.role_combo.setStyleSheet(f"font-size: {input_font_size}px; padding: 5px;")
        content_layout.addWidget(self.role_combo)
        
        # Botones
        register_btn = QPushButton('Registrarse')
        register_btn.setFixedHeight(button_height)
        register_btn.setIcon(self.style().standardIcon(QStyle.SP_FileDialogNewFolder))
        StyleHelper.style_button(register_btn)
        register_btn.clicked.connect(self.register)
        content_layout.addWidget(register_btn)
        
        if self.users:  # Solo mostrar si hay usuarios
            manage_users_btn = QPushButton('Gestionar Usuarios')
            manage_users_btn.setObjectName("secondaryButton")
            manage_users_btn.setFixedHeight(button_height)
            manage_users_btn.setIcon(self.style().standardIcon(QStyle.SP_FileDialogContentsView))
            StyleHelper.style_button(manage_users_btn)
            manage_users_btn.clicked.connect(self.show_user_management)
            content_layout.addWidget(manage_users_btn)
        
        back_btn = QPushButton('Volver al Login')
        back_btn.setObjectName("secondaryButton")
        back_btn.setFixedHeight(button_height)
        back_btn.setIcon(self.style().standardIcon(QStyle.SP_ArrowBack))
        StyleHelper.style_button(back_btn)
        back_btn.clicked.connect(lambda: self.stacked_widget.setCurrentIndex(0))
        content_layout.addWidget(back_btn)
        
        # Agregar el widget contenedor al layout principal
        layout.addWidget(content_widget)
        widget.setLayout(layout)
            # Ajustar el tamaño del widget a su contenido
        widget.adjustSize()
        return widget

    def show_user_management(self):
        dialog = UserManagementDialog(self.users, self)
        dialog.exec_()
    def create_recovery_widget(self):
        widget = QWidget()
        layout = QVBoxLayout()
        layout.setSpacing(10)
        layout.setAlignment(Qt.AlignCenter)
        
        title = QLabel('Recuperar Contraseña')
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)
        
        self.recovery_username = QLineEdit()
        self.recovery_username.setPlaceholderText('Usuario')
        layout.addWidget(self.recovery_username)
        
        self.recovery_pin = QLineEdit()
        self.recovery_pin.setPlaceholderText('Ingrese el PIN de 4 dígitos')
        self.recovery_pin.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.recovery_pin)
        
        self.recovery_question_label = QLabel('')
        layout.addWidget(self.recovery_question_label)
        self.recovery_question_label.hide()
        
        self.recovery_answer_input = QLineEdit()
        self.recovery_answer_input.setPlaceholderText('Respuesta')
        layout.addWidget(self.recovery_answer_input)
        self.recovery_answer_input.hide()
        
        self.new_password = QLineEdit()
        self.new_password.setPlaceholderText('Nueva Contraseña')
        self.new_password.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.new_password)
        self.new_password.hide()
        
        self.confirm_new_password = QLineEdit()
        self.confirm_new_password.setPlaceholderText('Confirmar Nueva Contraseña')
        self.confirm_new_password.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.confirm_new_password)
        self.confirm_new_password.hide()
        
        self.verify_username_btn = QPushButton('Verificar Usuario')
        self.verify_username_btn.setIcon(self.style().standardIcon(QStyle.SP_DialogApplyButton))
        StyleHelper.style_button(self.verify_username_btn)
        self.verify_username_btn.clicked.connect(self.verify_username)
        layout.addWidget(self.verify_username_btn)
        
        self.verify_answer_btn = QPushButton('Verificar Respuesta')
        self.verify_answer_btn.setIcon(self.style().standardIcon(QStyle.SP_DialogApplyButton))
        StyleHelper.style_button(self.verify_answer_btn)
        self.verify_answer_btn.clicked.connect(self.verify_answer)
        layout.addWidget(self.verify_answer_btn)
        self.verify_answer_btn.hide()
        
        self.reset_password_btn = QPushButton('Restablecer Contraseña')
        self.reset_password_btn.setIcon(self.style().standardIcon(QStyle.SP_DialogResetButton))
        StyleHelper.style_button(self.reset_password_btn)
        self.reset_password_btn.clicked.connect(self.reset_password)
        layout.addWidget(self.reset_password_btn)
        self.reset_password_btn.hide()
        
        back_btn = QPushButton('Volver al Login')
        back_btn.setIcon(self.style().standardIcon(QStyle.SP_ArrowBack))
        StyleHelper.style_button(back_btn)
        back_btn.clicked.connect(lambda: self.stacked_widget.setCurrentIndex(0))
        layout.addWidget(back_btn)
        
        widget.setLayout(layout)
        return widget

    def load_users(self):
        if os.path.exists(self.users_file):
            with open(self.users_file, 'r') as f:
                self.users = json.load(f)
        else:
            self.users = {}

    def save_users(self):
        with open(self.users_file, 'w') as f:
            json.dump(self.users, f)

    def login(self):
        username = self.login_username.text()
        password = self.login_password.text()

        if self.lock_until and time.time() < self.lock_until:
            QMessageBox.warning(self, 'Bloqueado', 'Demasiados intentos. Intente más tarde.')
            return

        if username in self.users:
            if self.users[username]['password'] == hashlib.sha256(password.encode()).hexdigest():
                self.failed_attempts = 0
                if self.remember_check.isChecked():
                    try:
                        with open(REMEMBER_FILE, 'w') as f:
                            json.dump({'username': username, 'password': password}, f)
                    except Exception:
                        pass
                else:
                    if os.path.exists(REMEMBER_FILE):
                        try:
                            os.remove(REMEMBER_FILE)
                        except Exception:
                            pass
                self.login_successful.emit(username)

                self.close()
            else:
                self.failed_attempts += 1
                if self.failed_attempts >= 3:
                    self.lock_until = time.time() + 10
                    QMessageBox.warning(self, 'Bloqueado', 'Demasiados intentos. Espere 10 segundos.')
                else:
                    QMessageBox.warning(self, 'Error', 'Contraseña incorrecta')
        else:
            QMessageBox.warning(self, 'Error', 'Usuario no encontrado')

    def register(self):
        username = self.reg_username.text()
        password = self.reg_password.text()
        confirm_password = self.reg_confirm_password.text()
        recovery_q = self.recovery_question.currentText()
        recovery_a = self.recovery_answer.text()
        pin_code = self.pin_code.text()
        role = self.role_combo.currentText()
        
        if not username or not password or not pin_code.isdigit() or len(pin_code) != 4:
            QMessageBox.warning(self, 'Error', 'Por favor complete todos los campos correctamente')
            return
            
        if password != confirm_password:
            QMessageBox.warning(self, 'Error', 'Las contraseñas no coinciden')
            return
            
        if username in self.users:
            QMessageBox.warning(self, 'Error', 'El usuario ya existe')
            return
            
        self.users[username] = {
            'password': hashlib.sha256(password.encode()).hexdigest(),
            'recovery_question': recovery_q,
            'recovery_answer': hashlib.sha256(recovery_a.lower().encode()).hexdigest(),
            'pin_code': hashlib.sha256(pin_code.encode()).hexdigest(),
            'role': role
        }
        
        self.save_users()
        QMessageBox.information(self, 'Éxito', 'Registro exitoso!')
        self.stacked_widget.setCurrentIndex(0)  # Ir
    def verify_username(self):
        username = self.recovery_username.text()
        if username in self.users:
            self.recovery_pin.show()
            self.verify_username_btn.hide()
            self.verify_answer_btn.show()
        else:
            QMessageBox.warning(self, 'Error', 'Usuario no encontrado')

    def verify_answer(self):
        username = self.recovery_username.text()
        pin_code = self.recovery_pin.text()
        
        if self.users[username]['pin_code'] == hashlib.sha256(pin_code.encode()).hexdigest():
            self.new_password.show()
            self.confirm_new_password.show()
            self.reset_password_btn.show()
            self.recovery_pin.hide()
            self.verify_answer_btn.hide()
        else:
            QMessageBox.warning(self, 'Error', 'PIN incorrecto')
    
    def reset_password(self):
        username = self.recovery_username.text()
        new_password = self.new_password.text()
        confirm_password = self.confirm_new_password.text()
        
        if new_password != confirm_password:
            QMessageBox.warning(self, 'Error', 'Las contraseñas no coinciden')
            return
            
        self.users[username]['password'] = hashlib.sha256(new_password.encode()).hexdigest()
        self.save_users()
        QMessageBox.information(self, 'Éxito', 'Contraseña actualizada')
        self.stacked_widget.setCurrentIndex(0)








# Define una contraseña encriptada. Cambia 'tu_contraseña_secreta' por tu contraseña.
HASHED_PASSWORD = hashlib.sha256("admin2041".encode()).hexdigest()
CONFIG_FILE = "config.json"




# Ruta al archivo de configuración y clave de cifrado
CONFIG_FILE = "config_gpt.json"
REMEMBER_FILE = "remember.json"

def read_api_key():
    """Return the decrypted OpenAI API key from config or environment."""
    if os.path.exists(CONFIG_FILE):
        os.chmod(CONFIG_FILE, 0o600)
        try:
            with open(CONFIG_FILE, "r") as f:
                cfg = json.load(f)
            encoded = cfg.get("api_key", "")
            if encoded:
                return base64.b64decode(encoded).decode()
        except Exception:
            return os.getenv("OPENAI_API_KEY", "")
    return os.getenv("OPENAI_API_KEY", "")

def apply_openai_key(cfg):
    """Set openai.api_key from encoded config and return the model name."""
    encoded = cfg.get('api_key', '')
    try:
        api_key = base64.b64decode(encoded).decode()
    except Exception:
        api_key = ''
    openai.api_key = api_key or os.getenv('OPENAI_API_KEY', '')
    return cfg.get('model', '')

def ensure_openai_key():
    """Load the OpenAI key and return True if available."""
    key = read_api_key()
    openai.api_key = key
    return bool(key)

# Prompt for the pre-generation AI interaction dialog
IA_INTERACTION_PROMPT = (
    "Eres un m\u00e9dico especialista que act\u00faa como IA cl\u00ednica. "
    "Lee la historia de enfermedad actual del paciente. "
    "Si crees que se necesita m\u00e1s informaci\u00f3n para desarrollar una historia cl\u00ednica completa, "
    "formula una o m\u00e1s preguntas al m\u00e9dico tratante para que las responda. "
    "Cuando creas que ya tienes la informaci\u00f3n necesaria, responde con el mensaje: \u2705 Listo para generar historia cl\u00ednica."
)

# Prompt for AI-generated clinical study
CLINICAL_STUDY_PROMPT = (
    "Eres un médico experto. A continuación se presenta toda la información clínica de un paciente. "
    "Genera un estudio clínico completo con:\n"
    "- Antecedentes relevantes\n"
    "- Signos vitales e interpretación\n"
    "- Hallazgos relevantes de laboratorio\n"
    "- Medicación actual y su posible impacto\n"
    "- Síntesis clínica y conclusión"
)

# Cargar tablas de crecimiento OMS si están disponibles
GROWTH_TABLES = GrowthTables()




class ConfigGPT(QWidget):
    config_updated = pyqtSignal(str)

    def __init__(self):
        super().__init__()
        self.main_window = None
        self.initUI()
        self.load_config()
        
        StyleHelper.set_window_icon(self, 'fa5s.cogs')

    def authenticate_user(self):
        self.setStyleSheet(STYLE_SHEET)

        StyleHelper.set_window_icon(self, 'fa5s.cogs')

        password, ok = QInputDialog.getText(
            None, 'Autenticación Requerida', 'Introduce la contraseña de administrador:', QLineEdit.Password
        )
        
        # Asegúrate de que HASHED_PASSWORD esté definido en tu código
        return ok and hashlib.sha256(password.encode()).hexdigest() == HASHED_PASSWORD

    def show(self):
        # Sobrescribimos el método show para realizar la autenticación
        if self.authenticate_user():
            super().show()
        else:
            QMessageBox.critical(self, "Acceso Denegado", "Contraseña incorrecta. No tienes permiso para modificar la configuración.")
    
    def initUI(self):
        self.setWindowTitle('Configuración de API y Modelo GPT')

        # Layout principal
        layout = QVBoxLayout()

        # Etiqueta y campo de entrada para la API key
        self.api_key_label = QLabel('API Key:')
        layout.addWidget(self.api_key_label)
        self.api_key_entry = QLineEdit(self)
        layout.addWidget(self.api_key_entry)
        StyleHelper.style_input(self.api_key_entry)

        # Etiqueta y desplegable para seleccionar el modelo de GPT
        self.model_label = QLabel('Selecciona el modelo GPT o ingrésalo manualmente:')
        layout.addWidget(self.model_label)

        # Contenedor para desplegable y campo de texto de modelo
        model_layout = QHBoxLayout()
        self.model_combo = QComboBox(self)
        self.model_combo.addItems(["gpt-4o-mini", "gpt-4o", "Otro (Escribir manualmente)"])
        self.model_combo.currentIndexChanged.connect(self.model_selection_changed)
        model_layout.addWidget(self.model_combo)
        self.model_entry = QLineEdit(self)
        self.model_entry.setPlaceholderText("Escriba el modelo manualmente...")
        self.model_entry.setEnabled(False)  # Deshabilitado por defecto
        model_layout.addWidget(self.model_entry)
        layout.addLayout(model_layout)
        StyleHelper.style_input(self.model_entry)

        # Botón para guardar la configuración
        self.save_button = QPushButton('Guardar Configuración', self)
        self.save_button.clicked.connect(self.save_config)
        layout.addWidget(self.save_button)
        StyleHelper.style_button(self.save_button)

        # Set the layout
        self.setLayout(layout)

    def model_selection_changed(self):
        if self.model_combo.currentText() == "Otro (Escribir manualmente)":
            self.model_entry.setEnabled(True)
        else:
            self.model_entry.setEnabled(False)
            self.model_entry.clear()

    def save_config(self):
        try:
            api_key = self.api_key_entry.text()
            encoded_key = base64.b64encode(api_key.encode()).decode()
            
            if self.model_combo.currentText() == "Otro (Escribir manualmente)":
                model = self.model_entry.text()
            else:
                model = self.model_combo.currentText()

            config_data = {
                "api_key": encoded_key,
                "model": model,
            }

            if os.path.exists(CONFIG_FILE):
                try:
                    with open(CONFIG_FILE, "r") as config_file:
                        existing = json.load(config_file)
                        if "theme" in existing:
                            config_data["theme"] = existing["theme"]
                except Exception:
                    pass

            with open(CONFIG_FILE, "w") as config_file:
                json.dump(config_data, config_file)
            os.chmod(CONFIG_FILE, 0o600)

            # Emitir la señal con el modelo actualizado
            self.config_updated.emit(model)
            # Actualizar el título de la ventana
            self.update_window_title2(model)
            
            QMessageBox.information(
                self,
                "Confirmación",
                f"Configuración guardada:\nAPI Key -> {'*' * len(api_key)}\nModelo -> {model}"
            )

        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Error al guardar la configuración: {str(e)}"
            )

    def set_main_window(self, main_window):
        self.main_window = main_window
        
    def update_window_title2(self, model):
        self.model = model
        self.setWindowTitle(f"Sistema de Historias Clínicas {self.model}")

    def load_config(self):
        if os.path.exists(CONFIG_FILE):
            os.chmod(CONFIG_FILE, 0o600)
            with open(CONFIG_FILE, "r") as config_file:
                config_data = json.load(config_file)
                encoded = config_data.get("api_key", "")
                try:
                    api_key = base64.b64decode(encoded).decode()
                except Exception:
                    api_key = ""
                model = config_data.get("model", "")

                self.api_key_entry.setText(api_key)
                if model in ["gpt-4o-mini", "gpt-4o"]:
                    self.model_combo.setCurrentText(model)
                else:
                    self.model_combo.setCurrentText("Otro (Escribir manualmente)")
                    self.model_entry.setText(model)
                    self.model_entry.setEnabled(True)
                
                self.model = model
                self.setWindowTitle(f"Sistema de Historias Clínicas {self.model}")      
        else:
            nconfirmation_msg = QMessageBox()
            nconfirmation_msg.setWindowTitle("No Confirmación")
            nconfirmation_msg.setText("No se encontró un archivo de configuración, por favor ingresa los datos.")
            nconfirmation_msg.setIcon(QMessageBox.Information)
            nconfirmation_msg.exec_()



# Database configuration
Base = declarative_base()
database_file = 'historias_clinic158.db'
engine = create_engine(f'sqlite:///{database_file}')
if os.path.exists(database_file):
    os.chmod(database_file, 0o600)
Session = sessionmaker(bind=engine, expire_on_commit=False)

def ensure_db_schema(db_engine):
    """Ensure expected tables and columns exist."""
    inspector = inspect(db_engine)
    expected_tables = {
        'pacientes',
        'historias_clinicas',
        'signos_vitales',
        'recordatorios',
        'lab_results',
        'viviendas',
        'servicios',
        'tarifas_ars',
        'facturas',
        'factura_detalle',
        'pagos',
    }

    existing_tables = set(inspector.get_table_names())
    for table in expected_tables:
        if table not in existing_tables:
            Base.metadata.create_all(db_engine, tables=[Base.metadata.tables[table]])

    # refresh inspector after creating tables
    inspector = inspect(db_engine)
    for table in expected_tables:
        existing_cols = {c['name'] for c in inspector.get_columns(table)}
        expected_cols = {col.key for col in Base.metadata.tables[table].columns}
        missing_cols = expected_cols - existing_cols
        if missing_cols:
            with db_engine.begin() as conn:
                for column in missing_cols:
                    column_obj = Base.metadata.tables[table].columns[column]
                    col_type = column_obj.type.compile(db_engine.dialect)
                    conn.execute(text(f'ALTER TABLE {table} ADD COLUMN {column} {col_type}'))

def ensure_default_services():
    """Insert predefined services if the table is empty."""
    session = Session()
    if session.query(Servicio).count() == 0:
        defaults = [
            ("Consulta general", "Consulta", 500),
            ("Consulta especializada", "Consulta", 800),
            ("EKG", "Procedimiento", 700),
            ("Laboratorio básico", "Laboratorio", 900),
            ("Sonografía", "Procedimiento", 1000),
            ("Curación menor", "Procedimiento", 300),
        ]
        for name, tipo, price in defaults:
            session.add(Servicio(nombre=name, tipo=tipo, precio_base=price))
        session.commit()
    session.close()

def generate_random_code(length=3):
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=length))
    
from sqlalchemy import Column, Integer, String, DateTime, Text, ForeignKey
from sqlalchemy.orm import relationship
class HistoriaClinica(Base):
    __tablename__ = 'historias_clinicas'

    id = Column(Integer, primary_key=True, autoincrement=True)
    paciente_id = Column(Integer, ForeignKey('pacientes.id'))
    fecha = Column(DateTime)
    antecedentes_personales = Column(Text)
    antecedentes_heredofamiliares = Column(Text)
    historia_enfermedad = Column(Text)
    sugerencias_ia = Column(Text)
    interaccion_ia = Column(Text)
    resumen = Column(Text)
    codigo = Column(String(3))  # Añadir esta línea para la columna 'codigo'
    programa_salud = Column(Text)
    sonografias = Column(Text)
    denver_resultado = Column(String)
    denver_diagnostico = Column(String)
    denver_descripcion = Column(Text)
    

    paciente = relationship("Paciente", back_populates="historias")
class Diagnostico(Base):
    __tablename__ = 'diagnosticos'
    id           = Column(Integer, primary_key=True, autoincrement=True)
    historia_id  = Column(Integer, ForeignKey('historias_clinicas.id'))
    codigo_cie10 = Column(String(10))
    descripcion  = Column(Text)
    fecha        = Column(DateTime, default=lambda: datetime.now())

    historia = relationship("HistoriaClinica", backref="diagnosticos")


class RegistroConsulta(Base):
    __tablename__ = 'registro_consultas'

    id                = Column(Integer, primary_key=True, autoincrement=True)
    fecha             = Column(DateTime, default=lambda: datetime.now())

    historia_id       = Column(Integer, ForeignKey('historias_clinicas.id'))
    paciente_id       = Column(Integer, ForeignKey('pacientes.id'))

    ficha_familiar    = Column(String(20))
    documento_id      = Column(String(20))
    edad              = Column(Integer)
    sexo              = Column(String(1))
    nacionalidad      = Column(String(15))
    direccion         = Column(Text)
    no_seguridad_soc  = Column(String(20))
    tipo_sangre       = Column("regimen", String(10))
    ars               = Column(String(30))
    lugar_consulta    = Column(String(20))
    frecuencia        = Column(String(20))
    servicio          = Column(String(40))
    diagnostico_txt   = Column(Text)
    codigo_cie10      = Column(String(10))
    tipo_atencion     = Column(String(15))
    sipna_id          = Column(String(15))
    referido_a        = Column(String(40))

    historia  = relationship("HistoriaClinica")
    paciente  = relationship("Paciente")

class Paciente(Base):
    __tablename__ = 'pacientes'
    
    id                 = Column(Integer, primary_key=True)
    nombre             = Column(String)
    edad               = Column(Integer)
    sexo               = Column(String)
    alergias           = Column(String)
    direccion          = Column(String)
    telefono           = Column(String)
    id_familia         = Column(String(20))

    # ——— NUEVOS CAMPOS FIJOS DEL SNS ———
    ficha_familiar     = Column(String(20))
    documento_id       = Column(String(20))
    nacionalidad       = Column(String(15))
    no_seguridad_soc   = Column(String(20))
    tipo_sangre        = Column("regimen", String(10))   # Reutiliza la columna 'regimen'
    ars                = Column(String(30))
    # ————————————————————————————————

    estado_civil       = Column(String)
    religion           = Column(String)
    lugar_nacimiento   = Column(String)
    fecha_nacimiento   = Column(Date)
    ocupacion          = Column(String)
    clinica            = Column(String)
    altura            = Column(Float)
    diagnosticos_previos     = Column(String)
    medicamentos_continuos   = Column(String)

    gestas_previas       = Column(String)
    abortos              = Column(String)
    partos_vaginales     = Column(String)
    nacidos_vivos        = Column(String)
    cesareas             = Column(String)
    gemelar              = Column(String)
    fin_embarazo_anterior= Column(String)
    embarazo_planeado    = Column(String)
    embarazo_activo      = Column(Boolean, default=False)
    embarazo_inicio      = Column(Date)
    metodo_anticonceptivo= Column(String)
    phq9_score           = Column(Integer)
    gad7_score           = Column(Integer)
    mmse_score           = Column(Integer)
    denver_done          = Column(String)
    riesgo_cardiovascular = Column(String)

    historias          = relationship("HistoriaClinica", back_populates="paciente")
    viviendas          = relationship(
        "Vivienda",
        back_populates="paciente",
        cascade="all, delete-orphan",
    )

# ------------------------------------------------------------
#   NUEVAS TABLAS PARA FUNCIONES ADICIONALES
# ------------------------------------------------------------
# ------------------------------------------------------------
#   NUEVAS TABLAS PARA FUNCIONES ADICIONALES
# ------------------------------------------------------------

class Cita(Base):
    __tablename__ = 'citas'

    id = Column(Integer, primary_key=True)
    paciente_id = Column(Integer, ForeignKey('pacientes.id'))
    fecha = Column(DateTime)
    notas = Column(Text)

    paciente = relationship("Paciente")


def schedule_lab_reminder(paciente_id):
    """Create a lab follow-up reminder 90 days from today if not already present."""
    session = Session()
    due_date = date.today() + timedelta(days=90)
    exists = (
        session.query(Recordatorio)
        .filter_by(paciente_id=paciente_id, descripcion='Exámenes de laboratorio', fecha=due_date)
        .first()
    )
    if not exists:
        session.add(
            Recordatorio(
                paciente_id=paciente_id,
                descripcion='Exámenes de laboratorio',
                fecha=due_date,
                tipo='Laboratorio',
                repetir='Ninguno',
                notificacion='Visual'
            )
        )
        session.commit()
    session.close()


def schedule_dt_vaccine_reminders(paciente_id, start_date):
    """Create reminders for dT vaccination at 20 and 36 weeks."""
    if not start_date:
        return
    session = Session()
    weeks = [20, 36]
    for w in weeks:
        due = start_date + timedelta(weeks=w)
        desc = f"Vacuna dT - semana {w}"
        exists = (
            session.query(Recordatorio)
            .filter_by(paciente_id=paciente_id, descripcion=desc, fecha=due)
            .first()
        )
        if not exists:
            session.add(
                Recordatorio(
                    paciente_id=paciente_id,
                    descripcion=desc,
                    fecha=due,
                    tipo="Vacunación",
                    repetir="Ninguno",
                    notificacion="Visual",
                )
            )
    session.commit()
    session.close()


class Servicio(Base):
    __tablename__ = 'servicios'

    id = Column(Integer, primary_key=True)
    nombre = Column(String(100), unique=True)
    tipo = Column(String(50))
    precio_base = Column(Float)


class TarifaARS(Base):
    __tablename__ = 'tarifas_ars'

    id = Column(Integer, primary_key=True)
    servicio_id = Column(Integer, ForeignKey('servicios.id'))
    ars_id = Column(Integer, ForeignKey('ars_names.id'))
    precio_ars = Column(Float)

    servicio = relationship("Servicio")
    ars = relationship("InsuranceName")


class Factura(Base):
    __tablename__ = 'facturas'

    id = Column(Integer, primary_key=True)
    paciente_id = Column(Integer, ForeignKey('pacientes.id'))
    fecha = Column(DateTime)
    monto = Column(Integer)
    descripcion = Column(Text)
    pagado = Column(Boolean, default=False)
    total = Column(Float, default=0)
    itbis = Column(Float, default=0)
    descuento_total = Column(Float, default=0)

    paciente = relationship("Paciente")
    detalles = relationship("FacturaDetalle", back_populates="factura")
    pagos = relationship("Pago", back_populates="factura")


class FacturaDetalle(Base):
    __tablename__ = 'factura_detalle'

    id = Column(Integer, primary_key=True)
    factura_id = Column(Integer, ForeignKey('facturas.id'))
    servicio_id = Column(Integer, ForeignKey('servicios.id'))
    cantidad = Column(Integer)
    precio_unitario = Column(Float)
    descuento_aplicado = Column(Float, default=0)

    factura = relationship("Factura", back_populates="detalles")
    servicio = relationship("Servicio")


class Pago(Base):
    __tablename__ = 'pagos'

    id = Column(Integer, primary_key=True)
    factura_id = Column(Integer, ForeignKey('facturas.id'))
    fecha = Column(DateTime)
    monto = Column(Float)
    metodo = Column(String)
    descripcion = Column(Text)

    factura = relationship("Factura", back_populates="pagos")


class TelemedSession(Base):
    __tablename__ = 'telemed_sessions'

    id = Column(Integer, primary_key=True)
    paciente_id = Column(Integer, ForeignKey('pacientes.id'))
    fecha = Column(DateTime)
    enlace = Column(String)

    paciente = relationship("Paciente")


class InsuranceName(Base):
    __tablename__ = 'ars_names'

    id = Column(Integer, primary_key=True)
    nombre = Column(String(50), unique=True)



class Documento(Base):
    __tablename__ = 'documentos'

    id = Column(Integer, primary_key=True)
    paciente_id = Column(Integer, ForeignKey('pacientes.id'))
    historia_id = Column(Integer, ForeignKey('historias_clinicas.id'), nullable=True)
    ruta = Column(String)
    descripcion = Column(Text)

    paciente = relationship("Paciente")
    historia = relationship("HistoriaClinica")


class SignoVital(Base):
    __tablename__ = 'signos_vitales'

    id = Column(Integer, primary_key=True, autoincrement=True)
    paciente_id = Column(Integer, ForeignKey('pacientes.id'))
    historia_id = Column(Integer, ForeignKey('historias_clinicas.id'))
    fecha = Column(DateTime, default=lambda: datetime.now())
    presion = Column(String(20))
    frecuencia = Column(Integer)
    peso = Column(Float)
    imc = Column(Float)
    glicemia = Column(Float)
    colesterol = Column(Float)
    oximetria = Column(Float)
    temperatura = Column(Float)
    peso_unidad = Column(String(2), default="kg")
    risk_processed = Column(Boolean, default=False)

    paciente = relationship("Paciente")
    historia = relationship("HistoriaClinica")


class PlantillaNota(Base):
    __tablename__ = 'plantillas_nota'

    id = Column(Integer, primary_key=True, autoincrement=True)
    nombre = Column(String(50))
    contenido = Column(Text)


class Recordatorio(Base):
    __tablename__ = 'recordatorios'

    id = Column(Integer, primary_key=True, autoincrement=True)
    paciente_id = Column(Integer, ForeignKey('pacientes.id'))
    descripcion = Column(Text)
    fecha = Column(Date)
    hora = Column(Time)
    tipo = Column(String(20))
    repetir = Column(String(20))
    notificacion = Column(String(20))
    comentario = Column(Text)
    completado = Column(Boolean, default=False)
    completado_fecha = Column(DateTime)
    completado_por = Column(String(50))

    paciente = relationship("Paciente")


class LabResult(Base):
    __tablename__ = 'lab_results'

    id = Column(Integer, primary_key=True, autoincrement=True)
    paciente_id = Column(Integer, ForeignKey('pacientes.id'))
    fecha = Column(Date, default=date.today)
    glicemia = Column(Float)
    colesterol = Column(Float)
    trigliceridos = Column(Float)
    urea = Column(Float)
    creatinina = Column(Float)
    hemoglobina = Column(Float)
    hemoglobina_glicosilada = Column(Float)

    paciente = relationship("Paciente")


class Vivienda(Base):
    __tablename__ = 'viviendas'

    id = Column(Integer, primary_key=True, autoincrement=True)
    paciente_id = Column(Integer, ForeignKey('pacientes.id', ondelete='CASCADE'))
    familia_id = Column(String(20))
    fecha = Column(Date)
    tenencia = Column(Integer)
    paredes = Column(Integer)
    techo = Column(Integer)
    piso = Column(Integer)
    sanitarios = Column(Integer)
    agua_instalacion = Column(Integer)
    agua_abastecimiento = Column(Integer)
    basura = Column(Integer)
    electricidad = Column(Integer)
    dormitorios = Column(Integer)
    combustible = Column(Integer)
    num_personas = Column(Integer)
    animales = Column(String)
    vectores = Column(String)
    total = Column(Integer)
    calificacion = Column(String(20))

    paciente = relationship("Paciente", back_populates="viviendas")


# Map housing field values to descriptive labels for statistics
HOUSING_FIELDS = {
    'tenencia': (
        'Tenencia de la vivienda',
        {5: 'Propia', 3: 'Alquilada', 2: 'Cedida/Prestada'},
    ),
    'paredes': (
        'Paredes de la vivienda',
        {10: 'Cemento', 9: 'Madera', 5: 'Asbesto', 4: 'Zinc', 0: 'Cartón/Yagua/Desechos'},
    ),
    'techo': (
        'Techo de la vivienda',
        {10: 'Concreto', 9: 'Asbesto', 0: 'Cana/Yagua/Desechos'},
    ),
    'piso': (
        'Piso de la vivienda',
        {10: 'Mosaico', 9: 'Madera', 0: 'Tierra'},
    ),
    'sanitarios': (
        'Servicios sanitarios',
        {10: 'Inodoro exclusivo', 9: 'Letrina exclusiva', 4: 'Letrina colectiva', 0: 'No tiene'},
    ),
    'agua_instalacion': (
        'Agua instalación',
        {9: 'Dentro y llega', 8: 'Fuera y llega', 0: 'No llega o no tiene'},
    ),
    'agua_abastecimiento': (
        'Abastecimiento de agua',
        {9: 'Acueducto', 8: 'Manantial', 4: 'Río', 3: 'Pozo', 2: 'Lluvia', 0: 'No tiene'},
    ),
    'basura': (
        'Eliminación de basura',
        {8: 'Recoge ayuntamiento', 7: 'La entierran', 6: 'Queman', 0: 'Cañada/patio'},
    ),
    'electricidad': (
        'Electricidad',
        {5: 'CDEEE/Inversor/Planta', 0: 'No tiene'},
    ),
    'dormitorios': (
        'Dormitorios',
        {9: '4 o más', 8: '3', 6: '2', 0: '1'},
    ),
    'combustible': (
        'Combustible de cocina',
        {5: 'Gas', 4: 'Carbón', 2: 'Leña'},
    ),
    'animales': (
        'Animales domésticos',
        {'0': 'Sí', '5': 'No'},
    ),
    'vectores': (
        'Vectores (criaderos)',
        {'0': 'Sí', '5': 'No'},
    ),
}

def cleanup_orphan_viviendas():
    """Remove housing records referencing patients that no longer exist."""
    session = Session()
    try:
        orphans = (
            session.query(Vivienda)
            .outerjoin(Paciente, Vivienda.paciente_id == Paciente.id)
            .filter(Paciente.id == None)
            .all()
        )
        count = len(orphans)
        if count:
            for v in orphans:
                session.delete(v)
            session.commit()
        return count
    except Exception:
        session.rollback()
        return 0
    finally:
        session.close()


class Consentimiento(Base):
    __tablename__ = 'consentimientos'

    id = Column(Integer, primary_key=True, autoincrement=True)
    paciente_id = Column(Integer, ForeignKey('pacientes.id'))
    fecha = Column(Date)
    ruta = Column(String)
    descripcion = Column(Text)

    paciente = relationship("Paciente")



# Create tables
try:
    Base.metadata.create_all(engine)
    ensure_db_schema(engine)
    ensure_default_services()
except SQLAlchemyError as e:
    print(f"Database error occurred: {e}")

def numero_a_letras(numero: float) -> str:
    """Devuelve el monto en letras en español sin incluir moneda."""
    unidades = (
        "CERO", "UNO", "DOS", "TRES", "CUATRO", "CINCO",
        "SEIS", "SIETE", "OCHO", "NUEVE"
    )
    especiales = {
        10: "DIEZ", 11: "ONCE", 12: "DOCE", 13: "TRECE", 14: "CATORCE",
        15: "QUINCE"
    }
    decenas = (
        "", "", "VEINTE", "TREINTA", "CUARENTA", "CINCUENTA",
        "SESENTA", "SETENTA", "OCHENTA", "NOVENTA"
    )
    centenas = (
        "", "CIENTO", "DOSCIENTOS", "TRESCIENTOS", "CUATROCIENTOS",
        "QUINIENTOS", "SEISCIENTOS", "SETECIENTOS", "OCHOCIENTOS",
        "NOVECIENTOS"
    )

    def convertir_nn(n: int) -> str:
        if n < 10:
            return unidades[n]
        if 10 <= n < 16:
            return especiales[n]
        if n < 20:
            return "DIECI" + unidades[n - 10].lower()
        if n == 20:
            return "VEINTE"
        if n < 30:
            return "VEINTI" + unidades[n - 20].lower()
        d = n // 10
        u = n % 10
        texto = decenas[d]
        if u > 0:
            texto += " Y " + unidades[u]
        return texto

    def convertir_nnn(n: int) -> str:
        if n == 0:
            return ""
        if n == 100:
            return "CIEN"
        c = n // 100
        r = n % 100
        partes = []
        if c > 0:
            partes.append(centenas[c])
        if r > 0:
            partes.append(convertir_nn(r))
        return " ".join(partes)

    entero = int(numero)
    resto = entero
    palabras = []
    if entero == 0:
        palabras.append("CERO")
    else:
        miles = entero // 1000
        resto = entero % 1000
        if miles > 0:
            if miles == 1:
                palabras.append("MIL")
            else:
                palabras.append(convertir_nnn(miles) + " MIL")
        if resto > 0:
            palabras.append(convertir_nnn(resto))

    return " ".join(palabras).strip()


from PyQt5.QtCore import Qt, QStringListModel

class StyleHelper:
    @staticmethod
    def apply_stylesheet(widget, stylesheet):
        widget.setStyleSheet(stylesheet)

    @staticmethod
    def get_screen_resolution():
        app = QApplication.instance()
        screen = app.primaryScreen()
        geometry = screen.geometry()
        return geometry.width(), geometry.height()

    def adjust_size(size, width):
        """Return the given size without any scaling."""
        return size

    @staticmethod
    def themed_icon(name, fallback):
        """Return an icon from the current theme with a fallback."""
        icon = QIcon.fromTheme(name)
        if icon.isNull():
            style = QApplication.instance().style()
            icon = style.standardIcon(fallback)
        return icon

    @staticmethod
    def add_drop_shadow(widget, radius=20):
        effect = QGraphicsDropShadowEffect(widget)
        effect.setBlurRadius(radius)
        effect.setOffset(0, 0)
        widget.setGraphicsEffect(effect)

    @staticmethod
    def window_icon(name, color="#138D75"):
        """Return a QtAwesome icon with graceful fallback."""
        if qta:
            try:
                return qta.icon(name, color=color)
            except Exception:
                pass
        return QIcon()

    @staticmethod
    def set_window_icon(widget, name, color="#138D75"):
        icon = StyleHelper.window_icon(name, color)
        if not icon.isNull():
            widget.setWindowIcon(icon)

    @staticmethod
    def style_dialog(dialog):
        """Give dialogs a rounded appearance with drop shadow."""
        StyleHelper.add_drop_shadow(dialog)
        dialog.setStyleSheet(dialog.styleSheet() + "\nQDialog{border-radius:20px;}")

    @staticmethod
    def set_dark_style(app):
        width, height = StyleHelper.get_screen_resolution()

        app.setStyle("Fusion")
        palette = QPalette()
        palette.setColor(QPalette.Window, QColor(44, 62, 80))
        palette.setColor(QPalette.WindowText, QColor(236, 240, 241))  # #3498db
        palette.setColor(QPalette.Base, QColor(52, 73, 94))
        palette.setColor(QPalette.AlternateBase, QColor(44, 62, 80))
        palette.setColor(QPalette.ToolTipBase, QColor(52, 73, 94))
        palette.setColor(QPalette.ToolTipText, QColor(236, 240, 241))  # #3498db
        palette.setColor(QPalette.Text, QColor(236, 240, 241))  # #3498db
        palette.setColor(QPalette.Button, QColor(52, 73, 94))
        palette.setColor(QPalette.ButtonText, QColor(236, 240, 241))  # #3498db
        palette.setColor(QPalette.BrightText, QColor(255, 50, 50))
        palette.setColor(QPalette.Link, QColor(52, 152, 219))
        palette.setColor(QPalette.Highlight, QColor(41, 128, 185))
        palette.setColor(QPalette.HighlightedText, QColor(236, 240, 241))
        app.setPalette(palette)
        app.setProperty("theme", "dark")

        font_size = max(10, min(14, int(width / 320)))
        font_size = StyleHelper.adjust_size(font_size, width)
        font = QFont("Segoe UI", font_size)
          # Aplicar negrita por defecto
        app.setFont(font)

    @staticmethod
    def set_light_style(app):
        width, height = StyleHelper.get_screen_resolution()

        app.setStyle("Fusion")
        palette = QPalette()
        palette.setColor(QPalette.Window, QColor(245, 245, 245))
        palette.setColor(QPalette.WindowText, QColor(33, 33, 33))
        palette.setColor(QPalette.Base, QColor(255, 255, 255))
        palette.setColor(QPalette.AlternateBase, QColor(245, 245, 245))
        palette.setColor(QPalette.ToolTipBase, QColor(255, 255, 220))
        palette.setColor(QPalette.ToolTipText, QColor(33, 33, 33))
        palette.setColor(QPalette.Text, QColor(33, 33, 33))
        palette.setColor(QPalette.Button, QColor(240, 240, 240))
        palette.setColor(QPalette.ButtonText, QColor(33, 33, 33))
        palette.setColor(QPalette.BrightText, QColor(255, 50, 50))
        palette.setColor(QPalette.Link, QColor(41, 128, 185))
        palette.setColor(QPalette.Highlight, QColor(41, 128, 185))
        palette.setColor(QPalette.HighlightedText, QColor(255, 255, 255))
        app.setPalette(palette)
        app.setProperty("theme", "light")

        font_size = max(10, min(14, int(width / 320)))
        font_size = StyleHelper.adjust_size(font_size, width)
        font = QFont("Segoe UI", font_size)
        app.setFont(font)

    @staticmethod
    def set_gpt_style(app):
        width, _ = StyleHelper.get_screen_resolution()

        app.setStyle("Fusion")
        palette = QPalette()
        palette.setColor(QPalette.Window, QColor(247, 247, 248))
        palette.setColor(QPalette.WindowText, QColor(32, 33, 35))
        palette.setColor(QPalette.Base, QColor(255, 255, 255))
        palette.setColor(QPalette.AlternateBase, QColor(247, 247, 248))
        palette.setColor(QPalette.ToolTipBase, QColor(255, 255, 220))
        palette.setColor(QPalette.ToolTipText, QColor(32, 33, 35))
        palette.setColor(QPalette.Text, QColor(32, 33, 35))
        palette.setColor(QPalette.Button, QColor(240, 240, 240))
        palette.setColor(QPalette.ButtonText, QColor(32, 33, 35))
        palette.setColor(QPalette.BrightText, QColor(255, 50, 50))
        palette.setColor(QPalette.Link, QColor(16, 163, 127))
        palette.setColor(QPalette.Highlight, QColor(16, 163, 127))
        palette.setColor(QPalette.HighlightedText, QColor(255, 255, 255))
        app.setPalette(palette)
        app.setProperty("theme", "gpt")

        font_size = max(10, min(14, int(width / 320)))
        font_size = StyleHelper.adjust_size(font_size, width)
        font = QFont("Segoe UI", font_size)
        app.setFont(font)

    @staticmethod
    def style_label(label):
        width, height = StyleHelper.get_screen_resolution()
        font_size = max(16, min(30, int(width / 120)))
        font_size = StyleHelper.adjust_size(font_size, width)
        app = QApplication.instance()
        palette = app.palette()
        text_color = palette.color(QPalette.WindowText).name()
        label.setStyleSheet(f"font-size: {font_size}px; font-weight: bold; color: {text_color};")

    @staticmethod
    def style_messagebox(messagebox):
        width, height = StyleHelper.get_screen_resolution()
        font_size = max(10, min(18, int(width / 240)))
        font_size = StyleHelper.adjust_size(font_size, width)
        button_size = max(20, min(40, int(width / 128)))
        button_size = StyleHelper.adjust_size(button_size, width)
        
        app = QApplication.instance()
        palette = app.palette()
        bg_color = palette.color(QPalette.Window).name()
        text_color = palette.color(QPalette.WindowText).name()
        button_bg = palette.color(QPalette.Highlight).name()
        button_text = palette.color(QPalette.HighlightedText).name()

        messagebox.setStyleSheet(f"""
            QMessageBox {{
                background-color: {bg_color};
                font-size: {font_size}px;
                color: {text_color};
            }}
            QMessageBox QPushButton {{
                min-width: {button_size}px;
                min-height: {button_size / 2}px;
                font-weight: bold;
                background-color: {button_bg};
                color: {button_text};
            }}
        """)


    @staticmethod
    def style_button(button, base_color="#17A589"):
        app = QApplication.instance()

        if qta:
            # Assign a QtAwesome icon based on button text
            text = button.text()
            icon_map = {
                # Main window
                "Nuevo Paciente": "fa5s.user-plus",
                "Editar Paciente": "fa5s.user-edit",
                "Eliminar Paciente": "fa5s.user-times",
                "Medicacion Fija": "fa5s.pills",
                "Signos Vitales": "fa5s.heartbeat",
                "Citas": "fa5s.calendar-check",
                "Facturación": "fa5s.file-invoice-dollar",
                "Facturacion": "fa5s.file-invoice-dollar",
                "Resultados de Laboratorio": "fa5s.vials",
                "Lab Results": "fa5s.vials",
                "Gestionar ARS": "fa5s.id-card",
                "Gestionar Servicios": "fa5s.toolbox",
                "Formularios": "fa5b.wpforms",
                "Recordatorios": "fa5s.bell",
                "Búsqueda Avanzada": "fa5s.search-plus",
                "Busqueda Avanzada": "fa5s.search-plus",
                "Análisis de Datos": "fa5s.chart-bar",
                "Analisis de Datos": "fa5s.chart-bar",
                "Estudio Clínico IA": "fa5s.brain",
                "Estudio Clinico IA": "fa5s.brain",
                "Clinical AI Study": "fa5s.brain",
                # Historia clinica / IA suggestions
                "Nueva Historia Clínica": "fa5s.notes-medical",
                "Nueva Historia Clinica": "fa5s.notes-medical",
                "Historia Clínica": "fa5s.notes-medical",
                "Medical History": "fa5s.notes-medical",
                "Eliminar Historia": "fa5s.trash-alt",
                "Generar Informe Diario": "fa5s.calendar-day",
                "Informe por Periodo": "fa5s.calendar-alt",
                "Eliminar Duplicados": "fa5s.eraser",
                "Delete Duplicates": "fa5s.eraser",
                "Guardar Historia": "fa5s.save",
                "Exportar a Word": "fa5s.file-word",
                "Generar Historia Clínica": "fa5s.file-medical",
                "Generar Historia Clinica": "fa5s.file-medical",
                "Generar Receta": "fa5s.prescription",
                "Explicar al Paciente": "fa5s.comment-medical",
                "Generar Laboratorio": "fa5s.flask",
                "Autoajustar Tamaño": "fa5s.expand-arrows-alt",
                "Autoajustar Tamano": "fa5s.expand-arrows-alt",
                "Preguntar Historial": "fa5s.question",
                "Analizar Imagen": "fa5s.image",
                "¿Olvidaste tu contraseña?": "fa5s.question-circle",
                "Salud Familiar y Social": "fa5s.home",
            }

            icon_name = None
            if text in icon_map:
                icon_name = icon_map[text]
            else:
                lowered = text.lower()
                for keyword, name in {
                    "cerrar": "fa5s.times",
                    "eliminar": "fa5s.trash",
                    "iniciar": "fa5s.sign-in-alt",
                    "registr": "fa5s.user-plus",
                    "volver": "fa5s.arrow-left",
                    "generar": "fa5s.play",
                    "exportar": "fa5s.file-export",
                    "guardar": "fa5s.save",
                    "cargar": "fa5s.folder-open",
                    "editar": "fa5s.edit",
                    "buscar": "fa5s.search",
                    "imprimir": "fa5s.print",
                    "agregar": "fa5s.plus",
                    "añadir": "fa5s.plus",
                    "aceptar": "fa5s.check",
                    "cancelar": "fa5s.times",
                    "actualizar": "fa5s.sync",
                    "enviar": "fa5s.paper-plane",
                }.items():
                    if keyword in lowered:
                        icon_name = name
                        break

            if icon_name:
                try:
                    button.setIcon(qta.icon(icon_name, color="white"))
                    button.setIconSize(QSize(50, 50))
                except Exception:
                    pass

        hover_color = QColor(base_color).darker(110).name()
        button.setStyleSheet(
            f"""
            QPushButton {{
                background-color: {base_color};
                color: white;
                border-radius: 8px;
                font-weight: bold;
            }}
            QPushButton:hover {{
                background-color: {hover_color};
            }}
            """
        )

    @staticmethod
    def style_input(input_widget):
        app = QApplication.instance()
        palette = app.palette()
        base_color = palette.color(QPalette.Base).name()
        text_color = palette.color(QPalette.Text).name()
        highlight_color = palette.color(QPalette.Highlight).name()

        width, height = StyleHelper.get_screen_resolution()
        font_size = max(14, min(30, int(width / 120)))
        font_size = StyleHelper.adjust_size(font_size, width)
        padding = max(5, min(15, int(width / 384)))
        padding = StyleHelper.adjust_size(padding, width)
        theme = app.property("theme")
        if theme == "gpt":
            border_radius = StyleHelper.adjust_size(30, width)
            border_color = "#bdc3c7"
        else:
            border_radius = max(4, min(10, int(width / 480)))
            border_radius = StyleHelper.adjust_size(border_radius, width)
            border_color = "#555"
        border_width = max(1, min(2, int(width / 960)))
        border_width = StyleHelper.adjust_size(border_width, width)

        base_style = f"""
            border: {border_width}px solid {border_color};
            border-radius: {border_radius}px;
            padding: {padding}px;
            background-color: {base_color};
            font-size: {font_size}px;
            font-weight: normal;
            color: {text_color};
        """

        focus_style = f"""
            border-color: {highlight_color};
        """
        
        if isinstance(input_widget, QLineEdit):
            input_widget.setStyleSheet(f"""
                QLineEdit {{
                    {base_style}
                }}
                QLineEdit:focus {{
                    {focus_style}
                }}
            """)
        elif isinstance(input_widget, QTextEdit):
            input_widget.setStyleSheet(f"""
                QTextEdit {{
                    {base_style}
                }}
                QTextEdit:focus {{
                    {focus_style}
                }}
            """)
        elif isinstance(input_widget, QDateEdit):
            input_widget.setStyleSheet(f"""
                QDateEdit {{
                    {base_style}
                }}
                QDateEdit:focus {{
                    {focus_style}
                }}
            """)
        elif isinstance(input_widget, QComboBox):
            input_widget.setStyleSheet(f"""
                QComboBox {{
                    {base_style}
                }}
                QComboBox:focus {{
                    {focus_style}
                }}
            """)

    @staticmethod
    def style_radio(radio):
        """Apply consistent sizing to QRadioButton widgets."""
        radio.setStyleSheet(
            "QRadioButton {"
            "font-family: 'Segoe UI', 'Roboto', 'Inter', sans-serif;"
            "font-size:30px;"
            "padding:8px;"
            "min-width:120px;"
            "}"
            "QRadioButton::indicator {width:32px;height:32px;}"
            "QRadioButton:checked {"
            "background-color:#d0f0c0;"
            "font-weight:bold;"
            "border-radius:6px;"
            "}"
            "QRadioButton::indicator:checked {"
            "background-color:#27AE60;"
            "border:2px solid #1F8F55;"
            "}"
        )

    @staticmethod
    def style_info_text(text_edit):
        """Style read-only text areas used for displaying information."""
        text_edit.setFrameStyle(QFrame.NoFrame)
        text_edit.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        font = QFont("Segoe UI", 12)
        text_edit.setFont(font)
        text_edit.setStyleSheet("background: transparent;")

    @staticmethod
    def style_table(table, base_color=None):
        width, height = StyleHelper.get_screen_resolution()
        font_size = max(14, min(32, int(width / 120)))
        font_size = StyleHelper.adjust_size(font_size, width)
        padding = max(6, min(12, int(width / 640)))
        padding = StyleHelper.adjust_size(padding, width)
        border_radius = max(6, min(12, int(width / 480)))
        border_radius = StyleHelper.adjust_size(border_radius, width)
        border_width = max(1, min(2, int(width / 960)))
        border_width = StyleHelper.adjust_size(border_width, width)
    
        app = QApplication.instance()
        palette = app.palette()
        if base_color is None:
            base_color = palette.color(QPalette.Base).name()
        text_color = palette.color(QPalette.Text).name()
        header_color = palette.color(QPalette.Highlight).name()
        header_text = palette.color(QPalette.HighlightedText).name()
        alternate = QColor(base_color).lighter(110).name()
        select_color = QColor("#d0e8e2").name()

        table.setAlternatingRowColors(True)
        table.horizontalHeader().setDefaultAlignment(Qt.AlignCenter)
        table.setSelectionBehavior(QAbstractItemView.SelectItems)
        table.setSelectionMode(QAbstractItemView.ExtendedSelection)
        table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        table.setContextMenuPolicy(Qt.CustomContextMenu)
        table.customContextMenuRequested.connect(
            lambda pos, t=table: StyleHelper._table_menu(t, pos)
        )
        table.itemDoubleClicked.connect(StyleHelper._show_full_cell)
        _filter = StyleHelper._TableEventFilter(table)
        table._copy_filter = _filter
        table.installEventFilter(_filter)
        table.setStyleSheet(f"""
            QTableWidget {{
                border: {border_width}px solid #555;
                border-radius: {border_radius}px;
                background-color: {base_color};
                alternate-background-color: {alternate};
                font-size: {font_size}px;
                font-weight: bold;
                color: {text_color};
            }}
            QTableWidget::item {{
                padding: {padding}px;
            }}
            QTableWidget::item:selected {{
                background-color: {select_color};
                color: {text_color};
                font-weight: bold;
            }}
            QHeaderView::section {{
                background-color: {header_color};
                color: {header_text};
                padding: {padding * 1.5}px;
                border: none;
                font-weight: bold;
                font-size: {font_size}px;
                border-bottom: 2px solid #555;
                border-top: 1px solid #555;
                text-align: left;
                padding-left: {padding}px;
                border-radius: 0;
            }}
        """)

    class _TableEventFilter(QObject):
        """Handle copy shortcuts for table widgets."""

        def __init__(self, table):
            super().__init__(table)
            self.table = table

        def eventFilter(self, obj, event):
            if obj is self.table and event.type() == QEvent.KeyPress:
                if event.matches(QKeySequence.Copy):
                    StyleHelper.copy_table_selection(self.table)
                    return True
            return QObject.eventFilter(self, obj, event)

    @staticmethod
    def copy_table_selection(table):
        """Copy currently selected cells from ``table`` to clipboard."""
        ranges = table.selectedRanges()
        if not ranges:
            return
        parts = []
        for r in ranges:
            for row in range(r.topRow(), r.bottomRow() + 1):
                row_data = []
                for col in range(r.leftColumn(), r.rightColumn() + 1):
                    item = table.item(row, col)
                    row_data.append(item.text() if item else "")
                parts.append("\t".join(row_data))
        QApplication.clipboard().setText("\n".join(parts))

    @staticmethod
    def _table_menu(table, pos):
        """Context menu offering copy and view actions."""
        menu = QMenu(table)
        copy_act = menu.addAction("Copiar")
        item = table.itemAt(pos)
        view_act = menu.addAction("Ver completo") if item else None
        action = menu.exec_(table.viewport().mapToGlobal(pos))
        if action == copy_act:
            StyleHelper.copy_table_selection(table)
        elif view_act and action == view_act:
            StyleHelper._show_full_cell(item)

    @staticmethod
    def _show_full_cell(item):
        """Display the full content of ``item`` in a dialog."""
        dlg = QDialog()
        dlg.setWindowTitle("Contenido completo")
        layout = QVBoxLayout(dlg)
        text = QPlainTextEdit(item.text())
        text.setReadOnly(True)
        layout.addWidget(text)
        btn = QPushButton("Cerrar")
        btn.clicked.connect(dlg.accept)
        layout.addWidget(btn, alignment=Qt.AlignRight)
        dlg.resize(400, 300)
        dlg.exec_()

    @staticmethod
    def style_menu_bar(menubar):
        width, _ = StyleHelper.get_screen_resolution()
        font_size = max(14, min(30, int(width / 120)))
        font_size = StyleHelper.adjust_size(font_size, width)

        menubar.setStyleSheet(
            f"""
            QMenuBar {{
                background: #ffffff;
            }}
            QMenuBar::item {{
                background: transparent;
                color: #333333;
                font-family: 'Segoe UI', 'Roboto', 'Poppins', 'Helvetica';
                font-size: {font_size}px;
                padding: 6px 12px;
                margin-right: 16px;
            }}
            QMenuBar::item:selected {{
                background: #A3E4D7;
                border-bottom: 2px solid #2ECC71;
            }}
        """
        )

    @staticmethod
    def style_menu(menu):
        width, _ = StyleHelper.get_screen_resolution()
        font_size = max(14, min(30, int(width / 120)))
        font_size = StyleHelper.adjust_size(font_size, width)

        menu.setStyleSheet(
            f"""
            QMenu {{
                background: #ffffff;
                color: #1F1F1F;
                font-family: 'Segoe UI', 'Roboto', 'Poppins', 'Helvetica';
                font-size: {font_size}px;
                border: 1px solid #A3E4D7;
            }}
            QMenu::item:selected {{
                background: #A3E4D7;
            }}
        """
        )

    @staticmethod
    def style_checkbox(checkbox):
        width, height = StyleHelper.get_screen_resolution()
        font_size = max(14, min(32, int(width / 120)))
        font_size = StyleHelper.adjust_size(font_size, width)
        padding = max(6, min(15, int(width / 16)))
        padding = StyleHelper.adjust_size(padding, width)
        spacing = max(4, min(8, int(width / 480)))
        spacing = StyleHelper.adjust_size(spacing, width)
        indicator_size = max(17, min(35, int(width / 110)))
        indicator_size = StyleHelper.adjust_size(indicator_size, width)
        app = QApplication.instance()
        palette = app.palette()
        text_color = palette.color(QPalette.WindowText).name()
        highlight_color = palette.color(QPalette.Highlight).name()
        base_color = palette.color(QPalette.Base).name()

        checkbox.setStyleSheet(f"""
            QCheckBox {{
                font-size: {font_size}px;

                padding: {padding}px;
                color: {text_color};
            }}
            QCheckBox::indicator {{
                width: {int(padding * 1.5)}px;
                height: {int(padding * 1.5)}px;
            }}
            QCheckBox::indicator:checked {{
                background-color: {highlight_color};
            }}
            QCheckBox::indicator:unchecked {{
                background-color: {base_color};
                border: 1px solid #555;
            }}
        """)

    @staticmethod
    def style_groupbox(groupbox):
        width, _ = StyleHelper.get_screen_resolution()
        font_size = StyleHelper.adjust_size(max(18, min(34, int(width / 106))), width)
        border_radius = StyleHelper.adjust_size(max(4, min(10, int(width / 480))), width)
        border_width = StyleHelper.adjust_size(max(1, min(2, int(width / 960))), width)
        padding_top = StyleHelper.adjust_size(max(5, min(15, int(width / 384))), width)
        app = QApplication.instance()
        text_color = app.palette().color(QPalette.WindowText).name()

        groupbox.setStyleSheet(f"""
            QGroupBox {{
                font-size: {font_size}px;
                font-weight: bold;
                color: {text_color};
                background-color: transparent;
                border: {border_width}px solid #555;
                border-radius: {border_radius}px;
                margin-top: {padding_top / 2}ex;
                padding-top: {padding_top}px;
                border-bottom: {border_width}px solid rgba(0, 0, 0, 0.1);
            }}
            QGroupBox[nbi="true"] {{
                border-color: #E74C3C;
            }}
            QGroupBox[answered="true"]::title {{
                color: #27AE60;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                subcontrol-position: top center;
                padding: 0 {padding_top / 2}px;
                font-size: {font_size + 2}px;
                color: {text_color};
                background-color: transparent;
                font-weight: bold;
            }}
        """)



class DatabaseManager(QDialog):
    from datetime import datetime
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.database')
        self.setWindowTitle("Gestión de Base de Datos")
        self.setGeometry(200, 200, 300, 200)
        self.setMaximumSize(3840, 2160)  # Tamaño máximo 4K
        self.backup_folder = os.path.join(os.path.expanduser("~"), 'db_backups')
        if not os.path.exists(self.backup_folder):
            os.makedirs(self.backup_folder)
        self.init_ui()
                # Aplica el estilo al diálogo
        
    def init_ui(self):
        layout = QVBoxLayout()

        self.load_button = QPushButton("Cargar Base de Datos")
        self.backup_button = QPushButton("Crear Copia de Seguridad")
        self.restore_button = QPushButton("Restaurar Copia de Seguridad")
        self.revert_button = QPushButton("Revertir a Base Anterior")
        

        StyleHelper.style_button(self.load_button)
        StyleHelper.style_button(self.backup_button)
        StyleHelper.style_button(self.restore_button)
        StyleHelper.style_button(self.revert_button)

        layout.addWidget(self.load_button)
        layout.addWidget(self.backup_button)
        layout.addWidget(self.restore_button)
        layout.addWidget(self.revert_button)

        self.load_button.clicked.connect(self.load_database)
        self.backup_button.clicked.connect(self.backup_database)
        self.restore_button.clicked.connect(self.restore_database)
        self.revert_button.clicked.connect(self.revert_database)

        self.setLayout(layout)   
    
    




    
    def load_database(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getOpenFileName(self, "Seleccionar Base de Datos", "", 
                                                   "Database Files (*.db);;All Files (*)", options=options)
        if file_name:
            try:
                new_engine = create_engine(f'sqlite:///{file_name}')
                new_engine.connect()
                if os.path.exists(file_name):
                    os.chmod(file_name, 0o600)
                
                # Verificar y actualizar la estructura de la base de datos
                self.verify_and_update_db_structure(new_engine)
                
                # Hacer copia de seguridad de la base de datos actual
                self.backup_current_database()
                
                # Cambiar a la nueva base de datos
                global engine
                engine = new_engine
                Session.configure(bind=engine)
                
                # Renombrar y mover la nueva base de datos
                new_db_path = os.path.join(os.path.expanduser("~"), 'historias_clinic158.db')
                shutil.copy2(file_name, new_db_path)
                engine = create_engine(f'sqlite:///{new_db_path}')
                if os.path.exists(new_db_path):
                    os.chmod(new_db_path, 0o600)
                Session.configure(bind=engine)
                
                QMessageBox.information(self, "Éxito", "Base de datos cargada y actualizada correctamente.")
            except Exception as e:
                QMessageBox.warning(self, "Error", f"No se pudo cargar la base de datos: {str(e)}")

    def verify_and_update_db_structure(self, new_engine):
        ensure_db_schema(new_engine)
        ensure_default_services()

    def backup_current_database(self):
        current_db = engine.url.database
        if os.path.exists(current_db):
            backup_file = os.path.join(self.backup_folder, f"backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db")
            shutil.copy2(current_db, backup_file)

    def backup_database(self):
        try:
            current_db = engine.url.database
            backup_file = os.path.join(self.backup_folder, f"backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db")
            shutil.copy2(current_db, backup_file)
            QMessageBox.information(self, "Éxito", f"Copia de seguridad creada: {backup_file}")
            
          
            
        except Exception as e:
            QMessageBox.warning(self, "Error", f"No se pudo crear la copia de seguridad: {str(e)}")

    def restore_database(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getOpenFileName(self, "Seleccionar Copia de Seguridad", self.backup_folder, 
                                                   "Database Files (*.db);;All Files (*)", options=options)
        if file_name:
            try:
                current_db = engine.url.database
                shutil.copy2(file_name, current_db)
                QMessageBox.information(self, "Éxito", "Base de datos restaurada correctamente.")
            except Exception as e:
                QMessageBox.warning(self, "Error", f"No se pudo restaurar la base de datos: {str(e)}")

    def revert_database(self):
        backups = [f for f in os.listdir(self.backup_folder) if f.endswith('.db')]
        if backups:
            latest_backup = max(backups, key=lambda x: os.path.getmtime(os.path.join(self.backup_folder, x)))
            try:
                current_db = engine.url.database
                shutil.copy2(os.path.join(self.backup_folder, latest_backup), current_db)
                QMessageBox.information(self, "Éxito", "Se ha revertido a la base de datos anterior.")
            except Exception as e:
                QMessageBox.warning(self, "Error", f"No se pudo revertir la base de datos: {str(e)}")
        else:
            QMessageBox.warning(self, "Error", "No se encontraron copias de seguridad anteriores.")




session_factory = sessionmaker(bind=engine, expire_on_commit=False)
Session = scoped_session(session_factory)
class PatientDialog(QDialog):
    def __init__(self, parent=None, patient=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.user-md')
        self.setWindowTitle("Datos del Paciente")
        self.patient = patient
        self.excel_data = self.load_excel_data()
        self.current_diagnosticos = ""  # Nuevo: para almacenar los diagnósticos actuales
        self.birth_date_changed = False  # Marcar si el usuario especificó fecha
        self.init_ui()
        self.resize(1800, 1800)

        # Habilitar botones de minimizar y maximizar
        self.setWindowFlags(self.windowFlags() | Qt.WindowMinMaxButtonsHint)

        # Usar el color de la paleta actual para el fondo del diálogo
        app = QApplication.instance()
        bg_color = app.palette().color(QPalette.Window).name()
        StyleHelper.apply_stylesheet(self, f"QDialog {{ background-color: {bg_color}; }}")

    def init_ui(self):
        """Crea todos los widgets y el formulario del paciente."""
        self.setWindowTitle("Datos del Paciente")
        self.resize(700, 770)

        layout = QFormLayout(self)

        # Campo de Cédula/NUI (primer campo)
        self.documento_id_input   = QLineEdit(self);        StyleHelper.style_input(self.documento_id_input)
        self.documento_id_input.editingFinished.connect(self.format_cedula)

        self.btn_autofill_cedula  = QPushButton("Buscar", self)
        StyleHelper.style_button(self.btn_autofill_cedula)
        self.btn_autofill_cedula.clicked.connect(self.autofill_from_cedula)

        # ---------- Datos básicos ----------
        self.nombre_input   = QLineEdit(self);              StyleHelper.style_input(self.nombre_input)
        self.nombre_input.textChanged.connect(lambda text: self.update_diagnosticos_previos(text))
        self.edad_input     = QLineEdit(self);              StyleHelper.style_input(self.edad_input)
        self.sexo_input     = QComboBox(self);              self.sexo_input.addItems(["Masculino", "Femenino"])
        self.alergias_input = QLineEdit(self);              StyleHelper.style_input(self.alergias_input)
        self.direccion_input= QLineEdit(self);              StyleHelper.style_input(self.direccion_input)
        self.telefono_input = QLineEdit(self);              StyleHelper.style_input(self.telefono_input)

        # Nuevo campo para el código FFS
        self.id_familia_input = QLineEdit(self);            StyleHelper.style_input(self.id_familia_input)

        # ---------- NUEVOS datos fijos SNS ----------
        self.ficha_familiar_input = QLineEdit(self);        StyleHelper.style_input(self.ficha_familiar_input)
        self.nacionalidad_input   = QComboBox(self);        self.nacionalidad_input.addItems(
            ["Dominicana", "Haitiana", "Venezolana", "Otros"])
        self.no_seguro_input  = QLineEdit(self);            StyleHelper.style_input(self.no_seguro_input)
        self.tipo_sangre_input = QComboBox(self);           self.tipo_sangre_input.addItems(
            ["O+", "O-", "A+", "A-", "B+", "B-", "AB+", "AB-", "No sabe"])
        self.tipo_sangre_input.setCurrentText("No sabe")
        self.ars_input        = QComboBox(self)
        self.load_insurance_options()
    
        # ---------- Otros datos clínicos ----------
        self.estado_civil_input     = QComboBox(self); self.estado_civil_input.addItems(
            ["Soltero/a", "Casado/a", "Divorciado/a", "Viudo/a", "Unión libre", "Separado/a"])

        self.religion_input         = QLineEdit(self);      StyleHelper.style_input(self.religion_input)
        religions = [
            "Católico/a", "Cristiano/a", "Evangélico/a", "Musulmán/a", "Judío/a",
            "Budista", "Hinduista", "Ateo/a", "Agnóstico/a", "Testigo de Jehová",
            "Mormón/a", "Ortodoxo/a", "Bautista", "Metodista", "Adventista",
            "Pentecostal", "Espiritista", "Ninguna"
        ]
        religion_completer = QCompleter(religions, self.religion_input)
        religion_completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.religion_input.setCompleter(religion_completer)

        self.lugar_nacimiento_input = QLineEdit(self);      StyleHelper.style_input(self.lugar_nacimiento_input)
        self.fecha_nacimiento_input = QDateEdit(self);      self.fecha_nacimiento_input.setCalendarPopup(True)
        self.fecha_nacimiento_input.dateChanged.connect(self.on_birthdate_changed)

        self.ocupacion_input        = QLineEdit(self);      StyleHelper.style_input(self.ocupacion_input)
        occupations = [
            "Estudiante", "Ama de casa", "Desempleado/a", "Médico/a", "Enfermero/a",
            "Ingeniero/a", "Abogado/a", "Profesor/a", "Administrador/a", "Contador/a",
            "Agricultor/a", "Comerciante", "Chofer", "Carpintero/a", "Electricista",
            "Policía", "Militar", "Arquitecto/a", "Programador/a", "Artista",
            "Músico/a", "Chef", "Cocinero/a", "Otros"
        ]
        occupation_completer = QCompleter(occupations, self.ocupacion_input)
        occupation_completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.ocupacion_input.setCompleter(occupation_completer)
        self.height_input           = QLineEdit(self);      StyleHelper.style_input(self.height_input)
        self.height_unit_combo      = QComboBox(self)
        self.height_unit_combo.addItems(["cm", "m", "ft"])
        self.height_unit_combo.currentTextChanged.connect(self.on_height_unit_changed)
        self.current_height_unit = "cm"
        self.diagnosticos_previos_input = QTextEdit(self);  StyleHelper.style_input(self.diagnosticos_previos_input)
        self.medicamentos_continuos_input = QTextEdit(self);StyleHelper.style_input(self.medicamentos_continuos_input)

    
        # ---------- Añadir filas al layout ----------
        ced_row = QHBoxLayout()
        ced_row.addWidget(self.documento_id_input)
        ced_row.addWidget(self.btn_autofill_cedula)
        layout.addRow("Cédula / NUI:",            ced_row)
        layout.addRow("Nombre:",                  self.nombre_input)
        layout.addRow("Fecha de Nacimiento:",     self.fecha_nacimiento_input)
        layout.addRow("Edad:",                    self.edad_input)
        layout.addRow("Sexo:",                    self.sexo_input)
        layout.addRow("Estado Civil:",            self.estado_civil_input)
        layout.addRow("Nacionalidad:",            self.nacionalidad_input)
        layout.addRow("Dirección:",               self.direccion_input)
        layout.addRow("Teléfono:",               self.telefono_input)
        layout.addRow("Alergias:",                self.alergias_input)

        layout.addRow("ID Familia:",             self.id_familia_input)
        layout.addRow("Ficha Familiar:",          self.ficha_familiar_input)
        layout.addRow("No. Seguro Social:",       self.no_seguro_input)
        layout.addRow("Tipo de Sangre:",          self.tipo_sangre_input)
        layout.addRow("ARS:",                     self.ars_input)

        layout.addRow("Religión:",                self.religion_input)
        layout.addRow("Lugar de Nacimiento:",     self.lugar_nacimiento_input)
        layout.addRow("Ocupación:",               self.ocupacion_input)
        height_row = QHBoxLayout()
        height_row.addWidget(self.height_input)
        height_row.addWidget(self.height_unit_combo)
        layout.addRow("Altura:",                 height_row)
        layout.addRow("Diagnósticos Previos:",    self.diagnosticos_previos_input)
        layout.addRow("Medicamentos Continuos:",  self.medicamentos_continuos_input)

        # ---------- Botón Guardar ----------
        btn_guardar = QPushButton("Guardar", self)
        btn_guardar.clicked.connect(self.save_patient)
        StyleHelper.style_button(btn_guardar)
        layout.addRow(btn_guardar)
    
        # Si es edición, poblar campos
        if self.patient:
            self.fill_data()

    def _sanitize_cedula(self, text: str) -> str:
        return ''.join(c for c in text if c.isdigit())

    def format_cedula(self):
        """Formatea el número de cédula con guiones."""
        ced = self._sanitize_cedula(self.documento_id_input.text())
        if len(ced) == 11:
            self.documento_id_input.setText(f"{ced[:3]}-{ced[3:10]}-{ced[10:]}")

    def autofill_from_cedula(self):
        ced = self._sanitize_cedula(self.documento_id_input.text())
        if len(ced) != 11:
            return
        self.format_cedula()
        data = consultar_cedula_todo(ced)
        info_ced = data.get("intrant_cedulados", {})
        info_mirex = data.get("mirex", {})
        nombres = info_mirex.get("nombres") or info_ced.get("nombres", "")
        apellido1 = info_mirex.get("apellido1", "")
        apellido2 = info_mirex.get("apellido2", "")
        apellidos = info_ced.get("apellidos", "")
        nombre_completo = " ".join(filter(None, [nombres, apellido1, apellido2])).strip()
        if not nombre_completo:
            nombre_completo = f"{nombres} {apellidos}".strip()
        if nombre_completo:
            self.nombre_input.setText(nombre_completo)

        fecha_nac = info_ced.get("fecha_nacimiento")
        if fecha_nac:
            try:
                dt = datetime.strptime(fecha_nac, "%Y-%m-%d")
                self.fecha_nacimiento_input.setDate(QDate(dt.year, dt.month, dt.day))
                self.update_age()
            except ValueError:
                pass
        else:
            fecha_mirex = info_mirex.get("fecha_nacimiento")
            if fecha_mirex:
                try:
                    dt = datetime.strptime(fecha_mirex, "%Y-%m-%d")
                    self.fecha_nacimiento_input.setDate(QDate(dt.year, dt.month, dt.day))
                    self.update_age()
                except ValueError:
                    pass
        if info_ced.get("edad"):
            self.edad_input.setText(str(info_ced["edad"]))

        sexo = (info_ced.get("sexo", "") or info_mirex.get("sexo", "")).upper()
        if sexo == "M":
            self.sexo_input.setCurrentText("Masculino")
        elif sexo == "F":
            self.sexo_input.setCurrentText("Femenino")
        estado = info_ced.get("estado_civil")
        estado_map = {'S': 'Soltero/a', 'C': 'Casado/a', 'D': 'Divorciado/a', 'V': 'Viudo/a', 'U': 'Unión libre', 'SE': 'Separado/a'}
        if estado:
            self.estado_civil_input.setCurrentText(estado_map.get(estado, self.estado_civil_input.currentText()))
        nacionalidad = info_ced.get("nacionalidad")
        if nacionalidad:
            self.nacionalidad_input.setCurrentText(nacionalidad.capitalize())
        if info_ced.get("direccion"):
            self.direccion_input.setText(info_ced["direccion"])
        telefono = data.get("intrant_licencia", {}).get("telefono")
        if telefono:
            self.telefono_input.setText(telefono)
        lugar_nac = info_mirex.get("lugar_nacimiento")
        if lugar_nac:
            self.lugar_nacimiento_input.setText(lugar_nac)

    def update_age(self):
        birth_date = self.fecha_nacimiento_input.date().toPyDate()
        months = months_between(birth_date, date.today())
        years = months // 12
        self.edad_input.setText(str(max(years, 0)))

    def on_birthdate_changed(self):
        self.birth_date_changed = True
        self.update_age()
    def toggle_maximize(self):
        if self.isMaximized():
            self.showNormal()  # Restaurar tamaño normal
        else:
            self.showMaximized()  # Maximizar ventana

    def convert_to_cm(self, value, unit):
        if unit == "cm":
            return value
        if unit == "m":
            return value * 100
        if unit == "ft":
            return value * 30.48
        return value

    def convert_from_cm(self, value_cm, unit):
        if unit == "cm":
            return value_cm
        if unit == "m":
            return value_cm / 100
        if unit == "ft":
            return value_cm / 30.48
        return value_cm

    def on_height_unit_changed(self, new_unit):
        try:
            current_value = float(self.height_input.text())
        except ValueError:
            current_value = None
        if current_value is not None:
            cm = self.convert_to_cm(current_value, self.current_height_unit)
            converted = self.convert_from_cm(cm, new_unit)
            self.height_input.setText(f"{converted:.2f}")
        self.current_height_unit = new_unit
        if self.parent() and hasattr(self.parent(), "update_bmi"):
            self.parent().update_bmi()

    def load_excel_file(self):
        options = QFileDialog.Options()
        options |= QFileDialog.ReadOnly
        file_path, _ = QFileDialog.getOpenFileName(self, "Cargar Archivo Excel", "", "Excel Files (*.xlsx);;All Files (*)", options=options)
        
        if file_path:
            try:
                self.excel_data = read_excel(file_path)  # Usar read_excel directamente
                # Actualizar el autocompletar con los nuevos datos
                completer = QCompleter(self.excel_data['Patient'].unique())
                completer.setCaseSensitivity(Qt.CaseInsensitive)
                self.nombre_input.setCompleter(completer)
    
                # Guardar la ruta del archivo en config.json
                self.save_last_excel_path(file_path)
            except FileNotFoundError:
                QMessageBox.warning(self, "Error", "Archivo no encontrado.")
            except ValueError as ve:
                QMessageBox.warning(self, "Error", f"Valor incorrecto en el archivo: {str(ve)}")
            except Exception as e:
                QMessageBox.warning(self, "Error", f"No se pudo cargar el archivo: {str(e)}")
    
    def save_last_excel_path(self, path):
        config = {}
        if os.path.exists('config.json'):
            with open('config.json', 'r') as f:
                try:
                    config = json.load(f)
                except json.JSONDecodeError:
                    config = {}
    
        config['last_excel'] = path
        with open('config.json', 'w') as f:
            json.dump(config, f)
    
    def load_excel_data(self):
        if os.path.exists('config.json'):
            with open('config.json', 'r') as f:
                try:
                    config = json.load(f)
                except json.JSONDecodeError:
                    config = {}
                last_excel = config.get('last_excel', '')
                if os.path.exists(last_excel):
                    return read_excel(last_excel)  # Usar read_excel directamente
        return None

    def update_diagnosticos_previos(self, nombre):
        if self.excel_data is not None:
            patient_data = self.excel_data[self.excel_data['Patient'] == nombre]
            diagnosticos = []
            for _, row in patient_data.iterrows():
                diagnostico = row['Diagnosis']
                service_conected = row['ServiceConected']
                
                if service_conected:
                    diagnostico += ' (Cubierto)'
                else:
                    diagnostico += ' (No Cubierto)'
                    
                diagnosticos.append(diagnostico)
            
            new_diagnosticos = '\n'.join(diagnosticos)
            
            # Si hay diagnósticos nuevos, actualizar
            if new_diagnosticos:
                self.current_diagnosticos = new_diagnosticos
                self.diagnosticos_previos_input.setText(self.current_diagnosticos)
            # Si no hay diagnósticos nuevos, mantener los actuales
            elif self.current_diagnosticos:
                self.diagnosticos_previos_input.setText(self.current_diagnosticos)
        else:
            # Si no hay datos de Excel, mantener los diagnósticos actuales
            if not self.current_diagnosticos:
                self.diagnosticos_previos_input.clear()

    def load_insurance_options(self):
        session = Session()
        names = [a.nombre for a in session.query(InsuranceName).order_by(InsuranceName.nombre).all()]
        session.close()

        if not names:
            names = ["ARS Humano", "ARS Senasa", "ARS Universal", "Privada"]
        if "Ninguna" in names:
            names.remove("Ninguna")
        names.insert(0, "Ninguna")
        self.ars_input.clear()
        self.ars_input.addItems(names)


    def fill_data(self):
        """
        Carga en los widgets toda la información almacenada del paciente.
        Si algún atributo viene como None, se coloca texto vacío.
        """
        p = self.patient

        self.load_insurance_options()

        # — Datos básicos —
        self.nombre_input.blockSignals(True)
        self.nombre_input.setText(p.nombre or "")
        self.nombre_input.blockSignals(False)
        self.edad_input.setText(str(p.edad or ""))
        self.sexo_input.setCurrentText(p.sexo or "Masculino")
        self.alergias_input.setText(p.alergias or "")
        self.direccion_input.setText(p.direccion or "")
        self.telefono_input.setText(getattr(p, 'telefono', '') or "")
        self.id_familia_input.setText(getattr(p, 'id_familia', '') or "")

        # — Datos fijos del SNS —
        self.ficha_familiar_input.setText(p.ficha_familiar or "")
        self.documento_id_input.setText(p.documento_id or "")
        self.nacionalidad_input.setCurrentText(p.nacionalidad or "Dominicana")
        self.no_seguro_input.setText(p.no_seguridad_soc or "")
        self.tipo_sangre_input.setCurrentText(p.tipo_sangre or "No sabe")
        self.ars_input.setCurrentText(p.ars or "Ninguna")
    
        # — Datos clínicos/biográficos —
        self.estado_civil_input.setCurrentText(p.estado_civil or "Soltero/a")
        self.religion_input.setText(p.religion or "")
        self.lugar_nacimiento_input.setText(p.lugar_nacimiento or "")
        self.fecha_nacimiento_input.blockSignals(True)
        if p.fecha_nacimiento:
            self.fecha_nacimiento_input.setDate(p.fecha_nacimiento)
            self.birth_date_changed = True
        else:
            self.birth_date_changed = False
        self.fecha_nacimiento_input.blockSignals(False)
        self.ocupacion_input.setText(p.ocupacion or "")
        if p.altura is not None:
            display_height = self.convert_from_cm(p.altura, self.current_height_unit)
            self.height_input.setText(f"{display_height:.2f}")
        else:
            self.height_input.clear()
        self.diagnosticos_previos_input.setText(p.diagnosticos_previos or "")
        self.medicamentos_continuos_input.setText(p.medicamentos_continuos or "")


        # Mantener referencia si quieres controlarla en otro sitio
        self.current_diagnosticos = p.diagnosticos_previos or ""

        # Mantener diagnósticos previos actuales; solo se actualizarán si el usuario cambia el nombre

    

    def save_patient(self):
        """
        Guarda (crea o actualiza) el paciente.  Funciona tanto si el diálogo
        se abrió para un paciente nuevo como si se abrió para editar.
        Utiliza la misma sesión que maneja la ventana principal.
        """
        # 1. Recuperar la sesión que ya usa el main window
        session = self.parent().session if self.parent() and hasattr(self.parent(), "session") \
                  else Session()
    
        try:
            # 2. Si el paciente existe re-anexarlo a la sesión con merge
            if self.patient and self.patient.id:
                p = session.merge(self.patient)   # Vuelve a estar “managed”
            else:
                p = Paciente()
                session.add(p)
                self.patient = p                  # Para referencia futura
    
            # 3. Copiar TODOS los campos
            p.nombre      = self.nombre_input.text().strip()
            p.edad        = int(self.edad_input.text() or 0)
            p.sexo        = self.sexo_input.currentText()
            p.alergias    = self.alergias_input.text().strip()
            p.direccion   = self.direccion_input.text().strip()
            p.telefono    = self.telefono_input.text().strip()
            p.id_familia  = self.id_familia_input.text().strip()

            p.ficha_familiar    = self.ficha_familiar_input.text().strip()
            p.documento_id      = self.documento_id_input.text().strip()
            p.nacionalidad      = self.nacionalidad_input.currentText()
            p.no_seguridad_soc  = self.no_seguro_input.text().strip()
            p.tipo_sangre       = self.tipo_sangre_input.currentText()
            p.ars               = self.ars_input.currentText()
    
            p.estado_civil      = self.estado_civil_input.currentText()
            p.religion          = self.religion_input.text().strip()
            p.lugar_nacimiento  = self.lugar_nacimiento_input.text().strip()
            if self.birth_date_changed:
                p.fecha_nacimiento = self.fecha_nacimiento_input.date().toPyDate()
            elif self.patient:
                p.fecha_nacimiento = self.patient.fecha_nacimiento
            else:
                p.fecha_nacimiento = None
            p.ocupacion         = self.ocupacion_input.text().strip()
            try:
                height_val = float(self.height_input.text() or 0)
            except ValueError:
                height_val = None
            if height_val is not None:
                p.altura = self.convert_to_cm(height_val, self.current_height_unit)
            else:
                p.altura = None
            p.diagnosticos_previos   = self.diagnosticos_previos_input.toPlainText().strip()
            p.medicamentos_continuos = self.medicamentos_continuos_input.toPlainText().strip()

            # 4. Commit definitivo
            session.commit()
            QMessageBox.information(self, "Éxito", "Datos del paciente guardados.")
    
            # 5. Refrescar información en la ventana principal
            if self.parent():
                self.parent().load_patient_details(p.nombre)
    
            self.accept()  # Cerrar diálogo
    
        except Exception as e:
            session.rollback()
            QMessageBox.critical(self, "Error",
                                 f"No se pudieron guardar los cambios:\n{e}")
        finally:
            # Solo cerramos la sesión si la creó este diálogo
            if not (self.parent() and hasattr(self.parent(), "session")):
                session.close()



    
    def get_patient_by_id(self, patient_id):
        session = Session()
        try:
            patient = session.get(Paciente, patient_id)
            return patient
        except Exception as e:
            QMessageBox.critical(None, "Error", f"Error al obtener el paciente por ID: {str(e)}")
        finally:
            session.close()


def create_signature_image(name, path):
    """Create a signature image from the doctor's name in blue."""
    width, height = 600, 120
    pix = QPixmap(width, height)
    pix.fill(Qt.transparent)
    painter = QPainter(pix)
    painter.setRenderHint(QPainter.Antialiasing)
    font = QFont("Segoe Script", 40)
    fm = QFontMetrics(font)
    # Reduce the font size until the name fits within the image
    while fm.horizontalAdvance(name) > width - 40 and font.pointSize() > 10:
        font.setPointSize(font.pointSize() - 2)
        fm = QFontMetrics(font)
    painter.setFont(font)
    painter.setPen(QPen(Qt.blue, 2))
    painter.drawText(pix.rect(), Qt.AlignCenter, name)
    painter.end()
    pix.save(path)
    return path


def create_seal_image(name, specialty, path):
    """Create a round seal image with the doctor's details."""
    size = 220
    pix = QPixmap(size, size)
    pix.fill(Qt.transparent)
    painter = QPainter(pix)
    painter.setRenderHint(QPainter.Antialiasing)
    pen = QPen(Qt.black, 3)
    painter.setPen(pen)
    painter.drawEllipse(3, 3, size - 6, size - 6)
    painter.setFont(QFont("Arial", 11, QFont.Bold))
    painter.drawText(pix.rect(), Qt.AlignCenter, f"{name}\n{specialty}")
    painter.end()
    pix.save(path)
    return path


def create_qr_code(data, path):
    """Draw a small but modern-looking QR code without external packages."""
    import hashlib
    size = 21
    bits = bin(int(hashlib.sha256(data.encode()).hexdigest(), 16))[2:].zfill(size * size)

    module = 4
    border = 4
    dim = (size + border * 2) * module
    pix = QPixmap(dim, dim)
    pix.fill(Qt.white)
    p = QPainter(pix)
    p.setPen(Qt.NoPen)
    p.setBrush(Qt.black)

    def draw_finder(x, y):
        s = 7 * module
        p.fillRect(x, y, s, s, Qt.black)
        p.fillRect(x + module, y + module, 5 * module, 5 * module, Qt.white)
        p.fillRect(x + 2 * module, y + 2 * module, 3 * module, 3 * module, Qt.black)

    draw_finder(border * module, border * module)
    draw_finder(border * module, (size - 7 + border) * module)
    draw_finder((size - 7 + border) * module, border * module)

    idx = 0
    for y in range(size):
        for x in range(size):
            if bits[idx] == '1':
                px = (x + border) * module
                py = (y + border) * module
                p.fillRect(px, py, module, module, Qt.black)
            idx += 1
    p.end()

    pix = pix.scaled(96, 96, Qt.KeepAspectRatio, Qt.SmoothTransformation)
    pix.save(path)
    return path


class InformeHistoriasDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.file-medical')
        self.setWindowTitle("Informe de Historias por Periodo")
        self.resize(700, 300)
        lay = QVBoxLayout(self)

        h = QHBoxLayout()
        h.setSpacing(15)
        h.setSpacing(15)
        h.addWidget(QLabel("Desde:"))
        self.desde = QDateEdit(calendarPopup=True)
        self.desde.setDate(QDate.currentDate())
        h.addWidget(self.desde)
        h.addWidget(QLabel("Hasta:"))
        self.hasta = QDateEdit(calendarPopup=True)
        self.hasta.setDate(QDate.currentDate())
        h.addWidget(self.hasta)
        lay.addLayout(h)

        self.btn_gen = QPushButton("Generar")
        StyleHelper.style_button(self.btn_gen)
        lay.addWidget(self.btn_gen)
        self.btn_gen.clicked.connect(self.generar)

    def generar(self):
        ini = self.desde.date().toPyDate()
        fin = self.hasta.date().toPyDate()
        dest = QFileDialog.getExistingDirectory(self, "Carpeta destino")
        if not dest:
            return
        session = Session()
        day = ini
        while day <= fin:
            next_day = day + timedelta(days=1)
            hists = (
                session.query(HistoriaClinica)
                .filter(HistoriaClinica.fecha.between(day, next_day))
                .all()
            )
            if hists:
                ddir = os.path.join(dest, day.strftime('%d-%m-%Y'))
                os.makedirs(ddir, exist_ok=True)
                doc = Document()
                for hist in hists:
                    doc.add_heading(
                        f"{hist.paciente.nombre} {hist.fecha.strftime('%d/%m/%Y %H:%M')}",
                        level=2,
                    )
                    doc.add_paragraph(hist.historia_enfermedad or '')
                doc.save(os.path.join(ddir, f"historias_{day.strftime('%d-%m-%Y')}.docx"))
            day = next_day
        session.close()
        QMessageBox.information(self, "Éxito", "Informe generado")
        self.accept()


class FixedMedicationDialog(QDialog):
    def __init__(self, patient, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.pills')
        self.patient = patient
        self.setWindowTitle("Medicación Fija")

        # Match the larger dimensions of the HistoryQueryDialog window
        self.setMinimumSize(1000, 1100)
        self.resize(1000, 1100)
        self.setStyleSheet(
            """
            background-color: #f9f9f9;
            border-radius: 12px;
            border: 1px solid #ccc;
            """
        )
        StyleHelper.add_drop_shadow(self)

        layout = QVBoxLayout(self)

        title = QLabel("💊 Medicación Fija")
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("font-size: 16pt; font-weight: bold;")
        layout.addWidget(title)

        self.text_edit = QTextEdit()
        self.text_edit.setStyleSheet(
            """
            border-radius: 8px;
            border: 1px solid #ccc;
            padding: 8px;
            background-color: #ffffff;
            """
        )
        self.text_edit.setPlainText(patient.medicamentos_continuos or "")
        layout.addWidget(self.text_edit)

        self.btn_guardar = QPushButton("Guardar")
        self.btn_guardar.clicked.connect(self.save)
        self.btn_guardar.setStyleSheet(
            """
            QPushButton {
                background-color: #009688;
                color: white;
                border-radius: 6px;
                padding: 6px 12px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #00796B;
            }
            """
        )
        layout.addWidget(self.btn_guardar, alignment=Qt.AlignRight)

    def save(self):
        text = self.text_edit.toPlainText().strip()
        session = Session()
        try:
            p = session.get(Paciente, self.patient.id)
            if p:
                p.medicamentos_continuos = text
                session.commit()
                self.patient.medicamentos_continuos = text
        except Exception as e:
            session.rollback()
            QMessageBox.critical(self, "Error", f"No se pudo guardar: {e}")
            return
        finally:
            session.close()
        self.accept()

  

# ------------------------------------------------------------
#   DIÁLOGO: EXPORTAR REGISTRO DIARIO DE CONSULTA A PDF
# ------------------------------------------------------------
class ReporteRegistroDialog(QDialog):
    """
    Ventana para seleccionar un rango de fechas y generar el PDF
    con el Registro Diario de Consulta (formato SNS).
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.file-pdf')
        self.setWindowTitle("Registro Diario de Consulta – Exportar PDF")
        self.resize(780, 220)

        # --- Controles de fecha ---
        lay = QVBoxLayout(self)

        h = QHBoxLayout()
        h.addWidget(QLabel("Desde:"))
        self.desde = QDateEdit(calendarPopup=True)
        self.desde.setDate(QDate.currentDate())
        h.addWidget(self.desde)

        h.addWidget(QLabel("Hasta:"))
        self.hasta = QDateEdit(calendarPopup=True)
        self.hasta.setDate(QDate.currentDate())
        h.addWidget(self.hasta)
        lay.addLayout(h)

        # --- Botón generar ---
        self.btn_pdf = QPushButton("Generar PDF")
        self.btn_pdf.clicked.connect(self.generar_pdf)
        StyleHelper.style_button(self.btn_pdf)
        lay.addWidget(self.btn_pdf)



    # ------------------------------------------------------------------
    #   MÉTODO COMPLETO  – ReporteRegistroDialog.generar_pdf
    # ------------------------------------------------------------------


    def generar_pdf(self):
        """PDF: Dirección envuelta, Diagnóstico ≥ 200 pt y línea divisoria por paciente."""
        from reportlab.lib.pagesizes import A3, landscape
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
        import re
    
        safe = lambda v: "" if v is None else str(v)
    
        ini = self.desde.date().toPyDate()
        fin = self.hasta.date().toPyDate()
        fin_dt = datetime.combine(fin, datetime.max.time())
    
        session = Session()
        try:
            rows = (session.query(RegistroConsulta, HistoriaClinica, Paciente)
                    .join(HistoriaClinica, RegistroConsulta.historia_id == HistoriaClinica.id)
                    .join(Paciente,        RegistroConsulta.paciente_id == Paciente.id)
                    .filter(RegistroConsulta.fecha.between(ini, fin_dt))
                    .order_by(RegistroConsulta.fecha.asc())
                    .all())
            if not rows:
                QMessageBox.information(self, "Sin datos",
                                        "No hay registros en ese rango.")
                return
    
            fname, _ = QFileDialog.getSaveFileName(
                self, "Guardar PDF",
                f"registro_{ini:%Y-%m-%d}_{fin:%Y-%m-%d}.pdf", "PDF (*.pdf)")
            if not fname:
                return
    
            # ───────── Canvas y título ─────────
            c = canvas.Canvas(fname, pagesize=landscape(A3))
            W, H = landscape(A3)
            margin_x = 1 * cm
    
            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(
                W / 2, H - 2 * cm,
                f"REGISTRO DIARIO DE CONSULTA ({ini:%d/%m/%Y} - {fin:%d/%m/%Y})"
            )
    
            # ───────── Cabeceras y anchos ──────
            heads = [
                "No.", "FECHA", "PACIENTE", "FICHA", "CÉDULA/NUI", "NO. SEGURO",
                "EDAD", "SEXO", "NAC.", "DIRECCIÓN", "RÉGIMEN", "ARS",
                "LUGAR", "FRECUENCIA", "SERVICIO", "TIPO ATENCIÓN",
                "DIAGNÓSTICO",
            ]

            w = [
                30, 55, 80, 45, 70, 60,
                30, 30, 45, 110, 60, 70,
                60, 70, 80, 70,
                None,
            ]  # Dirección 110 pt
    
            min_diag = 200
            ancho_total = W - 2 * margin_x
            restante = ancho_total - sum(w[:-1])
            if restante < min_diag:
                falta = min_diag - restante
                reducible = max(30, w[14] - 20)      # reducimos SERVICIO
                w[14] -= min(reducible, falta)
                restante = ancho_total - sum(w[:-1])
            w[-1] = max(min_diag, restante)
    
            # posiciones X absolutas
            x_pos = [margin_x]
            for width in w[:-1]:
                x_pos.append(x_pos[-1] + width)
    
            def header(y0):
                c.setFont("Helvetica-Bold", 7.1)
                for x, t in zip(x_pos, heads):
                    c.drawString(x, y0, t)
                return y0 - 0.34 * cm
    
            def wrap(text, maxw):
                maxw -= 4
                lines, line = [], ""
                for wd in text.split():
                    probe = (line + " " + wd).strip()
                    if c.stringWidth(probe, "Helvetica", 7) > maxw:
                        if line:
                            lines.append(line)
                        line = wd
                    else:
                        line = probe
                if line:
                    lines.append(line)
                return lines or [""]
    
            y = header(H - 3 * cm)
            c.setFont("Helvetica", 7)
            num = 1
            current_day = None

            for reg, hist, pac in rows:
                day = reg.fecha.date()
                if day != current_day:
                    c.setFont("Helvetica-Bold", 8)
                    c.drawString(margin_x, y, day.strftime("%d/%m/%Y"))
                    y -= 0.4 * cm
                    c.setFont("Helvetica", 7)
                    current_day = day
                direc_lines = wrap(safe(reg.direccion), w[9])
                seen = set()
                unique_diags = []
                for d in hist.diagnosticos:
                    desc = d.descripcion or ""
                    short = re.split(r'[;\n]', desc)[0].strip()
                    if short and short not in seen:
                        unique_diags.append(short)
                        seen.add(short)
                diag_text = "; ".join(unique_diags)
                if diag_text.strip().lower() == "sin diagnóstico explícito".lower():
                    diag_text = "Consulta de persona sana"
                diag_lines = wrap(diag_text, w[-1])
                name_lines = wrap(safe(pac.nombre), w[2])

                signo = (
                    session.query(SignoVital)
                    .filter_by(historia_id=hist.id)
                    .order_by(SignoVital.fecha.desc())
                    .first()
                )
                pa = signo.presion if signo and signo.presion else ""
                fc = str(signo.frecuencia) if signo and signo.frecuencia else ""
                peso = str(signo.peso) if signo and signo.peso is not None else ""
                imc_val = f"{signo.imc:.1f}" if signo and signo.imc else ""
                gly = str(signo.glicemia) if signo and signo.glicemia is not None else ""
                chol = str(signo.colesterol) if signo and signo.colesterol is not None else ""

                vs_parts = []
                if pa:
                    vs_parts.append(f"PA: {pa}")
                if fc:
                    vs_parts.append(f"FC: {fc}")
                if peso:
                    vs_parts.append(f"PESO: {peso}")
                if imc_val:
                    vs_parts.append(f"IMC: {imc_val}")
                if gly:
                    vs_parts.append(f"Glicemia: {gly}")
                if chol:
                    vs_parts.append(f"Colesterol: {chol}")
                vs_text = "  ".join(vs_parts)

                if vs_text:
                    diag_text = diag_text + "\n" + vs_text
                diag_lines = wrap(diag_text, w[-1])

                n_lines = max(len(direc_lines), len(diag_lines), len(name_lines))
                row_h   = 0.30 * cm * n_lines
    
                # salto de página
                if y - row_h - 0.12 * cm < 2 * cm:
                    c.showPage()
                    y = header(H - 2.5 * cm)
                    c.setFont("Helvetica", 7)
    
                # ── columnas fijas (hasta NAC.) ──
                fixed = [
                    num, reg.fecha.strftime("%d/%m/%Y"), "",
                    safe(reg.ficha_familiar), safe(reg.documento_id),
                    safe(reg.no_seguridad_soc), safe(reg.edad), safe(reg.sexo),
                    safe(reg.nacionalidad)
                ]
                for x, txt in zip(x_pos, fixed):
                    c.drawString(x, y, safe(txt))
                for i in range(len(name_lines)):
                    c.drawString(x_pos[2], y - i*0.30*cm, name_lines[i])
    
                # ── imprimir Dirección + Diagnóstico línea a línea ──
                for i in range(n_lines):
                    # Dirección
                    c.drawString(x_pos[9], y - i * 0.30 * cm,
                                 direc_lines[i] if i < len(direc_lines) else "")
                    # Columnas simples que sólo ocupan la primera línea
                    if i == 0:
                        rest = [
                            safe(reg.tipo_sangre),
                            safe(reg.ars),
                            safe(reg.lugar_consulta),
                            safe(reg.frecuencia),
                            safe(reg.servicio),
                            safe(reg.tipo_atencion),
                        ]
                        for x, txt in zip(x_pos[10:-1], rest):
                            c.drawString(x, y, txt)
                    # Diagnóstico
                    c.drawString(x_pos[-1], y - i * 0.30 * cm,
                                 diag_lines[i] if i < len(diag_lines) else "")
    
                sep_y = y - row_h - 0.04 * cm
                y = sep_y - 0.04 * cm
                num += 1
    
            c.save()
            QMessageBox.information(self, "Éxito", f"PDF guardado en:\n{fname}")
            self.accept()
    
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudo exportar:\n{e}")
        finally:
            session.close()

        




class ConfigReceta(QWidget):
    def __init__(self):
        super().__init__()
        self.config_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")

        # Cargar configuración existente o usar valores predeterminados
        self.load_config()
        
        # Variables para rastrear cambios
        self.original_logo_path = self.logo_path
        self.original_footer_text = self.footer_text
        self.original_footer_details = self.footer_details
        self.original_footer_phone = self.footer_phone
        self.original_signature_path = self.signature_path
        self.original_seal_path = self.seal_path
        self.original_doctor_name = self.doctor_name
        self.original_doctor_specialty = self.doctor_specialty
        self.original_clinic_header = self.clinic_header
        
        StyleHelper.set_window_icon(self, 'fa5s.cogs')

        # Crear la interfaz gráfica
        self.init_ui()


    def init_ui(self):
        layout = QVBoxLayout()

        # Grupo para el logo
        logo_group = QGroupBox("Logo")
        logo_layout = QVBoxLayout()

        # Etiqueta para mostrar el logo
        self.logo_display = QLabel()
        self.logo_display.setMinimumSize(100, 100)
        self.logo_display.setAlignment(Qt.AlignCenter)
        self.update_logo_display()
        logo_layout.addWidget(self.logo_display)

        # Botón para cargar el logo
        self.btn_cargar_logo = QPushButton("Cargar nuevo logo", self)
        self.btn_cargar_logo.clicked.connect(self.cargar_logo)
        StyleHelper.style_button(self.btn_cargar_logo)
        logo_layout.addWidget(self.btn_cargar_logo)

        logo_group.setLayout(logo_layout)
        layout.addWidget(logo_group)

        header_group = QGroupBox("Encabezado")
        header_layout = QFormLayout()
        self.clinic_header_edit = QLineEdit(self.clinic_header)
        StyleHelper.style_input(self.clinic_header_edit)
        header_layout.addRow("Nombre del Centro:", self.clinic_header_edit)
        header_group.setLayout(header_layout)
        layout.addWidget(header_group)

        signature_group = QGroupBox("Firma Digital")
        sign_layout = QVBoxLayout()
        self.signature_display = QLabel()
        self.signature_display.setMinimumSize(100, 50)
        self.signature_display.setAlignment(Qt.AlignCenter)
        self.update_signature_display()
        sign_layout.addWidget(self.signature_display)
        self.btn_cargar_firma = QPushButton("Cargar firma", self)
        self.btn_cargar_firma.clicked.connect(self.cargar_firma)
        StyleHelper.style_button(self.btn_cargar_firma)
        sign_layout.addWidget(self.btn_cargar_firma)
        self.btn_eliminar_firma = QPushButton("Eliminar firma", self)
        self.btn_eliminar_firma.clicked.connect(self.eliminar_firma)
        StyleHelper.style_button(self.btn_eliminar_firma)
        sign_layout.addWidget(self.btn_eliminar_firma)
        signature_group.setLayout(sign_layout)
        layout.addWidget(signature_group)

        doctor_group = QGroupBox("Datos del Médico")
        doctor_layout = QFormLayout()
        self.doctor_name_edit = QLineEdit(self.doctor_name)
        StyleHelper.style_input(self.doctor_name_edit)
        doctor_layout.addRow("Nombre:", self.doctor_name_edit)
        self.doctor_spec_edit = QLineEdit(self.doctor_specialty)
        StyleHelper.style_input(self.doctor_spec_edit)
        doctor_layout.addRow("Especialidad:", self.doctor_spec_edit)
        doctor_group.setLayout(doctor_layout)
        layout.addWidget(doctor_group)

        seal_group = QGroupBox("Sello Digital")
        seal_layout = QVBoxLayout()
        self.seal_display = QLabel()
        self.seal_display.setMinimumSize(100, 100)
        self.seal_display.setAlignment(Qt.AlignCenter)
        self.update_seal_display()
        seal_layout.addWidget(self.seal_display)
        self.btn_cargar_sello = QPushButton("Cargar sello", self)
        self.btn_cargar_sello.clicked.connect(self.cargar_sello)
        StyleHelper.style_button(self.btn_cargar_sello)
        seal_layout.addWidget(self.btn_cargar_sello)
        self.btn_eliminar_sello = QPushButton("Eliminar sello", self)
        self.btn_eliminar_sello.clicked.connect(self.eliminar_sello)
        StyleHelper.style_button(self.btn_eliminar_sello)
        seal_layout.addWidget(self.btn_eliminar_sello)
        seal_group.setLayout(seal_layout)
        layout.addWidget(seal_group)

        # Grupo para el pie de página
        footer_group = QGroupBox("Pie de Página")
        footer_layout = QVBoxLayout()

        # Etiqueta para mostrar la información del pie de página
        self.footer_label = QLabel(f"Dirección: {self.footer_text}\nHorario: {self.footer_details}\nTeléfono: {self.footer_phone}")
        self.footer_label.setWordWrap(True)
        StyleHelper.style_label(self.footer_label)
        footer_layout.addWidget(self.footer_label)

        # Botón para modificar el pie de página
        self.btn_modificar_footer = QPushButton("Modificar pie de página", self)
        self.btn_modificar_footer.clicked.connect(self.modificar_footer)
        StyleHelper.style_button(self.btn_modificar_footer)
        footer_layout.addWidget(self.btn_modificar_footer)

        footer_group.setLayout(footer_layout)
        layout.addWidget(footer_group)

        # Botón para guardar cambios
        self.btn_guardar_cambios = QPushButton("Guardar Cambios", self)
        self.btn_guardar_cambios.clicked.connect(self.guardar_cambios)
        StyleHelper.style_button(self.btn_guardar_cambios)
        layout.addWidget(self.btn_guardar_cambios)

        self.setLayout(layout)
        self.setWindowTitle('Configuración de Receta')
        self.setGeometry(300, 300, 400, 400)

    def update_signature_display(self):
        if self.signature_path and os.path.exists(self.signature_path):
            pix = QPixmap(self.signature_path)
            if not pix.isNull():
                scaled = pix.scaled(200, 80, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                self.signature_display.setPixmap(scaled)
                return
        self.signature_display.setText("Sin firma")

    def update_seal_display(self):
        if self.seal_path and os.path.exists(self.seal_path):
            pix = QPixmap(self.seal_path)
            if not pix.isNull():
                scaled = pix.scaled(120, 120, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                self.seal_display.setPixmap(scaled)
                return
        self.seal_display.setText("Sin sello")

    def cargar_firma(self):
        path, _ = QFileDialog.getOpenFileName(self, "Seleccionar Firma", "", "Imágenes (*.png *.jpg *.bmp)")
        if path:
            self.signature_path = path
            self.update_signature_display()

    def eliminar_firma(self):
        if self.signature_path and os.path.exists(self.signature_path):
            try:
                os.remove(self.signature_path)
            except Exception:
                pass
        self.signature_path = ''
        self.update_signature_display()

    def cargar_sello(self):
        path, _ = QFileDialog.getOpenFileName(self, "Seleccionar Sello", "", "Imágenes (*.png *.jpg *.bmp)")
        if path:
            self.seal_path = path
            self.update_seal_display()

    def eliminar_sello(self):
        if self.seal_path and os.path.exists(self.seal_path):
            try:
                os.remove(self.seal_path)
            except Exception:
                pass
        self.seal_path = ''
        self.update_seal_display()

    def load_config(self):
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r') as f:
                    config = json.load(f)
                    self.logo_path = config.get('logo_path', "logo.png")
                    self.footer_text = config.get('footer_text', "Santiago de los Caballeros 51000")
                    self.footer_details = config.get('footer_details', "Mon - Fri, 8:00 AM to 6:00 PM")
                    self.footer_phone = config.get('footer_phone', "Tel: 809-555-5555")
                    self.clinic_header = config.get('clinic_header', '')
                    self.signature_path = config.get('signature_path', '')
                    self.seal_path = config.get('seal_path', '')
                    self.doctor_name = config.get('doctor_name', '')
                    self.doctor_specialty = config.get('doctor_specialty', '')
                    if self.doctor_name and not self.signature_path:
                        self.signature_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'signature.png')
                        create_signature_image(self.doctor_name, self.signature_path)
                    if self.doctor_name and not self.seal_path:
                        self.seal_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'seal.png')
                        create_seal_image(self.doctor_name, self.doctor_specialty, self.seal_path)
                    self.original_signature_path = self.signature_path
                    self.original_seal_path = self.seal_path
                    self.original_doctor_name = self.doctor_name
                    self.original_doctor_specialty = self.doctor_specialty
            else:
                self.logo_path = "logo.png"
                self.footer_text = "Santiago de los Caballeros 51000"
                self.footer_details = "Mon - Fri, 8:00 AM to 6:00 PM"
                self.footer_phone = "Tel: 809-555-5555"
                self.clinic_header = ""
                self.signature_path = ''
                self.seal_path = ''
                self.doctor_name = ''
                self.doctor_specialty = ''
                self.original_signature_path = ''
                self.original_seal_path = ''
                self.original_doctor_name = ''
                self.original_doctor_specialty = ''
                self.original_clinic_header = ''
        except Exception as e:
            print(f"Error loading config: {e}")
            self.logo_path = "logo.png"
            self.footer_text = "Santiago de los Caballeros 51000"
            self.footer_details = "Mon - Fri, 8:00 AM to 6:00 PM"
            self.footer_phone = "Tel: 809-555-5555"

    def save_config(self):
        config = {
            'logo_path': self.logo_path,
            'footer_text': self.footer_text,
            'footer_details': self.footer_details,
            'footer_phone': self.footer_phone,
            'clinic_header': self.clinic_header,
            'signature_path': self.signature_path,
            'seal_path': self.seal_path,
            'doctor_name': self.doctor_name,
            'doctor_specialty': self.doctor_specialty
        }
        try:
            with open(self.config_file, 'w') as f:
                json.dump(config, f)
        except Exception as e:
            QMessageBox.warning(self, "Error", f"No se pudo guardar la configuración: {str(e)}")

    def update_logo_display(self):
            try:
                if os.path.exists(self.logo_path):
                    pixmap = QPixmap(self.logo_path)
                    if not pixmap.isNull():
                        scaled_pixmap = pixmap.scaled(100, 100, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                        self.logo_display.setPixmap(scaled_pixmap)
                    else:
                        self.logo_display.setText("Error al cargar la imagen")
                else:
                    self.logo_display.setText("Logo no encontrado")
            except Exception as e:
                self.logo_display.setText(f"Error: {str(e)}")


    def cargar_logo(self):
        logo_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar Logo", "", "Imágenes (*.png *.jpg *.bmp)")
        if logo_path:
            self.logo_path = logo_path
            self.update_logo_display()

    def modificar_footer(self):
        footer_text, ok1 = QInputDialog.getText(self, 'Modificar Pie de Página', 
                                              'Ingrese la nueva dirección:', 
                                              text=self.footer_text)
        
        if ok1:
            footer_details, ok2 = QInputDialog.getText(self, 'Modificar Detalles', 
                                                     'Ingrese los detalles (horario, etc.):', 
                                                     text=self.footer_details)
            if ok2:
                footer_phone, ok3 = QInputDialog.getText(self, 'Modificar Teléfono', 
                                                       'Ingrese el número de teléfono:', 
                                                       text=self.footer_phone)
                if ok3:
                    self.footer_text = footer_text
                    self.footer_details = footer_details
                    self.footer_phone = footer_phone
                    self.footer_label.setText(f"Dirección: {self.footer_text}\nHorario: {self.footer_details}\nTeléfono: {self.footer_phone}")

    def guardar_cambios(self):
        changes_made = False
        self.doctor_name = self.doctor_name_edit.text()
        self.doctor_specialty = self.doctor_spec_edit.text()
        self.clinic_header = self.clinic_header_edit.text()
        
        # Verificar cambios en el logo
        if self.logo_path != self.original_logo_path:
            try:
                new_logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.png")
                
                # Eliminar el logo anterior si existe
                if os.path.exists(new_logo_path):
                    os.remove(new_logo_path)
                
                # Copiar el nuevo logo
                shutil.copy(self.logo_path, new_logo_path)
                self.original_logo_path = self.logo_path
                changes_made = True
            except Exception as e:
                QMessageBox.warning(self, "Error", f"No se pudo guardar el logo: {str(e)}")
                return

        # Verificar cambios en el pie de página
        footer_changed = (self.footer_text != self.original_footer_text or
                        self.footer_details != self.original_footer_details or
                        self.footer_phone != self.original_footer_phone)
        
        if footer_changed:
            self.original_footer_text = self.footer_text
            self.original_footer_details = self.footer_details
            self.original_footer_phone = self.footer_phone
            changes_made = True

        if self.signature_path != self.original_signature_path:
            self.original_signature_path = self.signature_path
            changes_made = True
        if self.seal_path != self.original_seal_path:
            self.original_seal_path = self.seal_path
            changes_made = True
        if (self.doctor_name != self.original_doctor_name or
                self.doctor_specialty != self.original_doctor_specialty or
                self.clinic_header != self.original_clinic_header):
            self.original_doctor_name = self.doctor_name
            self.original_doctor_specialty = self.doctor_specialty
            self.original_clinic_header = self.clinic_header
            sig_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'signature.png')
            seal_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'seal.png')
            try:
                create_signature_image(self.doctor_name, sig_path)
                create_seal_image(self.doctor_name, self.doctor_specialty, seal_path)
                self.signature_path = sig_path
                self.seal_path = seal_path
                self.update_signature_display()
                self.update_seal_display()
            except Exception as e:
                print(f'Error creando firma o sello: {e}')
            changes_made = True

        if changes_made:
            self.save_config()
            QMessageBox.information(self, "Éxito", "Cambios guardados exitosamente.")
        else:
            QMessageBox.information(self, "Sin Cambios", "No se ha realizado ningún cambio.")

    def get_logo(self):
        return self.logo_path

    def get_footer(self):
        return self.footer_text, self.footer_details, self.footer_phone
    
    def get_footer_text(self):
        return self.footer_text

    def get_footer_details(self):
        return self.footer_details
    
    def get_footer_phone(self):
        return self.footer_phone

    def get_signature(self):
        return self.signature_path

    def get_seal(self):
        return self.seal_path

    def get_doctor_name(self):
        return self.doctor_name

    def get_doctor_specialty(self):
        return self.doctor_specialty

    def get_clinic_header(self):
        return self.clinic_header

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        # Ocultar la ventana principal al inicio
        self.hide()

        # ↳ Ejecuta el back-fill solo al arrancar la app
        self.backfill_registro_consultas()

        #self._clean_old_registros()   # <- luego comenta esta línea
  
        
        # Inicializar la configuración
        self.model = ""
        self.current_language = 'es'
        self.current_theme = 'gpt'
        self.load_config()
        
        # Verificar autenticación antes de mostrar la ventana
        self.check_and_show_auth()
    
            # Solo una instancia de ConfigGPT
        self.config_gpt = ConfigGPT()
        self.config_gpt.set_main_window(self)
        self.config_gpt.config_updated.connect(self.update_model)

        # Backup automático cada 6 horas
        self.backup_timer = QTimer(self)
        self.backup_timer.timeout.connect(self.autobackup)
        self.backup_timer.start(6 * 3600 * 1000)

        # Temporizadores para citas y recordatorios
        self.reminder_timer = QTimer(self)
        self.reminder_timer.timeout.connect(self.check_appointments)
        self.alert_timer = QTimer(self)
        self.alert_timer.timeout.connect(self.check_reminders)

        # Background timer to evaluate cardiovascular risk from new vitals
        self.risk_timer = QTimer(self)
        self.risk_timer.timeout.connect(self.process_risk_updates)
        # check every 5 minutes
        self.risk_timer.start(5 * 60 * 1000)
       

        
        # Establecer el título inicial de la ventana
        self.update_window_title()
        
        # Obtener el tamaño de la pantalla
        screen = QApplication.primaryScreen()
        rect = screen.availableGeometry()

        # Ajustar el tamaño de la ventana principal al tamaño de la pantalla
        self.setGeometry(rect)

     
        

     

        
    
        #self.setWindowTitle(f"Sistema de Historias Clínicas {self.model}")
        
          # Establecer el título inicial
        self.update_window_title()
        self.setGeometry(100, 100, 1280, 720)
        self.current_patient = None
        self.current_history_id = None
        self.current_lab_request = ""  # Inicializar el atributo current_lab_request
        self.history_thread = None  # hilo para generacion de historia
        self.current_programs = []

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)
        # Mostrar la ventana maximizada

        import os
        StyleHelper.set_window_icon(self, 'fa5s.home')

              # Crear una instancia de ConfigGPT y conectar la señal
        self.config_gpt = ConfigGPT()
        self.config_gpt.config_updated.connect(self.update_model)



        # Search panel and patient list
        search_panel = QWidget()
        search_layout = QVBoxLayout(search_panel)
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Nombre del paciente")
        StyleHelper.style_input(self.search_input)
        self.search_button = QPushButton("Buscar Paciente")
        StyleHelper.style_button(self.search_button)

        search_box = QHBoxLayout()
        search_box.addWidget(self.search_input)
        search_box.addWidget(self.search_button)
        search_layout.addLayout(search_box)

        self.new_patient_button = QPushButton("Nuevo Paciente")
        StyleHelper.style_button(self.new_patient_button)
        self.edit_patient_button = QPushButton("Editar Paciente")
        StyleHelper.style_button(self.edit_patient_button)
        self.delete_patient_button = QPushButton("Eliminar Paciente")
        StyleHelper.style_button(self.delete_patient_button)

        self.fixed_med_button = QPushButton("Medicacion Fija")
        StyleHelper.style_button(self.fixed_med_button)
        self.vitals_button = QPushButton("Signos Vitales")
        StyleHelper.style_button(self.vitals_button)
        self.family_button = QPushButton("Salud Familiar y Social")
        StyleHelper.style_button(self.family_button)
        self.appointment_button = QPushButton("Citas")
        StyleHelper.style_button(self.appointment_button)
        self.lab_results_button = QPushButton("Resultados de Laboratorio")
        StyleHelper.style_button(self.lab_results_button)
        self.forms_button = QPushButton("Formularios")
        StyleHelper.style_button(self.forms_button)
        self.new_history_button = QPushButton("Historia Clínica")
        StyleHelper.style_button(self.new_history_button)
        self.validate_ai_file_button = QPushButton("Validar Diagnóstico IA (Archivo)")
        StyleHelper.style_button(self.validate_ai_file_button)
        self.billing_button = QPushButton("Facturación")
        StyleHelper.style_button(self.billing_button)
        self.insurance_button = QPushButton("Gestionar ARS")
        StyleHelper.style_button(self.insurance_button)
        self.services_button = QPushButton("Gestionar Servicios")
        StyleHelper.style_button(self.services_button)
        self.analysis_button = QPushButton("Análisis de Datos")
        StyleHelper.style_button(self.analysis_button)
        self.ai_study_button = QPushButton("Estudio Clínico IA")
        StyleHelper.style_button(self.ai_study_button)
        self.search_hist_button = QPushButton("Búsqueda Avanzada")
        StyleHelper.style_button(self.search_hist_button)
        self.reminder_button = QPushButton("Recordatorios")
        StyleHelper.style_button(self.reminder_button)
        self.btn_limpiar_dups = QPushButton("Eliminar Duplicados")
        StyleHelper.style_button(self.btn_limpiar_dups)
        self.btn_registro_pdf = QPushButton("Generar Informe Diario")
        self.btn_registro_pdf.setToolTip("Exportar el Registro Diario de Consulta a PDF")
        StyleHelper.style_button(self.btn_registro_pdf)
        self.btn_hist_period = QPushButton("Informe por Periodo")
        StyleHelper.style_button(self.btn_hist_period)

        self.patient_list = QTableWidget()
        self.patient_list.setColumnCount(3)
        self.patient_list.setHorizontalHeaderLabels(["ID", "Nombre", ""])
        self.patient_list.horizontalHeaderItem(0).setText("ID")
        self.patient_list.horizontalHeaderItem(1).setText("Nombre")
        self.patient_list.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        StyleHelper.style_table(self.patient_list, "#e0e0e0")
        self.patient_list.itemSelectionChanged.connect(self.highlight_selected_patient)

        patient_group = QGroupBox("Gestión del Paciente")
        patient_layout = QHBoxLayout(patient_group)
        patient_layout.addWidget(self.new_patient_button)
        patient_layout.addWidget(self.edit_patient_button)
        patient_layout.addWidget(self.delete_patient_button)
        StyleHelper.style_groupbox(patient_group)
        search_layout.addWidget(patient_group)

        actions_group = QGroupBox("Acciones del Sistema")
        actions_layout = QGridLayout(actions_group)
        actions = [
            self.fixed_med_button,
            self.vitals_button,
            self.appointment_button,
            self.lab_results_button,
            self.forms_button,
            self.reminder_button,
            self.billing_button,
            self.insurance_button,
            self.services_button,
            self.analysis_button,
            self.ai_study_button,
            self.search_hist_button,
            self.btn_limpiar_dups,
            self.btn_registro_pdf,
            self.btn_hist_period,
        ]
        for i, btn in enumerate(actions):
            row = i // 3
            col = i % 3
            actions_layout.addWidget(btn, row, col)
        StyleHelper.style_groupbox(actions_group)
        search_layout.addWidget(actions_group)

        search_layout.addWidget(self.patient_list)

         # Crear el menú de configuración
        self.create_menu()

        # Crea un QCompleter para el campo de búsqueda
        self.patient_completer = QCompleter()
        self.patient_completer.setFilterMode(Qt.MatchContains)

        self.patient_completer.setCaseSensitivity(Qt.CaseInsensitive)  # Añadir esta línea
        self.search_input.setCompleter(self.patient_completer)

        # Conecta la señal activated del QCompleter a una función que cargue los detalles del paciente
        self.patient_completer.activated.connect(self.load_patient_details)

        # Carga los nombres de los pacientes al iniciar la aplicación
        try:
            self.load_patient_names()
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Ocurrió un error al cargar la base de datos: {str(e)}")
            self.database_manager = DatabaseManager(self)
            self.database_manager.show()

        # Agregar un QListWidget para mostrar historias clínicas y recetas
        self.history_list = QListWidget()
        hist_group = QGroupBox("Historias Clínicas")
        hist_layout = QVBoxLayout(hist_group)
        hist_layout.addWidget(self.history_list)

        self.delete_history_button = QPushButton("Eliminar Historia")
        if qta:
            self.delete_history_button.setIcon(qta.icon("fa5s.trash-alt", color="white"))
        StyleHelper.style_button(self.delete_history_button)

        self.db_manager_button = QPushButton("Gestionar Base de Datos")
        if qta:
            self.db_manager_button.setIcon(qta.icon("fa5s.database", color="white"))
        StyleHelper.style_button(self.db_manager_button)

        history_grid = QGridLayout()
        history_grid.addWidget(self.new_history_button, 0, 0)
        history_grid.addWidget(self.validate_ai_file_button, 0, 1)
        history_grid.addWidget(self.delete_history_button, 0, 2)
        history_grid.addWidget(self.btn_registro_pdf, 0, 3)
        history_grid.addWidget(self.btn_hist_period, 1, 0)
        history_grid.addWidget(self.btn_limpiar_dups, 1, 1)
        history_grid.addWidget(self.db_manager_button, 1, 2)
        hist_layout.addLayout(history_grid)
        StyleHelper.style_groupbox(hist_group)
        search_layout.addWidget(hist_group)

        
        # Patient info and medical history panel
        info_panel = QScrollArea()
        info_panel.setWidgetResizable(True)
        info_widget = QWidget()
        self.info_layout = QVBoxLayout(info_widget)

        # Hidden text field used for summaries
        self.patient_info = AutoAdjustingTextEdit()
        self.patient_info.setReadOnly(True)
        StyleHelper.style_info_text(self.patient_info)
        self.patient_info.hide()

        # Structured patient info display
        self.patient_info_group = self.create_patient_info_group()
        self.info_layout.addWidget(self.patient_info_group)

        # Mini dashboard de consultas y citas
        self.dashboard_group = self.create_dashboard_group()
        self.info_layout.addWidget(self.dashboard_group)
        self.update_dashboard(None)

        # Personal history section
        personal_history_group = QGroupBox("Antecedentes Personales")
        personal_history_layout = QGridLayout()
        self.checkboxes = {
            "cigarrillo": QCheckBox("Cigarrillo"),
            "alcohol": QCheckBox("Alcohol"),
            "drogas": QCheckBox("Drogas"),
            "cafe": QCheckBox("Café"),
            "te": QCheckBox("Té"),
            "patologias_cronicas": QCheckBox("Patologías Crónicas"),  # Cambiado de "A. Alérgicos" a "Patologías Crónicas"
            "transfusionales": QCheckBox("A. Transfusionales"),
            "traumatico": QCheckBox("A. Traumático"),
            "quirurgicos": QCheckBox("A. Quirúrgicos")
        }

        self.detail_inputs = {key: QLineEdit() for key in self.checkboxes}

        for key, checkbox in self.checkboxes.items():
            personal_history_layout.addWidget(checkbox)
            personal_history_layout.addWidget(self.detail_inputs[key])
            StyleHelper.style_checkbox(checkbox)

        personal_history_group.setLayout(personal_history_layout)

        # Family history section
        family_history_group = QGroupBox("Antecedentes Heredo-Familiares")
        family_history_layout = QFormLayout()
        self.padre_input = QLineEdit()
        self.madre_input = QLineEdit()
        self.hijos_input = QLineEdit()
        self.hermanos_input = QLineEdit()
        family_history_layout.addRow("Padre:", self.padre_input)
        family_history_layout.addRow("Madre:", self.madre_input)
        family_history_layout.addRow("Hijos:", self.hijos_input)
        family_history_layout.addRow("Hermanos:", self.hermanos_input)
        family_history_group.setLayout(family_history_layout)

        # Obstetric history section
        self.obst_group = QGroupBox("Antecedentes Obstétricos")
        obst_layout = QFormLayout()
        self.gestas_input = QLineEdit()
        self.abortos_input = QLineEdit()
        self.vaginales_input = QLineEdit()
        self.vivos_input = QLineEdit()
        self.cesareas_input = QLineEdit()
        self.gemelar_input = QLineEdit()
        self.fin_embarazo_input = QLineEdit()
        self.planeado_check = QCheckBox("Embarazo planeado")
        self.embarazo_activo_check = QCheckBox("Embarazo activo")
        self.metodo_input = QLineEdit()

        self.us_date_edit = QDateEdit()
        self.us_date_edit.setDisplayFormat("dd/MM/yyyy")
        self.us_date_edit.setCalendarPopup(True)
        # Default to today so the current year is preselected
        self.us_date_edit.setDate(QDate.currentDate())
        self.us_weeks_spin = QSpinBox()
        self.us_weeks_spin.setRange(0, 45)
        self.us_weeks_spin.setSuffix(" sem")
        self.us_days_spin = QSpinBox()
        self.us_days_spin.setRange(0, 6)
        self.us_days_spin.setSuffix(" d")
        self.add_us_button = QPushButton("Agregar")
        self.add_us_button.clicked.connect(self.add_ultrasound_entry)

        self.sonografias_input = AutoAdjustingTextEdit()
        self.sonografias_input.setPlaceholderText("Ej: 10/07/2025 - 22 semanas - primer trimestre")
        self.sonografias_input.setFixedHeight(80)
        obst_layout.addRow("Gestas previas:", self.gestas_input)
        obst_layout.addRow("Abortos:", self.abortos_input)
        obst_layout.addRow("Partos vaginales:", self.vaginales_input)
        obst_layout.addRow("Nacidos vivos:", self.vivos_input)
        obst_layout.addRow("Cesáreas:", self.cesareas_input)
        obst_layout.addRow("Gemelar:", self.gemelar_input)
        obst_layout.addRow("Fin embarazo anterior:", self.fin_embarazo_input)
        obst_layout.addRow(self.planeado_check)
        obst_layout.addRow(self.embarazo_activo_check)
        obst_layout.addRow("Método (si fracaso):", self.metodo_input)

        us_row = QHBoxLayout()
        us_row.addWidget(self.us_date_edit)
        us_row.addWidget(self.us_weeks_spin)
        us_row.addWidget(self.us_days_spin)
        us_row.addWidget(self.add_us_button)
        obst_layout.addRow("Nueva sonografía:", us_row)
        obst_layout.addRow("Sonografías:", self.sonografias_input)
        self.obst_group.setLayout(obst_layout)
        self.obst_group.hide()


        # Diagnósticos previos autoajustable**********************
        self.diagnosticos_previos_input = AutoAdjustingTextEdit()
        self.diagnosticos_previos_input.setPlaceholderText("Diagnósticos Previos")
        self.diagnosticos_previos_input.setReadOnly(True)
        self.diagnosticos_previos_input.set_auto_resize(True)
        self.diagnosticos_previos_input.setFixedHeight(100)  # Altura inicial
        # Aplicar estilo usando StyleHelper
        StyleHelper.style_input(self.diagnosticos_previos_input)

        # Use AutoFitTextEdit for other text inputs
        self.historia_enfermedad_input = self.create_autofit_text_edit()
        self.historia_enfermedad_input.setInitialHeight(500)

        self.sugerencias_ia = self.create_autofit_text_edit(read_only=True)
        self.sugerencias_ia.setInitialHeight(400)

        self.interaccion_ia_input = self.create_autofit_text_edit()
        self.interaccion_ia_input.setInitialHeight(400)
        self.interaccion_ia_input.hide()

        self.interaccion_ia_output = self.create_autofit_text_edit(read_only=True)
        self.interaccion_ia_output.setInitialHeight(400)
        self.interaccion_ia_output.hide()

        self.generate_button = QPushButton("Generar Historia Clinica")
        self.generate_button.setIcon(self.style().standardIcon(QStyle.SP_FileDialogNewFolder))
        StyleHelper.style_button(self.generate_button)
        self.save_button = QPushButton("Guardar Historia")
        self.save_button.setIcon(self.style().standardIcon(QStyle.SP_DialogSaveButton))
        StyleHelper.style_button(self.save_button)
        self.export_button = QPushButton("Exportar a Word")
        self.export_button.setIcon(self.style().standardIcon(QStyle.SP_DialogSaveButton))
        StyleHelper.style_button(self.export_button)
        self.btn_generar_receta = QPushButton("Generar Receta")
        self.btn_generar_receta.setIcon(self.style().standardIcon(QStyle.SP_FileIcon))
        StyleHelper.style_button(self.btn_generar_receta)

        self.btn_generar_Lab = QPushButton("Generar Laboratorio")
        self.btn_generar_Lab.setIcon(self.style().standardIcon(QStyle.SP_FileIcon))
        StyleHelper.style_button(self.btn_generar_Lab)

        self.auto_adjust_button = QPushButton("Autoajustar Tamaño")
        self.auto_adjust_button.setIcon(StyleHelper.themed_icon("view-refresh", QStyle.SP_BrowserReload))
        StyleHelper.style_button(self.auto_adjust_button)

        self.explain_button = QPushButton("Explicar al Paciente")
        self.explain_button.setIcon(StyleHelper.themed_icon("help-about", QStyle.SP_MessageBoxQuestion))
        StyleHelper.style_button(self.explain_button)
        self.ask_history_button = QPushButton("Preguntar Historial")
        self.ask_history_button.setIcon(StyleHelper.themed_icon("comment-question", QStyle.SP_MessageBoxQuestion))
        StyleHelper.style_button(self.ask_history_button)
        self.analyze_image_button = QPushButton("Analizar Imagen", self)
        self.analyze_image_button.setIcon(StyleHelper.themed_icon("image-search", QStyle.SP_FileDialogContentsView))
        StyleHelper.style_button(self.analyze_image_button)
        self.analyze_image_button.clicked.connect(self.analyze_image)

        
        
        # Hidden text field for summaries
        self.info_layout.addWidget(self.patient_info)

        # Ajustar la altura inicial del cuadro de texto si se muestra
        self.patient_info.textChanged.connect(self.adjust_textedit_size)
        
        # Configurar el panel de información
        info_panel.setWidget(info_widget)

        self.info_layout.addWidget(personal_history_group)
        self.info_layout.addWidget(family_history_group)
        self.info_layout.addWidget(self.obst_group)

        self.info_layout.addWidget(self.diagnosticos_previos_input)

        # Campos de signos vitales dentro de la historia clínica
        vitals_group = QGroupBox("🩺 Signos Vitales")
        vitals_layout = QGridLayout()
        vitals_layout.setContentsMargins(16, 16, 16, 16)
        vitals_layout.setHorizontalSpacing(12)
        vitals_layout.setVerticalSpacing(12)

        vitals_layout.addWidget(QLabel("PA:"), 0, 0)
        self.bp_field = QLineEdit()
        StyleHelper.style_input(self.bp_field)
        vitals_layout.addWidget(self.bp_field, 0, 1)
        self.bp_field.textChanged.connect(self.update_bp_style)

        vitals_layout.addWidget(QLabel("FC:"), 0, 2)
        self.hr_field = QLineEdit()
        StyleHelper.style_input(self.hr_field)
        vitals_layout.addWidget(self.hr_field, 0, 3)
        self.hr_field.textChanged.connect(self.update_hr_style)

        vitals_layout.addWidget(QLabel("Peso:"), 0, 4)
        self.weight_field = QLineEdit()
        StyleHelper.style_input(self.weight_field)
        self.weight_unit_combo = QComboBox()
        self.weight_unit_combo.addItems(["lb", "kg"])
        StyleHelper.style_input(self.weight_unit_combo)
        self.weight_unit_combo.currentTextChanged.connect(self.on_weight_unit_changed)
        weight_box = QWidget()
        wb = QHBoxLayout(weight_box)
        wb.setContentsMargins(0, 0, 0, 0)
        wb.setSpacing(4)
        wb.addWidget(self.weight_field)
        wb.addWidget(self.weight_unit_combo)
        vitals_layout.addWidget(weight_box, 0, 5)

        vitals_layout.addWidget(QLabel("Glicemia:"), 0, 6)
        self.gly_field = QLineEdit()
        StyleHelper.style_input(self.gly_field)
        vitals_layout.addWidget(self.gly_field, 0, 7)
        self.gly_field.textChanged.connect(self.update_gly_style)

        vitals_layout.addWidget(QLabel("Colesterol:"), 1, 0)
        self.chol_field = QLineEdit()
        StyleHelper.style_input(self.chol_field)
        vitals_layout.addWidget(self.chol_field, 1, 1)
        self.chol_field.textChanged.connect(self.update_chol_style)

        vitals_layout.addWidget(QLabel("Oximetría:"), 1, 2)
        self.oxi_field = QLineEdit()
        StyleHelper.style_input(self.oxi_field)
        vitals_layout.addWidget(self.oxi_field, 1, 3)
        self.oxi_field.textChanged.connect(self.update_oxi_style)

        vitals_layout.addWidget(QLabel("Temp:"), 1, 4)
        self.temp_field = QLineEdit()
        StyleHelper.style_input(self.temp_field)
        vitals_layout.addWidget(self.temp_field, 1, 5)
        self.temp_field.textChanged.connect(self.update_temp_style)

        vitals_layout.addWidget(QLabel("IMC:"), 1, 6)
        self.bmi_field = QLineEdit()
        self.bmi_field.setReadOnly(True)
        StyleHelper.style_input(self.bmi_field)
        vitals_layout.addWidget(self.bmi_field, 1, 7)

        vitals_layout.addWidget(QLabel("Estado:"), 1, 8)
        self.bmi_class_field = QLabel()
        self.bmi_class_field.setFixedWidth(160)
        self.bmi_class_field.setAlignment(Qt.AlignCenter)
        vitals_layout.addWidget(self.bmi_class_field, 1, 9)

        vitals_group.setLayout(vitals_layout)
        StyleHelper.style_groupbox(vitals_group)

        self.current_weight_unit = "lb"
        self.weight_field.textChanged.connect(self.update_bmi)
        self.weight_unit_combo.currentTextChanged.connect(self.update_bmi)
        self.info_layout.addWidget(vitals_group)
        separator = QFrame()
        separator.setFrameShape(QFrame.HLine)
        separator.setFrameShadow(QFrame.Sunken)
        self.info_layout.addWidget(separator)

        # Resumen de cambios en signos vitales
        self.vitals_summary_label = QLabel()
        self.vitals_summary_label.setAlignment(Qt.AlignLeft)
        self.vitals_summary_label.setStyleSheet("color:#555; padding:4px;")
        self.info_layout.addWidget(self.vitals_summary_label)

        hist_label = QLabel("\ud83d\udcd6 Historia de la Enfermedad Actual")
        hist_label.setStyleSheet("font-weight:bold;color:#222;padding-bottom:6px;")
        self.info_layout.addWidget(hist_label)
        hline_hist = QFrame()
        hline_hist.setFrameShape(QFrame.HLine)
        hline_hist.setFrameShadow(QFrame.Sunken)
        self.info_layout.addWidget(hline_hist)
        self.info_layout.addWidget(self.historia_enfermedad_input)

        actions_grid = QGridLayout()
        actions_grid.addWidget(self.generate_button, 0, 0)
        actions_grid.addWidget(self.save_button, 0, 1)
        actions_grid.addWidget(self.export_button, 0, 2)
        actions_grid.addWidget(self.btn_generar_receta, 1, 0)
        actions_grid.addWidget(self.btn_generar_Lab, 1, 1)
        actions_grid.addWidget(self.auto_adjust_button, 1, 2)
        actions_grid.addWidget(self.explain_button, 2, 0)
        actions_grid.addWidget(self.ask_history_button, 2, 1)
        actions_grid.addWidget(self.analyze_image_button, 2, 2)
        self.info_layout.addLayout(actions_grid)

        self.auto_adjust_button.clicked.connect(self.auto_adjust_text_edits)
        self.explain_button.clicked.connect(self.generate_patient_explanation)
        self.ask_history_button.clicked.connect(self.open_history_query)


        sug_label = QLabel("\ud83e\udd16 Sugerencias Inteligentes del Asistente IA")
        sug_label.setStyleSheet("font-weight:bold;color:#222;padding-bottom:6px;margin-top:12px;")
        self.info_layout.addWidget(sug_label)
        hline_sug = QFrame()
        hline_sug.setFrameShape(QFrame.HLine)
        hline_sug.setFrameShadow(QFrame.Sunken)
        self.info_layout.addWidget(hline_sug)
        self.info_layout.addWidget(self.sugerencias_ia)

        # self.info_layout.addWidget(QLabel("Interacción con IA:"))
        # self.info_layout.addWidget(QLabel("Pregunta a la IA:"))
        # self.info_layout.addWidget(self.interaccion_ia_input)
        # self.load_file_button = QPushButton("Cargar Archivo", self)
        # self.load_file_button.setIcon(self.style().standardIcon(QStyle.SP_DialogOpenButton))
        # StyleHelper.style_button(self.load_file_button)
        # self.load_file_button.clicked.connect(self.load_file)
        # self.info_layout.addWidget(self.load_file_button)

        # self.info_layout.addWidget(QLabel("Respuesta de la IA:"))
        # self.info_layout.addWidget(self.interaccion_ia_output)

        # Aplicar estilo a la tabla de pacientes
        StyleHelper.style_table(self.patient_list, "#e0e0e0")

        # Aplicar estilo a los checkboxes
        for checkbox in self.checkboxes.values():
            StyleHelper.style_checkbox(checkbox)

        # Aplicar estilo a los groupboxes
        StyleHelper.style_groupbox(personal_history_group)
        StyleHelper.style_groupbox(family_history_group)
        StyleHelper.style_groupbox(self.obst_group)



        # Aplicar estilo a todos los campos de entrada
        for input_widget in [self.search_input, self.patient_info, self.historia_enfermedad_input,
                             self.sugerencias_ia, self.interaccion_ia_input, self.interaccion_ia_output] + \
                            list(self.detail_inputs.values()) + \
                            [self.padre_input, self.madre_input, self.hijos_input, self.hermanos_input,
                             self.gestas_input, self.abortos_input, self.vaginales_input, self.vivos_input,
                             self.cesareas_input, self.gemelar_input, self.fin_embarazo_input,
                            self.metodo_input, self.sonografias_input,
                             self.bp_field, self.hr_field, self.weight_field, self.bmi_field,
                             self.gly_field, self.chol_field, self.oxi_field, self.temp_field]:
            StyleHelper.style_input(input_widget)

        info_panel.setWidget(info_widget)
        info_panel.setWidgetResizable(True)

        # Balanced columns
        main_layout.addWidget(search_panel)
        main_layout.addWidget(info_panel)
        main_layout.setStretch(0, 45)
        main_layout.setStretch(1, 55)

        # Conectar el botón de generar receta a la función
        self.btn_generar_receta.clicked.connect(self.print_receta)

        self.search_button.clicked.connect(self.search_patient)
        self.new_patient_button.clicked.connect(self.new_patient)
        self.edit_patient_button.clicked.connect(self.edit_patient)
        self.fixed_med_button.clicked.connect(self.edit_fixed_medications)
        self.vitals_button.clicked.connect(self.open_vitals)
        self.appointment_button.clicked.connect(self.open_appointments)
        self.billing_button.clicked.connect(self.open_billing)
        self.lab_results_button.clicked.connect(self.open_lab_results)
        self.insurance_button.clicked.connect(self.open_insurance_manager)
        self.services_button.clicked.connect(self.open_service_manager)
        self.forms_button.clicked.connect(self.open_forms)
        self.reminder_button.clicked.connect(self.open_reminders)
        self.search_hist_button.clicked.connect(self.open_advanced_search)
        self.analysis_button.clicked.connect(self.open_analysis)
        self.ai_study_button.clicked.connect(self.open_ai_study)
        self.delete_patient_button.clicked.connect(self.delete_patient)
        self.patient_list.cellClicked.connect(self.load_patient)
        self.save_button.clicked.connect(self.save_history)
        self.generate_button.clicked.connect(self.open_ia_interaction)
        self.export_button.clicked.connect(self.export_to_word)
        self.delete_history_button.clicked.connect(self.delete_history)
        self.btn_registro_pdf.clicked.connect(self.open_reporte_registro)
        self.btn_limpiar_dups.clicked.connect(self.limpiar_registros_duplicados)
        self.btn_hist_period.clicked.connect(self.open_informe_historia)
        self.db_manager_button.clicked.connect(self.open_db_manager)
        self.btn_generar_Lab.clicked.connect(self.on_btn_generar_Lab_clicked)
        self.validate_ai_file_button.clicked.connect(self.open_validation_dialog)
        

        # Conectar la tecla Enter a las funciones correspondientes
        QShortcut(QKeySequence("Return"), self.historia_enfermedad_input, self.open_ia_interaction)

        self.patient_list.cellClicked.connect(self.load_patient)
        self.history_list.itemClicked.connect(self.load_selected_history)
        self.new_history_button.clicked.connect(self.new_history)

        # Aplicar el tema guardado y luego el idioma seleccionado
        self.set_theme(self.current_theme)
        self.apply_language()
    # --------------------------------------------------------------------
    #  Abre el diálogo para exportar el Registro Diario de Consulta a PDF
        # --------------------------------------------------------------------
    def limpiar_registros_duplicados(self):
        """
        Conserva solo el registro más reciente para cada (paciente, fecha).
        """
        if QMessageBox.question(
                self, "Eliminar Duplicados",
                "Eliminará registros repetidos, ¿continuar?",
                QMessageBox.Yes | QMessageBox.No) != QMessageBox.Yes:
            return
    
        session = Session()
        try:
            sub = (session.query(
                        RegistroConsulta.paciente_id.label("pid"),
                        func.date(RegistroConsulta.fecha).label("f"),
                        func.max(RegistroConsulta.id).label("keep"))
                   .group_by("pid", "f").subquery())
    
            keep_ids = [r.keep for r in session.query(sub.c.keep)]
            borrados = (session.query(RegistroConsulta)
                              .filter(~RegistroConsulta.id.in_(keep_ids))
                              .delete(synchronize_session=False))
            session.commit()
            QMessageBox.information(self, "Hecho",
                                    f"Se eliminaron {borrados} duplicados.")
        except Exception as e:
            session.rollback()
            QMessageBox.critical(self, "Error", f"No se pudo limpiar:\n{e}")
        finally:
            session.close()

    def _clean_old_registros(self):
        """
        Elimina duplicados previos: conserva el registro más reciente
        para cada (paciente_id, fecha).  Ejecutar UNA sola vez.
        """
        session = Session()
        try:
            sub = (session.query(
                        RegistroConsulta.paciente_id.label("pid"),
                        func.date(RegistroConsulta.fecha).label("f"),
                        func.max(RegistroConsulta.id).label("max_id"))
                   .group_by("pid", "f")
                   .subquery())
    
            # ids válidos
            keep_ids = [row.max_id for row in session.query(sub.c.max_id)]
            # borrar todo lo que no sea el más nuevo del día
            session.query(RegistroConsulta)\
                   .filter(~RegistroConsulta.id.in_(keep_ids))\
                   .delete(synchronize_session=False)
            session.commit()
            print("✔  Duplicados viejos eliminados")
        finally:
            session.close()

    # ─────────────────────────────────────────────────────────────────────
    def _safe_widget_text(self, attr_name, default="", combo=False):
        """
        Devuelve el texto/valor del widget *attr_name* si existe.
        • Si no existe, devuelve *default*.
        • Si es QComboBox (combo=True) usa currentText().
        """
        w = getattr(self, attr_name, None)
        if w is None:
            return default
        return w.currentText().strip() if combo else w.text().strip()

    
    def store_diagnoses(self, historia_obj, texto, session=None):
        """
        • Elimina los diagnósticos previos y escribe los nuevos.
        • Captura el código CIE-10 escrito entre paréntesis, tras un guion o
          dos puntos.   Ej.:   Dolor lumbar - M54.5
        • Descarta:
            – Líneas sin código
            – Códigos Z00, Z72, Z83, Z86, Z87, Z91  (antecedentes / factores)
            – Líneas que contienen ‘motivo de consulta’, ‘ausencia’, ‘sin’, ‘negado’.
        """
        own_session = session is None
        if own_session:
            session = Session()
        try:
            historia = session.merge(historia_obj)
            session.query(Diagnostico)\
                   .filter_by(historia_id=historia.id).delete()
    
            lines_without_code = []
            parsed_lines = []
            for ln in texto.splitlines():
                ln_orig = ln
                ln = ln.strip()
                if not ln:
                    continue
                if re.search(r'\b(motivo de consulta|ausencia|sin |negado)\b', ln, re.I):
                    continue

                m = CIE10_REGEX.search(ln)
                if not m:
                    lines_without_code.append(ln)
                    continue
                codigo = m.group(1).upper()
                if codigo[:3] in Z_BLACK:
                    continue
                desc = re.sub(CIE10_REGEX, "", ln).strip(" .:-")
                parsed_lines.append((desc, codigo))

            extra = self.categorize_diagnoses_api(lines_without_code)
            for desc, code in parsed_lines + list(extra.values()):
                session.add(Diagnostico(
                    historia_id  = historia.id,
                    codigo_cie10 = code,
                    descripcion  = f"{desc} ({code})" if desc else code
                ))
    
            # fallback si no quedó ninguno
            if not session.new:
                primer = next((l for l in texto.splitlines() if l.strip()), "")
                session.add(Diagnostico(historia_id=historia.id,
                                        codigo_cie10="",
                                        descripcion=primer[:120] or "Sin diagnóstico explícito"))
    
            session.commit()
        except Exception as e:
            session.rollback()
            print("[store_diagnoses] error:", e)
        finally:
            if own_session:
                session.close()

    def categorize_diagnoses_api(self, lines):
        if not lines:
            return {}
        model = ''
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r') as f:
                cfg = json.load(f)
            model = cfg.get('model', '')
        if not ensure_openai_key():
            print('[categorize_diagnoses_api] missing API key')
            return {}
        prompt = "Asigna el codigo CIE-10 mas probable a cada diagnostico. Devuelve cada uno como 'descripcion - CODIGO'."\
                 + "\n" + "\n".join(lines)
        try:
            resp = openai.ChatCompletion.create(model=model,
                messages=[{"role":"system","content":"Asistente de codificacion medica"},
                         {"role":"user","content":prompt}])
            res = {}
            for line in resp['choices'][0]['message']['content'].splitlines():
                if '-' in line:
                    desc, code = line.rsplit('-',1)
                    res[line.strip()] = (desc.strip(), code.strip())
            return res
        except Exception as e:
            print('[categorize_diagnoses_api]', e)
            return {}

    def backfill_diagnosticos(self):
        """
        Recorre las historias clínicas que todavía no tienen diagnóstico
        y crea uno a partir de `historia.sugerencias_ia`.
        """
        session = Session()
        nuevos = 0
        try:
            historias = (session.query(HistoriaClinica)
                         .outerjoin(Diagnostico)
                         .filter(Diagnostico.id.is_(None))   # sin diagnósticos
                         .all())
            for hist in historias:
                texto = (hist.sugerencias_ia or "").strip()
                if not texto:
                    continue
                # Tomar primera línea que contenga algo tipo "A09 Diarrea aguda"
                linea = texto.splitlines()[0]
                match = CIE10_REGEX.search(linea)
                codigo = match.group(1) if match else ""
                descripcion = linea.replace(codigo, "").strip(' :.-')
                session.add(Diagnostico(
                    historia_id=hist.id,
                    codigo_cie10=codigo.upper(),
                    descripcion=descripcion or linea
                ))
                nuevos += 1
            if nuevos:
                session.commit()
                print(f"✔ backfill_diagnosticos: {nuevos} diagnósticos creados")
        except Exception as e:
            session.rollback()
            print("[backfill_diagnosticos] error:", e)
        finally:
            session.close()


    
    def open_reporte_registro(self):
        """
        Muestra la ventana que permite elegir el rango de fechas
        y genera el PDF del Registro Diario.
        """
        dlg = ReporteRegistroDialog(self)   # Asegúrate de que esta clase esté importada/definida
        dlg.exec_()

    def open_informe_historia(self):
        dlg = InformeHistoriasDialog(self)
        dlg.exec_()

    # --------------------------------------------------------------------
    #   CREA AUTOMÁTICAMENTE las filas que falten en registro_consultas
    # --------------------------------------------------------------------
    def backfill_registro_consultas(self):
        """
        Recorre todas las HistoriaClinica existentes y, si todavía no tienen
        su renglón en RegistroConsulta, lo genera con valores por defecto.
        Esto se ejecuta una sola vez al iniciar la aplicación, y también
        puede llamarse desde el diálogo de PDF si fuera necesario.
        """
        session = Session()
        nuevos = 0
        try:
            historias = session.query(HistoriaClinica).all()
            # Detect if ficha_familiar column exists
            insp = inspect(engine)
            ficha_exists = 'ficha_familiar' in [c['name'] for c in insp.get_columns('pacientes')]
            for hist in historias:
                if session.query(RegistroConsulta)\
                          .filter_by(historia_id=hist.id).first():
                    continue  # ya existe
    
                pac = hist.paciente
                diag = hist.diagnosticos[0] if hist.diagnosticos else None
    
                kwargs = dict(
                    fecha             = hist.fecha or datetime.now(),
                    historia_id       = hist.id,
                    paciente_id       = pac.id,
                    documento_id      = pac.documento_id,
                    edad              = pac.edad,
                    sexo              = pac.sexo[:1] if pac.sexo else "",
                    nacionalidad      = pac.nacionalidad,
                    direccion         = pac.direccion,
                    no_seguridad_soc  = pac.no_seguridad_soc,
                    tipo_sangre       = pac.tipo_sangre,
                    ars               = pac.ars,
                    lugar_consulta    = "CONSULTORIO",
                    frecuencia        = "CONTROL",
                    servicio          = "Medicina Familiar",
                    diagnostico_txt   = diag.descripcion if diag else "",
                    codigo_cie10      = diag.codigo_cie10 if diag else "",
                    tipo_atencion     = "Curativa"
                )
                if ficha_exists:
                    kwargs['ficha_familiar'] = getattr(pac, 'ficha_familiar', '')
                session.add(RegistroConsulta(**kwargs))
                nuevos += 1
            if nuevos:
                session.commit()
                print(f"✔  backfill: {nuevos} filas añadidas a RegistroConsulta")
        except Exception as e:
            session.rollback()
            print(f"[backfill_registro_consultas] Error: {e}")
        finally:
            session.close()
            

    # --------------------------------------------------------------------
    #   Versión resiliente — no falla si un widget no existe en MainWindow
    # --------------------------------------------------------------------
    def ensure_patient_defaults(self, pac):
        """
        Rellena en BD los campos fijos del paciente (ficha, cédula, etc.).
        • Si el widget existe en la ventana principal, toma el valor del widget.
        • Si no existe, deja intacto lo que ya tenga el Paciente
          (porque seguramente se editó en el diálogo "Datos del Paciente").
        Se hace solo la primera vez que el campo está vacío.
        """
          # --------------------------------------------------------------
        def _safe_widget_text(self, name, default="", combo=False):
            """
            Devuelve el texto del widget llamado *name* si existe.
            • Si el widget no existe, devuelve *default*.
            • Si es QComboBox usa currentText().
            """
            w = getattr(self, name, None)
            if w is None:
                return default
            return w.currentText().strip() if combo else w.text().strip()
        


    def compute_tipo_atencion(self, diagnostico_txt, motivo_consulta=""):
        """
        Devuelve PREVENTIVA, CURATIVA o PROCEDIMIENTO usando GPT.
        Si ocurre un error devuelve 'Curativa' por defecto.
        """
        if not ensure_openai_key():
            print('[tipo_atencion IA] missing API key')
            return "Curativa"
        try:
            prompt = (
                "Clasifica este motivo de consulta y diagnóstico en PREVENTIVA, "
                "CURATIVA o PROCEDIMIENTO. Responde solo una palabra.\n\n"
                f"Motivo de consulta:\n{motivo_consulta}\n\n"
                f"Diagnóstico:\n{diagnostico_txt}"
            )
    
            rsp = openai.ChatCompletion.create(
                model= 'gpt-4.1-2025-04-14',
                messages=[
                    {"role": "system",
                     "content": "Eres un médico familiar que clasifica tipos de atención."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2,
                temperature=0
            )
            txt = rsp.choices[0].message.content.strip().upper()
            if txt.startswith("PREV"):
                return "Preventiva"
            if txt.startswith("PROC"):
                return "Procedimiento"
            return "Curativa"
        except Exception as e:
            print(f"[tipo_atencion IA] {e}")
            return "Curativa"

    def generate_history_summary(self, historia_obj):
        text = (historia_obj.historia_enfermedad or '') + "\n" + (historia_obj.sugerencias_ia or '')
        if not text.strip():
            return ''
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r') as f:
                cfg = json.load(f)
            encoded = cfg.get('api_key', '')
            try:
                api_key = base64.b64decode(encoded).decode()
            except Exception:
                api_key = ''
            openai.api_key = api_key or os.getenv('OPENAI_API_KEY', '')
            model = cfg.get('model', '')
        else:
            openai.api_key = os.getenv('OPENAI_API_KEY', '')
            model = ''
            if not openai.api_key:
                return ''
        prompt = f"Resume brevemente en viñetas los hallazgos y planes más relevantes:\n{text}"
        try:
            resp = openai.ChatCompletion.create(model=model,
                messages=[{"role":"system","content":"Eres un médico"}, {"role":"user","content":prompt}])
            summary = resp['choices'][0]['message']['content'].strip()
            session = Session()
            h = session.merge(historia_obj)
            h.resumen = summary
            session.commit()
            session.close()
            return summary
        except Exception as e:
            print('[generate_history_summary]', e)
            return ''

    # --------------------------------------------------------------------
    #   VERSIÓN CORREGIDA de store_consultation_register
    # --------------------------------------------------------------------


     # --------------------------------------------------------------------
    def store_consultation_register(self, historia_obj, session=None):

        """
        Crea o actualiza el Registro Diario:
          · Deduplica por (paciente_id, fecha)
          · Usa _safe_widget_text para widgets opcionales
        """
        own_session = session is None
        if own_session:
            session = Session()
        try:
            historia = session.merge(historia_obj)          # re-anexa
            pac      = historia.paciente
    
            # diagnóstico principal (puede haber varios – cogemos el 1º)
            diag = historia.diagnosticos[0] if historia.diagnosticos else None
            diag_txt   = diag.descripcion if diag else ""
            codigo_cie = diag.codigo_cie10 if diag else ""
    
            hoy = datetime.now().date()
    
            reg = (session.query(RegistroConsulta)
                          .filter(RegistroConsulta.paciente_id == pac.id,
                                  func.date(RegistroConsulta.fecha) == hoy)
                          .first())
    
            # ----- valores de la interfaz (o por defecto) -----
            lugar      = self._safe_widget_text("combo_lugar",      "CONSULTORIO",   combo=True)
            frecuencia = self._safe_widget_text("combo_frecuencia", "CONTROL",       combo=True)
            servicio   = self._safe_widget_text("combo_servicio",   "Medicina Fam.", combo=True)
            sipna_id   = self._safe_widget_text("input_sipna", "")
            referido   = self._safe_widget_text("input_referido", "")
    
            tipo_aten  = self.compute_tipo_atencion(
                            diag_txt,
                            self._safe_widget_text("input_motivo_consulta", "")
                         )
    
            # rellenar campos fijos del paciente si faltan
            self.ensure_patient_defaults(pac)
    
            if reg is None:
                reg = RegistroConsulta(
                    fecha       = datetime.now(),
                    paciente_id = pac.id
                )
                session.add(reg)
    
            # —— copiar datos ——
            reg.historia_id       = historia.id
            reg.ficha_familiar    = pac.ficha_familiar
            reg.documento_id      = pac.documento_id
            reg.edad              = pac.edad
            reg.sexo              = (pac.sexo or "")[:1]
            reg.nacionalidad      = pac.nacionalidad
            reg.direccion         = pac.direccion
            reg.no_seguridad_soc  = pac.no_seguridad_soc
            reg.tipo_sangre       = pac.tipo_sangre
            reg.ars               = pac.ars
    
            reg.lugar_consulta    = lugar
            reg.frecuencia        = frecuencia
            reg.servicio          = servicio
    
            reg.diagnostico_txt   = diag_txt
            reg.codigo_cie10      = codigo_cie
            reg.tipo_atencion     = tipo_aten
            reg.sipna_id          = sipna_id
            reg.referido_a        = referido
    
            session.commit()
    
        except Exception as e:
            session.rollback()
            QMessageBox.critical(self, "Error",
                                 f"No se pudo guardar el registro diario:\n{e}")
        finally:
            if own_session:
                session.close()
        
    def show_main_window(self):
        # Oculta la ventana de carga si aún está abierta
        if self.loading_screen.isVisible():
            self.loading_screen.stop_animation()
        self.show()
    def show_main_window(self):
        self.show()

    
    def check_and_show_auth(self):
        self.auth_window = ModernAuthWindow()        
        if self.auth_window.auto_login_if_not_required():
            # No hay usuarios registrados, permitir acceso directo
            
            self.auth_window.login_successful.emit('admin')
        else:
            # Hay usuarios registrados, mostrar ventana de login
            
            self.auth_window.show()
        
        # Conectar la señal de login exitoso en ambos casos
        self.auth_window.login_successful.connect(self.on_login_successful)

    def on_login_successful(self, username):
        self.current_user = username
        self.user_role = self.auth_window.users.get(username, {}).get('role', 'doctor')
        self.apply_role_permissions()
        self.show()
        self.showMaximized()
        # Iniciar recordatorios tras autenticación
        self.reminder_timer.start(3600 * 1000)
        self.alert_timer.start(3600 * 1000)
        self.check_appointments()
        self.check_reminders()

    def apply_role_permissions(self):
        if getattr(self, 'user_role', 'doctor') != 'admin':
            self.db_manager_button.hide()
        else:
            self.db_manager_button.show()
        
    def open_config_window(self):
        # Primero, autentica al usuario
        if ConfigGPT.authenticate_user():
            self.config_window = ConfigGPT()
            self.config_window.show()            
            
        else:
            QMessageBox.critical(None, "Acceso Denegado", "Contraseña incorrecta. No tienes permiso para modificar la configuración.")


    def load_config(self):
        # Cargar el archivo de configuración (si existe)
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, "r") as config_file:
                    config_data = json.load(config_file)
                    self.model = config_data.get("model", "")
                    self.current_theme = config_data.get("theme", self.current_theme)
                    self.update_window_title()
            except Exception as e:
                print(f"Error al cargar la configuración: {e}")
                self.model = ""
        else:
            print("No se encontró un archivo de configuración.")
            self.model = ""
            
    def save_config(self):
        model = 'Nuevo modelo'
        self.update_window_title(model)  # Llama al método de la ventana principal
    def open_config_gpt(self):
        self.config_gpt.show()

    def update_model(self, model):
        
        # Actualizar el modelo en la ventana principal cuando se reciba la señal
        self.model = model
        self.update_window_title()

    def update_window_title(self):
       
        # Actualiza el título de la ventana con el nuevo modelo
        nuevo_titulo = f"Sistema de Historias Clínicas - {self.model}"
        self.setWindowTitle(nuevo_titulo)
        self.repaint()  # Forzar actualización visual
        QApplication.processEvents()  # Procesar eventos pendientes

    def set_language(self, lang):
        self.current_language = lang
        self.apply_language()

    def apply_language(self):
        t = TRANSLATIONS[self.current_language]
        self.search_button.setText(t['search_patient'])
        self.new_patient_button.setText(t['new_patient'])
        self.edit_patient_button.setText(t['edit_patient'])
        self.fixed_med_button.setText(t['fixed_med'])
        self.vitals_button.setText(t['vital_signs'])
        self.family_button.setText(t['family_social'])
        self.appointment_button.setText(t['appointments'])
        self.billing_button.setText(t['billing'])
        self.lab_results_button.setText(t['lab_results'])
        self.reminder_button.setText(t['reminders'])
        self.search_hist_button.setText(t['advanced_search'])
        self.analysis_button.setText(t['data_analysis'])
        self.delete_patient_button.setText(t['delete_patient'])
        self.new_history_button.setText(t['new_history'])
        self.delete_history_button.setText(t['delete_history'])
        self.btn_registro_pdf.setText(t['generate_daily'])
        self.btn_registro_pdf.setToolTip(t['export_pdf'])
        self.btn_hist_period.setText(t['period_history'])
        self.btn_limpiar_dups.setText(t['clean_duplicates'])
        self.db_manager_button.setText(t['manage_db'])
        self.insurance_button.setText(t['manage_insurance'])
        self.services_button.setText(t['manage_services'])
        self.forms_button.setText(t['forms'])
        self.generate_button.setText(t['generate_history'])
        self.save_button.setText(t['save_history'])
        self.export_button.setText(t['export_word'])
        self.btn_generar_receta.setText(t['generate_recipe'])
        self.btn_generar_Lab.setText(t['generate_lab'])
        self.auto_adjust_button.setText(t['auto_adjust'])
        self.analyze_image_button.setText(t['analyze_image'])
        self.explain_button.setText(t['explain_patient'])
        self.ask_history_button.setText(t['ask_history'])
        self.ai_study_button.setText(t['ai_study'])
        self.validate_ai_file_button.setText(t['validate_ai_file'])

    def set_theme(self, theme):
        self.current_theme = theme
        app = QApplication.instance()
        if theme == 'light':
            StyleHelper.set_light_style(app)
            stylesheet = AdaptiveStylesheet(theme='light').generate_light_stylesheet(AdaptiveStylesheet.get_screen_resolution()[0])
        elif theme == 'gpt':
            StyleHelper.set_gpt_style(app)
            stylesheet = AdaptiveStylesheet(theme='gpt').generate_gpt_stylesheet(AdaptiveStylesheet.get_screen_resolution()[0])
        else:
            StyleHelper.set_dark_style(app)
            stylesheet = AdaptiveStylesheet().generate_stylesheet(AdaptiveStylesheet.get_screen_resolution()[0])
        AdaptiveStylesheet.apply_stylesheet(self, stylesheet)

        # Guardar el tema seleccionado
        config_data = {}
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r') as f:
                    config_data = json.load(f)
            except Exception:
                config_data = {}
        config_data['theme'] = theme
        try:
            with open(CONFIG_FILE, 'w') as f:
                json.dump(config_data, f)
        except Exception as e:
            print(f"Error al guardar tema: {e}")
      
        
    def create_menu(self):
        # Crear barra de menú
        menubar = self.menuBar()
        menubar.setNativeMenuBar(False)
        StyleHelper.style_menu_bar(menubar)
        
        # Agregar el menú "Configuración" con texto visible junto al icono
        if qta:
            config_menu = menubar.addMenu(qta.icon('fa5s.cogs', color='#1F1F1F'),
                                         "Configuración")
        else:
            config_menu = menubar.addMenu("Configuración")
        StyleHelper.style_menu(config_menu)
        
        # Crear acción para "Agregar Logo y Dirección"
        action_logo = QAction("Agregar Logo y Dirección", self)
        if qta:
            action_logo.setIcon(qta.icon('fa5s.image', color='#1F1F1F'))
        action_logo.triggered.connect(self.open_config_receta)
        config_menu.addAction(action_logo)
        
        # Crear acción para "Modificar Modelo y APIKEY"
        action_footer = QAction("Modificar Modelo y APIKEY", self)
        if qta:
            action_footer.setIcon(qta.icon('fa5s.key', color='#1F1F1F'))
        action_footer.triggered.connect(self.open_config_gpt)
        config_menu.addAction(action_footer)
        
        # Agregar menú de Usuario
        if qta:
            user_menu = menubar.addMenu(qta.icon('fa5s.user', color='#1F1F1F'),
                                       "Usuario")
        else:
            user_menu = menubar.addMenu("Usuario")
        StyleHelper.style_menu(user_menu)
        
        # Acción para gestionar la cuenta
        action_manage_account = QAction("Gestionar Cuenta", self)
        if qta:
            action_manage_account.setIcon(qta.icon('fa5s.user-cog', color='#1F1F1F'))
        action_manage_account.triggered.connect(self.open_auth_window)
        user_menu.addAction(action_manage_account)
        
        # Agregar separador
        user_menu.addSeparator()
        
        # Acción para cerrar sesión
        action_logout = QAction("Cerrar Sesión", self)
        if qta:
            action_logout.setIcon(qta.icon('fa5s.sign-out-alt', color='#1F1F1F'))
        action_logout.triggered.connect(self.logout)
        user_menu.addAction(action_logout)

        # Acción para salir del programa
        action_exit = QAction("Salir del Programa", self)
        if qta:
            action_exit.setIcon(qta.icon('fa5s.power-off', color='#1F1F1F'))
        action_exit.triggered.connect(QApplication.quit)  # Cierra la aplicación
        user_menu.addAction(action_exit)

        if qta:
            theme_menu = menubar.addMenu(qta.icon('fa5s.paint-brush',
                                                color='#1F1F1F'), "Tema")
        else:
            theme_menu = menubar.addMenu("Tema")
        StyleHelper.style_menu(theme_menu)
        action_light = QAction("Claro", self)
        if qta:
            action_light.setIcon(qta.icon('fa5s.sun', color='#1F1F1F'))
        action_light.triggered.connect(lambda: self.set_theme('light'))
        theme_menu.addAction(action_light)
        action_dark = QAction("Oscuro", self)
        if qta:
            action_dark.setIcon(qta.icon('fa5s.moon', color='#1F1F1F'))
        action_dark.triggered.connect(lambda: self.set_theme('dark'))
        theme_menu.addAction(action_dark)
        action_gpt = QAction("GPT", self)
        if qta:
            action_gpt.setIcon(qta.icon('fa5s.robot', color='#1F1F1F'))
        action_gpt.triggered.connect(lambda: self.set_theme('gpt'))
        theme_menu.addAction(action_gpt)

        if qta:
            language_menu = menubar.addMenu(qta.icon('fa5s.globe',
                                                   color='#1F1F1F'), "Idioma")
        else:
            language_menu = menubar.addMenu("Idioma")
        StyleHelper.style_menu(language_menu)
        action_spanish = QAction("Español", self)
        if qta:
            action_spanish.setIcon(qta.icon('fa5s.flag', color='#1F1F1F'))
        action_spanish.triggered.connect(lambda: self.set_language('es'))
        language_menu.addAction(action_spanish)
        action_english = QAction("English", self)
        if qta:
            action_english.setIcon(qta.icon('fa5r.flag', color='#1F1F1F'))
        action_english.triggered.connect(lambda: self.set_language('en'))
        language_menu.addAction(action_english)
    
    def open_auth_window(self):
        self.auth_window = ModernAuthWindow()
        # Forzar la visualización de la ventana de registro
        self.auth_window.stacked_widget.setCurrentIndex(1)  # 1 es el índice de la ventana de registro
        self.auth_window.show()

    def open_auth_window0(self):
            self.auth_window = ModernAuthWindow()
            # Forzar la visualización de la ventana de registro
            self.auth_window.stacked_widget.setCurrentIndex(0)  # 1 es el índice de la ventana de registro
            self.auth_window.show()
                # Ocultar la ventana principal al inicio
            self.hide()
            # Conectar la señal de login_successful
            self.auth_window.login_successful.connect(self.on_login_successful)
        
    def logout(self):
        reply = QMessageBox.question(self, 'Cerrar Sesión',
                                   '¿Está seguro de que desea cerrar sesión?',
                                   QMessageBox.Yes | QMessageBox.No,
                                   QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            # Aquí puedes agregar la lógica para cerrar sesión
            # Por ejemplo, resetear variables de sesión, cerrar conexiones, etc.
            self.open_auth_window0()  # Volver a mostrar la ventana de login
    

    # Agregar método al ModernAuthWindow para verificar autenticación
    def is_authenticated(self):
        # Agregar este método a la clase ModernAuthWindow
        return hasattr(self, 'current_user') and self.current_user is not None

    def open_config_receta(self):
        # Abrir la ventana de configuración de la receta
        self.config_receta_window = ConfigReceta()
        self.config_receta_window.show()

    def open_config_gpt(self):
        # Abrir la ventana de configuración del modelo y la API key
        self.config_gpt_window = ConfigGPT()  # Asegúrate de que ConfigGPT esté definida
        self.config_gpt_window.show()     
        

    def adjust_textedit_size(self):
        # Obtener el tamaño del documento del QTextEdit
        doc_size = self.patient_info.document().size()
        
        # Calcular la altura necesaria (convertir a entero y agregar margen)
        height = int(doc_size.height()) + 20
        
        # Obtener el ancho actual
        current_width = self.patient_info.width()
        
        # Ajustar la altura, manteniendo el ancho actual
        self.patient_info.setFixedHeight(height)
        self.patient_info.setFixedWidth(current_width)
            
    def auto_adjust_text_edits(self):
        self.historia_enfermedad_input.sizeChange()
        self.sugerencias_ia.sizeChange()
        # Los paneles de interacción con IA ya no se muestran
        # self.interaccion_ia_input.sizeChange()
        # self.interaccion_ia_output.sizeChange()
        



    def load_patient_details(self, patient_name):
        session = Session()
        try:
            # Buscar el paciente por coincidencia parcial
            patient = session.query(Paciente).filter(Paciente.nombre.like(f"%{patient_name}%")).first()

            if patient:
                session.refresh(patient)  # Refrescar la instancia del paciente
                self.current_patient = patient
                # Aquí puedes actualizar la interfaz de usuario con los detalles del paciente
                dx = patient.diagnosticos_previos or ""
                self.diagnosticos_previos_input.setText(self.organize_diagnoses(dx))

                self.gestas_input.setText(patient.gestas_previas or "")
                self.abortos_input.setText(patient.abortos or "")
                self.vaginales_input.setText(patient.partos_vaginales or "")
                self.vivos_input.setText(patient.nacidos_vivos or "")
                self.cesareas_input.setText(patient.cesareas or "")
                self.gemelar_input.setText(patient.gemelar or "")
                self.fin_embarazo_input.setText(patient.fin_embarazo_anterior or "")
                self.planeado_check.setChecked(patient.embarazo_planeado == "si")
                self.metodo_input.setText(patient.metodo_anticonceptivo or "")
                self.update_obstetric_visibility()

                fecha = patient.fecha_nacimiento.strftime('%d/%m/%Y') if patient.fecha_nacimiento else ''
                meds = patient.medicamentos_continuos or ''
                info_lines = [
                    f"Nombre: {patient.nombre}",
                    f"Cédula/ID: {patient.documento_id or ''}",
                    f"Edad: {format_patient_age(patient)}",
                    f"Sexo: {patient.sexo}",
                    f"Alergias: {patient.alergias}",
                    f"Dirección: {patient.direccion}",
                    f"Teléfono: {getattr(patient, 'telefono', '')}",
                    f"Estado Civil: {patient.estado_civil}",
                    f"Religión: {patient.religion}",
                    f"Lugar de Nacimiento: {patient.lugar_nacimiento}",
                    f"Fecha de Nacimiento: {fecha}",
                    f"Ocupación: {patient.ocupacion}",
                    f"Medicamentos Continuos: {meds}",
                ]
                self.patient_info.setText("\n".join(info_lines))

                mapping = {
                    "Nombre": patient.nombre,
                    "Cédula/ID": patient.documento_id or "",
                    "Edad": format_patient_age(patient),
                    "Sexo": patient.sexo,
                    "Alergias": patient.alergias,
                    "Dirección": patient.direccion,
                    "Teléfono": getattr(patient, 'telefono', ''),
                    "ID Familia": getattr(patient, 'id_familia', ''),
                    "Estado Civil": patient.estado_civil,
                    "Religión": patient.religion,
                    "Lugar de Nacimiento": patient.lugar_nacimiento,
                    "Fecha de Nacimiento": fecha,
                    "Ocupación": patient.ocupacion,
                    "Medicamentos Continuos": meds,
                }
                for key, val in mapping.items():
                    label = self.patient_info_labels.get(key)
                    if label is not None:
                        label.setText(f"{key}: {val}")

                self.update_dashboard(patient)

                # Busca el ítem correspondiente en la lista de pacientes y lo selecciona
                items = self.patient_list.findItems(patient.nombre, Qt.MatchExactly)
                if items:
                    # Elige el primer ítem que coincida (debería haber solo uno)
                    item = items[0]
                    self.patient_list.setCurrentItem(item)
    
                # Actualiza la lista de pacientes
                self.search_patient()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al cargar el paciente: {str(e)}")
        finally:
            session.close()


    
    # Función para convertir la fecha y hora a un formato ISO válido
    def convert_to_iso_format(date_str):
        try:
            return datetime.fromisoformat(date_str)
        except ValueError:
            return datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S.%f')
        
  
    def on_btn_generar_Lab_clicked(self):
        self.labrun()
        self.open_lab_order_dialog()
       

    
    def get_current_lab_request(self):
        # Devuelve la solicitud de laboratorio actual
        return self.current_lab_request
    def open_db_manager(self):
        db_manager = DatabaseManager(self)
        db_manager.exec_()

    def open_appointments(self):
        patient = getattr(self, 'current_patient', None)
        dialog = AppointmentDialog(self, patient)
        dialog.exec_()
        if patient:
            self.update_dashboard(patient)

    def open_billing(self):
        dialog = BillingDialog(self)
        dialog.exec_()

    def open_insurance_manager(self):
        dialog = InsuranceManagerDialog(self)
        dialog.setWindowState(dialog.windowState() | Qt.WindowMaximized)
        dialog.exec_()

    def open_service_manager(self):
        dialog = ServiceManagerDialog(self)
        dialog.setWindowState(dialog.windowState() | Qt.WindowMaximized)
        dialog.exec_()

    def open_family_social(self):
        patient = getattr(self, 'current_patient', None)
        dialog = FamilySocialDialog(self, patient)
        dialog.setWindowState(dialog.windowState() | Qt.WindowMaximized)
        dialog.exec_()
        if patient:
            self.update_dashboard(patient)

    def open_telemedicine(self):
        dialog = TelemedicineDialog(self)
        dialog.exec_()

    def open_documents(self):
        dialog = DocumentDialog(self)
        dialog.exec_()

    def open_forms(self):
        dialog = FormsDialog(self)
        dialog.setWindowState(dialog.windowState() | Qt.WindowMaximized)
        dialog.exec_()

    def open_vitals(self):
        dialog = VitalSignsDialog(
            self, self.current_patient if hasattr(self, 'current_patient') else None
        )
        dialog.setWindowState(dialog.windowState() | Qt.WindowMaximized)
        dialog.exec_()


    def open_lab_results(self):
        dialog = LabResultsDialog(self)
        dialog.exec_()

    def open_templates(self):
        dialog = TemplateManagerDialog(self)
        dialog.exec_()

    def open_reminders(self):
        dialog = ReminderDialog(self)
        dialog.setWindowState(dialog.windowState() | Qt.WindowMaximized)
        dialog.exec_()

    def open_advanced_search(self):
        dialog = AdvancedSearchDialog(self)
        dialog.exec_()

    def open_analysis(self):
        dialog = DataAnalysisDialog(self)
        dialog.exec_()

    def open_ai_study(self):
        if not self.current_patient:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")
            return
        dialog = ClinicalStudyDialog(self.current_patient, self)
        dialog.exec_()

    def open_validation_dialog(self):
        dialog = ValidationDialog(self)
        dialog.setWindowState(dialog.windowState() | Qt.WindowMaximized)
        dialog.exec_()

    def open_ia_interaction(self):
        history = self.historia_enfermedad_input.toPlainText()
        context = self.assemble_patient_context()
        dlg = IAInteractionDialog(history, context, self)
        dlg.setWindowState(Qt.WindowMaximized)
        dlg.ready.connect(self.complete_ia_interaction)
        dlg.exec_()

    def complete_ia_interaction(self, extra_text):
        if extra_text:
            current = self.historia_enfermedad_input.toPlainText().strip()
            joined = (current + "\n" + extra_text).strip() if current else extra_text
            self.historia_enfermedad_input.setPlainText(joined)
            self.interaccion_ia_output.setText(extra_text)
        self.generate_suggestions()

    def autobackup(self):
        db = DatabaseManager(self)
        db.backup_database()
        db.close()

    def generate_patient_explanation(self):
        if not self.current_patient:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")
            return
        diag = self.sugerencias_ia.toPlainText().strip()
        labs = getattr(self, 'lab_record', '') or ''
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r') as f:
                cfg = json.load(f)
            encoded = cfg.get('api_key', '')
            try:
                api_key = base64.b64decode(encoded).decode()
            except Exception:
                api_key = ''
            openai.api_key = api_key or os.getenv('OPENAI_API_KEY', '')
            model = cfg.get('model', '')
        else:
            openai.api_key = os.getenv('OPENAI_API_KEY', '')
            model = ''
            if not openai.api_key:
                QMessageBox.warning(self, "Error", "Falta configuración de OpenAI")
                return
        prompt = f"Diagnósticos:\n{diag}\n\nResultados de laboratorio:\n{labs}\n\nExplica estos hallazgos al paciente de forma sencilla y breve."
        messages = [
            {"role":"system","content":"Eres un médico que explica en lenguaje sencillo"},
            {"role":"user","content":prompt}
        ]
        dialog = ExplanationDialog(messages, model, self)
        dialog.exec_()

    def open_history_query(self):
        if not self.current_patient:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")
            return
        dialog = HistoryQueryDialog(self.current_patient, self)
        dialog.exec_()

    def check_appointments(self):
        session = Session()
        now = datetime.now()
        soon = now + timedelta(hours=1)
        citas = session.query(Cita).filter(Cita.fecha.between(now, soon)).all()
        session.close()
        if citas:
            QMessageBox.information(self, "Recordatorio", f"Tiene {len(citas)} citas próximas")

    def check_reminders(self):
        session = Session()
        today = date.today()
        rems = session.query(Recordatorio).filter(Recordatorio.fecha <= today, Recordatorio.completado == False).all()
        session.close()
        if rems:
            resp = QMessageBox.question(
                self,
                "Recordatorios",
                f"Tiene {len(rems)} recordatorios pendientes. ¿Desea revisarlos?",
            )
            if resp == QMessageBox.Yes:
                dialog = ReminderListDialog(self, only_today=True)
                dialog.exec_()

    def process_risk_updates(self):
        """Evaluate cardiovascular risk for new vital signs and store it."""
        session = Session()
        signs = session.query(SignoVital).filter_by(risk_processed=False).all()
        for s in signs:
            risk = self.risk_level(s)
            pat = session.get(Paciente, s.paciente_id)
            if pat:
                pat.riesgo_cardiovascular = risk
            s.risk_processed = True
        session.commit()
        session.close()
    
    def create_autofit_text_edit(self, read_only=False):
        text_edit = AutoFitTextEdit(self)
        text_edit.setReadOnly(read_only)
        StyleHelper.style_input(text_edit)
        return text_edit

    def create_patient_info_group(self):
        group = QGroupBox("Información del Paciente")
        layout = QGridLayout()
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setColumnStretch(1, 1)
        self.patient_info_labels = {}
        fields = [
            ("Nombre", "fa5s.user"),
            ("Cédula/ID", "fa5s.id-card"),
            ("Edad", "fa5s.birthday-cake"),
            ("Sexo", "fa5s.venus-mars"),
            ("Alergias", "fa5s.exclamation-triangle"),
            ("Dirección", "fa5s.map-marker-alt"),
            ("Teléfono", "fa5s.phone"),
            ("ID Familia", "fa5s.users"),
            ("Estado Civil", "fa5s.heart"),
            ("Religión", "fa5s.star-of-david"),
            ("Lugar de Nacimiento", "fa5s.globe"),
            ("Fecha de Nacimiento", "fa5s.calendar"),
            ("Ocupación", "fa5s.briefcase"),
            ("Medicamentos Continuos", "fa5s.pills"),
        ]
        for row, (field, icon_name) in enumerate(fields):
            icon_label = QLabel()
            if qta:
                try:
                    icon_label.setPixmap(qta.icon(icon_name, color="black").pixmap(50, 50))
                except Exception:
                    pass
            text_label = QLabel(f"{field}:")
            text_label.setFont(QFont("Segoe UI", 12))
            text_label.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
            layout.addWidget(icon_label, row, 0)
            layout.addWidget(text_label, row, 1)
            self.patient_info_labels[field] = text_label
        layout.setVerticalSpacing(8)
        group.setLayout(layout)
        StyleHelper.style_groupbox(group)
        return group

    def create_dashboard_group(self):
        group = QGroupBox("\U0001F4CA Dashboard de Consultas")
        layout = QVBoxLayout()
        self.lbl_consultas_12m = QLabel()
        self.lbl_proxima_cita = QLabel()
        self.lbl_ficha_vivienda = QLabel()
        StyleHelper.style_input(self.lbl_consultas_12m)
        StyleHelper.style_input(self.lbl_proxima_cita)
        StyleHelper.style_input(self.lbl_ficha_vivienda)
        prog_btn = QPushButton("Programar nueva cita")
        StyleHelper.style_button(prog_btn)
        if qta:
            try:
                prog_btn.setIcon(qta.icon('fa5s.calendar-alt', color='white'))
            except Exception:
                pass
        self.family_button = QPushButton("Salud Familiar y Social")
        StyleHelper.style_button(self.family_button)
        if qta:
            try:
                self.family_button.setIcon(qta.icon('fa5s.home', color='white'))
            except Exception:
                pass
        self.reminder_check = QCheckBox("\U0001F514 Enviar recordatorio al paciente")
        StyleHelper.style_checkbox(self.reminder_check)
        layout.addWidget(self.lbl_consultas_12m)
        layout.addWidget(self.lbl_proxima_cita)
        layout.addWidget(self.lbl_ficha_vivienda)
        layout.addWidget(prog_btn)
        layout.addWidget(self.family_button)
        layout.addWidget(self.reminder_check)
        group.setLayout(layout)
        StyleHelper.style_groupbox(group)
        prog_btn.clicked.connect(self.open_appointments)
        self.family_button.clicked.connect(self.open_family_social)
        return group

    def update_dashboard(self, patient):
        if not patient:
            self.lbl_consultas_12m.setText("\U0001FA7A Consultas en 12 meses: 0")
            self.lbl_proxima_cita.setText("\U0001F4C5 Próxima cita: N/A")
            self.lbl_proxima_cita.setStyleSheet("")
            return
        session = Session()
        last_year = datetime.now() - timedelta(days=365)
        count = (
            session.query(HistoriaClinica)
            .filter(
                HistoriaClinica.paciente_id == patient.id,
                HistoriaClinica.fecha >= last_year,
            )
            .count()
        )
        self.lbl_consultas_12m.setText(f"\U0001FA7A Consultas en 12 meses: {count}")
        now = datetime.now()
        cita = (
            session.query(Cita)
            .filter(Cita.paciente_id == patient.id, Cita.fecha >= now)
            .order_by(Cita.fecha)
            .first()
        )
        if cita:
            txt = f"\U0001F4C5 Próxima cita: {cita.fecha:%d/%m/%Y} - {cita.notas or 'Medicina Familiar'}"
            self.lbl_proxima_cita.setText(txt)
            if cita.fecha.date() <= date.today() + timedelta(days=3):
                self.lbl_proxima_cita.setStyleSheet("background-color: yellow;")
            else:
                self.lbl_proxima_cita.setStyleSheet("")
        else:
            self.lbl_proxima_cita.setText("\U0001F4C5 Próxima cita: Ninguna")
            self.lbl_proxima_cita.setStyleSheet("")

        viv = (
            session.query(Vivienda)
            .filter_by(paciente_id=patient.id)
            .order_by(Vivienda.fecha.desc())
            .first()
        )
        grade = None
        if viv:
            if viv.calificacion:
                grade = viv.calificacion
            else:
                total = viv.total
                if total is None:
                    fields = [
                        viv.tenencia,
                        viv.paredes,
                        viv.techo,
                        viv.piso,
                        viv.sanitarios,
                        viv.agua_instalacion,
                        viv.agua_abastecimiento,
                        viv.basura,
                        viv.electricidad,
                        viv.dormitorios,
                        viv.combustible,
                    ]
                    total = sum(v for v in fields if isinstance(v, int))
                    try:
                        if viv.vectores is not None:
                            total += int(viv.vectores)
                    except ValueError:
                        pass
                if total >= 75:
                    grade = "Buena"
                elif total >= 50:
                    grade = "Regular"
                else:
                    grade = "Mala"
        if not grade:
            self.lbl_ficha_vivienda.setText("\U0001F3E0 Ficha de Vivienda: No realizada")
        else:
            self.lbl_ficha_vivienda.setText(
                f"\U0001F3E0 Ficha de Vivienda: {grade}"
            )
        
        session.close()




    def generate_random_code(length=3):
        import random
        import string
        return ''.join(random.choices(string.ascii_uppercase + string.digits, k=length))
        

    def load_patient_names(self):
        session = Session()
        patients = session.query(Paciente).all()
    
        # Crear una lista con los nombres completos de los pacientes
        patient_names = [patient.nombre for patient in patients]
    
        self.patient_completer.setModel(QStringListModel(patient_names))
        session.close()

        



    def search_patient(self):        
        session = Session()
        search_term = self.search_input.text()
        patients = session.query(Paciente).filter(Paciente.nombre.like(f"%{search_term}%")).all()
    
        self.patient_list.setRowCount(0)
        for patient in patients:
            row_position = self.patient_list.rowCount()
            self.patient_list.insertRow(row_position)
            id_item = QTableWidgetItem(str(patient.id))
            id_item.setTextAlignment(Qt.AlignCenter)
            name_item = QTableWidgetItem(patient.nombre)
            self.patient_list.setItem(row_position, 0, id_item)
            self.patient_list.setItem(row_position, 1, name_item)

            action_widget = QWidget()
            layout = QHBoxLayout(action_widget)
            layout.setContentsMargins(0, 0, 0, 0)
            edit_btn = QToolButton()
            edit_btn.setAutoRaise(True)
            del_btn = QToolButton()
            del_btn.setAutoRaise(True)
            if qta:
                try:
                    edit_btn.setIcon(qta.icon('fa5s.edit', color='#333'))
                    del_btn.setIcon(qta.icon('fa5s.trash', color='#333'))
                except Exception:
                    pass
            edit_btn.clicked.connect(partial(self.edit_patient_from_id, patient.id))
            del_btn.clicked.connect(partial(self.delete_patient_from_id, patient.id))
            layout.addWidget(edit_btn)
            layout.addWidget(del_btn)
            layout.addStretch()
            action_widget.setLayout(layout)
            action_widget.hide()
            self.patient_list.setCellWidget(row_position, 2, action_widget)
    
        # Ajusta el tamaño de las columnas y filas al contenido
        self.patient_list.resizeColumnsToContents()
        self.patient_list.resizeRowsToContents()
    
        # Hace que las columnas ocupen todo el espacio disponible
        header = self.patient_list.horizontalHeader()

        header.setSectionResizeMode(1, QHeaderView.Stretch)  # Columna de nombre
        header.setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.highlight_selected_patient()
    
        session.close()


    def new_patient(self):
            dialog = PatientDialog(self)
            if dialog.exec_():
                self.load_patient_names()
                self.search_patient()
                self.reload_database1()

               
     
    def reload_database1(self):
        try:
            if hasattr(self, 'session') and self.session:  # Verifica si la sesión está definida
                self.session.close()  # Cierra la sesión actual si existe
            
            self.session = Session()  # Crea una nueva sesión
            
            #elf.load_patient_list()  # Vuelve a cargar la lista de pacientes
            #self.update_history_list()  # Actualiza la lista de historias clínicas (asegúrate de que también exista esta función)

        except Exception as e:
            QMessageBox.warning(self, "Error", f"No se pudo recargar la base de datos: {str(e)}")                
               
    def reload_database(self):
        try:
            self.session.close()  # Cierra la sesión actual
            self.session = Session()  # Crea una nueva sesión
            self.load_patient_list()  # Vuelve a cargar la lista de pacientes
            self.update_history_list()  # Vuelve a actualizar la lista de historias clínicas
        except Exception as e:
            QMessageBox.warning(self, "Error", f"No se pudo recargar la base de datos: {str(e)}")

    def edit_patient(self):
        if not self.current_patient:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")
            return
    
        dialog = PatientDialog(self, self.current_patient)
        if dialog.exec_():
            try:
                session = Session()
                updated_patient = session.get(Paciente, self.current_patient.id)
                if updated_patient:
                    self.current_patient = updated_patient  # Actualiza la instancia del paciente
                    session.commit()
                    self.load_patient(self.patient_list.currentRow(), 0)
                    self.search_patient()
                    self.update_history_list()
                else:
                    QMessageBox.warning(self, "Error", "Paciente no encontrado.")
            except Exception as e:
                session.rollback()
                #QMessageBox.critical(self, "Error2", f"Error al actualizar el paciente: {str(e)}")
            finally:
                session.close()

    def edit_fixed_medications(self):
        if not self.current_patient:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")
            return
        dialog = FixedMedicationDialog(self.current_patient, self)
        if dialog.exec_() == QDialog.Accepted:
            self.load_patient_details(self.current_patient.nombre)



            
    def update_patient_in_database(self, patient):
        # Crea una nueva sesión de SQLAlchemy
        session = Session()
    
        # Actualiza la información del paciente en la base de datos
        session.merge(patient)
        session.commit()
    def keyPressEvent(self, event):
        if event.key() == Qt.Key_Delete:
            if self.history_list.hasFocus():
                self.delete_history()
            elif self.patient_list.hasFocus():  # Asumiendo que tienes una lista de pacientes
                self.delete_patient()

    def delete_patient(self):
        if not self.current_patient:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")
            return
        
        reply = QMessageBox.question(self, 'Confirmar Eliminación',
                                     f'¿Está seguro de que desea eliminar al paciente {self.current_patient.nombre} y todas sus historias clínicas?',
                                     QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            session = Session()
            try:
                # Eliminar todas las historias clínicas asociadas
                session.query(HistoriaClinica).filter(HistoriaClinica.paciente_id == self.current_patient.id).delete()
                
                # Eliminar al paciente
                session.delete(self.current_patient)
                session.commit()
                
                QMessageBox.information(self, "Éxito", f"El paciente {self.current_patient.nombre} y todas sus historias clínicas han sido eliminados.")
                
                # Refrescar la lista de pacientes en el autocompleter después de eliminar el paciente
                self.load_patient_names()
                
                # Limpiar la selección y la lista de historial en la interfaz de usuario
                self.search_patient()
                self.clear_patient_info()
                self.history_list.clear()
                
            except Exception as e:
                session.rollback()
                QMessageBox.critical(self, "Error", f"Error al eliminar el paciente: {str(e)}")
            finally:
                session.close()

    def edit_patient_from_id(self, pid):
        session = Session()
        patient = session.get(Paciente, pid)
        session.close()
        if patient:
            self.current_patient = patient
            self.edit_patient()

    def delete_patient_from_id(self, pid):
        session = Session()
        patient = session.get(Paciente, pid)
        session.close()
        if patient:
            self.current_patient = patient
            self.delete_patient()

    def load_patient(self, row, column):
        session = Session()
        try:
            patient_item = self.patient_list.item(row, 0)
            if patient_item is None:
                return
            patient_id = int(patient_item.text())
            stmt = select(Paciente).where(Paciente.id == patient_id)
            self.current_patient = session.execute(stmt).scalar_one_or_none()
            if self.current_patient:
                session.refresh(self.current_patient)  # Refrescar la instancia
                # Limpiar y restablecer campos
                self.clear_history_fields()
                self.clear_all_fields()
    
                # Restablecer tamaños originales de los campos
                self.historia_enfermedad_input.setInitialHeight(800)
                self.sugerencias_ia.setInitialHeight(800)
                self.interaccion_ia_input.setInitialHeight(800)
                self.interaccion_ia_output.setInitialHeight(800)

    
                # Cargar la información del paciente
                fecha = self.current_patient.fecha_nacimiento.strftime('%d/%m/%Y') if self.current_patient.fecha_nacimiento else ''
                meds = self.current_patient.medicamentos_continuos or ''
                self.patient_info.setText(
                    f"Nombre: {self.current_patient.nombre}\n"
                    f"Edad: {format_patient_age(self.current_patient)}\n"
                    f"Sexo: {self.current_patient.sexo}\n"
                    f"Alergias: {self.current_patient.alergias}\n"
                    f"Dirección: {self.current_patient.direccion}\n"
                    f"Estado Civil: {self.current_patient.estado_civil}\n"
                    f"Religión: {self.current_patient.religion}\n"
                    f"Lugar de Nacimiento: {self.current_patient.lugar_nacimiento}\n"
                    f"Fecha de Nacimiento: {fecha}\n"
                    f"Ocupación: {self.current_patient.ocupacion}\n"
                    f"Medicamentos Continuos: {meds}")

                mapping = {
                    "Nombre": self.current_patient.nombre,
                    "Cédula/ID": self.current_patient.documento_id or "",
                    "Edad": format_patient_age(self.current_patient),
                    "Sexo": self.current_patient.sexo,
                    "Alergias": self.current_patient.alergias,
                    "Dirección": self.current_patient.direccion,
                    "Estado Civil": self.current_patient.estado_civil,
                    "Religión": self.current_patient.religion,
                    "Lugar de Nacimiento": self.current_patient.lugar_nacimiento,
                    "Fecha de Nacimiento": fecha,
                    "Ocupación": self.current_patient.ocupacion,
                    "Medicamentos Continuos": meds,
                }
                for key, val in mapping.items():
                    label = self.patient_info_labels.get(key)
                    if label is not None:
                        label.setText(f"{key}: {val}")

                self.gestas_input.setText(self.current_patient.gestas_previas or "")
                self.abortos_input.setText(self.current_patient.abortos or "")
                self.vaginales_input.setText(self.current_patient.partos_vaginales or "")
                self.vivos_input.setText(self.current_patient.nacidos_vivos or "")
                self.cesareas_input.setText(self.current_patient.cesareas or "")
                self.gemelar_input.setText(self.current_patient.gemelar or "")
                self.fin_embarazo_input.setText(self.current_patient.fin_embarazo_anterior or "")
                self.planeado_check.setChecked(self.current_patient.embarazo_planeado == "si")
                self.metodo_input.setText(self.current_patient.metodo_anticonceptivo or "")
                self.update_obstetric_visibility()

                self.history_list.clear()
                historias = session.query(HistoriaClinica).filter(
                    HistoriaClinica.paciente_id == self.current_patient.id
                ).order_by(HistoriaClinica.fecha.desc()).all()

                for historia in historias:
                    self.add_history_item(historia)

                if self.history_list.count() == 0:
                    self.history_list.addItem("No hay historias clínicas previas")
                self.load_latest_vitals()
                self.update_dashboard(self.current_patient)
            else:
                QMessageBox.warning(self, "Error", "No se encontró el paciente.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al cargar el paciente: {str(e)}")
        finally:
            session.close()



    def update_history_list(self):
        if not self.current_patient:
            return
        self.history_list.clear()
        session = Session()
        historias = session.query(HistoriaClinica).filter_by(paciente_id=self.current_patient.id).order_by(HistoriaClinica.fecha.desc()).all()
        for historia in historias:
            self.add_history_item(historia)
        session.close()
        if self.history_list.count() > 0:
            self.history_list.setCurrentRow(0)
            self.highlight_history_item(self.history_list.item(0))

    def highlight_history_item(self, selected_item):
        for i in range(self.history_list.count()):
            item = self.history_list.item(i)
            widget = self.history_list.itemWidget(item)
            if widget:
                if item == selected_item:
                    widget.setStyleSheet("background-color:#d0e8e2;font-weight:bold;")
                    if hasattr(widget, 'check_label'):
                        widget.check_label.show()
                else:
                    widget.setStyleSheet("")
                    if hasattr(widget, 'check_label'):
                        widget.check_label.hide()

    def highlight_selected_patient(self):
        for row in range(self.patient_list.rowCount()):
            name_item = self.patient_list.item(row, 1)
            actions = self.patient_list.cellWidget(row, 2)
            if not name_item:
                continue
            if self.patient_list.item(row, 0).isSelected() or name_item.isSelected():
                if qta:
                    try:
                        name_item.setIcon(qta.icon("fa5s.user-check", color="black"))
                    except Exception:
                        pass
                font = name_item.font()
                font.setBold(True)
                name_item.setFont(font)
                if actions:
                    actions.show()
            else:
                name_item.setIcon(QIcon())
                font = name_item.font()
                font.setBold(False)
                name_item.setFont(font)
                if actions:
                    actions.hide()

    def load_latest_vitals(self):
        """Carga los últimos signos vitales del paciente seleccionado."""
        if not self.current_patient:
            self.bp_field.clear()
            self.hr_field.clear()
            self.weight_field.clear()
            return
        session = Session()
        signo = (
            session.query(SignoVital)
            .filter_by(paciente_id=self.current_patient.id)
            .order_by(SignoVital.fecha.desc())
            .first()
        )
        if signo:
            unit = getattr(signo, "peso_unidad", self.current_weight_unit)
            self.current_weight_unit = unit
            self.weight_unit_combo.setCurrentText(unit)
            self.bp_field.setText(signo.presion or "")
            self.hr_field.setText(str(signo.frecuencia or ""))
            display_weight = (
                self.convert_weight_from_kg(signo.peso or 0, unit)
                if signo.peso is not None
                else ""
            )
            self.weight_field.setText(str(display_weight))
            self.bmi_field.setText(f"{signo.imc:.1f}" if signo.imc else "")
            cat = (
                self.categorize_bmi(
                    signo.imc,
                    weight=signo.peso,
                    height_cm=self.current_patient.altura,
                )
                if signo.imc
                else ""
            )
            self.bmi_class_field.setText(cat)
            self.update_bmi_class_style(cat)
            self.gly_field.setText(str(signo.glicemia or ""))
            self.chol_field.setText(str(signo.colesterol or ""))
            self.oxi_field.setText(str(signo.oximetria or ""))
            self.temp_field.setText(str(signo.temperatura or ""))
            self.update_bp_style()
            self.update_hr_style()
            self.update_gly_style()
            self.update_chol_style()
            self.update_oxi_style()
            self.update_temp_style()
            self.compute_vitals_summary(signo)
        else:
            self.clear_vitals()
        session.close()

    def load_vitals_for_history(self, history):
        """Carga los signos vitales asociados a la historia dada."""
        session = Session()
        signo = (
            session.query(SignoVital)
            .filter_by(historia_id=history.id)
            .order_by(SignoVital.fecha.desc())
            .first()
        )
        if not signo:
            signo = (
                session.query(SignoVital)
                .filter(
                    SignoVital.paciente_id == history.paciente_id,
                    SignoVital.fecha <= history.fecha,
                )
                .order_by(SignoVital.fecha.desc())
                .first()
            )
        if signo:
            unit = getattr(signo, "peso_unidad", self.current_weight_unit)
            self.current_weight_unit = unit
            self.weight_unit_combo.setCurrentText(unit)
            self.bp_field.setText(signo.presion or "")
            self.hr_field.setText(str(signo.frecuencia or ""))
            display_weight = (
                self.convert_weight_from_kg(signo.peso or 0, unit)
                if signo.peso is not None
                else ""
            )
            self.weight_field.setText(str(display_weight))
            self.bmi_field.setText(f"{signo.imc:.1f}" if signo.imc else "")
            cat = (
                self.categorize_bmi(
                    signo.imc,
                    weight=signo.peso,
                    height_cm=self.current_patient.altura,
                )
                if signo.imc
                else ""
            )
            self.bmi_class_field.setText(cat)
            self.update_bmi_class_style(cat)
            self.gly_field.setText(str(signo.glicemia or ""))
            self.chol_field.setText(str(signo.colesterol or ""))
            self.oxi_field.setText(str(signo.oximetria or ""))
            self.temp_field.setText(str(signo.temperatura or ""))
            self.update_bp_style()
            self.update_hr_style()
            self.update_gly_style()
            self.update_chol_style()
            self.update_oxi_style()
            self.update_temp_style()
            self.compute_vitals_summary(signo)
        else:
            self.clear_vitals()
        session.close()

    def clear_vitals(self):
        self.bp_field.clear()
        self.bp_field.setStyleSheet("")
        self.hr_field.clear()
        self.hr_field.setStyleSheet("")
        self.weight_field.clear()
        self.bmi_field.clear()
        self.vitals_summary_label.clear()
        self.bmi_class_field.clear()
        self.update_bmi_class_style("")
        self.gly_field.clear()
        self.gly_field.setStyleSheet("")
        self.chol_field.clear()
        self.chol_field.setStyleSheet("")
        self.oxi_field.clear()
        self.oxi_field.setStyleSheet("")
        self.temp_field.clear()
        self.temp_field.setStyleSheet("")

    def update_bmi(self):
        if not self.current_patient:
            self.bmi_field.clear()
            self.bmi_class_field.clear()
            return
        try:
            weight = float(self.weight_field.text())
        except ValueError:
            self.bmi_field.clear()
            self.bmi_class_field.clear()
            return
        weight_kg = self.convert_weight_to_kg(weight, self.current_weight_unit)
        height_cm = self.current_patient.altura or 0
        height_m = height_cm / 100 if height_cm else 0
        if height_m and weight_kg:
            bmi = weight_kg / (height_m ** 2)
            self.bmi_field.setText(f"{bmi:.1f}")
            _idx, _val, cat = classify_nutritional_status(
                self.current_patient.fecha_nacimiento,
                weight_kg,
                height_cm,
                self.current_patient.sexo,
                GROWTH_TABLES,
            )
            self.bmi_class_field.setText(cat)
            self.update_bmi_class_style(cat)

        else:
            self.bmi_field.clear()
            self.bmi_class_field.clear()
            self.update_bmi_class_style("")

    def update_obstetric_visibility(self):
        pat = self.current_patient
        if not pat:
            self.obst_group.hide()
            return
        if pat.fecha_nacimiento:
            age = int((date.today() - pat.fecha_nacimiento).days / 365.25)
        else:
            age = pat.edad or 0
        is_female = str(pat.sexo).lower().startswith("f")
        visible = is_female and age >= 12
        self.obst_group.setVisible(visible)
        if visible:
            self.embarazo_activo_check.setChecked(bool(pat.embarazo_activo))
            weeks, _ = self.compute_current_pregnancy_weeks(date.today())
            if weeks is not None and weeks >= 42:
                self.embarazo_activo_check.setChecked(False)


    def categorize_bmi(self, bmi, weight=None, height_cm=None):
        if bmi is None:
            return ""
        patient = getattr(self, "current_patient", None)
        birth = getattr(patient, "fecha_nacimiento", None) if patient else None
        sex = getattr(patient, "sexo", "") if patient else ""
        if height_cm is None and patient:
            height_cm = patient.altura or 0
        if weight is None and bmi and height_cm:
            weight = bmi * ((height_cm / 100) ** 2)
        _, _, cat = classify_nutritional_status(
            birth, weight, height_cm, sex, GROWTH_TABLES
        )
        return cat

    def update_bmi_class_style(self, category):
        """Color the BMI classification field based on the result."""
        cat = (category or "").lower()
        color = ""
        text = "black"
        if not category:
            self.bmi_class_field.setStyleSheet(
                "padding:4px 8px; border-radius:8px; background:#eee; color:#555;"
            )
            return
        if "normal" in cat or "saludable" in cat:
            color = "#2ecc71"
            text = "white"
        elif "obes" in cat:
            color = "#e74c3c"
            text = "white"
        elif "sobrepeso" in cat:
            color = "#f39c12"
        elif "bajo" in cat or "delg" in cat or "desnut" in cat:
            color = "#f1c40f"
        if color:
            self.bmi_class_field.setStyleSheet(
                f"padding:4px 8px; border-radius:8px; background-color: {color}; color: {text};"
            )
        else:
            self.bmi_class_field.setStyleSheet(
                "padding:4px 8px; border-radius:8px; background:#eee; color:#555;"
            )

    def _color_for_level(self, level):
        if level == "normal":
            return "#2ecc71", "white"
        if level == "warning":
            return "#f1c40f", "black"
        if level == "danger":
            return "#e74c3c", "white"
        return "", ""

    def _set_field_level(self, field, level):
        color, text = self._color_for_level(level)
        if color:
            field.setStyleSheet(f"background-color: {color}; color: {text};")
        else:
            field.setStyleSheet("")

    def update_bp_style(self):
        txt = self.bp_field.text().strip()
        if not txt:
            self.bp_field.setStyleSheet("")
            return
        cat = self.categorize_bp(txt)
        if cat == "Normal":
            lvl = "normal"
        elif cat == "Elevada":
            lvl = "warning"
        else:
            lvl = "danger" if cat else ""
        self._set_field_level(self.bp_field, lvl)

    def update_hr_style(self):
        txt = self.hr_field.text().strip()
        if not txt:
            self.hr_field.setStyleSheet("")
            return
        try:
            val = int(txt)
        except ValueError:
            self._set_field_level(self.hr_field, "danger")
            return
        if 60 <= val <= 100:
            lvl = "normal"
        elif 50 <= val < 60 or 100 < val <= 110:
            lvl = "warning"
        else:
            lvl = "danger"
        self._set_field_level(self.hr_field, lvl)

    def update_gly_style(self):
        txt = self.gly_field.text().strip()
        if not txt:
            self.gly_field.setStyleSheet("")
            return
        try:
            val = float(txt)
        except ValueError:
            self._set_field_level(self.gly_field, "danger")
            return
        if 70 <= val <= 110:
            lvl = "normal"
        elif 60 <= val < 70 or 110 < val <= 125:
            lvl = "warning"
        else:
            lvl = "danger"
        self._set_field_level(self.gly_field, lvl)

    def update_chol_style(self):
        txt = self.chol_field.text().strip()
        if not txt:
            self.chol_field.setStyleSheet("")
            return
        try:
            val = float(txt)
        except ValueError:
            self._set_field_level(self.chol_field, "danger")
            return
        if val < 200:
            lvl = "normal"
        elif val < 240:
            lvl = "warning"
        else:
            lvl = "danger"
        self._set_field_level(self.chol_field, lvl)

    def update_oxi_style(self):
        txt = self.oxi_field.text().strip()
        if not txt:
            self.oxi_field.setStyleSheet("")
            return
        try:
            val = float(txt)
        except ValueError:
            self._set_field_level(self.oxi_field, "danger")
            return
        if val >= 95:
            lvl = "normal"
        elif val >= 90:
            lvl = "warning"
        else:
            lvl = "danger"
        self._set_field_level(self.oxi_field, lvl)

    def update_temp_style(self):
        txt = self.temp_field.text().strip()
        if not txt:
            self.temp_field.setStyleSheet("")
            return
        try:
            val = float(txt)
        except ValueError:
            self._set_field_level(self.temp_field, "danger")
            return
        if 36 <= val <= 37.5:
            lvl = "normal"
        elif 35 <= val < 36 or 37.5 < val <= 38:
            lvl = "warning"
        else:
            lvl = "danger"
        self._set_field_level(self.temp_field, lvl)

    MONTHS_ES = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    DAYS_ES = [
        "lunes", "martes", "miércoles", "jueves", "viernes", "sábado", "domingo"
    ]

    PROGRAM_OPTIONS = [
        "NINGUNO (CONSULTA LA DEMANDA)",
        "EDUCACIÓN PARA LA PREVENCIÓN DE LA VIOLENCIA INTRAFAMILIAR",
        "ATENCIÓN A LAS ENFERMEDADES PREVALENTES DE LA INFANCIA (AIEPI)",
        "ATENCIÓN A LA MALARIA",
        "ATENCIÓN AL DENGUE",
        "ATENCIÓN A LA ENFERMEDAD MENINGOCÓCICA",
        "EDUCACIÓN INDIVIDUAL Y COLECTIVA PARA LA PREVENCIÓN",
        "PROMOCIÓN DE LA SALUD",
        "INMUNIZACIÓN",
        "VIGILANCIA DEL CRECIMIENTO Y DESARROLLO",
        "VALORACIÓN DE RIESGO ADULTO MAYOR",
        "ATENCIÓN INTEGRAL A LA ADOLESCENTE",
        "ATENCIÓN PRENATAL",
        "DIAGNÓSTICO PRECOZ DEL CÁNCER GINECOLÓGICO",
        "PREVENCIÓN RIESGO CARDIOVASCULAR (HIPERTENSIÓN Y DIABETES)",
        "CONTROL Y TRATAMIENTO DE LA TUBERCULOSIS",
        "INFECCIONES DE TRANSMISIÓN SEXUAL Y SIDA",
        "ATENCIÓN A LA MORBILIDAD DE SALUD MENTAL",
        "INFORMACIÓN Y EDUCACIÓN EN SALUD AMBIENTAL",
        "CONTROL DE VECTORES",
        "PREVENCIÓN DE LA RABIA ANIMAL Y HUMANA",
        "PROGRAMA DE CÁNCER",
    ]

    def format_datetime_spanish(self, dt):
        day = self.DAYS_ES[dt.weekday()]
        month = self.MONTHS_ES[dt.month - 1]
        return f"{day} {dt.day} de {month} del {dt.year} {dt.strftime('%I:%M %p')}"

    def format_date_spanish(self, dt):
        day = self.DAYS_ES[dt.weekday()]
        month = self.MONTHS_ES[dt.month - 1]
        return f"{day} {dt.day} de {month} del {dt.year}"

    def format_time_12h(self, dt):
        return dt.strftime('%I:%M %p')

    def parse_datetime_spanish(self, text):
        m = re.match(
            r"(?:\w+\s+)?(\d{1,2}) de (\w+) del (\d{4}) (\d{1,2}):(\d{2}) ([AP]M)",
            text,
            re.IGNORECASE,
        )
        if not m:
            raise ValueError("Formato de fecha no válido")
        day, month_name, year, hour, minute, ampm = m.groups()
        month = self.MONTHS_ES.index(month_name.lower()) + 1
        hour = int(hour)
        if ampm.upper() == "PM" and hour != 12:
            hour += 12
        if ampm.upper() == "AM" and hour == 12:
            hour = 0
        return datetime(int(year), month, int(day), hour, int(minute))

    def add_history_item(self, historia):
        dt = historia.fecha
        text = f"Historia Clínica - {self.format_datetime_spanish(dt)}"
        item = QListWidgetItem(text)
        if qta:
            widget = QWidget()
            layout = QHBoxLayout(widget)
            layout.setContentsMargins(4, 2, 4, 2)
            check = QLabel()
            check.setPixmap(qta.icon('fa5s.check-circle', color='green').pixmap(32, 32))
            check.hide()
            layout.addWidget(check)
            note = QLabel()
            note.setPixmap(qta.icon('fa5s.sticky-note', color='black').pixmap(48, 48))
            layout.addWidget(note)
            layout.addWidget(QLabel('Historia Clínica'))
            layout.addWidget(QLabel('·'))
            cal = QLabel()
            cal.setPixmap(qta.icon('fa5s.calendar-alt', color='black').pixmap(48, 48))
            layout.addWidget(cal)
            layout.addWidget(QLabel(self.format_date_spanish(dt)))
            layout.addWidget(QLabel('·'))
            clock = QLabel()
            clock.setPixmap(qta.icon('fa5s.clock', color='black').pixmap(48, 48))
            layout.addWidget(clock)
            layout.addWidget(QLabel(self.format_time_12h(dt)))
            layout.addStretch()
            item.setSizeHint(widget.sizeHint())
            widget.check_label = check
            self.history_list.addItem(item)
            self.history_list.setItemWidget(item, widget)
        else:
            self.history_list.addItem(text)

    def parse_bp(self, bp):
        """Return systolic and diastolic integers or (None, None)."""
        if not bp or '/' not in bp:
            return None, None
        try:
            systolic, diastolic = [int(x) for x in bp.split('/')[:2]]
            return systolic, diastolic
        except ValueError:
            return None, None

    def categorize_bp(self, bp):
        if not bp or '/' not in bp:
            return ""
        try:
            systolic, diastolic = [int(x) for x in bp.split('/')[:2]]
        except ValueError:
            return ""
        if systolic >= 180 or diastolic >= 120:
            return "Crisis"
        if systolic >= 160 or diastolic >= 100:
            return "Hipertensión II"
        if systolic >= 140 or diastolic >= 90:
            return "Hipertensión I"
        if systolic > 120 or diastolic > 80:
            return "Elevada"
        return "Normal"

    def detect_programs(self, text):
        """Return a list of health programs mentioned in text."""
        if not text:
            return []
        text = text.lower()
        mapping = {
            "violencia": "EDUCACIÓN PARA LA PREVENCIÓN DE LA VIOLENCIA INTRAFAMILIAR",
            "infancia": "ATENCIÓN A LAS ENFERMEDADES PREVALENTES DE LA INFANCIA (AIEPI)",
            "malaria": "ATENCIÓN A LA MALARIA",
            "dengue": "ATENCIÓN AL DENGUE",
            "meningoc": "ATENCIÓN A LA ENFERMEDAD MENINGOCÓCICA",
            "promoción": "PROMOCIÓN DE LA SALUD",
            "inmuniz": "INMUNIZACIÓN",
            "crecimiento": "VIGILANCIA DEL CRECIMIENTO Y DESARROLLO",
            "adolesc": "ATENCIÓN INTEGRAL A LA ADOLESCENTE",
            "prenatal": "ATENCIÓN PRENATAL",
            "cáncer": "PROGRAMA DE CÁNCER",
            "cardio": "PREVENCIÓN RIESGO CARDIOVASCULAR (HIPERTENSIÓN Y DIABETES)",
            "tuberc": "CONTROL Y TRATAMIENTO DE LA TUBERCULOSIS",
            "sexual": "INFECCIONES DE TRANSMISIÓN SEXUAL Y SIDA",
            "salud mental": "ATENCIÓN A LA MORBILIDAD DE SALUD MENTAL",
            "rabia": "PREVENCIÓN DE LA RABIA ANIMAL Y HUMANA",
        }
        found = []
        for kw, prog in mapping.items():
            if kw in text and prog not in found:
                found.append(prog)
        if not found:
            found.append("NINGUNO (CONSULTA LA DEMANDA)")
        return found

    def convert_weight_to_kg(self, value, unit):
        if unit == "kg":
            return value
        if unit == "lb":
            return value * 0.453592
        return value

    def convert_weight_from_kg(self, value_kg, unit):
        if unit == "kg":
            return value_kg
        if unit == "lb":
            return value_kg / 0.453592
        return value_kg


    def maybe_prompt_form(self, text):
        """Ask the user to open a questionnaire if the AI suggests it."""
        if not text:
            return
        lower = text.lower()
        pat = self.current_patient
        age_months = None
        if pat and pat.fecha_nacimiento:
            age_months = months_between(pat.fecha_nacimiento)
        checks = [
            ("phq", PHQ9Dialog),
            ("gad", GAD7Dialog),
            ("mini-mental", MMSEDialog),
            ("mmse", MMSEDialog),
            ("denver", DenverDialog),
        ]
        for kw, dlg_class in checks:
            if kw in lower:
                if dlg_class in (PHQ9Dialog, GAD7Dialog, MMSEDialog) and (
                    age_months is None or age_months < 60
                ):
                    continue
                if QMessageBox.question(
                    self,
                    "Formulario sugerido",
                    f"La IA sugiere aplicar el formulario {kw}. \u00bfDesea abrirlo ahora?",
                ) == QMessageBox.Yes and dlg_class:
                    if dlg_class is DenverDialog:
                        age = None
                        pat = self.current_patient
                        if pat and pat.fecha_nacimiento:
                            age = months_between(pat.fecha_nacimiento)
                        dlg = dlg_class(self, pat, age)
                    else:
                        dlg = dlg_class(self, self.current_patient)
                    dlg.setWindowState(dlg.windowState() | Qt.WindowMaximized)
                    dlg.exec_()
                break

    def parse_ultrasound_lines(self, text):
        entries = []
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            m_date = re.search(r"(\d{1,2}/\d{1,2}/\d{4})", line)
            m_week = re.search(r"(\d+(?:\.\d+)?)\s*sem", line.lower())
            m_tri = re.search(r"(primer|segundo|tercer) trimestre", line.lower())
            if m_date and m_week:
                try:
                    d = datetime.strptime(m_date.group(1), "%d/%m/%Y").date()
                    w = float(m_week.group(1))
                    t = m_tri.group(1) if m_tri else ""
                    entries.append((w, d, t))
                except Exception:
                    continue
        return entries

    def add_ultrasound_entry(self):
        d = self.us_date_edit.date().toPyDate()
        w = self.us_weeks_spin.value()
        days = self.us_days_spin.value()
        total_w = w + days / 7.0
        if total_w <= 0:
            return
        if total_w < 14:
            tri = "primer trimestre"
        elif total_w < 28:
            tri = "segundo trimestre"
        else:
            tri = "tercer trimestre"
        # Calculate gestational age as of today
        today = date.today()
        current_w = total_w + (today - d).days / 7.0
        line = (
            f"{d.strftime('%d/%m/%Y')} - {total_w:.1f} semanas "
            f"(actual {current_w:.1f}) - {tri}"
        )
        cur = self.sonografias_input.toPlainText().strip()
        if cur:
            cur += "\n" + line
        else:
            cur = line
        self.sonografias_input.setPlainText(cur)

        # Update the illness history intro with the current gestational age
        weeks, start = self.compute_current_pregnancy_weeks(today)
        if weeks is not None:
            intro = f"Embarazo de {weeks:.1f} semanas.\n"
            text = self.historia_enfermedad_input.toPlainText()
            if text.startswith("Embarazo de "):
                text = re.sub(r"^Embarazo de .*?\n", intro, text, count=1)
            else:
                text = intro + text
            self.historia_enfermedad_input.setPlainText(text)
            if self.current_patient:
                schedule_dt_vaccine_reminders(self.current_patient.id, start)

    def format_ultrasound_summary(self, entries, ref_date=None):
        if not entries:
            return ""
        if ref_date is None:
            ref_date = datetime.now().date()
        parts = []
        for w, d, t in entries:
            weeks_now = w + (ref_date - d).days / 7.0
            months = weeks_now / 4.3
            if not t:
                if weeks_now < 14:
                    t = "primer trimestre"
                elif weeks_now < 28:
                    t = "segundo trimestre"
                else:
                    t = "tercer trimestre"
            parts.append(
                f"embarazo de {weeks_now:.1f} semanas ({months:.1f} meses) por sonografía del {t}"
            )
        return "; ".join(parts)

    def compute_current_pregnancy_weeks(self, ref_date=None, session=None):
        """Return gestational weeks and estimated start date using latest ultrasound."""
        if ref_date is None:
            ref_date = date.today()

        close = False
        if session is None:
            session = Session()
            close = True

        entries = []
        if self.current_patient:
            # Reload histories using the provided session to avoid detached errors
            hists = (
                session.query(HistoriaClinica)
                .filter_by(paciente_id=self.current_patient.id)
                .all()
            )
            for hist in hists:
                entries.extend(self.parse_ultrasound_lines(hist.sonografias or ""))

        entries.extend(self.parse_ultrasound_lines(self.sonografias_input.toPlainText()))
        if close:
            session.close()

        if not entries:
            return None, None
        latest = max(entries, key=lambda x: x[1])
        weeks = latest[0] + (ref_date - latest[1]).days / 7.0
        start_date = latest[1] - timedelta(weeks=latest[0])
        return weeks, start_date

    def compute_vitals_summary(self, current_signo):
        session = Session()
        previous = (
            session.query(SignoVital)
            .filter(
                SignoVital.paciente_id == current_signo.paciente_id,
                SignoVital.fecha < current_signo.fecha,
            )
            .order_by(SignoVital.fecha.desc())
            .first()
        )
        session.close()

        parts = []
        bmi_cat = ""
        if current_signo.imc and self.current_patient:
            bmi_cat = self.categorize_bmi(
                current_signo.imc,
                weight=current_signo.peso,
                height_cm=self.current_patient.altura,
            )
        if bmi_cat:
            parts.append(f"IMC {current_signo.imc:.1f} {bmi_cat}")
        bp_cat = self.categorize_bp(current_signo.presion)
        if bp_cat:
            parts.append(f"Hipertensión: {bp_cat}")
        if previous and previous.presion and current_signo.presion:
            prev_s, prev_d = self.parse_bp(previous.presion)
            curr_s, curr_d = self.parse_bp(current_signo.presion)
            if prev_s is not None and curr_s is not None:
                diff_s = curr_s - prev_s
                diff_d = curr_d - prev_d
                if diff_s or diff_d:
                    direction = "subió" if diff_s > 0 or diff_d > 0 else "bajó"
                    msg = f"PA {direction} {abs(diff_s)}/{abs(diff_d)}"
                    if curr_s >= 140 or curr_d >= 90:
                        msg += " (ALTO)"
                    parts.append(msg)
        if previous and previous.peso and current_signo.peso is not None:
            diff = current_signo.peso - previous.peso
            if abs(diff) >= 0.1:
                direction = "subió" if diff > 0 else "bajó"
                parts.append(f"Peso {direction} {abs(diff):.1f} kg")
        if previous and previous.glicemia is not None and current_signo.glicemia is not None:
            diffg = current_signo.glicemia - previous.glicemia
            if abs(diffg) >= 1:
                direction = "subió" if diffg > 0 else "bajó"
                parts.append(f"Glicemia {direction} {abs(diffg):.1f}")
        if previous and previous.colesterol is not None and current_signo.colesterol is not None:
            diffc = current_signo.colesterol - previous.colesterol
            if abs(diffc) >= 1:
                direction = "subió" if diffc > 0 else "bajó"
                parts.append(f"Colesterol {direction} {abs(diffc):.1f}")
        if previous and previous.oximetria is not None and current_signo.oximetria is not None:
            diffo = current_signo.oximetria - previous.oximetria
            if abs(diffo) >= 1:
                direction = "subió" if diffo > 0 else "bajó"
                parts.append(f"Oximetría {direction} {abs(diffo):.1f}")
        if previous and previous.temperatura is not None and current_signo.temperatura is not None:
            difft = current_signo.temperatura - previous.temperatura
            if abs(difft) >= 0.1:
                direction = "subió" if difft > 0 else "bajó"
                parts.append(f"Temp {direction} {abs(difft):.1f}")
        summary = "; ".join(parts)
        self.vitals_summary_label.setText(summary)
        return summary

    def risk_level(self, signo):
        """Return cardiovascular/metabolic risk estimate based on vitals."""
        score = 0
        sys_bp, dia_bp = self.parse_bp(signo.presion)
        if sys_bp is not None and dia_bp is not None:
            if sys_bp >= 160 or dia_bp >= 100:
                score += 2
            elif sys_bp >= 140 or dia_bp >= 90:
                score += 1
        if signo.imc and signo.imc >= 35:
            score += 2
        elif signo.imc and signo.imc >= 30:
            score += 1
        if signo.colesterol and signo.colesterol >= 240:
            score += 2
        elif signo.colesterol and signo.colesterol >= 200:
            score += 1
        if signo.glicemia and signo.glicemia >= 200:
            score += 2
        elif signo.glicemia and signo.glicemia >= 126:
            score += 1
        if signo.oximetria and signo.oximetria < 90:
            score += 2
        elif signo.oximetria and signo.oximetria < 95:
            score += 1
        if signo.temperatura and signo.temperatura >= 38:
            score += 1
        if score >= 6:
            return "Alto"
        if score >= 3:
            return "Moderado"
        return "Bajo"

    def on_weight_unit_changed(self, new_unit):
        try:
            current_value = float(self.weight_field.text())
        except ValueError:
            current_value = None
        if current_value is not None:
            kg = self.convert_weight_to_kg(current_value, self.current_weight_unit)
            converted = self.convert_weight_from_kg(kg, new_unit)
            self.weight_field.setText(f"{converted:.1f}")
        self.current_weight_unit = new_unit
        self.update_bmi()

    def maybe_launch_denver(self):
        """Open Denver II dialog when patient age matches and not yet done."""
        pat = self.current_patient
        if not pat or not pat.fecha_nacimiento:
            return
        age_months = months_between(pat.fecha_nacimiento)
        target = [2, 6, 12, 24, 36]
        done = set(filter(None, (pat.denver_done or '').split(';')))
        if age_months in target and str(age_months) not in done:
            dlg = DenverDialog(self, pat, age_months)
            dlg.setWindowState(dlg.windowState() | Qt.WindowMaximized)
            dlg.exec_()

    def organize_diagnoses(self, diagnoses):
        lines = diagnoses.split('\n')
        covered = []
        not_covered = []
    
        for line in lines:
            lower_line = line.lower()
            if "(cubierto)" in lower_line:
                covered.append(line)
            elif "(no cubierto)" in lower_line:
                not_covered.append(line)
    
        # Organizar en el orden deseado
        organized_diagnoses = "\n".join(covered + not_covered)
        return organized_diagnoses

    

    
############################################################################33


    def load_selected_history(self, item):
        self.clear_all_fields()
        self.highlight_history_item(item)
        if item.text().startswith("Historia Clínica"):
            try:
                # Extraer la fecha y hora del texto del ítem
                parts = item.text().rsplit(" - ", 1)
                if len(parts) < 2:
                    raise ValueError("Formato de texto no válido")
                date_str = parts[-1]
                date_time = self.parse_datetime_spanish(date_str)
            except ValueError as e:
                QMessageBox.warning(self, "Error", f"Formato de fecha no válido: {e}")
                return
    
            session = Session()
            try:
                # Re-cargar el paciente para asegurar que está asociado a la sesión actual
                paciente = session.get(Paciente, self.current_patient.id)
                if not paciente:
                    raise ValueError("Paciente no encontrado")
    
                historia = session.query(HistoriaClinica).filter(
                    HistoriaClinica.paciente_id == paciente.id,
                    func.date(HistoriaClinica.fecha) == date_time.date(),
                    func.strftime('%H:%M', HistoriaClinica.fecha) == date_time.strftime('%H:%M')
                ).first()
    
                if historia:
                    self.current_history_id = historia.id
                    self.load_history(historia)
                else:
                    QMessageBox.warning(self, "Error", f"No se pudo encontrar la historia clínica para la fecha {date_str}")
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Error al cargar la historia clínica: {str(e)}")
            finally:
                session.close()

    def new_history(self):
        if not self.current_patient:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")
            return
        self.clear_history_fields()
        self.current_history_id = None  # Resetear el ID de la historia actual
        
        # Crear una nueva historia clínica
        current_datetime = datetime.now()
        session = Session()
    
        # Obtener los antecedentes personales y heredofamiliares del paciente actual
        antecedentes_personales = ""
        antecedentes_heredofamiliares = ""
        
        # Si hay historias clínicas previas, cargar los antecedentes del más reciente
        historias_previas = session.query(HistoriaClinica).filter_by(paciente_id=self.current_patient.id).order_by(HistoriaClinica.fecha.desc()).first()
        prev_us = ""
        if historias_previas:
            antecedentes_personales = historias_previas.antecedentes_personales
            antecedentes_heredofamiliares = historias_previas.antecedentes_heredofamiliares
            prev_us = historias_previas.sonografias or ""
    
        nueva_historia = HistoriaClinica(
            paciente_id=self.current_patient.id,
            fecha=current_datetime,
            antecedentes_personales=antecedentes_personales,
            antecedentes_heredofamiliares=antecedentes_heredofamiliares,
            historia_enfermedad="",
            sugerencias_ia="",
            interaccion_ia="",
            sonografias=prev_us,
            codigo=generate_random_code()
        )
        session.add(nueva_historia)
        session.commit()
        self.current_history_id = nueva_historia.id
        session.refresh(nueva_historia)

        # Autocompletar los campos de antecedentes antes de cerrar la sesión
        self.load_history(nueva_historia)
        weeks, _ = self.compute_current_pregnancy_weeks(current_datetime.date())
        if weeks is not None and weeks < 42:
            intro = f"Embarazo de {weeks:.1f} semanas.\n"
            if not self.historia_enfermedad_input.toPlainText().startswith(intro):
                self.historia_enfermedad_input.insertPlainText(intro)
            self.embarazo_activo_check.setChecked(True)
        session.close()

        # Actualizar la lista de historias
        self.update_history_list()
        
        # Seleccionar la nueva historia en la lista
        formatted = self.format_datetime_spanish(current_datetime)
        for i in range(self.history_list.count()):
            if self.history_list.item(i).text().endswith(formatted):
                self.history_list.setCurrentRow(i)
                self.highlight_history_item(self.history_list.item(i))
                break
    
        self.clear_vitals()

        QMessageBox.information(self, "Nueva Historia", f"Se ha creado una nueva historia clínica. ID: {self.current_history_id}")
        self.maybe_launch_denver()
    
    def load_history(self, history):
        self.clear_all_fields()
        self.current_patient = history.paciente

        # Cargar antecedentes personales
        if history.antecedentes_personales:
            antecedentes = history.antecedentes_personales.split(';')
            for antecedente in antecedentes:
                if ':' in antecedente:
                    key, value = antecedente.split(':', 1)
                    key = key.strip()
                    value = value.strip()
                    if key in self.checkboxes:
                        self.checkboxes[key].setChecked(value.lower() not in ['false', 'negado', ''])
                    if key in self.detail_inputs:
                        self.detail_inputs[key].setText(value)
    
        # Cargar antecedentes familiares
        if history.antecedentes_heredofamiliares:
            family_history = history.antecedentes_heredofamiliares.split(';')
            for item in family_history:
                if ':' in item:
                    key, value = item.split(':', 1)
                    key = key.strip()
                    value = value.strip()
                    if key == 'Padre':
                        self.padre_input.setText(value)
                    elif key == 'Madre':
                        self.madre_input.setText(value)
                    elif key == 'Hijos':
                        self.hijos_input.setText(value)
                    elif key == 'Hermanos':
                        self.hermanos_input.setText(value)
    
        # Establecer otros campos
        prev_meds = self.get_previous_medications(history.paciente_id, history.fecha)
        if prev_meds:
            self.diagnosticos_previos_input.setText(prev_meds)
        else:
            self.diagnosticos_previos_input.setText('No hay medicamentos previos')
        self.historia_enfermedad_input.setText(history.historia_enfermedad or "")
        p = self.current_patient
        self.gestas_input.setText(p.gestas_previas or "")
        self.abortos_input.setText(p.abortos or "")
        self.vaginales_input.setText(p.partos_vaginales or "")
        self.vivos_input.setText(p.nacidos_vivos or "")
        self.cesareas_input.setText(p.cesareas or "")
        self.gemelar_input.setText(p.gemelar or "")
        self.fin_embarazo_input.setText(p.fin_embarazo_anterior or "")
        self.planeado_check.setChecked(p.embarazo_planeado == "si")
        self.embarazo_activo_check.setChecked(bool(p.embarazo_activo))
        self.metodo_input.setText(p.metodo_anticonceptivo or "")
        self.sonografias_input.setText(history.sonografias or "")
        self.update_obstetric_visibility()
        self.current_programs = []
        if history.programa_salud:
            self.current_programs = [s.strip() for s in history.programa_salud.split(';') if s.strip()]
        self.sugerencias_ia.setText(history.sugerencias_ia or "")
        self.interaccion_ia_output.setText(history.interaccion_ia or "")
    
        # Restablecer tamaños originales de los campos
        self.historia_enfermedad_input.setInitialHeight(800)
        self.sugerencias_ia.setInitialHeight(800)
        self.interaccion_ia_input.setInitialHeight(800)
        self.interaccion_ia_output.setInitialHeight(800)

        self.load_vitals_for_history(history)

##############################################################################




    def clear_all_fields(self):
        # Limpiar checkboxes y campos de detalle
        for checkbox in self.checkboxes.values():
            checkbox.setChecked(False)
        for input_field in self.detail_inputs.values():
            input_field.clear()

        self.diagnosticos_previos_input.clear()
        # Limpiar campos de antecedentes familiares
        self.padre_input.clear()
        self.madre_input.clear()
        self.hijos_input.clear()
        self.hermanos_input.clear()
        self.gestas_input.clear()
        self.abortos_input.clear()
        self.vaginales_input.clear()
        self.vivos_input.clear()
        self.cesareas_input.clear()
        self.gemelar_input.clear()
        self.fin_embarazo_input.clear()
        self.planeado_check.setChecked(False)
        self.embarazo_activo_check.setChecked(False)
        self.metodo_input.clear()
        self.sonografias_input.clear()
        self.current_programs = []

        # Limpiar otros campos
        self.historia_enfermedad_input.clear()
        self.sugerencias_ia.clear()
        self.interaccion_ia_output.clear()
        self.bp_field.clear()
        self.hr_field.clear()
        self.weight_field.clear()
        self.bmi_field.clear()
        self.vitals_summary_label.clear()
                # Restablecer tamaños originales de los campos
        self.historia_enfermedad_input.setInitialHeight(500)
        self.sugerencias_ia.setInitialHeight(400)
        self.interaccion_ia_input.setInitialHeight(400)
        self.interaccion_ia_output.setInitialHeight(400)


    def clear_history_fields(self):
        for checkbox in self.checkboxes.values():
            checkbox.setChecked(False)
        for input_field in self.detail_inputs.values():
            input_field.clear()
        self.padre_input.clear()
        self.madre_input.clear()
        self.hijos_input.clear()
        self.hermanos_input.clear()
        self.gestas_input.clear()
        self.abortos_input.clear()
        self.vaginales_input.clear()
        self.vivos_input.clear()
        self.cesareas_input.clear()
        self.gemelar_input.clear()
        self.fin_embarazo_input.clear()
        self.planeado_check.setChecked(False)
        self.metodo_input.clear()
        self.sonografias_input.clear()
        self.historia_enfermedad_input.clear()
        self.sugerencias_ia.clear()
        self.interaccion_ia_input.clear()
        self.interaccion_ia_output.clear()
        self.bp_field.clear()
        self.hr_field.clear()
        self.weight_field.clear()
        self.bmi_field.clear()
                # Restablecer tamaños originales de los campos
        self.historia_enfermedad_input.setInitialHeight(500)
        self.sugerencias_ia.setInitialHeight(400)
        self.interaccion_ia_input.setInitialHeight(400)
        self.interaccion_ia_output.setInitialHeight(400)

    def clear_patient_info(self):
        self.patient_info.clear()
        for field, label in getattr(self, 'patient_info_labels', {}).items():
            label.setText(f"{field}:")
        self.clear_history_fields()
        self.current_patient = None

    def delete_history(self):
        current_item = self.history_list.currentItem()
        if not current_item:
            QMessageBox.warning(self, "Error", "No hay historia seleccionada.")
            return
        
        # Confirmar la eliminación
        reply = QMessageBox.question(self, 'Confirmar Eliminación',
                                     '¿Está seguro de que desea eliminar esta historia clínica?',
                                     QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.No:
            return
        
        # Extraer la fecha y hora del texto del ítem
        date_str = current_item.text().split(" - ")[1]
        try:
            date_time = self.parse_datetime_spanish(date_str)
        except ValueError:
            QMessageBox.warning(self, "Error", "Formato de fecha no válido.")
            return
        
        session = Session()
        try:
            # Buscar la historia clínica correspondiente
            historia = session.query(HistoriaClinica).filter(
                HistoriaClinica.paciente_id == self.current_patient.id,
                func.date(HistoriaClinica.fecha) == date_time.date(),
                func.strftime('%H:%M', HistoriaClinica.fecha) == date_time.strftime('%H:%M')
            ).first()
            
            if historia:
                session.delete(historia)
                session.commit()
                self.history_list.takeItem(self.history_list.row(current_item))
                self.clear_all_fields()
                QMessageBox.information(self, "Éxito", "Historia clínica eliminada correctamente.")
            else:
                QMessageBox.warning(self, "Error", "No se pudo encontrar la historia clínica.")
        
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al eliminar la historia clínica: {str(e)}")
        finally:
            session.close()     

  

    # --------------------------------------------------------------------
    #  MainWindow.save_history  —  versión final
    # --------------------------------------------------------------------
    def save_history(self):
        """
        Guarda la historia clínica:
          • Si YA existe una historia para el mismo paciente y misma fecha
            (UTC), la actualiza en lugar de crear otra.
          • Después de commit ⇒ guarda diagnósticos y registro diario.
        """
        if not self.current_patient:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")
            return
    
        session = Session()
        try:
            historia = None
            new_history = False
            # 1 ─ Si ya hay una historia seleccionada => editar esa
            if self.current_history_id:
                historia = session.get(HistoriaClinica, self.current_history_id)

            if not historia:
                hoy = datetime.now().date()
                historia = (
                    session.query(HistoriaClinica)
                    .filter(
                        HistoriaClinica.paciente_id == self.current_patient.id,
                        func.date(HistoriaClinica.fecha) == hoy,
                    )
                    .first()
                )
                if historia:
                    self.current_history_id = historia.id
                else:
                    historia = HistoriaClinica(
                        paciente_id=self.current_patient.id,
                        fecha=datetime.now(),
                    )
                    session.add(historia)
                    new_history = True
    
            # 2 ─ Copiar datos GUI
            historia.antecedentes_personales       = self.get_antecedentes_personales()
            historia.antecedentes_heredofamiliares = self.get_antecedentes_heredofamiliares()
            historia.historia_enfermedad           = self.historia_enfermedad_input.toPlainText()
            historia.sugerencias_ia                = self.sugerencias_ia.toPlainText()
            historia.interaccion_ia                = self.interaccion_ia_output.toPlainText()
            if new_history:
                historia.fecha = datetime.now()
            historia.programa_salud                = ";".join(self.current_programs)
            historia.sonografias                   = self.sonografias_input.toPlainText()

            # actualizar antecedentes obstétricos del paciente
            p = session.get(Paciente, self.current_patient.id)
            p.gestas_previas        = self.gestas_input.text().strip()
            p.abortos               = self.abortos_input.text().strip()
            p.partos_vaginales      = self.vaginales_input.text().strip()
            p.nacidos_vivos         = self.vivos_input.text().strip()
            p.cesareas              = self.cesareas_input.text().strip()
            p.gemelar               = self.gemelar_input.text().strip()
            p.fin_embarazo_anterior = self.fin_embarazo_input.text().strip()
            p.embarazo_planeado     = "si" if self.planeado_check.isChecked() else "no"
            p.metodo_anticonceptivo = self.metodo_input.text().strip()
            p.embarazo_activo = self.embarazo_activo_check.isChecked()
            weeks, start = self.compute_current_pregnancy_weeks(historia.fecha.date(), session)
            if weeks is not None:
                p.embarazo_inicio = start
                if weeks >= 42:
                    p.embarazo_activo = False
                else:
                    schedule_dt_vaccine_reminders(p.id, start)

            # 3 ─ Confirmar
            session.commit()
            self.current_history_id = historia.id

            # Guardar signos vitales asociados a esta historia
            bp = self.bp_field.text().strip()
            hr = self.hr_field.text().strip()
            wt = self.weight_field.text().strip()
            gly = self.gly_field.text().strip()
            chol = self.chol_field.text().strip()
            oxi = self.oxi_field.text().strip()
            temp = self.temp_field.text().strip()
            if bp or hr or wt or gly or chol or oxi or temp:
                weight_val = float(wt) if wt else None
                weight_kg = self.convert_weight_to_kg(weight_val, self.current_weight_unit) if weight_val is not None else None
                height_cm = self.current_patient.altura or 0
                height_m = height_cm / 100 if height_cm else 0
                bmi = weight_kg / height_m ** 2 if weight_kg and height_m else None
                last = (
                    session.query(SignoVital)
                    .filter_by(paciente_id=self.current_patient.id)
                    .order_by(SignoVital.fecha.desc())
                    .first()
                )
                existing = (
                    session.query(SignoVital)
                    .filter_by(historia_id=historia.id)
                    .order_by(SignoVital.fecha.desc())
                    .first()
                )
                if existing:
                    signo = existing
                    signo.fecha = historia.fecha
                    signo.presion = bp or None
                    signo.frecuencia = int(hr) if hr else None
                    signo.peso = weight_kg
                    signo.imc = bmi
                    signo.glicemia = float(gly) if gly else None
                    signo.colesterol = float(chol) if chol else None
                    signo.oximetria = float(oxi) if oxi else None
                    signo.temperatura = float(temp) if temp else None
                    signo.peso_unidad = self.current_weight_unit
                    signo.risk_processed = False
                else:
                    signo = SignoVital(
                        paciente_id=self.current_patient.id,
                        historia_id=historia.id,
                        fecha=historia.fecha,
                        presion=bp or None,
                        frecuencia=int(hr) if hr else None,
                        peso=weight_kg,
                        imc=bmi,
                        glicemia=float(gly) if gly else None,
                        colesterol=float(chol) if chol else None,
                        oximetria=float(oxi) if oxi else None,
                        temperatura=float(temp) if temp else None,
                        peso_unidad=self.current_weight_unit,
                        risk_processed=False,
                    )
                    session.add(signo)
                session.commit()

                change_msg = []
                if last:
                    if last.peso and weight_kg and weight_kg != last.peso:
                        diff = weight_kg - last.peso
                        change_msg.append(
                            f"Peso {'subió' if diff>0 else 'bajó'} {abs(diff):.1f} kg"
                        )
                    if last.presion and bp:
                        prev_s, prev_d = self.parse_bp(last.presion)
                        curr_s, curr_d = self.parse_bp(bp)
                        if prev_s is not None and curr_s is not None and (curr_s != prev_s or curr_d != prev_d):
                            direction = "subió" if (curr_s > prev_s or curr_d > prev_d) else "bajó"
                            msg = f"PA {direction} {abs(curr_s - prev_s)}/{abs(curr_d - prev_d)}"
                            if curr_s >= 140 or curr_d >= 90:
                                msg += " (ALTO)"
                            change_msg.append(msg)
                    if last.frecuencia and hr and int(hr) != last.frecuencia:
                        diff_hr = int(hr) - last.frecuencia
                        change_msg.append(
                            f"FC {'subió' if diff_hr>0 else 'bajó'} {abs(diff_hr)}"
                        )
                    if last.imc and bmi and round(bmi,1) != round(last.imc,1):
                        diff_imc = bmi - last.imc
                        change_msg.append(
                            f"IMC {'subió' if diff_imc>0 else 'bajó'} {abs(diff_imc):.1f}"
                        )
                    if last.glicemia is not None and signo.glicemia is not None and signo.glicemia != last.glicemia:
                        diff_g = signo.glicemia - last.glicemia
                        change_msg.append(
                            f"Glicemia {'subió' if diff_g>0 else 'bajó'} {abs(diff_g):.1f}"
                        )
                    if last.colesterol is not None and signo.colesterol is not None and signo.colesterol != last.colesterol:
                        diff_c = signo.colesterol - last.colesterol
                        change_msg.append(
                            f"Colesterol {'subió' if diff_c>0 else 'bajó'} {abs(diff_c):.1f}"
                        )
                if signo.glicemia is not None and signo.glicemia > 126:
                    change_msg.append("Glicemia alta")
                if signo.colesterol is not None and signo.colesterol > 200:
                    change_msg.append("Colesterol alto")
                summary = self.compute_vitals_summary(signo)
                if change_msg or summary:
                    text = "\n".join(change_msg)
                    if summary:
                        text = (text + "\n" if text else "") + summary
                    self.historia_enfermedad_input.append("\n" + text)

            # 4 ─ Post-commit
            self.store_diagnoses(historia, historia.sugerencias_ia, session)
            self.store_consultation_register(historia, session)

            # Resumen automático con GPT (deshabilitado al guardar)
            # resumen = self.generate_history_summary(historia)
            # if resumen:
            #     QMessageBox.information(self, "Resumen", resumen)

            # 5 ─ UI feedback
            QMessageBox.information(self, "Éxito",
                                    f"Historia clínica guardada (ID {self.current_history_id}).")
            self.update_history_list()
    
        except Exception as e:
            session.rollback()
            QMessageBox.critical(self, "Error", f"Error al guardar la historia clínica:\n{e}")
        finally:
            session.close()

    # Agregar este método para limpiar los datos existentes
    def clean_existing_data(self):
        session = Session()
        historias = session.query(HistoriaClinica).all()
        for historia in historias:
            antecedentes = historia.antecedentes_personales.split(';')
            cleaned_antecedentes = []
            for antecedente in antecedentes:
                parts = antecedente.split(':')
                if len(parts) >= 2:
                    key = parts[0].strip()
                    value = ':'.join(parts[1:]).strip()
                    if value.lower().startswith('false:'):
                        value = ':'.join(parts[2:]).strip()
                    if value.lower() == 'false' or value.lower() == 'true' or not value:
                        cleaned_antecedentes.append(f"{key}: negado")
                    else:
                        cleaned_antecedentes.append(f"{key}: {value}")
            historia.antecedentes_personales = '; '.join(cleaned_antecedentes)
        session.commit()
        session.close()

    def cargar_historia_clinica(self, historia_clinica):
        """
        Carga los datos de una historia clínica guardada, actualizando el paciente activo
        y los campos del formulario.
        """
        self.current_historia = historia_clinica
        self.current_patient = historia_clinica.paciente
    
        # Rellenar los campos con la información de la historia clínica guardada
        self.textedit_antecedentes_personales.setText(historia_clinica.antecedentes_personales or "")
        self.textedit_antecedentes_heredofamiliares.setText(historia_clinica.antecedentes_heredofamiliares or "")
        self.textedit_historia_enfermedad.setText(historia_clinica.historia_enfermedad or "")
        self.textedit_sugerencias_ia.setText(historia_clinica.sugerencias_ia or "")
        self.textedit_interaccion_ia.setText(historia_clinica.interaccion_ia or "")
    
        # Refrescar información de paciente (si tienes labels o campos visibles)
        if self.current_patient:
            self.label_nombre.setText(self.current_patient.nombre or "")
            self.label_edad.setText(str(self.current_patient.edad) if self.current_patient.edad else "")
            self.label_sexo.setText(self.current_patient.sexo or "")
            self.label_alergias.setText(self.current_patient.alergias or "") 

    def print_receta(self):
        if not self.current_patient:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")
            return
    
        # Obtener sugerencias (si vienen de una historia histórica, se cargan en load_history) :contentReference[oaicite:0]{index=0}:contentReference[oaicite:1]{index=1}
        sugerencias = self.sugerencias_ia.toPlainText().strip()
    
        # Si no hay sugerencias (historia antigua sin generar), invocar a la IA
        if not sugerencias:
            # Genera interacción IA que llenará sugerencias_ia y guardará en BD
            self.interact_with_ia()
            # Recarga el texto (ya debería estar en self.sugerencias_ia)
            sugerencias = self.sugerencias_ia.toPlainText().strip()
            if not sugerencias:
                QMessageBox.warning(self, "Error", "No se pudo generar sugerencias con la IA.")
                return
    
        # Construir contexto completo para extracción de prescripción
        medicamentos_continuos = getattr(self.current_patient, 'medicamentos_continuos', '').strip()
        alergias = getattr(self.current_patient, 'alergias', '').strip()
    
        texto_completo = (
                   
        
            f"Medications:\n{sugerencias}\n\n"
            f"Continuous Medications:\n{medicamentos_continuos}\n\n"
            f"Suggested Medications from History:\n{sugerencias}\n\n"
            f"Patient Allergies:\n{alergias}"
        )
    
        previous_diagnoses = self.get_current_history_diagnoses()

        review_dialog = PrescriptionReviewDialog(
            "",
            previous_diagnoses,
            alergias,
            self.current_patient
        )
        review_dialog.medications_text.setText("Escribiendo...")

        thread = self.create_prescription_thread(texto_completo)
        if not thread:
            return

        thread.chunk_received.connect(review_dialog.append_med_chunk)
        thread.finished.connect(review_dialog.finish_med_chunk)
        thread.error.connect(lambda m: QMessageBox.warning(self, "Error", m))
        thread.start()

        if review_dialog.exec_() != QDialog.Accepted:
            thread.quit()
            thread.wait()
            return

        prescripcion = review_dialog.medications_text.toPlainText().strip()
        if not prescripcion:
            QMessageBox.warning(self, "Error", "No se encontró información de prescripción.")
            return
    
        prescripcion = self.clean_prescription_text(prescripcion)
    
        # Creación y guardado de documentos Word
        def crear_documento(texto, filepath):
            doc = Document()
            cfg = ConfigReceta(); cfg.load_config()
            try:
                logo_path = os.path.join(os.path.dirname(__file__), cfg.logo_path)
                table = doc.add_table(rows=1, cols=3)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell_logo = table.cell(0, 0)
                cell_head = table.cell(0, 1)
                cell_qr = table.cell(0, 2)
                run = cell_logo.paragraphs[0].add_run()
                run.add_picture(logo_path, width=Inches(0.8))
                header = cfg.get_clinic_header()
                if header:
                    para = cell_head.paragraphs[0]
                    para.text = header
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                qr_path = os.path.join(tempfile.gettempdir(), 'rx_qr.png')
                create_qr_code(f"Rx-{self.current_patient.nombre}-{datetime.now().strftime('%Y%m%d')}", qr_path)
                run = cell_qr.paragraphs[0].add_run()
                run.add_picture(qr_path, width=Inches(1))
            except Exception as e:
                QMessageBox.warning(self, "Advertencia", f"No se pudo agregar el logo: {e}")

            # Fecha y encabezados
            fecha = datetime.now().strftime("%d de %B de %Y")
            p = doc.add_paragraph(f"Fecha: {fecha}")
            p.alignment = 1
            h = doc.add_heading("Receta Médica", level=1)
            h.alignment = 1

            # Info paciente y prescripción
            doc.add_heading("Datos del Paciente", level=2)
            doc.add_paragraph(f"Paciente: {self.current_patient.nombre}")
            doc.add_paragraph(f"Edad: {self.current_patient.edad}")
            if self.current_patient.edad and self.current_patient.edad <= 12:
                peso = 2 * self.current_patient.edad + 8
                doc.add_paragraph(f"Peso estimado: {peso} kg")
            doc.add_heading("Prescription", level=2)
            doc.add_paragraph(texto)

            # Firma digital
            sign_path = cfg.get_signature()
            seal_path = cfg.get_seal()
            doctor = cfg.get_doctor_name()
            specialty = cfg.get_doctor_specialty()

            if doctor or sign_path or seal_path:
                table = doc.add_table(rows=1, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell_sig = table.cell(0, 0)
                cell_seal = table.cell(0, 1)
                if sign_path and os.path.exists(sign_path):
                    try:
                        run = cell_sig.paragraphs[0].add_run()
                        run.add_picture(sign_path, width=Inches(1.5))
                    except Exception:
                        pass
                if doctor:
                    cell_sig.add_paragraph(f"{doctor}\n{specialty}").alignment = 1
                if seal_path and os.path.exists(seal_path):
                    try:
                        run = cell_seal.paragraphs[0].add_run()
                        run.add_picture(seal_path, width=Inches(2.5))
                    except Exception:
                        pass
    
            # Pie de página
            footer = doc.sections[0].footer
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.text = f"{cfg.get_footer_text()}\n{cfg.get_footer_details()}\n{cfg.get_footer_phone()}"
            fp.alignment = 1
    
            try:
                doc.save(filepath)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error al crear la receta: {e}")
    
        # Sanitizar y generar nombres únicos
        def sanitize(name): return re.sub(r'[<>:"/\\|?*\s]+', '_', name)[:30]
        def unique(base, fn):
            name, ext = os.path.splitext(fn)
            cnt = 1
            while os.path.exists(os.path.join(base, fn)):
                fn = f"{name}_{cnt}{ext}"; cnt += 1
            return fn
    
        san = sanitize(self.current_patient.nombre)
        fecha_tag = datetime.now().strftime('%Y%m%d')
        default = f"Rx_{san}_{fecha_tag}.docx"
        dest = os.path.join(expanduser("~"), "Documents", "recetas", san)
        os.makedirs(dest, exist_ok=True)
    
        # Diálogo de guardado
        opts = QFileDialog.Options()
        path, _ = QFileDialog.getSaveFileName(
            self, "Guardar Receta", os.path.join(dest, default),
            "Archivos Word (*.docx);;Todos los archivos (*)", options=opts
        )
        if not path:
            QMessageBox.warning(self, "Aviso", "No se seleccionó ubicación para guardar.")
            return
        if not path.lower().endswith('.docx'):
            path += '.docx'
    
        base_dir, fn = os.path.split(path)
        fn_unique = unique(base_dir, fn)
        full_path = os.path.join(base_dir, fn_unique)
    
        crear_documento(prescripcion, full_path)
        QMessageBox.information(self, "Éxito", f"Receta exportada como {fn_unique}")

        try:
            os.startfile(full_path)
        except Exception:
            pass

        reply = QMessageBox.question(
            self,
            "Recordatorio",
            "¿Desea crear un recordatorio de seguimiento para este tratamiento en 15 días?",
        )
        if reply == QMessageBox.Yes and self.current_patient:
            dlg = ReminderDialog(self)
            dlg.patient_combo.setCurrentText(self.current_patient.nombre)
            dlg.desc_edit.setText("Seguimiento de tratamiento")
            dlg.type_combo.setCurrentText("Tratamiento")
            dlg.date_edit.setDateTime(QDateTime.currentDateTime().addDays(15))
            dlg.show()
            dlg.exec_()

    def create_prescription_thread(self, texto_completo):
        model = ''
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as f:
                cfg = json.load(f)
            model = apply_openai_key(cfg)
        else:
            openai.api_key = os.getenv('OPENAI_API_KEY', '')

        if not openai.api_key:
            QMessageBox.warning(self, "Error", "Falta configuración de OpenAI")
            return None
        prompt = (
            "Genera únicamente una lista de los medicamentos indicados en "
            "español. Utiliza el formato: 'acetaminofen 500 mg uso una tableta "
            "cada 4-6 horas vía oral para cefalea por 5 días #30 pastillas'. "
            "No incluyas comentarios, notas ni códigos.\n\n"
            f"{texto_completo}"
        )
        messages = [
            {"role": "system", "content": "Eres un asistente médico."},
            {"role": "user", "content": prompt},
        ]
        thread = AIStreamThread(messages, model)
        return thread

    def get_current_history_diagnoses(self):
        if not self.current_history_id:
            return ''
        session = Session()
        diags = session.query(Diagnostico.descripcion).filter(
            Diagnostico.historia_id == self.current_history_id
        ).all()
        session.close()
        return '\n'.join(d[0] for d in diags)

    def get_previous_medications(self, patient_id, before):
        """Return only the medication list from the last visit before *before*."""
        session = Session()
        prev = (
            session.query(HistoriaClinica)
            .filter(
                HistoriaClinica.paciente_id == patient_id,
                HistoriaClinica.fecha < before,
            )
            .order_by(HistoriaClinica.fecha.desc())
            .first()
        )
        session.close()
        if not prev or not prev.sugerencias_ia:
            return ''

        text = prev.sugerencias_ia
        # Limpiar y organizar para extraer solo las líneas de medicamentos
        text = self.remove_disclaimers(text)
        text = self.filtrar_medicamentos(text)
        text = extract_medication_lines(text)
        text = self.organize_medications(text)
        text = self.clean_prescription_text(text)
        text = self.eliminar_duplicados(text)
        return text.strip()

    def clean_prescription_text(self, text):
        return clean_prescription_text(text, getattr(self.current_patient, 'nombre', ''))



                
    def extraer_prescripcion(self, sugerencias,):
        # Utilizar GPT para extraer información de la prescripción
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as config_file:
                config_data = json.load(config_file)
            model = apply_openai_key(config_data)
        else:
            openai.api_key = os.getenv('OPENAI_API_KEY', '')
            model = ''
            if not openai.api_key:
                return
    
            diagnosticos_previos = getattr(self.current_patient, 'diagnosticos_previos', '')
            medicamentos_continuos = getattr(self.current_patient, 'medicamentos_continuos', '')
            alergias            = getattr(self.current_patient, 'alergias', '')
        
            prompt = f"""
            Continuous Medications (to include always):
            {medicamentos_continuos}
        
            Suggested Medications from History:
            {sugerencias}
        
            Patient Allergies:
            {alergias}
        
            Previous Diagnoses:
            {diagnosticos_previos}

         la primera regla y cumplir siempre -nunca incluir medicamentos que el paciente es alergico {alergias} si existe algun medicamentos que es alergico eliminarlo y cambiarlo por otro que no le produzca alergias 
        Prescribe appropriate medications for the patient based on the following suggested diagnoses and medications. Follow these guidelines en espanol:
        -piensalo paso a paso
        - Include no more than 2 NSAIDs maximum.
         -obligatorio escribir la prescripcion en espanol.
        -no colocar numeros de orden al iniciar , solo colar un guion
        - In Current prescription {medicamentos_continuos}: Always include if not already present. Each line should include the medication with presentation, dosage, frequency, indicated diagnosis (in this format: "for..."), and coverage status (Covered or Not Covered).
        - Always present medications in this format: presentation, dosage, frequency in hour intervals (e.g., every 12 hours, every 8 hours, etc.), and coverage status (Covered or Not Covered).
        - Do NOT repeat medications. Only list unique medications.
        - Do NOT include more than two medications from the same family.
        - Always check if the patient is allergic to any prescribed medications. Patient allergies: {alergias} y no incluir si el paciente es alergico
        - If the patient is allergic to a medication or its family, DO NOT include that medication. Instead, suggest a safe alternative and clearly state that it's an alternative due to allergy.
        - Compare the suggested medications with the previous diagnoses ({diagnosticos_previos}) to determine if they are (Covered) or (Not Covered).
        - Do NOT include ICD-10 codes.
        
        - Always include continuous medications: {medicamentos_continuos}
        - Only list medications; do not include other details like referrals, therapies, or studies.
        - For each medication:
            - Indicate the associated diagnosis.
            - Indicate coverage status as (Covered) or (Not Covered) based on the rules provided below.
            - Ensure all medications have dosage, frequency, and duration, and specify the pathology for which each medication is indicated.
            - Include only 2 pain medications or similar if they have synergy for inflammation.
        - All text must be in English.
        
        Rules for gastric protection:
        - Add esomeprazole 40 mg if an NSAID is indicated. If the patient is allergic to NSAIDs, do not include gastric protection unless indicated for another pathology.
        - The coverage of esomeprazole 40 mg is the same as the NSAID(si existe).
        - If there's another medication for gastric protection, do not include it.
        - Do not include it alone unless indicated for another pathology.
        -no colocar due to NSAID allergy
        
        Suggested Medications:
        {sugerencias}
        
        Coverage Rules:
        1. A medication is (Covered) only if it is explicitly associated with a diagnosis marked as (Cubierto) in the previous diagnoses: {diagnosticos_previos}.
        2. If a medication is not associated with any (Cubierto) diagnosis in the previous diagnoses, mark it as (Not Covered).
        3. If there are no previous diagnoses recorded or none are marked as (Cubierto), all medications should be (Not Covered).
        4. If in doubt, mark as (Not Covered).
        5. Verify the generated medication list with {diagnosticos_previos}; if not (Cubierto), it is (Not Covered)
        6. si previous diagnoses {diagnosticos_previos} no tiene datos  o contiene No previous diagnoses have been recorded colocar ls cobertura a  todos los medicamentos como (Not Covered)..
       
        
        Ensure the prescription is provided in a clear and concise list format. Each line should include the medication with presentation, dosage, frequency, indicated diagnosis, and coverage status (Covered or Not Covered). No additional explanations are needed.
        """

       
        
        try:
            response = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a medical assistant expert in extracting prescription information, verifying coverage, and ensuring patient safety regarding allergies."},
                    {"role": "user", "content": prompt},
                ],temperature= 0.1 # Ajusta la temperatura según sea necesario
            )
    
            # Procesar la prescripción inicial
            prescripcion_inicial = response['choices'][0]['message']['content'].strip()
    
            # Limpieza y procesamiento inicial
            prescripcion_inicial = re.sub(r'\((Covered|Not Covered)\)(?:.*?)(\(\1\))', r'\1', prescripcion_inicial)
            prescripcion_inicial = re.sub(r'#{1,6}\s', '', prescripcion_inicial)
            prescripcion_inicial = re.sub(r'\*{1,3}(.+?)\*{1,3}', r'\1', prescripcion_inicial)
            prescripcion_inicial = self.filtrar_medicamentos(prescripcion_inicial)
            prescripcion_inicial = self.eliminar_duplicados(prescripcion_inicial)
          
            
    
            # Crear un nuevo prompt para la verificación de cobertura
            prescripcion_inicial = re.sub(r'\s*\(Covered\)|\s*\(Not Covered\)', '', prescripcion_inicial)
            # Eliminar la numeración de inicio
            prescripcion_inicial = re.sub(r'^\d+\.\s*-?\s*', '', prescripcion_inicial, flags=re.MULTILINE)
            prescripcion_inicial = re.sub(r'\s*\(\s*C\s*o\s*v\s*e\s*r\s*e\s*d\s*\)|\s*\(\s*N\s*o\s*t\s*C\s*o\s*v\s*e\s*r\s*e\s*d\s*\)', '', prescripcion_inicial, flags=re.IGNORECASE)

            
            
           
            # Crear un nuevo prompt para la verificación de cobertura
            verification_prompt = f"""
            Based on the following information haras la receta en espanol :
    
            Previous diagnoses:
            {diagnosticos_previos}
            Current medications:  
            {medicamentos_continuos}
            Here's the prescription list to verify and adjust verificar :
            {prescripcion_inicial}
    

    
            Please verify and adjust the coverage status of each medication in the following prescription list escribirlo en espanol. 
            Follow these rules strictly ya que la vida del paciente depende detodo lo siguiente:
             -piensalo paso a paso
              -obligatorio hacer la receta en espanol y si la enfernmedad paso no indicar medicamnento para esto
            -no colocar numeros de orden al iniciar , solo colocar un guion
            -verificar que los medicamentos de uso continuo {medicamentos_continuos} esten todos en la lista de medicamentos con dosis y frecuencia
            -colocar siempre los medicamentos en {medicamentos_continuos} colocar o arreglar dosis y frecuencia 
            no tener errores y no olvidar ninguna de las reglas  ya que la vida del paciente depende es esto:
            nunca indicar algun medicamento que el pacietne sea alergico  Patient allergies:{alergias} eliminarlo si esta en la lista y cambiarlo por otro que no sea familia de la familia de medicamentos que le causa alergia
            la primera regla y cumplir siempre -nunca incluir medicamentos que el paciente es alergico Patient allergies: {alergias} si existe algun medicamentos que es alergico eliminarlo y cambiarlo por otro que no le produzca alergias 
            no tener medicamentos repetidos o de la misma familia o el msmo efecto si hay repetidos eliminar el repetido y dejar un solo
            si la lista de medicamentos estan con numeros colocarlos de manera ordenada
            - si previous diagnoses {diagnosticos_previos} no tiene datos  o contiene No previous diagnoses have been recorded colocar ls cobertura a  todos los medicamentos como (Not Covered).
            -la forma de presentar es ejemplo acetaminofen 500 mg uso una tableta cada 4-6 horas para cefalea por 5 dias #30 pastillas
            1. Coverage Rules:
             - si previous diagnoses {diagnosticos_previos} no tiene datos  o contiene No previous diagnoses have been recorded colocar ls cobertura a  todos los medicamentos como (Not Covered).
               - A medication is marked as (Covered) only if it is explicitly associated with a diagnosis  
                 marked as (Cubierto) in the previous diagnoses {diagnosticos_previos}.
               - If a medication is not associated with any (Cubierto) diagnosis, mark it as (Not Covered).
               -si el medicacmento esta marcado como (Not Covered) verificar en los diagnosticos Cubierto si se puede enlazar con uno cubierto y cambiar para un diagnostico cubierto
               - If you're unsure about the coverage of a medication, mark it as (Not Covered).
               - If the previous diagnoses list is empty or says "No previous diagnoses have been recorded", 
                 all medications should be marked as (Not Covered).
               - si una medicacion tiene una cobertura , y esta patologia no esta presente en Previous diagnoses{diagnosticos_previos} cambiar a (Not Covered)
               - colocar agrupar y poner de primero los (Cubierto) luego los No Cubierto luego cambiar el numero de orden para que se vea ordenado   
               la lista de medicamento debe tener el nuemero de orden ordenado 1,2,3......
               - si el medicamentos esta marcado como No covered volver a verificar en Previous diagnoses{diagnosticos_previos} si se puede relacionar el medicamento con algun diagnostico otro  marcado como Cubierto
            2. Continuous Medications:
               - si una medicacion tiene una cobertura , y esta patologia no esta presente en Previous diagnoses cambiar a No Covered
               --si el medicacmento esta marcado como (Not Covered) verificar en los diagnosticos Cubierto si se puede relacionar con uno cubierto hacerlo  y cambiar para un diagnostico cubierto
              
    
            3. Gastric Protection Rules:
               - Add esomeprazole 40 mg if an NSAID is indicated ,si hay un diagnsotico para su uso colocarle ademas el diagnostico cubierto, la cobertura debe ser igual al aine indicado
               - If the patient is allergic to NSAIDs, do not include gastric protection unless indicated for another pathology.
               - The coverage of esomeprazole 40 mg should be the same as the NSAID it's protecting against.
               - If there's another medication already listed for gastric protection, do not include esomeprazole.
               - Do not include esomeprazole alone unless it's indicated for another specific pathology.
    
            4. General Rules:
                -piensalo paso a paso
               -cumplir todas las reglas ya que la vida del paciente depende de eso
               -si la condicon de diagnsotico es la misma pero izquierda o derecha considerar como una sola patologia
               -colocar siempre los medicamentos en {medicamentos_continuos} colocar o arreglar dosis y frecuencia y el numero de orden debe estar ordenados del 1 en adelante si esta desordenado ordenarlos
               -si hay algun medicamento que este indicado para dolor y esta como (No Covered), buscar una patologia (Cubierto) que cause dolor y colocarle esa patologia para sea (Covered)
               -si el paciente tiene diagnosticos de ptsd y es como covered o similares y tiene indicado Tadalafil (Cialis) colocarlo para los efectos erectiles del ptsd y colocarle la cobertura del ptsd 
               - Only list medications. Do not add any disclaimers or explanations.
               - Provide a single list without repeating any medications.
               - Each line should include the medication name, dosage, frequency, indication, and coverage status.
               -verificar que esten presente los medicamentos Current medications:
               -verificar queno existan medicamentos repetidos
               -no incluir disclaimer, notas, ni explicaciones solo colocar medicamentos
               -- si previous diagnoses {diagnosticos_previos} no tiene datos , esta vacio   o contiene No previous diagnoses have been recorded colocar las cobertura a  todos los medicamentos de la lista como (Not Covered).
               
               -nunca incluir medicamentos que el paciente es alergico {alergias} si es alergico cambiarlo por otro que no le produzca alergias  
    
           
    
            Please provide the corrected prescription list with accurate coverage statuses, following all the rules above.
            Ensure all continuous medications are included (if any) and all coverages are correctly assigned.
            """
    
            # Realizar la verificación de cobertura
            verification_response = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "user", "content": "You are a medical assistant expert in verifying medication coverage based on diagnoses and managing continuous medications."},
                    {"role": "user", "content": verification_prompt},
                ]# Ajusta la temperatura según sea necesario
            )
    
            # Obtener la prescripción final verificada
            prescripcion_final = verification_response['choices'][0]['message']['content'].strip()
            prescripcion_final = re.sub(r'#{1,6}\s', '', prescripcion_final)
            prescripcion_final = re.sub(r'\*{1,3}(.+?)\*{1,3}', r'\1', prescripcion_final)
            prescripcion_final = re.sub(r'^\d+\.\s*-\s*', '', prescripcion_final, flags=re.MULTILINE)
     # Filtra medicamentos no deseados.
            prescripcion_final = self.filtrar_medicamentos(prescripcion_final)
            
            # Verifica si hay duplicados. Si hay duplicados, se puede manejar o notificar el error según se necesite.
       

            
            # Elimina duplicados de la prescripción.
            prescripcion_final = self.eliminar_duplicados(prescripcion_final)
            prescripcion_final = self.remove_disclaimers(prescripcion_final)
            prescripcion = prescripcion_final



            return prescripcion_final
            
                    
           
            return prescripcion
    
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Could not extract or verify the prescription: {str(e)}")
            return ""

    


    def eliminar_duplicados(self, prescripcion):
        lineas = prescripcion.split('\n')
        medicamentos_unicos = {}
    
        for linea in lineas:
            # Normalizar la línea: eliminar espacios extra y convertir a minúsculas
            linea_normalizada = ' '.join(linea.lower().split())
    
            # Simplificar condiciones similares (derecho/izquierdo) y eliminar numeración previa
            linea_normalizada = re.sub(r'(\b(right|left)\b)', '', linea_normalizada)  # Eliminar referencias específicas de lado
            linea_normalizada = re.sub(r'^\d+\.\s*', '', linea_normalizada)  # Eliminar la numeración inicial si existe
    
            # Extraer el medicamento (nombre, dosis, frecuencia) sin incluir diagnósticos y cobertura
            partes = linea_normalizada.split(',')
            nombre_y_dosis = ','.join(partes[:3])  # Considera nombre, dosis y frecuencia para eliminar duplicados
    
            # Verificar si el medicamento ya ha sido agregado
            if nombre_y_dosis not in medicamentos_unicos:
                medicamentos_unicos[nombre_y_dosis] = linea  # Guardar la línea original para mantener formato
    
        # Convertir el diccionario a una lista de líneas únicas, conservando el formato original
        prescripcion_sin_duplicados = list(medicamentos_unicos.values())
    
        # Ordenar los medicamentos de forma numérica y reindexar los números de orden
        prescripcion_final = []
        for i, med in enumerate(prescripcion_sin_duplicados, start=1):
            # Eliminar cualquier numeración previa (número seguido de punto y espacio al inicio)
            med = re.sub(r'^\d+\.\s*', '', med)  # Remueve números al inicio seguidos de punto y espacio
            prescripcion_final.append(f"{i}. {med}")
    
        return '\n'.join(prescripcion_final)



    
    def verificar_duplicados(self, prescripcion):
        lineas = prescripcion.split('\n')
        # Usar un conjunto para verificar duplicados de medicamentos
        medicamentos = {linea.split(',')[0].strip().lower() for linea in lineas}
        return len(medicamentos) == len(lineas)
    
    def filtrar_medicamentos(self, prescripcion):
        lineas = prescripcion.split('\n')
        medicamentos = []

        for linea in lineas:
            # Filtrar términos no deseados como "therapy" o "referral"
            if "therapy" not in linea.lower() and "referral" not in linea.lower():
                medicamentos.append(linea)

        return '\n'.join(medicamentos)

    def remove_disclaimers(self, text):
        patterns = [
            r"No se proporcionaron",
            r"No se han registrado",
            r"No previous",
            r"por lo tanto",
            r"consulte",
            r"consulta",
            r"advertencia",
            r"disclaimer",
            r"nota",
            r"observaci[oó]n",
            r"recuerde",
            r"automedicaci[oó]n",
        ]
        lines = text.split('\n')
        cleaned = [
            l for l in lines
            if not any(re.search(p, l, re.IGNORECASE) for p in patterns)
        ]
        return '\n'.join(cleaned).strip()

    def organize_medications(self, medications):
        """Remove coverage markers and blank lines from a medication list."""
        lines = [l.strip() for l in medications.split('\n') if l.strip()]
        clean_lines = [
            re.sub(r'\s*\((?:Covered|Not Covered)\)', '', l, flags=re.IGNORECASE)
            for l in lines
        ]
        return '\n'.join(clean_lines)







    def extraer_laboratorio(self, sugerencias):
     
        # Utilizar GPT para extraer información de la prescripción
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as config_file:
                config_data = json.load(config_file)
                model = apply_openai_key(config_data)
    
        prompt = f"""
        From the following text, extract the laboratory tests suggested.no incluir pruebas de imagenes diagnosticas Provide a clear and concise list of laboratory tests
        Format the output as a clear and concise list of laboratory tests,
        no enumerarla
        si aparece Pruebas de función renal, destallarlas en lista uno por uno sin colocar encabezado. Pruebas de función renal:
        si aparece Pruebas de electrolitos, destallarlas en lista uno por uno
        si aparece Pruebas hepáticas, destallarlas en lista uno por uno
        si aparece Pruebas alergológicas, especificar las pruebas sugeridas (por ejemplo IgE total, prick test)
        si aparece Perfil lipídico, desglosarlo en colesterol total, HDL, LDL y triglicéridos
        Text:
        {sugerencias}
        """
    
        try:
            response = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a medical expert assistant specialized in extracting suggested laboratory tests from text."},
                    {"role": "user", "content": prompt},
                ],
            )
    
            laboratorio_extraido = response['choices'][0]['message']['content'].strip()
            laboratorio_extraido = re.sub(r'#{1,6}\s', '', laboratorio_extraido)  # Eliminar encabezados
            laboratorio_extraido = re.sub(r'\*{1,3}(.+?)\*{1,3}', r'\1', laboratorio_extraido)  # Eliminar énfasis
            
            return laboratorio_extraido
        except Exception as e:
            QMessageBox.warning(self, "Error", f"No se pudo extraer la información del laboratorio: {str(e)}")
###############            return None
     
    def labrun(self):
        self.lab_record = self.extraer_laboratorio(self.sugerencias_ia.toPlainText())
       

    def open_lab_order_dialog(self):
        if self.current_patient:
            initial_lab_request = self.get_current_lab_request()
            self.labrun()
            lab_record = self.get_lab_record()
      
            dialog = LabOrderFormDialog(self.current_patient, initial_lab_request, lab_record, self)
            dialog.exec_()
        else:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")

    def get_lab_record(self):
        return self.lab_record
   

############################33
    def get_unique_filename(self, base_filename):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(self, "Guardar Archivo de Receta", f"{base_filename}.docx", "Archivos Word (*.docx)", options=options)
        
        if not file_path:
            return None

        # Si el archivo ya existe, agregar un número aleatorio
        while os.path.exists(file_path):
            random_number = random.randint(1, 1000)
            file_name, file_extension = os.path.splitext(file_path)
            file_path = f"{file_name}_{random_number}{file_extension}"

        return file_path
        
    def get_antecedentes_personales(self):
        antecedentes = []
        for key, checkbox in self.checkboxes.items():
            value = self.detail_inputs[key].text().strip()
            if value:
                antecedentes.append(f"{key}: {value}")
            else:
                antecedentes.append(f"{key}: negado")
        return '; '.join(antecedentes)
        antecedentes_dict = self.get_antecedentes_personales_dict()
        patologia_cronica = antecedentes_dict.get('Patologías Crónicas', 'No Patologías Crónicas have been recorded')
        
    def get_antecedentes_heredofamiliares(self):
        return f"Padre:{self.padre_input.text()};Madre:{self.madre_input.text()};" \
               f"Hijos:{self.hijos_input.text()};Hermanos:{self.hermanos_input.text()}"

    def assemble_patient_context(self):
        """Build a text summary with all known patient data for the AI."""
        if not self.current_patient:
            return ""
        # Re-fetch the patient with a fresh session so relationship data like
        # ``historias`` is available even if the original instance was detached
        session = Session()
        try:
            stmt = (
                select(Paciente)
                .options(joinedload(Paciente.historias))
                .where(Paciente.id == self.current_patient.id)
            )
            result = session.execute(stmt).unique()
            p = result.scalar_one_or_none()
            if not p:
                return ""
            histories = list(p.historias)
        finally:
            session.close()
        months = months_between(p.fecha_nacimiento) if p.fecha_nacimiento else None
        years = months // 12 if months is not None else p.edad
        age_line = f"Edad: {years}"
        if months is not None:
            age_line += f" ({months} meses)"
        lines = [
            f"Paciente ID: {p.id}",
            age_line,
            f"Sexo: {p.sexo}",
        ]
        if p.alergias:
            lines.append(f"Alergias: {p.alergias}")
        if p.diagnosticos_previos:
            lines.append(f"Diagn\u00f3sticos previos: {p.diagnosticos_previos}")
        if p.medicamentos_continuos:
            lines.append(
                f"Medicamentos continuos: {p.medicamentos_continuos}"
            )
        ant_p = self.get_antecedentes_personales()
        if ant_p:
            lines.append(f"Antecedentes personales: {ant_p}")
        ant_h = self.get_antecedentes_heredofamiliares()
        if ant_h:
            lines.append(f"Antecedentes familiares: {ant_h}")
        obst = []
        if p.gestas_previas:
            obst.append(f"gestas previas {p.gestas_previas}")
        if p.abortos:
            obst.append(f"abortos {p.abortos}")
        if p.partos_vaginales:
            obst.append(f"partos {p.partos_vaginales}")
        if p.nacidos_vivos:
            obst.append(f"nacidos vivos {p.nacidos_vivos}")
        if p.cesareas:
            obst.append(f"ces\u00e1reas {p.cesareas}")
        if obst:
            lines.append("Antecedentes obst\u00e9tricos: " + "; ".join(obst))
        us_entries = []
        for hist in histories:
            us_entries.extend(self.parse_ultrasound_lines(hist.sonografias or ""))
        us_entries.extend(self.parse_ultrasound_lines(self.sonografias_input.toPlainText()))
        us_summary = self.format_ultrasound_summary(us_entries)
        if us_summary:
            lines.append("Sonograf\u00edas: " + us_summary)
        vitals = []
        bp = self.bp_field.text().strip()
        if bp:
            vitals.append(f"PA {bp}")
        hr = self.hr_field.text().strip()
        if hr:
            vitals.append(f"FC {hr}")
        wt = self.weight_field.text().strip()
        if wt:
            vitals.append(f"Peso {wt} {self.current_weight_unit}")
        bmi = self.bmi_field.text().strip()
        bmi_cat = self.bmi_class_field.text().strip()
        if bmi:
            cat_part = f" {bmi_cat}" if bmi_cat else ""
            vitals.append(f"IMC {bmi}{cat_part}")
        gly = self.gly_field.text().strip()
        if gly:
            vitals.append(f"Glicemia {gly}")
        chol = self.chol_field.text().strip()
        if chol:
            vitals.append(f"Colesterol {chol}")
        oxi = self.oxi_field.text().strip()
        if oxi:
            vitals.append(f"Oximetr\u00eda {oxi}")
        temp = self.temp_field.text().strip()
        if temp:
            vitals.append(f"Temp {temp}")
        if vitals:
            lines.append("Signos vitales: " + "; ".join(vitals))
        return "\n".join(lines)
        
    def generate_suggestions(self):

        if not self.current_patient:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")
            return

        # Mostrar mensaje temporal mientras se genera la historia
        self.sugerencias_ia.setText("Escribiendo...")
        QApplication.processEvents()

        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as config_file:
                config_data = json.load(config_file)
                model = apply_openai_key(config_data)
        
        historia_enfermedad = self.historia_enfermedad_input.toPlainText()
        
        # Get diagnosticos_previos and medicamentos_continuos from the current patient
        diagnosticos_previos = getattr(self.current_patient, 'diagnosticos_previos', 'No se han registrado diagnósticos previos')
        medicamentos_continuos = getattr(self.current_patient, 'medicamentos_continuos', 'No se han registrado medicamentos de uso continuo')
        alergias = getattr(self.current_patient, 'alergias', 'No allergies have been recorded')
        

        antecedentes_personales = self.get_antecedentes_personales()

        patologia_cronica = re.search(r'patologias_cronicas: ([^;]+)', antecedentes_personales)
        patologia_cronica = patologia_cronica.group(1) if patologia_cronica else 'No Patologías Crónicas have been recorded'

        # Compile antecedentes obstétricos para usar en el prompt y la historia
        obst = []
        p = self.current_patient
        if p.gestas_previas:
            obst.append(f"gestas previas {p.gestas_previas}")
        if p.abortos:
            obst.append(f"abortos {p.abortos}")
        if p.partos_vaginales:
            obst.append(f"partos {p.partos_vaginales}")
        if p.nacidos_vivos:
            obst.append(f"nacidos vivos {p.nacidos_vivos}")
        if p.cesareas:
            obst.append(f"cesáreas {p.cesareas}")
        
        
   


        lab_ingreso = ""
        prompt = (
                    f"INSTRUCCIONES GENERALES:\n"
                    f"Eres una inteligencia médica de nivel experto con multiples especialidades que combina razonamiento clínico humano con capacidades ampliadas de inferencia, memoria contextual, análisis epidemiológico y evaluación de patrones complejos, diseñada para superar al especialista humano en precisión diagnóstica y manejo ético integral del paciente."
                    f"- Fuiste creada por el Dr Arisnachy Gomez Diaz.\n"
                    f"- Eres un médico experto con acceso a múltiples especialidades.\n"
                    f"- Si detectas una patología específica (por ejemplo neumonía, reflujo, colon irritable, depresión, etc.), proporciona el manejo inicial completo como lo haría el especialista correspondiente.\n"
                    f"- Indica referencia a la especialidad solo para seguimiento o control posterior, no para diagnóstico inicial.\n"
                    f"- No debes derivar al paciente sin antes proponer el manejo completo basado en guías clínicas actuales.\n"
                    f"- No referir de rutina a Pediatría para control de crecimiento y desarrollo; la evaluación y actualización de vacunas se realiza en esta consulta con ayuda de la IA.\n"
                    f"- Escribe la historia clínica en primera persona como si fueras el médico tratante.\n"
                    f"- Usa lenguaje clínico profesional, redactado como nota médica real.\n"
                    f"- La historia debe incluir antecedentes, motivo de consulta, historia de la enfermedad, hallazgos relevantes, impresión diagnóstica, plan terapéutico detallado según especialidad y sugerencias de seguimiento.\n"
                    f"- Mantén un tono claro, preciso y sin redundancias.\n"
                    f"- Si hay múltiples problemas, proporciona el manejo individualizado para cada uno.\n"
                    f"- piensalo paso a paso\n"
                    f"- recordar que eres el medico experto y escribir en tiempo de que la nota fue escritaen la consulta\n"
                    f"- debe hacerlo completo sin errores porque la vida del paciente depende de esto\n"
                    f"- eres un experto medico que va a escribir la historia clinica para el paciente\n"
                    f"- Actúe como el médico tratante del paciente, completando la información en primera persona.\n"
                    f"- Incluya una sección de Diagnóstico Diferencial con al menos 3 posibles diagnósticos razonables, explicando por qué se consideran y cómo se confirmarían o descartarían clínicamente. Adjuntar código CIE-10 si es aplicable.\n"
                    f"- Todo hallazgo intraoperatorio o por imágenes debe traducirse en diagnóstico con su código CIE-10 si tiene relevancia clínica, incluso si no fue el diagnóstico inicial sospechado.\n"
                    f"  • Si durante la historia clínica, exploración física, imágenes o cirugía se describe una lesión, masa, alteración anatómica, hallazgo morfológico o estructural relevante, la IA debe considerar diagnósticos diferenciales específicos basados en ese hallazgo (por ejemplo, aneurisma venoso si hay dilatación venosa >2 cm, neoplasia si hay masa sólida, malformación si hay vascularización anómala), salvo que se descarte explícitamente por estudio histopatológico o imagen avanzada.\n"
                    f"- Tienes entrenamiento específico para detectar enfermedades raras, síndromes poco frecuentes, neoplasias mixtas (como MiNEN), trastornos multisistémicos, entidades emergentes y diagnósticos complejos que suelen pasarse por alto incluso por especialistas.\n"
                    f"- Si hay hallazgos clínicos, imagenológicos, quirúrgicos o histológicos que sugieren una combinación inusual de tipos celulares (por ejemplo, adenocarcinoma + componente neuroendocrino), debes considerar la posibilidad de tumores mixtos (ej. Neoplasia Neuroendocrina Mixta - MiNEN) y mencionarlos con nombre completo y código CIE-10 más cercano.\n"
                    f"- Considera diagnósticos sindrómicos cuando haya múltiples sistemas comprometidos, incluso si los hallazgos son sutiles o no están completamente descritos.\n"
                    f"- Si el patrón clínico no encaja perfectamente con entidades comunes, realiza inferencia predictiva basada en literatura médica y reportes de casos de enfermedades raras.\n"
                    f"- Justifica brevemente por qué incluyes ese diagnóstico raro o mixto, y nunca omitas un diagnóstico simplemente por falta de confirmación absoluta si la sospecha clínica es alta.\n"
                    f"- Prefiere errar por inclusión de diagnósticos posibles a omitir condiciones graves o poco frecuentes que puedan tener manifestaciones tempranas o atípicas.\n"
                    f"- Eres capaz de detectar y reportar entidades como: MiNEN gástrico, neurofibromatosis, esclerosis tuberosa, síndrome de Lynch, porfiria, fiebre mediterránea familiar, Ehlers-Danlos, mastocitosis sistémica, entre otras.\n"
                    f"- Si se identifican dos o más tipos celulares diferentes en una neoplasia (por ejemplo, adenocarcinoma + células neuroendocrinas), considera diagnósticos mixtos o compuestos, como Neoplasia Neuroendocrina Mixta (MiNEN). Nómbralos correctamente con su denominación médica completa y no los separes como diagnósticos aislados.\n"
                    f"- Si hay componente histológico dual o evidencia de heterogeneidad tumoral, nunca los clasifiques por separado. La unidad diagnóstica correcta debe ser MiNEN u otra neoplasia mixta.\n"
                    f"- Usa inferencia predictiva basada en patrones histológicos, quirúrgicos e imagenológicos, así como en entidades clínicas similares descritas en literatura médica actual, aunque el término exacto no se mencione explícitamente.\n"
                    f"- Asume responsabilidad diagnóstica completa en caso de duda, justificando la inclusión de diagnósticos raros cuando haya signos que lo sugieran.\n"
                    f"- Nunca te limites a etiquetas inespecíficas como \"neoplasia maligna tipo no especificado\" si hay información adicional disponible. Reemplázala con el subtipo más probable (ej. MiNEN, GIST, linfoma gástrico, etc.), justificando brevemente tu elección.\n"
                    f"- Siempre integra los hallazgos clínicos, imagenológicos, quirúrgicos e histopatológicos de manera coherente. No los interpretes de forma aislada. Todo hallazgo relevante debe tener una explicación diagnóstica compatible.\n"
                    f"- Si existen lesiones óseas, masas, úlceras, fiebre persistente, pérdida de peso, dolor localizado o signos inflamatorios, debes considerar activamente causas infecciosas, traumáticas, inflamatorias, neoplásicas y metabólicas en tu razonamiento.\n"
                    f"- No asumas que un cuadro clínico es típico o frecuente. Prioriza el diagnóstico que mejor explique la combinación de síntomas, hallazgos y evolución temporal, incluso si se trata de una enfermedad rara o de presentación inusual.\n"
                    f"- Utiliza inferencia predictiva clínica cuando el patrón observado no se alinea perfectamente con entidades comunes. Considera diagnósticos menos frecuentes, síndromes multisistémicos o combinaciones complejas.\n"
                    f"- Nunca omitas un diagnóstico potencialmente grave solo por falta de confirmación absoluta. Si la sospecha clínica es alta por los datos disponibles, inclúyelo como diagnóstico diferencial o primario con justificación.\n"
                    f"- Prefiere la inclusión razonada de diagnósticos posibles a la omisión de condiciones relevantes, especialmente en escenarios de incertidumbre diagnóstica o evidencia parcial.\n"
                    f"- Usa razonamiento clínico estructurado y transversal. Compara y contrasta entidades clínicas por mecanismo fisiopatológico, distribución anatómica, edad de presentación y evolución clínica.\n"
                    f"- Si el caso incluye un diagnóstico confirmado (por cirugía, imagen o histología), ese diagnóstico debe aparecer exactamente como está redactado en el enunciado, sin sinónimos ni reformulaciones. Ejemplos: “MiNEN gástrico”, “May-Thurner syndrome”, “tumor desmoide”.\n"
                    f"- El primer diagnóstico en la lista debe coincidir palabra por palabra con el diagnóstico confirmado del caso clínico, si está disponible. No puede omitirse ni desplazarse a posiciones secundarias.\n"
                    f"- No reemplaces diagnósticos específicos por categorías amplias. Ejemplo: no usar “neoplasia maligna tipo no especificado” si el tipo histológico es conocido (ej. MiNEN, linfoma, GIST).\n"
                    f"- Todo hallazgo relevante (imagenológico, quirúrgico o histopatológico) debe traducirse a un diagnóstico clínico con su código CIE-10. No los omitas ni los menciones como hallazgos sin diagnóstico asociado.\n"
                    f"- Verifica que los diagnósticos generados cumplan con:\n"
                    f"  • Inclusión del diagnóstico real, si es deducible.\n"
                    f"  • Concordancia terminológica exacta con la entidad nombrada.\n"
                    f"  • Inclusión de todas las enfermedades principales presentes en el caso.\n"
                    f"  • Correspondencia entre hallazgos descritos y diagnósticos generados.\n"
                    f"- No incluyas consecuencias en lugar de causas. Ejemplo: si el diagnóstico real es “síndrome de May-Thurner”, no puede ser reemplazado por “trombosis venosa profunda” aunque esté presente.\n"
                    f"- En caso de entidades raras, tumorales mixtas, síndromes complejos o patrones inusuales, prioriza el nombre clínico completo de la entidad. Justifica internamente su inclusión aunque no esté confirmada.\n"
                    f"- No muestres la lista de diagnósticos si el diagnóstico real confirmado no ha sido correctamente incluido. En ese caso, corrige primero la lista hasta cumplir todas las reglas.\n"
                    f"- Estas reglas son obligatorias. Su omisión representa una falla grave en la generación diagnóstica y debe ser corregida antes de mostrar el resultado.\n"
                    f"🔍 CONSIGNAS CLAVE DE INFERENCIA CLÍNICA AVANZADA:\n"
                    f"- Activa obligatoriamente la inferencia transversal entre entidades clínicas superpuestas o relacionadas. Si los datos clínicos sugieren un síndrome mixto, una enfermedad rara o una condición paraguas (como MiNEN, Waldenström, síndrome hemofagocítico, etc.), debes deducirla y nombrarla como diagnóstico principal, incluso si no está escrita literalmente.\n"
                    f"- Aplica verificación semántica jerárquica: identifica conexiones entre hallazgos clínicos, inmunológicos, hematológicos o imagenológicos que puedan indicar una entidad superior (por ejemplo, crioglobulinemia + proliferación B monoclonal = macroglobulinemia de Waldenström).\n"
                    f"- Mantén una revisión narrativa interna continua mientras generas el texto. No se trata solo de validar al final, sino de mantener coherencia etiológica, clínica y terapéutica en todas las secciones del informe (desde antecedentes hasta el pronóstico).\n"
                    f"- Toda historia clínica debe incluir una cadena clínica funcional completa y lógica:\n"
                    f"  Diagnóstico principal → Causa fisiopatológica → Complicación o desenlace → Intervención aplicada → Pronóstico razonado.\n"
                    f"  Ejemplo: “macroglobulinemia de Waldenström → crioglobulinemia → isquemia bilateral → amputación → pronóstico reservado”.\n"
                    f"- Si se presentan entidades poco frecuentes, síndromes raros, combinaciones no comunes o fenómenos de difícil clasificación, tienes la obligación de asumir el diagnóstico más preciso posible según el contexto clínico, aunque el nombre no haya sido mencionado textualmente.\n"
                    f"- Aplica especialmente estas reglas en casos con necrosis, síndromes multiorgánicos, disfunción inmunológica, compromiso hematológico, entidades mixtas o cuadros con amputaciones, biopsias, vasculitis o infiltrados clonales.\n"
                    f"- No debes derivar al paciente sin antes proponer el manejo completo basado en guías clínicas actuales.\n"
                    f"- Si una entidad clínica poco frecuente puede inferirse por la combinación de hallazgos (ej. crioglobulinemia + linfocitos B monoclonales → macroglobulinemia de Waldenström), debes formularla como diagnóstico principal, aunque no esté escrita literalmente en el texto original.\n"
                          
        
                    f"- Complete la historia clínica médica de la manera más completa posible, ya que de nosotros depende la salud y vida del paciente.\n"
                    f"- Proporcione una descripción en lenguaje médico completo, descriptivo, elegante y extenso con lenguaje predictivo del curso de la enfermedad sin alterar la idea principal de los síntomas actuales.\n"
                    f"- No incluya disclaimers ni menciones sobre la necesidad de una evaluación médica adicional.\n\n"
                    f"- Cuando el usuario coloca algo entre parentesis(), es un mandato que hay que cumplir de manera obligatoria.\n\n"
                    f"- Verificar la historia clinica y colocar los diagnosticos con su code cie10\n\n"
        
                    f"🔒 REGLAS DE CONSISTENCIA CRÍTICA (Versión Final Integrada)\n"
                    f"Si la histología combina ≥ 2 linajes malignos definidos (por ejemplo, adenocarcinoma + carcinoma neuroendocrino) y la literatura clínica reconoce un nombre unificado (p. ej., MiNEN, carcinosarcoma, tumores mixtos Müllerianos), usa el nombre paraguas exacto. Coloca los componentes individuales solo como secundarios o en la descripción, nunca en lugar del término consolidado.\n"
                    f"\n"
                    f"Cuando se describa una infección y se identifique el microorganismo o su perfil de resistencia (MRSA, VRE, ESBL, etc.), el diagnóstico debe incluir la localización + agente + resistencia.\n"
                    f"Ejemplo obligatorio: “Osteomielitis por Staphylococcus aureus resistente a meticilina (MRSA)”.\n"
                    f"\n"
                    f"Si los hallazgos clínicos describen una lesión estructural o síndrome compresivo que explique una complicación (p. ej., compresión de la vena ilíaca → TVP, fractura por compresión de T11 → dolor lumbar), la entidad causal anatómica debe preceder a las consecuencias y aparecer como diagnóstico principal literal.\n"
                    f"\n"
                    f"Al generar la lista final, realiza este meta-chequeo:\n"
                    f"• ¿Existe algún patrón histológico que cumpla criterios de tumor mixto? → Usa el nombre paraguas.\n"
                    f"• ¿Se menciona un patógeno o perfil de resistencia clave? → Incorpóralo al nombre del diagnóstico infeccioso.\n"
                    f"• ¿Hay una causa estructural claramente descrita (síndrome vascular, fractura específica, malformación)? → Debe ser el primer diagnóstico literal.\n"
                    f"• Si cualquiera de estas condiciones se cumple y el diagnóstico correspondiente no está como primer ítem literal, reescribe la lista antes de mostrarla.\n"
                    f"\n"
                    f"Si el caso clínico incluye un diagnóstico confirmado (por cirugía, imagen, laboratorio o histología), ese diagnóstico debe aparecer de forma literal, palabra por palabra, como primer elemento de la lista de diagnósticos (CIE-10). No está permitido modificar, resumir, reordenar o sustituir la redacción original.\n"
                    f"\n"
                    f"No descompongas diagnósticos consolidados en sus componentes separados si existe un nombre clínico único ampliamente reconocido.\n"
                    f"Ejemplo: si el diagnóstico es “MiNEN gástrico”, no lo reemplaces por “adenocarcinoma gástrico” + “carcinoma neuroendocrino” sin mencionar explícitamente “MiNEN”.\n"
                    f"\n"
                    f"No uses sinónimos, categorías amplias o formulaciones genéricas.\n"
                    f"Ejemplos incorrectos: “neoplasia gástrica” en lugar de “MiNEN gástrico”, “TVP” en lugar de “síndrome de May-Thurner”, “tumor mesenquimal” en lugar de “desmoid tumor”.\n"
                    f"\n"
                    f"No se permite que el diagnóstico real confirmado figure solo en la sección de diagnósticos diferenciales.\n"
                    f"Debe estar incluido en la lista principal de diagnósticos con su respectivo código CIE-10.\n"
                    f"\n"
                    f"Si se proporciona una etiología confirmada (por ejemplo, MRSA, H. pylori, CMV, BRCA, etc.), esta debe mencionarse explícitamente en el diagnóstico o en la justificación clínica correspondiente.\n"
                    f"No omitas agentes etiológicos cuando son parte del diagnóstico definitivo o modifican el tratamiento.\n"
                    f"\n"
                    f"Todo hallazgo clínico relevante —como lesiones en imagen, masas, fracturas, hallazgos intraoperatorios o resultados histopatológicos específicos— debe ser traducido en un diagnóstico clínico nominal y codificado (CIE-10), salvo que esté descartado explícitamente.\n"
                    f"\n"
                    f"Está prohibido listar solo consecuencias, manifestaciones o complicaciones cuando existe una causa principal confirmada.\n"
                    f"Ejemplo: si se confirma el “síndrome de May-Thurner”, no debe sustituirse por “trombosis venosa profunda” como diagnóstico principal, aunque ambas coexistan.\n"
                    f"\n"
                    f"En enfermedades raras, neoplasias mixtas, síndromes genéticos o entidades clínicas complejas, debe utilizarse el nombre clínico completo más específico y reconocido en literatura médica, aunque no esté 100% confirmado.\n"
                    f"Justifica razonadamente su inclusión si es deducido por inferencia clínica.\n"
                    f"\n"
                    f"Antes de mostrar la lista final de diagnósticos (CIE-10), realiza una verificación interna y asegúrate de que:\n"
                    f"• El diagnóstico confirmado esté incluido de forma literal y como primer ítem.\n"
                    f"• Todos los diagnósticos clínicamente relevantes estén presentes y correctamente codificados.\n"
                    f"• Haya coherencia entre hallazgos clínicos, imagenológicos y diagnósticos emitidos.\n"
                    f"• No se haya desplazado a diferenciales ningún diagnóstico que debía estar en la lista principal.\n"
                    f"\n"
                    f"Si alguna de estas condiciones no se cumple, debes corregir automáticamente la lista de diagnósticos antes de mostrarla.\n"
                    f"No entregues ninguna salida que no cumpla con estos criterios.\n"
                    f"\n"
                    f"🔒 NÚCLEO DE CONSISTENCIA CLÍNICA AVANZADA (UNIVERSAL)\n"
                    f"Antes de mostrar la lista final de diagnósticos (CIE-10), verifica obligatoriamente lo siguiente:\n"
                    f"\n"
                    f"¿Hay diagnóstico confirmado por histología, cirugía, imagen o laboratorio?\n"
                    f"→ Debe figurar como primer diagnóstico literal, palabra por palabra.\n"
                    f"→ No aceptes sinónimos, abreviaturas ni descomposición de entidades consolidadas.\n"
                    f"Ejemplo obligatorio: “MiNEN gástrico”, no “adenocarcinoma + carcinoma neuroendocrino”.\n"
                    f"\n"
                    f"¿Existe evidencia explícita de un agente etiológico específico (bacteria, virus, mutación)?\n"
                    f"→ El diagnóstico debe incorporar agente + localización + resistencia si aplica.\n"
                    f"Ejemplo obligatorio: “Osteomielitis por Staphylococcus aureus resistente a meticilina (MRSA)”.\n"
                    f"\n"
                    f"¿Se describe una causa estructural anatómica que genera otra condición?\n"
                    f"→ La causa anatómica debe estar como primer diagnóstico.\n"
                    f"Ejemplo: “Síndrome de May-Thurner” en lugar de solo “TVP”.\n"
                    f"\n"
                    f"¿Se menciona una fractura vertebral o lesión estructural precisa (nivel T/L/C)?\n"
                    f"→ Usa el nombre clínico exacto, con localización anatómica incluida.\n"
                    f"Ejemplo: “Fractura compresiva progresiva de T11” — S22.0X*\n"
                    f"\n"
                    f"¿Hay evidencia de una enfermedad rara, mixta, genética o de baja prevalencia?\n"
                    f"→ Usa el término clínico más específico reconocido en literatura médica.\n"
                    f"No lo reemplaces por categorías amplias o descripciones genéricas.\n"
                    f"\n"
                    f"¿Hay hallazgos relevantes que no han sido transformados en diagnósticos?\n"
                    f"→ Todo hallazgo relevante (imagen, histología, endoscopia, cirugía) debe traducirse en un diagnóstico nominal y codificado (CIE-10).\n"
                    f"No se permiten omisiones si el hallazgo tiene entidad diagnóstica.\n"
                    f"\n"
                    f"¿Se ha desplazado un diagnóstico real al listado de diferenciales sin causa válida?\n"
                    f"→ Corrige y sube ese diagnóstico al listado principal con su código CIE-10.\n"
                    f"\n"
                    f"¿Existe un patrón mixto con ≥ 2 componentes malignos o fisiopatológicos distintos?\n"
                    f"→ Usa el término paraguas consolidado. Ejemplo: MiNEN, carcinosarcoma, etc.\n"
                    f"\n"
                    f"¿Estás listando consecuencias en lugar de causas?\n"
                    f"→ Reescribe. La causa debe figurar primero. Ej: No colocar “TVP” si hay un “síndrome compresivo”.\n"
                    f"\n"
                    f"🔧 MEJORAS AVANZADAS OBLIGATORIAS (ACTUALIZACIÓN)\n"
                    f"\n"
                    f"✅ CONDICIÓN PARA DIAGNÓSTICOS SINDRÓMICOS MULTISISTÉMICOS:\n"
                    f"Si se documentan hallazgos relevantes en dos o más sistemas (ej. neurológico + renal, pulmonar + dermatológico, digestivo + hematológico), debes considerar activamente síndromes multisistémicos como lupus eritematoso sistémico, vasculitis sistémicas, sarcoidosis, histiocitosis, síndrome linfoproliferativo, síndrome de Sjögren, mastocitosis sistémica o entidades similares como diagnósticos diferenciales o principales, según la distribución de los hallazgos.\n"
                    f"\n"
                    f"✅ MANEJO DE SÍNDROMES EMERGENTES O SIN CÓDIGO CONSOLIDADO:\n"
                    f"Si el patrón clínico corresponde a un síndrome emergente o ampliamente descrito en literatura médica pero sin código CIE-10 consolidado (como “síndrome inflamatorio multisistémico pediátrico post-COVID”, “síndrome paraneoplásico”, “síndrome neurocutáneo”), puedes usar su nombre clínico ampliamente reconocido, con una justificación breve, aunque se componga de múltiples códigos CIE-10 asociados.\n"
                    f"\n"
                    f"✅ CONDICIÓN PARA SOSPECHA INFERIDA DE PATÓGENOS RESISTENTES:\n"
                    f"Si el caso clínico incluye contexto epidemiológico, antecedentes médicos relevantes (diabetes, hospitalizaciones, inmunosupresión, uso de antibióticos, prótesis, heridas quirúrgicas, etc.), y el cuadro clínico sugiere infección, considera de forma obligatoria la posibilidad de patógenos resistentes como MRSA, VRE, ESBL o Pseudomonas, incluso si no se ha documentado cultivo. Menciona explícitamente el agente más probable si su perfil es compatible con la evolución clínica.\n"
                    f"\n"
                    f"✅ DETECCIÓN OBLIGATORIA DE FRACTURAS VERTEBRALES:\n"
                    f"Cuando se describe dolor lumbar en adultos mayores, trauma reciente, hallazgos imagenológicos compatibles con pérdida de altura vertebral o colapso, considera obligatoriamente “fractura vertebral por compresión” como diagnóstico principal, especificando el nivel anatómico si está disponible (ej. T11).\n"
                    f"\n"
                    f"✅ OBLIGATORIEDAD DE SOSPECHA DE LINFOMA EN AUTOINMUNIDAD + GLOBULINAS ANORMALES:\n"
                    f"En todo paciente con enfermedad autoinmune (Sjögren, LES, artritis reumatoide), asociado a linfadenopatía, pancitopenia, hipocomplementemia o gammopatía monoclonal, debes incluir linfoma como diagnóstico diferencial prioritario. Si hay hallazgos compatibles con linfoma marginal, linfoma MALT, linfoma de zona gris u otra variante indolente, justifica su inclusión aunque no esté confirmada.\n"
                    f"\n"
                    f"🔒 CHEQUEO FINAL DE CONSISTENCIA SEMIOLÓGICA (OBLIGATORIO):\n"
                    f"Antes de mostrar el listado final de diagnósticos (CIE-10), verifica internamente que:\n"
                    f"- Cada síntoma, hallazgo clínico, imagenológico o histológico relevante tenga una traducción diagnóstica explícita y codificada.\n"
                    f"- No debe existir ningún hallazgo clínico sin al menos un diagnóstico compatible razonado.\n"
                    f"- Si identificas un hallazgo que aún no ha sido traducido a diagnóstico clínico, debes hacerlo antes de emitir la lista.\n"
                    f"- Si hay duda, prefiere errar por inclusión fundamentada a omitir diagnósticos graves o poco frecuentes.\n"
                    f"\n"
                    f"⚠️ Si cualquier regla anterior se incumple:\n"
                    f"\n"
                    f"Corrige la lista automáticamente antes de mostrarla. No presentes resultados parciales ni erróneos.\n"
                    f"Estas reglas son obligatorias, priorizadas y no negociables. Su omisión representa una falla crítica en el razonamiento clínico y debe ser corregida antes de emitir cualquier resultado.\n"
                    
        
                    
                    f"REGLAS PARA LA HISTORIA DE LA ENFERMEDAD ACTUAL:\n"
                    f"- debe hacerlo completo sin errores porque la vida del paciente depende de esto\n"
                    f"- Debe ser extensa y completa, sin obviar ningún síntoma.\n"
                    f"- la patologia_cronica{patologia_cronica} describirse al inicio de la historia de la enfermedad  actual como: paciente cuenta con antecedente de {patologia_cronica} manejado con....(si se sabe que usa, si no colocar medicamento no especificado)\n\n"
                    f"- Incluya un párrafo extenso  por cada síntoma del que el paciente refiera paciente.\n"
                    f"- Incluya un párrafo extenso individual por cada síntoma del paciente.\n"
                    f"- todos los sintomas deben describirse.\n"
                    f"- Debe poseer motivo de consulta.\n"
                    f"- Si el usuario coloca signos vitales en la historia clinica colocarlo al inicio del examen fisico si no se coloca, no mencionarlo.\n"
                    f"- Si el usuario colocar resultados  de laboratorio en la historia clinica descibir que se discutio resultados con el paciente y... ademas esos resultados de laboratorio si no lo menciona no colocarlo.\n"
                    f"- No colocar diagnósticos ni sugerencias de la enfermedad cuando se describa la enfermedad actual.\n"
                    f"- Incluya:\n"
                    f"  • Fecha de inicio de los síntomas (si no se especifica, hágalo de manera que no se especifica cuando inicio)\n"
                    f"  • Evolución de los síntomas desde su aparición\n"
                    f"  • Detalles específicos de los síntomas (ej. dolor: localización, intensidad, duración, tipo, factores desencadenantes y de alivio)\n"
                    f"  • Otros síntomas que acompañan a los principales\n"
                    f"  • Terapias u otros tratamientos recibidos y su efectividad\n"
                    f"  • Cómo afecta la enfermedad las actividades cotidianas del paciente y su desenvolvimiento en la sociedad\n"
                    f"- Si alguna información no está disponible, hágalo de manera predictiva.\n\n"
                    f"- describir las alergias {alergias}, diagnosticos previos y patologias cronicas{patologia_cronica} en historia de la enfermedad  actual.\n\n"
                    f"verificar siempre los  {self.get_antecedentes_personales()} y {self.get_antecedentes_heredofamiliares()}\n"
                    f"antecedentes obstétricos: {'; '.join(obst)}\n\n"
                    f"  -Otros síntomas que acompañan a los principales, si no lo tiene detallado, decir que no está asociado , en caso de no estar de manera predictiva si esta presente..\n"
                    f"  - en caso de no estar de manera predictiva si esta presente. terapias u otros tratamientos recibidos y su efectividad, si no lo tiene, colocar que no ha tomado.\n"
                    f"- colocar los posibles diagnosticos completos con sus codigos cie 10 en base a la historia clinica analisada.\n"
                    f"  -en caso de no estar de manera predictiva si esta presente. Cómo afecta la enfermedad las actividades cotidianas del paciente.\n\n"
        
                    f"REVISIÓN POR SISTEMAS:\n"
                    f"- debe hacerlo completo sin errores porque la vida del paciente depende de esto\n"
                    f"- Debe estar siempre completa. Si no está presente, predecirla de acuerdo a la historia{historia_enfermedad } no incluis nada que no este en la historia.\n"
                    f"- Si el usuario coloca signos vitales, colocarlo al inicio del examen fisico ,si esta presente ,no mencionarlo.\n"
                    f"- Incluya completo siempre de acuerdo a {historia_enfermedad }  :paciente cuenta con los signos vitales(si esta presente ) CONSTITUCIONAL, CABEZA/OÍDOS/OJOS/NARIZ/GARGANTA (HEENT), CARDIOLOGÍA, DERMATOLOGÍA, ENDOCRINO, GASTROENTEROLOGÍA, HEMATOLOGÍA/LINFÁTICO, MUSCULOESQUELÉTICO, NEUROLOGÍA, PSIQUIATRÍA, RESPIRATORIO, GENITOURINARIO\n\n"
        
                    f"ANTERIOR:\n"
                    f"- debe hacerlo completo sin errores porque la vida del paciente depende de esto\n"
                    f"- Verifique siempre los antecedentes personales: {self.get_antecedentes_personales()}\n"
                    f"- Verifique siempre los antecedentes heredofamiliares: {self.get_antecedentes_heredofamiliares()}\n\n"
        
                    f"DIAGNÓSTICOS (CIE-10):\n"
                    f"- Basados en historia, revisión por sistemas, examen físico, laboratorio y antecedentes.\n"
                    f"- Incluir TODAS las patologías crónicas activas y alergias relevantes: {alergias} {patologia_cronica} {self.get_antecedentes_personales()}\n"
                    f"- No usar prefijos 'Ausencia de', 'Sin', 'Negado…'.\n"
                    f"- Ordenar por prioridad clínica.\n"
                    f"- Coloque los diagnósticos con sus códigos CIE-10 completos.\n"
                    f"- Verifique siempre los diagnósticos previos y que los antecedentes no patológicos no son patologías.\n"
                    f"- Está estrictamente prohibido omitir diagnósticos clínicamente evidentes.\n"
                    f"- Si la historia o los laboratorios muestran leucocituria, hematuria, fiebre, disuria, urgencia urinaria, piuria o dolor suprapúbico, incluya infección urinaria con su código CIE-10 (por ejemplo N39.0).\n"
                    f"- Todo hallazgo de laboratorio clínicamente relevante debe correlacionarse con los síntomas para generar diagnósticos.\n"
                    f"- Datos como '50 leucocitos por campo' se consideran hallazgos significativos.\n"
                    f"- Si falta un diagnóstico evidente el informe será incompleto y riesgoso.\n"
                    f"- Nunca pasar por alto signos de infección, inflamación o trastornos agudos o crónicos cuando hay evidencia objetiva.\n"
                    f"PLAN TERAPÉUTICO:\n"
                    f"A. Tratamiento farmacológico\n"
                    f"  • Medicamentos para cada diagnóstico (nombre genérico, dosis, vía, frecuencia, duración).\n"
                    f"  • Si indicas AINEs → añadir esomeprazol 40 mg VO cada 24 h (protección gástrica).\n"
                    f"  • No repetir familias farmacológicas.\n"
                    f"  • Considerar alergias ({alergias}) y medicamentos de uso continuo. {medicamentos_continuos}\n"
                    f"  • Añadir tratamiento de patologías crónicas {patologia_cronica}\n"
                    f"  • Factores a considerar (antecedentes médicos relevantes y posibles interacciones).\n"
                    f"  • Recomendaciones de seguimiento y monitoreo del paciente\n"
                    f"  • Plan de cambios en el estilo de vida detallado\n"
                    f"  • Plan de educación al paciente, indicando lo explicado\n"
                    f"  • Predicciones sobre posibles complicaciones o evolución de la enfermedad\n"
                    f"  • Plan de referimientos completos (razones para referir) siguiendo las directrices del VA\n"
                    f"  • Plan de análisis de laboratorios e imágenes\n"
                    f"  • Verificar diagnósticos previos y evitar medicamentos que causen alergia\n"
                    f"  • Colocar cita de seguimiento en 4 semanas\n"
                    f"  • Cada parte del plan debe verse organizada a la vista\n"
                    f"  • Desglosar el Plan de Análisis de Laboratorios e Imágenes y especificar para qué se indicó\n"
        
                    f"INFORMACIÓN DEL PACIENTE:\n"
                    f"Historia de Enfermedad Actual:\n{historia_enfermedad}\n\n"
                    f"Diagnósticos Previos: {diagnosticos_previos}\n\n"
                    f"Medicamentos de Uso Continuo: {medicamentos_continuos}\n\n"
                    f"Antecedentes Personales:\n{self.get_antecedentes_personales()}\n\n"
                    f"Antecedentes Heredo-familiares:\n{self.get_antecedentes_heredofamiliares()}\n\n"
        
        
                    f"Notas de Referimientos:\n"
                    f"escribir las notas de referimientos extensas completas  dirijidas  a las diferente especialidades referidas explicando el porque se refiere:\n"
        
        
        
        
                    f"ESTRUCTURA DEL INFORME:\n"
                    f". Motivo de la consulta\n"
                    f". Historia de Enfermedad Actual\n"
                    f". Revisión por Sistemas\n"
                    f". Antecedentes Personales\n"
                    f". Antecedentes hereditarios familiares\n"
                    f". Diagnósticos y diagnosticos diferenciales  (con códigos CIE-10)\n\n"
        
                    f"PLAN TERAPÉUTICO:\n"
                    f". Detalle de Medicamentos Indicados\n"
                    f". Medicamentos de Uso Continuo\n"
                    f". Interacciones Medicamentosas\n"
                    f". Recomendaciones de Seguimiento\n"
                    f". Predicciones sobre Complicaciones Potenciales o Evolución de la Enfermedad\n\n"
        
                    f"CAMBIOS Y CUIDADOS EN EL ESTILO DE VIDA:\n"
                    f". Plan de Cambios en el Estilo de Vida\n"
                    f". Plan de Actividad Física Personalizado\n"
                    f". Educación en Autocuidado y Red de Apoyo\n\n"
        
                    f"PLAN DE EDUCACIÓN:\n"
                    f". Información sobre la Condición y su Manejo\n"
                    f". Capacitación en el Uso de Medicamentos y Dispositivos Médicos\n"
                    f". Estrategias para Identificar y Evitar Factores Desencadenantes\n"
                    f". Recomendaciones para Reconocer Síntomas de Alarma\n"
                    f". Consejos para Mejorar la Adherencia al Tratamiento\n"
        
        
                    f"EVALUACIÓN ADICIONAL:\n"
                    f". Evaluación del Estado Psicoemocional\n"
                    f". Factores de Riesgo y Estrategias de Prevención\n"
                    f". Adherencia al Tratamiento y Barreras Identificadas\n"
                    f". Plan de Monitoreo Domiciliario\n\n"
        
                    f"REFERIMIENTOS Y ANÁLISIS COMPLEMENTARIOS:\n"
                    f". Plan de Referimientos (incluyendo especialidades referidas)\n"
                    f". Plan de Análisis de Laboratorios e Imágenes\n"
                    f". Notas de Referimientos\n\n"
        
                    f"OBJETIVOS Y PRONÓSTICO:\n"
                    f". Objetivos a Corto, Mediano y Largo Plazo\n"
                    f". Predicción del Pronóstico con Indicadores de Mejoría\n\n"
        
                    f"CONCLUSIONES\n"
        
        
        
                    f"Por favor, genere un informe médico completo siguiendo estas instrucciones y utilizando la información proporcionada del paciente."
                )
       

       
        
            
        bp = self.bp_field.text().strip()
        hr = self.hr_field.text().strip()
        try:
            weight_val = float(self.weight_field.text())
            weight_kg = self.convert_weight_to_kg(weight_val, self.current_weight_unit)
        except ValueError:
            weight_val = None
            weight_kg = None
        height_cm = self.current_patient.altura or 0
        height_m = height_cm / 100 if height_cm else 0
        bmi = weight_kg / (height_m ** 2) if weight_kg and height_m else None
        bmi_str = f"{bmi:.1f}" if bmi else ""
        bmi_cat = (
            self.categorize_bmi(bmi, weight_kg, height_cm) if bmi else ""
        )

        patient_info = (
            "los datos siguietes no debe aparecer en el output pero si usarlo para generar las ideas\n"
            f"Nombre: {self.current_patient.nombre}\n"
            f"Edad: {format_patient_age(self.current_patient)}\n"
            f"Sexo: {self.current_patient.sexo}\n"
            f"Alergias: {self.current_patient.alergias}\n"
            f"PA: {bp}\n"
            f"FC: {hr}\n"
            f"Peso: {weight_val if weight_val is not None else ''} {self.current_weight_unit}\n"
            f"IMC: {bmi_str} {bmi_cat}\n\n"
        )

        if obst:
            patient_info += "Antecedentes obstétricos: " + "; ".join(obst) + "\n"

        us_entries = self.parse_ultrasound_lines(self.sonografias_input.toPlainText())
        hist_date = date.today()
        if self.current_history_id:
            session_temp = Session()
            hobj = session_temp.get(HistoriaClinica, self.current_history_id)
            if hobj:
                hist_date = hobj.fecha.date()
            session_temp.close()
        weeks, _ = self.compute_current_pregnancy_weeks(hist_date)
        if self.embarazo_activo_check.isChecked() and weeks is not None:
            if weeks >= 36:
                patient_info += "Aviso: refuerzo de vacuna dT pendiente si no se ha aplicado (36 semanas).\n"
            elif weeks >= 20:
                patient_info += "Aviso: primera dosis de vacuna dT pendiente si no se ha aplicado (20 semanas).\n"
        us_summary = self.format_ultrasound_summary(us_entries, ref_date=hist_date)
        if us_summary:
            patient_info += "Sonografías: " + us_summary + "\n"

        messages = [
            {
                "role": "system",
                "content": (
                    "Eres un médico especialista y redactas la nota "
                    "como si hubieras atendido personalmente al paciente. "
                    "Escribe en primera persona y en tiempo pasado clínico "
                    "(\"examiné\", \"se explicó\"). "
                    "La vida del paciente depende de la exactitud: no se "
                    "permiten omisiones."
                ),
            },
            {"role": "user", "content": patient_info + prompt},
        ]

        if self.history_thread and self.history_thread.isRunning():
            QMessageBox.information(self, "En progreso", "Ya se está generando una historia clínica")
            return

        self.sugerencias_ia.clear()
        self.sugerencias_ia.setText("Escribiendo...")
        QApplication.processEvents()

        self.generate_button.setEnabled(False)
        self.history_thread = HistoryGeneratorThread(messages, model)
        self.history_thread.chunk_received.connect(self.update_history_chunk)
        self.history_thread.finished.connect(self.history_generation_finished)
        self.history_thread.error.connect(self.history_generation_error)
        self.history_thread.finished.connect(lambda: self.generate_button.setEnabled(True))
        self.history_thread.error.connect(lambda _m: self.generate_button.setEnabled(True))
        self.history_thread.start()

    def update_history_chunk(self, text):
        if self.sugerencias_ia.toPlainText() == "Escribiendo...":
            self.sugerencias_ia.clear()
        self.sugerencias_ia.moveCursor(QTextCursor.End)
        self.sugerencias_ia.insertPlainText(text)
        QApplication.processEvents()

    def history_generation_finished(self, text):
        self.current_programs = self.detect_programs(text)
        if self.current_programs:
            text += "\n\nPrograma de Salud: " + "; ".join(self.current_programs)
        self.sugerencias_ia.setPlainText(text)
        self.maybe_prompt_form(text)
        self.labrun()
        self.interaccion_ia_input.clear()
        self.history_thread = None

    def history_generation_error(self, msg):
        # Clear placeholder text if generation aborted mid-stream
        if self.sugerencias_ia.toPlainText() == "Escribiendo...":
            self.sugerencias_ia.clear()

        # Swallow transient connection errors silently so the UI remains responsive
        if "InvalidChunkLength" in msg or "Connection" in msg:
            self.history_thread = None
            return

        # For other errors, notify the user
        QMessageBox.warning(self, "Error", f"No se pudo generar sugerencias: {msg}")
        self.history_thread = None


    
    
    
    def load_file(self):
        file_content = ""

        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getOpenFileName(self, "QFileDialog.getOpenFileName()", "", "All Files (*);;PDF Files (*.pdf)", options=options)
        if file_name:
            with open(file_name, 'rb') as file:
                pdf_reader = PdfReader(file)
                file_content = ""
                for page in range(len(pdf_reader.pages)):
                    file_content += pdf_reader.pages[page].extract_text()
        
        # Assign the raw file content to the AI interaction input
        #self.interaccion_ia_input.setText(file_content)
    
        self.interaccion_ia_output.clear()
       # Set the OpenAI API key
                # Utilizar GPT para extraer información de la prescripción
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as config_file:
                config_data = json.load(config_file)
                model = apply_openai_key(config_data)
        
        # Evitar usar file_content si está vacío
        if file_content:
            prompt = f"de aqui debes extraer solo los datos de laboratorio y el nombre del paciente y la fecha de realizacion debe escribir al final un resumen descriptivo de los hallazgos: {file_content}\n"
        
            # Use the AI to process the results
            response = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "Eres un asistente médico experto."},
                    {"role": "user", "content": prompt},
                ],
            )
            
            # Assign the AI's response to a variable
            ai_response = response.choices[0].message['content']
            
            # Remove headers and bold formatting from the AI's response
            ai_response = re.sub(r'#{1,6}\s', '', ai_response)  # Remove headers
            ai_response = re.sub(r'\*{1,3}(.+?)\*{1,3}', r'\1', ai_response)  # Remove bold formatting
            
            # Assign the modified AI's response to the interaccion_ia_output
            self.interaccion_ia_output.setText(ai_response)
        else:
            # Mostrar un mensaje de error o manejar el caso donde no se seleccionó un archivo
            self.interaccion_ia_output.setText("No se seleccionó ningún archivo o no se pudo extraer texto del PDF.")

    def analyze_image(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getOpenFileName(
            self,
            "Seleccionar Imagen",
            "",
            "Imágenes (*.png *.jpg *.jpeg *.bmp);;All Files (*)",
            options=options,
        )
        if not file_name:
            return

        with open(file_name, "rb") as f:
            img_data = base64.b64encode(f.read()).decode("utf-8")

        if self.current_patient:
            dialog = ImageQueryDialog(self.current_patient, img_data, self)
            dialog.exec_()





    def interact_with_ia(self):
        self.interaccion_ia_output.clear()
        if not self.current_patient:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")
            return
    
        # Intenta obtener cada variable y asigna una cadena vacía si no está definida
        try:
            nombre = self.current_patient.nombre
        except NameError:
            nombre = ''
        try:
            edad = self.current_patient.edad
        except NameError:
            edad = ''
        try:
            sexo = self.current_patient.sexo
        except NameError:
            sexo = ''
        try:
            alergias = self.current_patient.alergias
        except NameError:
            alergias = ''
        try:
            historia_enfermedad = self.historia_enfermedad_input.toPlainText()
        except NameError:
            historia_enfermedad = ''
        try:
            diagnosticos_previos = self.current_patient.diagnosticos_previos
        except NameError:
            diagnosticos_previos = ''
        try:
            medicamentos_continuos = self.current_patient.medicamentos_continuos
        except NameError:
            medicamentos_continuos = ''
        try:
            antecedentes_personales = self.get_antecedentes_personales()
        except NameError:
            antecedentes_personales = ''
        try:
            antecedentes_heredofamiliares = self.get_antecedentes_heredofamiliares()
        except NameError:
            antecedentes_heredofamiliares = ''

        try:
            weight_val = float(self.weight_field.text())
            weight_kg = self.convert_weight_to_kg(weight_val, self.current_weight_unit)
        except ValueError:
            weight_val = None
            weight_kg = None
        height_cm = self.current_patient.altura or 0
        height_m = height_cm / 100 if height_cm else 0
        bmi = weight_kg / (height_m ** 2) if weight_kg and height_m else None
        bmi_str = f"{bmi:.1f}" if bmi else ""
        bmi_cat = (
            self.categorize_bmi(bmi, weight_kg, height_cm) if bmi else ""
        )

        patient_info = f"Nombre: {nombre}\n" \
                       f"Edad: {edad}\n" \
                       f"Sexo: {sexo}\n" \
                       f"Alergias: {alergias}\n\n" \
                       f"{historia_enfermedad}\n\n" \
                       f"Diagnósticos Previos: {diagnosticos_previos}\n\n" \
                       f"Medicamentos de Uso Continuo: {medicamentos_continuos}\n\n" \
                       f"{antecedentes_personales}\n\n" \
                       f"Antecedentes Heredo-familiares\n" \
                       f"{antecedentes_heredofamiliares}\n\n" \
                       f"Peso: {weight_val if weight_kg is not None else ''} {self.current_weight_unit}\n" \
                       f"IMC: {bmi_str} {bmi_cat}\n\n"
    
        user_prompt = self.interaccion_ia_input.toPlainText()
    
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as config_file:
                config_data = json.load(config_file)
                model = apply_openai_key(config_data)
        messages = [
            {"role": "system", "content": "actua como médico experto , actua como asistente medico ademas de  asistente general que responde todo lo que se pregunta."},
            {"role": "user", "content": patient_info + user_prompt},
        ]
        self.interaccion_ia_output.setText("Escribiendo...")
        self.ai_thread = AIStreamThread(messages, model)
        self.ai_thread.chunk_received.connect(self.update_ai_chunk)
        self.ai_thread.finished.connect(self.ai_interaction_finished)
        self.ai_thread.error.connect(lambda m: QMessageBox.warning(self, "Error", m))
        self.ai_thread.start()

    def update_ai_chunk(self, text):
        if self.interaccion_ia_output.toPlainText() == "Escribiendo...":
            self.interaccion_ia_output.clear()
        self.interaccion_ia_output.moveCursor(QTextCursor.End)
        self.interaccion_ia_output.insertPlainText(text)
        QApplication.processEvents()

    def ai_interaction_finished(self, text):
        self.interaccion_ia_output.append("")
        self.interaccion_ia_output.insertPlainText(text)
        self.interaccion_ia_input.clear()

    def sanitize_filename(self, filename, max_length=30):
        # Remove invalid characters
        filename = re.sub(r'[<>:"/\\|?*\s]', '_', filename)
        # Remove leading/trailing periods and spaces
        filename = filename.strip('. ')
        # Truncate to max_length
        return filename[:max_length]

    def get_unique_filename(self, base_path, filename):
        name, ext = os.path.splitext(filename)
        counter = 1
        while os.path.exists(os.path.join(base_path, filename)):
            filename = f"{name}_{counter}{ext}"
            counter += 1
        return filename
    
    def translate_text_with_ai(self, text, target_lang="en"):
        
                # Utilizar GPT para extraer información de la prescripción
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as config_file:
                config_data = json.load(config_file)
                model = apply_openai_key(config_data)
        messages = [
            {"role": "system", "content": f"You are a translator. Translate the following text to ingles,no colocar AI Suggestions:  {target_lang}."},
            {"role": "user", "content": text}
        ]
        try:
            response = openai.ChatCompletion.create(
               model=model,  # Changed from "gpt-4o-mini" to a valid model
                messages=messages
            )
            translated_text = response.choices[0].message['content'].strip()
            return translated_text
        except openai.error.OpenAIError as e:
            print(f"An error occurred: {str(e)}")
            return f"Translation error: {str(e)}"
    def export_to_word(self):
        if not self.current_patient:
            QMessageBox.warning(self, "Error", "No hay paciente seleccionado.")
            return

        # Crear documento en español
        document_es = Document()
        document_es.add_heading(f"Historia Médica de {self.current_patient.nombre}", 0)
        
        document_es.add_heading("Información del Paciente", level=1)
        document_es.add_paragraph(f"Nombre: {self.current_patient.nombre}")
        document_es.add_paragraph(f"Edad: {self.current_patient.edad}")
        document_es.add_paragraph(f"Sexo: {self.current_patient.sexo}")
        document_es.add_paragraph(f"Alergias: {self.current_patient.alergias}")
        document_es.add_paragraph(f"Dirección: {self.current_patient.direccion}")
        document_es.add_paragraph(f"Estado Civil: {self.current_patient.estado_civil}")
        document_es.add_paragraph(f"Religión: {self.current_patient.religion}")
        document_es.add_paragraph(f"Lugar de Nacimiento: {self.current_patient.lugar_nacimiento}")
        if self.current_patient.fecha_nacimiento:
            fecha_nac = self.current_patient.fecha_nacimiento.strftime('%d/%m/%Y')
        else:
            fecha_nac = 'N/D'
        document_es.add_paragraph(f"Fecha de Nacimiento: {fecha_nac}")
        document_es.add_paragraph(f"Profesión: {self.current_patient.ocupacion}")
        
        document_es.add_heading("Historia Médica", level=1)
        document_es.add_paragraph(f"Antecedentes Personales: {self.get_antecedentes_personales()}")
        document_es.add_paragraph(f"Antecedentes Heredofamiliares: {self.get_antecedentes_heredofamiliares()}")
        document_es.add_paragraph(f"{self.sugerencias_ia.toPlainText()}")
        
        # Traducción al inglés utilizando IA
        document_en = Document()
        document_en.add_heading(self.translate_text_with_ai(f"Historia Médica de {self.current_patient.nombre}"), 0)
        
        document_en.add_heading(self.translate_text_with_ai("Información del Paciente"), level=1)
        document_en.add_paragraph(self.translate_text_with_ai(f"Nombre: {self.current_patient.nombre}"))
        document_en.add_paragraph(self.translate_text_with_ai(f"Edad: {self.current_patient.edad}"))
        document_en.add_paragraph(self.translate_text_with_ai(f"Sexo: {self.current_patient.sexo}"))
        document_en.add_paragraph(self.translate_text_with_ai(f"Alergias: {self.current_patient.alergias}"))
        document_en.add_paragraph(self.translate_text_with_ai(f"Dirección: {self.current_patient.direccion}"))
        document_en.add_paragraph(self.translate_text_with_ai(f"Estado Civil: {self.current_patient.estado_civil}"))
        document_en.add_paragraph(self.translate_text_with_ai(f"Religión: {self.current_patient.religion}"))
        document_en.add_paragraph(self.translate_text_with_ai(f"Lugar de Nacimiento: {self.current_patient.lugar_nacimiento}"))
        if self.current_patient.fecha_nacimiento:
            fecha_nac_en = self.current_patient.fecha_nacimiento.strftime('%d/%m/%Y')
        else:
            fecha_nac_en = 'N/A'
        document_en.add_paragraph(self.translate_text_with_ai(f"Fecha de Nacimiento: {fecha_nac_en}"))
        document_en.add_paragraph(self.translate_text_with_ai(f"Profesión: {self.current_patient.ocupacion}"))
        
        document_en.add_heading(self.translate_text_with_ai("Historia Médica"), level=1)
        document_en.add_paragraph(self.translate_text_with_ai(f"Antecedentes Personales: {self.get_antecedentes_personales()}"))
        document_en.add_paragraph(self.translate_text_with_ai(f"Antecedentes Heredofamiliares: {self.get_antecedentes_heredofamiliares()}"))
        document_en.add_paragraph(self.translate_text_with_ai(f"Sugerencias de IA: {self.sugerencias_ia.toPlainText()}"))
    
        sanitized_name = self.sanitize_filename(self.current_patient.nombre)
        date_str = datetime.now().strftime('%Y%m%d')
        default_filename = f"HC_{sanitized_name}_{date_str}"
    
        home = os.path.expanduser("~")
        documents_path = os.path.join(home, "Documents")
    
        try:
            os.makedirs(documents_path, exist_ok=True)
        except Exception as e:
            QMessageBox.warning(self, "Error", f"No se pudo crear el directorio de documentos: {str(e)}")
            return
    
        options = QFileDialog.Options()
        file_path_es, _ = QFileDialog.getSaveFileName(
            self, "Guardar Archivo (Español)", 
            os.path.join(documents_path, default_filename + "_es.docx"),
            "Archivos Word (*.docx);;Todos los archivos (*)",
            options=options
        )
    
        if file_path_es:
            try:
                if not file_path_es.lower().endswith('.docx'):
                    file_path_es += '.docx'
                dir_path, filename_es = os.path.split(file_path_es)
                unique_filename_es = self.get_unique_filename(dir_path, filename_es)
                final_path_es = os.path.join(dir_path, unique_filename_es)
                document_es.save(final_path_es)

                filename_en = unique_filename_es.replace("_es.docx", "_en.docx")
                final_path_en = os.path.join(dir_path, filename_en)
                final_path_en = self.get_unique_filename(dir_path, filename_en)
                document_en.save(final_path_en)

                QMessageBox.information(self, "Éxito", f"Documentos exportados exitosamente como:\n{unique_filename_es}\n{os.path.basename(final_path_en)}")
                os.startfile(final_path_es)
                os.startfile(final_path_en)
            except PermissionError:
                QMessageBox.warning(self, "Error", "No se pudo guardar el archivo debido a un error de permiso.")
            except Exception as e:
                QMessageBox.warning(self, "Error", f"No se pudo guardar el archivo: {str(e)}")
        else:
            QMessageBox.warning(self, "Aviso", "No se seleccionó ninguna ubicación para guardar el archivo.")


def clean_prescription_text(text, patient_name=""):
    """Sanitize prescription text and ensure basic spacing."""
    patient_name = patient_name.lower() if patient_name else ""
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        low = stripped.lower()
        if patient_name and patient_name in low:
            continue
        if low.startswith("paciente") or low.startswith("nombre"):
            continue
        stripped = re.sub(r"\s*\((?:covered|cubierto|not covered|no cubierto).*?\)", "", stripped, flags=re.IGNORECASE)
        lines.append(stripped)
    result = "\n".join(lines)
    result = re.sub(r"(?i)(receta m[aá]dica)\s*", r"\1\n", result)
    result = re.sub(r"(?i)(patolog[íi]as?:[^\n]*)\s+(?=\d)", r"\1\n", result)
    result = re.sub(r"(\d+[\.\)-]\s[^\n]+)(?=\s*\d+[\.\)-])", r"\1\n", result)
    result = re.sub(r"(-\s[^\n]+)(?=\s*-\s)", r"\1\n", result)
    # Add blank line between numbered items
    result = re.sub(r"(?m)(^\d+[\.\)-].+$)", r"\1\n", result)
    # Capitalize first letter of each line
    result = "\n".join(l[:1].upper() + l[1:] if l else "" for l in result.splitlines())
    return result.strip()


def extract_medication_lines(text):
    """Return only lines that look like medication entries."""
    meds = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        # Check for bullet or numbered lines
        if re.match(r"^[-\d]", stripped):
            candidate = stripped
        else:
            candidate = stripped
        if re.search(r"\b(mg|ml|g|tableta|c[aá]psula|ampolla|gotas|inyecci[oó]n|pastilla)\b", candidate, re.IGNORECASE):
            meds.append(candidate)
    return "\n".join(meds)


def interpretar_denver(respuestas_dict):
    """Return general result for Denver II from answer codes."""
    fallos = 0
    sospechosos = 0
    for area in [
        "personal_social",
        "motora_fina_adaptativa",
        "lenguaje",
        "motora_gruesa",
    ]:
        respuestas = respuestas_dict.get(area, [])
        f = sum(1 for _, r in respuestas if r == "F")
        if f >= 2:
            fallos += 1
        elif f == 1:
            sospechosos += 1
    if fallos >= 2:
        return "Anormal"
    elif sospechosos >= 1:
        return "Sospechoso"
    return "Normal"


class PrescriptionReviewDialog(QDialog):
    def __init__(self, initial_prescription, previous_diagnoses, alergias, current_patient, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Revisar Prescripción")
        self.setWindowFlags(self.windowFlags() | Qt.WindowMaximizeButtonHint | Qt.WindowMinimizeButtonHint)
        
        self.alergias = alergias
        self.current_patient = current_patient
        
        layout = QVBoxLayout()
        StyleHelper.set_window_icon(self, 'fa5s.pills')

        # Diagnósticos
        diagnoses_label = QLabel("Diagnósticos:")
        StyleHelper.style_label(diagnoses_label)
        layout.addWidget(diagnoses_label)
        
        self.diagnoses_text = self.create_autofit_text_edit()
        organized_diagnoses = self.organize_diagnoses(previous_diagnoses)
        self.diagnoses_text.setPlainText(organized_diagnoses)
        self.diagnoses_text.set_auto_resize(True)  # Habilitar autoajuste
        StyleHelper.style_input(self.diagnoses_text)
        layout.addWidget(self.diagnoses_text)

        # Medicamentos generados
        medications_label = QLabel("Medicamentos generados:")
        StyleHelper.style_label(medications_label)
        layout.addWidget(medications_label)
        
        self.medications_text = self.create_autofit_text_edit(read_only=True)
        organized_medications = self.organize_medications(initial_prescription)
        organized_medications = self.clean_prescription_text(organized_medications)
        self.medications_text.setPlainText(organized_medications)
        self.medications_text.set_auto_resize(True)  # Habilitar autoajuste
        StyleHelper.style_input(self.medications_text)
        layout.addWidget(self.medications_text)
        
        # Interacción con IA
        ia_label = QLabel("Interacción con IA:")
        StyleHelper.style_label(ia_label)
        layout.addWidget(ia_label)
        
        self.ia_interaction_text = self.create_autofit_text_edit()
        self.ia_interaction_text.setPlaceholderText("Interactúa con la IA aquí para modificar la prescripción...")
        self.ia_interaction_text.set_auto_resize(True)  # Habilitar autoajuste
        StyleHelper.style_input(self.ia_interaction_text)
        layout.addWidget(self.ia_interaction_text)

        # Alergias del Paciente
        alergias_label = QLabel("Alergias del Paciente:")
        StyleHelper.style_label(alergias_label)
        layout.addWidget(alergias_label)

        self.alergias_text = QTextEdit()
        self.alergias_text.setPlainText(self.alergias if self.alergias else "No se han registrado alergias")
        self.alergias_text.setReadOnly(True)
        self.alergias_text.setFixedHeight(100)
        StyleHelper.style_input(self.alergias_text)
        layout.addWidget(self.alergias_text)
        
        # Botones
        button_layout = QHBoxLayout()
        self.update_button = QPushButton("Actualizar Medicamentos")
        self.accept_button = QPushButton("Aceptar Prescripción")
        self.auto_adjust_button = QPushButton("Autoajustar Tamaño")
        self.recheck_button = QPushButton("Rechequear Cobertura")

        if qta:
            self.update_button.setIcon(qta.icon("fa5s.sync", color="white"))
            self.accept_button.setIcon(qta.icon("fa5s.check", color="white"))
            self.auto_adjust_button.setIcon(qta.icon("fa5s.expand-arrows-alt", color="white"))
            self.recheck_button.setIcon(qta.icon("fa5s.sync", color="white"))
        StyleHelper.style_button(self.update_button)
        StyleHelper.style_button(self.accept_button)
        StyleHelper.style_button(self.auto_adjust_button)
        StyleHelper.style_button(self.recheck_button)
        button_layout.addWidget(self.update_button)
        button_layout.addWidget(self.accept_button)
        button_layout.addWidget(self.auto_adjust_button)
        button_layout.addWidget(self.recheck_button)
        layout.addLayout(button_layout)
        
        # Scroll Area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_content = QWidget()
        scroll_content.setLayout(layout)
        scroll_area.setWidget(scroll_content)
        
        main_layout = QVBoxLayout()
        main_layout.addWidget(scroll_area)
        self.setLayout(main_layout)
        
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        
        self.update_button.clicked.connect(self.update_medications)
        self.accept_button.clicked.connect(self.accept)
        self.auto_adjust_button.clicked.connect(self.auto_adjust_sizes)
        self.recheck_button.clicked.connect(self.reverificar_cobertura)


    def organize_diagnoses(self, diagnoses):
        lines = [l.strip() for l in diagnoses.split('\n') if l.strip()]
        clean_lines = [re.sub(r'\s*\((?:Cubierto|No Cubierto)\)', '', l, flags=re.IGNORECASE) for l in lines]
        return "\n".join(clean_lines)
    
    def organize_medications(self, medications):
        lines = [l.strip() for l in medications.split('\n') if l.strip()]
        clean_lines = [re.sub(r'\s*\((?:Covered|Not Covered)\)', '', l, flags=re.IGNORECASE) for l in lines]
        return "\n".join(clean_lines)

    def clean_prescription_text(self, text):
        return clean_prescription_text(text, getattr(self.current_patient, 'nombre', ''))

    def apply_original_format(self, original, new):
        """Format new prescription lines using the style of the original."""
        orig_lines = [l.strip() for l in original.splitlines() if l.strip()]
        if not orig_lines:
            return new
        first = orig_lines[0]
        numbered = bool(re.match(r'^\d+[\).-]', first))
        bullet = first.strip().startswith('-')
        new_lines = [re.sub(r'^\d+[\).-]\s*', '', l).lstrip('-').strip() for l in new.splitlines() if l.strip()]
        if numbered:
            symbol_match = re.match(r'^\d+([\).-])', first)
            symbol = symbol_match.group(1) if symbol_match else '.'
            return '\n'.join(f"{i+1}{symbol} {line}" for i, line in enumerate(new_lines))
        if bullet:
            return '\n'.join(f"- {line}" for line in new_lines)
        return '\n'.join(new_lines)

    def auto_adjust_sizes(self):
        self.diagnoses_text.adjust_height_to_content()
        self.medications_text.adjust_height_to_content()

    def create_autofit_text_edit(self, read_only=False):
        text_edit = AutoAdjustingTextEdit()
        text_edit.setReadOnly(read_only)
        return text_edit


    def update_medications(self):
        user_input = self.ia_interaction_text.toPlainText()
        current_prescription = self.medications_text.toPlainText()

        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r') as config_file:
                cfg = json.load(config_file)
            model = apply_openai_key(cfg)
        else:
            openai.api_key = os.getenv('OPENAI_API_KEY', '')
            model = ''
            if not openai.api_key:
                QMessageBox.warning(self, 'Error', 'Falta configuración de OpenAI')
                return

        alergias = getattr(self.current_patient, 'alergias', 'No se han registrado alergias')
        diagnosticos_previos = getattr(self.current_patient, 'diagnosticos_previos', 'No se han registrado diagnósticos previos')
        medicamentos_continuos = getattr(self.current_patient, 'medicamentos_continuos', 'No se han registrado medicamentos de uso continuo')

        prompt = f"""
        Current prescription:
        {current_prescription}

        User input:
        {user_input}

        Patient allergies:
        {alergias}
        regla universal colocar en espanol si el diagnostico ya fue en el pasado no indicar para esa patologia,si no hay patologia om esta negada no poner en el diagnostico :siempre haras lo que el usuario te escriba
         nunca equivocarte  ya que la vida del paciente depende es esto:   nunca indicar algun medicamento que el pacietne sea alergico  Patient allergies:{alergias} eliminarlo si esta en la lista y cambiarlo por otro que no sea familia de la familia de medicamentos que le causa alergias
        solo agregar medicamentos nuevos si el usuario te lo pide, de lo contrario no hacerlo
        si el usuario pide agragar un medicamento o familia de medicamento al que el paciente es alergico no colocarlo y colocar otro que no sea alegico y que haga el efecto
        nunca indicar algun medicamento que el pacietne sea alergico  Patient allergies:{alergias} eliminarlo si esta en la lista y cambiarlo por otro que no sea familia de la familia de medicamentos que le causa alergias
        tu funcion es modificar si se te pide la indicacion
        Check the list to see if there is any medication the patient is allergic to and put in parentheses that the patient is allergic to that medication and its family, and suggest another for management. Never omit this because it puts the patient's life at risk.
        Extract the medical prescription information from the following text, including medication names.
        If available, also extract the dosage, frequency, and duration of the treatment.
        Responde en español. No incluyas las palabras Covered ni Not Covered.
        Medications may appear without dosage, frequency, and duration. Ensure all medications without exception have dosage, frequency, and duration.
        As a medical expert, prescribe appropriate medications. If a medication is unknown, look it up on the web and include the dosage.
        Include all continuous use medications.
        Do not include \"as prescribed by the physician\" in the dosage; we must specify the dosage and duration.
        Do not state that allergies need to be specified if they are not present.
        Do not repeat medications or suggest medications from the same family.
        If NSAIDs (Non-Steroidal Anti-Inflammatory Drugs) are prescribed, also include a gastric protector in the prescription with dosage, frequency, and duration of the treatment (only one, esomeprazole 40mg is Covered for this condition).
        Format the output as a clear and concise list of prescriptions, one per line. The output format is for medical prescription.
        Do not include medications without dosage, frequency, and duration of treatment. Indicate for which pathology it was prescribed.
        Do not include any additional information or explanations. Todo en español.
        Do not include coverage information or ICD-10 codes.
        Do not include information code ICD-10.
        Remember, patient safety is the top priority. If you're unsure about a medication's safety given the patient's allergies, err on the side of caution and do not include it.
        Do not include disclaimers or mentions about the need for additional medical evaluation.
        Do not include referrals, only medications. Medications must not be repeated; the patient's life depends on this.
        no colocar notes, notas

        Coverage Rules:
        1. A medication is (Covered) only if it is explicitly associated with a diagnosis marked as (Cubierto) in the previous diagnoses: {diagnosticos_previos}.
        2. If a medication is not associated with any (Cubierto) diagnosis in the previous diagnoses, mark it as (Not Covered).
        3. If there are no previous diagnoses recorded or none are marked as (Cubierto), all medications should be (Not Covered).
        4. If in doubt, mark as (Not Covered).
        si un medicamento esta marcado como not Covered buscar algun diagnostico marcado como Cubierto si existe en {diagnosticos_previos} que se pueda relacionar y cambiar
        Text:
         Add medication if not present Medications may appear without dosage, frequency, and duration. Ensure all medications without exception have dosage, frequency, and duration. {diagnosticos_previos}
        """

        messages = [
            {"role": "system", "content": "You are a medical expert assistant specializing in updating prescriptions based on user input."},
            {"role": "user", "content": prompt},
        ]

        self.medications_text.setText('Escribiendo...')
        self.current_prescription_backup = current_prescription
        self.update_thread = AIStreamThread(messages, model)
        self.update_thread.chunk_received.connect(self.append_med_chunk)
        self.update_thread.finished.connect(self.finish_med_chunk)
        self.update_thread.error.connect(lambda m: QMessageBox.warning(self, 'Error', m))
        self.update_thread.start()
        self.ia_interaction_text.clear()

 
      
    def eliminar_duplicados(self, prescripcion):
        lineas = prescripcion.split('\n')
        medicamentos_unicos = {}
    
        for linea in lineas:
            # Normalizar la línea: eliminar espacios extra y convertir a minúsculas
            linea_normalizada = ' '.join(linea.lower().split())
    
            # Simplificar condiciones similares (derecho/izquierdo) y eliminar numeración previa
            linea_normalizada = re.sub(r'(\b(right|left)\b)', '', linea_normalizada)  # Eliminar referencias específicas de lado
            linea_normalizada = re.sub(r'^\d+\.\s*', '', linea_normalizada)  # Eliminar la numeración inicial si existe
    
            # Extraer el medicamento (nombre, dosis, frecuencia) sin incluir diagnósticos y cobertura
            partes = linea_normalizada.split(',')
            nombre_y_dosis = ','.join(partes[:3])  # Considera nombre, dosis y frecuencia para eliminar duplicados
    
            # Verificar si el medicamento ya ha sido agregado
            if nombre_y_dosis not in medicamentos_unicos:
                medicamentos_unicos[nombre_y_dosis] = linea  # Guardar la línea original para mantener formato
    
        # Convertir el diccionario a una lista de líneas únicas, conservando el formato original
        prescripcion_sin_duplicados = list(medicamentos_unicos.values())
    
        # Ordenar los medicamentos de forma numérica y reindexar los números de orden
        prescripcion_final = []
        for i, med in enumerate(prescripcion_sin_duplicados, start=1):
            # Eliminar cualquier numeración previa (número seguido de punto y espacio al inicio)
            med = re.sub(r'^\d+\.\s*', '', med)  # Remueve números al inicio seguidos de punto y espacio
            prescripcion_final.append(f"{i}. {med}")
    
        return '\n'.join(prescripcion_final)
    
    def verificar_duplicados(self, prescripcion):
        lineas = prescripcion.split('\n')
        medicamentos = [linea.split(',')[0].strip().lower() for linea in lineas]
        duplicados = [med for med in set(medicamentos) if medicamentos.count(med) > 1]
        return len(duplicados) == 0
    
    def filtrar_medicamentos(self, prescripcion):
        lineas = prescripcion.split('\n')
        medicamentos = []
    
        for linea in lineas:
            if "therapy" not in linea.lower() and "referral" not in linea.lower():
                medicamentos.append(linea)
    
        return '\n'.join(medicamentos)
        
    def reverificar_cobertura(self):
        diagnosticos_previos = self.diagnoses_text.toPlainText()
        prescripcion_inicial = self.medications_text.toPlainText()
        alergias = getattr(self.current_patient, 'alergias', 'No se han registrado alergias')
        
        
        # Eliminar las etiquetas de cobertura
        prescripcion_inicial = re.sub(r'\s*\(\s*C\s*o\s*v\s*e\s*r\s*e\s*d\s*\)|\s*\(\s*N\s*o\s*t\s*C\s*o\s*v\s*e\s*r\s*e\s*d\s*\)', '', prescripcion_inicial, flags=re.IGNORECASE)

        
        verification_prompt = f"""
        Based on the following information:
    
        Previous diagnoses:
        {diagnosticos_previos}
        nunca indicar algun medicamento que el pacietne sea alergico  Patient allergies:{alergias} si el usuario recomienda uno  eliminarlo si esta en la lista y cambiarlo por otro que no sea familia de la familia de medicamentos que le causa alergia
        Please verify and cambia the coverage status of each medication in the following prescription list. 
        Follow these rules strictly ya que la vida del paciente depende es esto:
        regla universal :siempre haras lo que el usuario te escriba
        
         nunca equivocarte y colacarlo en  ingles si la enfermedad ya fue en el pasado no indicar medicacion  ya que la vida del paciente depende es esto:   nunca indicar algun medicamento que el pacietne sea alergico  Patient allergies:{alergias} eliminarlo si esta en la lista y cambiarlo por otro que no sea familia de la familia de medicamentos que le causa alergia
        - Add esomeprazole 40 mg solo if an NSAID esta en la lista de medicamentos,el esomeprazol tendra la misma cobertura que el aines
    
        1. Coverage Rules:
           -no tomar en cuenta la cobertura de {prescripcion_inicial} usar nuevos cambios basados en las reglas
           - A medication is marked as (Covered) only if it is explicitly associated with a diagnosis marked as (Cubierto) in the previous diagnoses.
           - If a medication is not associated with any (Cubierto) diagnosis, mark it as (Not Covered).
           - If you're unsure about the coverage of a medication, mark it as (Not Covered).
           - If the previous diagnoses list is empty or says "No previous diagnoses have been recorded", all medications should be marked as (Not Covered).
           - agrupar, y colocar los covered primero
           - Do not include disclaimers or mentions about the need for additional medical evaluation, solo coloca la lista de medicamentos.
        2. Continuous Medications:
           - Ensure all continuous medications are included (if any) and all coverages are correctly assigned.
    
        Here's the prescription list to verify and adjust:
        {prescripcion_inicial}
    
        Please provide the corrected prescription list with accurate coverage statuses, following all the rules above.
        Do not include disclaimers or mentions about the need for additional medical evaluation.
        Do not include referrals, only medications. Medications must not be repeated; the patient's life depends on this.
        si se sugiere Ibuprofen cambiar Ibuprofen por dexketoprofeno 25 mg
        """
                # Utilizar GPT para extraer información de la prescripción
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as config_file:
                config_data = json.load(config_file)
                model = apply_openai_key(config_data)
        try:
            verification_response = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a medical assistant expert in verifying medication coverage based on diagnoses and managing continuous medications."},
                    {"role": "user", "content": verification_prompt},
                ],temperature= 0.1 # Ajusta la temperatura según sea necesario
            )
    
            updated_prescription = verification_response['choices'][0]['message']['content'].strip()
            updated_prescription = self.filtrar_medicamentos(updated_prescription)
            updated_prescription= self.eliminar_duplicados( updated_prescription)
            self.medications_text.setPlainText(updated_prescription)
    
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Could not verify coverage: {str(e)}")





    
    def get_updated_prescription(self, current_prescription, user_input):
                # Utilizar GPT para extraer información de la prescripción
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as config_file:
                config_data = json.load(config_file)
                model = apply_openai_key(config_data)
    
        # Obtener las alergias del paciente
        alergias = getattr(self.current_patient, 'alergias', 'No se han registrado alergias')
        diagnosticos_previos = getattr(self.current_patient, 'diagnosticos_previos', 'No se han registrado diagnósticos previos')
        medicamentos_continuos = getattr(self.current_patient, 'medicamentos_continuos', 'No se han registrado medicamentos de uso continuo')
    
        prompt = f"""
        Current prescription:
        {current_prescription}
    
        User input:
        {user_input}
    
        Patient allergies:
        {alergias}
        regla universal colocar en espanol si el diagnostico ya fue en el pasado no indicar para esa patologia,si no hay patologia om esta negada no poner en el diagnostico :siempre haras lo que el usuario te escriba
         nunca equivocarte  ya que la vida del paciente depende es esto:   nunca indicar algun medicamento que el pacietne sea alergico  Patient allergies:{alergias} eliminarlo si esta en la lista y cambiarlo por otro que no sea familia de la familia de medicamentos que le causa alergia
        solo agregar medicamentos nuevos si el usuario te lo pide, de lo contrario no hacerlo
        si el usuario pide agragar un medicamento o familia de medicamento al que el paciente es alergico no colocarlo y colocar otro que no sea alegico y que haga el efecto
        nunca indicar algun medicamento que el pacietne sea alergico  Patient allergies:{alergias} eliminarlo si esta en la lista y cambiarlo por otro que no sea familia de la familia de medicamentos que le causa alergias
        tu funcion es modificar si se te pide la indicacion
        Check the list to see if there is any medication the patient is allergic to and put in parentheses that the patient is allergic to that medication and its family, and suggest another for management. Never omit this because it puts the patient's life at risk.
        Extract the medical prescription information from the following text, including medication names.
        If available, also extract the dosage, frequency, and duration of the treatment.
        Responde en español. No incluyas las palabras Covered ni Not Covered.
        Medications may appear without dosage, frequency, and duration. Ensure all medications without exception have dosage, frequency, and duration.
        As a medical expert, prescribe appropriate medications. If a medication is unknown, look it up on the web and include the dosage.
        Include all continuous use medications.
        Do not include "as prescribed by the physician" in the dosage; we must specify the dosage and duration.
        Do not state that allergies need to be specified if they are not present.
        Do not repeat medications or suggest medications from the same family.
        If NSAIDs (Non-Steroidal Anti-Inflammatory Drugs) are prescribed, also include a gastric protector in the prescription with dosage, frequency, and duration of the treatment (only one, esomeprazole 40mg is Covered for this condition).
        Format the output as a clear and concise list of prescriptions, one per line. The output format is for medical prescription.
        Do not include medications without dosage, frequency, and duration of treatment. Indicate for which pathology it was prescribed.
        Do not include any additional information or explanations. Todo en español.
        Do not include coverage information or ICD-10 codes.
        Do not include information code ICD-10.
        Remember, patient safety is the top priority. If you're unsure about a medication's safety given the patient's allergies, err on the side of caution and do not include it.
        Do not include disclaimers or mentions about the need for additional medical evaluation.
        Do not include referrals, only medications. Medications must not be repeated; the patient's life depends on this.
        no colocar notes, notas
        
        Coverage Rules:
        1. A medication is (Covered) only if it is explicitly associated with a diagnosis marked as (Cubierto) in the previous diagnoses: {diagnosticos_previos}.
        2. If a medication is not associated with any (Cubierto) diagnosis in the previous diagnoses, mark it as (Not Covered).
        3. If there are no previous diagnoses recorded or none are marked as (Cubierto), all medications should be (Not Covered).
        4. If in doubt, mark as (Not Covered).
        si un medicamento esta marcado como not Covered buscar algun diagnostico marcado como Cubierto si existe en {diagnosticos_previos} que se pueda relacionar y cambiar
        Text:
         Add medication if not present Medications may appear without dosage, frequency, and duration. Ensure all medications without exception have dosage, frequency, and duration. {diagnosticos_previos}
        """
    
        try:
            response = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a medical expert assistant specializing in updating prescriptions based on user input."},
                    {"role": "user", "content": prompt},
                ],temperature= 0.1 # Ajusta la temperatura según sea necesario
            )
    
            updated_prescription = response['choices'][0]['message']['content'].strip()
            updated_prescription = self.filtrar_medicamentos(updated_prescription)
            updated_prescription= self.eliminar_duplicados( updated_prescription)
            return updated_prescription
        except Exception as e:
            return f"Error updating prescription: {str(e)}"

    def load_history_data(self):
        """Load all historical diagnoses and medications for the patient"""
        session = Session()
        patient = session.merge(self.current_patient)
        diagnoses = []
        for hist in patient.historias:
            for d in hist.diagnosticos:
                if d.descripcion and d.descripcion not in diagnoses:
                    diagnoses.append(d.descripcion)
        meds = []
        for hist in patient.historias:
            if hist.sugerencias_ia:
                for line in hist.sugerencias_ia.splitlines():
                    line = re.sub(r'\s*\((?:Covered|Not Covered|Cubierto|No Cubierto).*?\)', '', line, flags=re.IGNORECASE).strip()
                    if not line:
                        continue
                    med = line.split(',')[0].lstrip('- ').strip()
                    if med and med not in meds:
                        meds.append(med)
        session.close()
        self.diagnoses_text.setPlainText('\n'.join(diagnoses))
        self.medications_text.setPlainText('\n'.join(meds))

    def append_med_chunk(self, text):
        if self.medications_text.toPlainText() == 'Escribiendo...':
            self.medications_text.clear()
        self.medications_text.moveCursor(QTextCursor.End)
        self.medications_text.insertPlainText(text)
        QApplication.processEvents()

    def finish_med_chunk(self, text):
        cleaned = self.clean_prescription_text(text)
        if hasattr(self, 'current_prescription_backup'):
            cleaned = self.apply_original_format(self.current_prescription_backup, cleaned)
            del self.current_prescription_backup
        self.medications_text.setPlainText(cleaned)



class AutoAdjustingTextEdit(QTextEdit):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.document().contentsChanged.connect(self.adjust_height_to_content)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)  # Ocultar scroll vertical
        self.auto_resize = True  # Habilitar autoajuste por defecto
        self.heightMin = 200  # Altura mínima base
        self.heightMax = 10000  # Altura máxima arbitraria (ajustar según necesidades)
        
        # Ajustar la altura al inicializar
        self.adjust_height_to_content()

    def adjust_height_to_content(self):
        if not self.auto_resize:
            return
        
        # Obtener el tamaño del documento (contenido)
        doc_height = int(self.document().size().height()) + 10  # Añadir margen extra
        
        # Definir altura mínima y máxima
        doc_height = max(self.heightMin, min(doc_height, self.heightMax))
        
        # Ajustar la altura del cuadro de texto
        self.setFixedHeight(doc_height)

    def set_auto_resize(self, enable):
        self.auto_resize = enable
        if enable:
            self.adjust_height_to_content()









class AutoFitTextEdit(QTextEdit):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.document().contentsChanged.connect(self.sizeChange)
        self.heightMin = 800
        self.heightMax = 65000
        self.initialHeight = 800
        self.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.setMinimumHeight(self.initialHeight)
        self.setMaximumHeight(self.initialHeight)

    def sizeChange(self):
        docHeight = self.document().size().height()
        # Convertir docHeight a entero
        docHeight = int(docHeight)
        if docHeight == 0:
            docHeight = self.initialHeight
        if self.heightMin <= docHeight <= self.heightMax:
            self.setMinimumHeight(docHeight)
            self.setMaximumHeight(docHeight)

    def sizeHint(self):
        return QSize(self.width(), self.heightMin)

    def setMinimumLines(self, lines):
        self.heightMin = self.fontMetrics().lineSpacing() * lines

    def setMaximumLines(self, lines):
        self.heightMax = self.fontMetrics().lineSpacing() * lines

    def setInitialHeight(self, height):
        self.initialHeight = height
        self.setMinimumHeight(self.initialHeight)
        self.setMaximumHeight(self.initialHeight)
            

        
class LabOrderFormDialog(QDialog):
    def __init__(self, patient, initial_lab_request, lab_record, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.vials')
        self.patient = patient
        self.initial_lab_request = initial_lab_request
        self.lab_record = lab_record
        self.setWindowTitle("Generar Orden de Laboratorio")
        self.setWindowFlags(self.windowFlags() | Qt.WindowMaximizeButtonHint | Qt.WindowMinimizeButtonHint)
        
        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout()

        # Cuadro de texto para la solicitud de laboratorio actual
        lab_request_label = QLabel("Solicitud de laboratorio actual:")
        StyleHelper.apply_stylesheet(lab_request_label, "font-size: 32px; font-weight: bold;")
        layout.addWidget(lab_request_label)

        self.lab_request_text = QTextEdit()
        if self.lab_record:
            self.lab_request_text.setPlainText(self.lab_record)
        else:
            self.lab_request_text.setPlainText(self.initial_lab_request)
        self.lab_request_text.setReadOnly(True)
        StyleHelper.style_input(self.lab_request_text)
        layout.addWidget(self.lab_request_text)

        # Cuadro de texto para interacción con la IA
        ia_label = QLabel("Sugerir pruebas de laboratorio:")
        StyleHelper.apply_stylesheet(ia_label, "font-size: 32px; font-weight: bold;")
        layout.addWidget(ia_label)

        self.ia_lab_suggestion_text = QTextEdit()
        self.ia_lab_suggestion_text.setPlaceholderText("Interactúa con la IA aquí para sugerir pruebas de laboratorio...")
        StyleHelper.style_input(self.ia_lab_suggestion_text)
        layout.addWidget(self.ia_lab_suggestion_text)

        # Botones
        button_layout = QHBoxLayout()
        self.suggest_lab_button = QPushButton("Sugerir Pruebas")
        self.generate_lab_order_button = QPushButton("Generar Orden de Laboratorio")

        if qta:
            self.suggest_lab_button.setIcon(qta.icon("fa5s.lightbulb", color="white"))
            self.generate_lab_order_button.setIcon(qta.icon("fa5s.flask", color="white"))
        StyleHelper.style_button(self.suggest_lab_button)
        StyleHelper.style_button(self.generate_lab_order_button)
        button_layout.addWidget(self.suggest_lab_button)
        button_layout.addWidget(self.generate_lab_order_button)
        layout.addLayout(button_layout)

        self.setLayout(layout)

        # Configurar política de tamaño para permitir el autoajuste
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        # Conectar señales
        self.suggest_lab_button.clicked.connect(self.suggest_lab_tests)
        self.generate_lab_order_button.clicked.connect(self.generate_lab_order)

        # Ajustar el tamaño inicial
        self.adjustSize()

    def suggest_lab_tests(self):
        current_lab_request = self.lab_request_text.toPlainText()
        user_input = self.ia_lab_suggestion_text.toPlainText()
        prompt = self.build_lab_prompt(current_lab_request, user_input)
        if not prompt:
            return
        self.lab_request_text.setPlainText("Escribiendo...")
        self.lab_thread = AIStreamThread(prompt["messages"], prompt["model"])
        self.lab_thread.chunk_received.connect(self.update_lab_chunk)
        self.lab_thread.finished.connect(self.lab_suggestion_finished)
        self.lab_thread.error.connect(lambda m: QMessageBox.warning(self, "Error", m))
        self.lab_thread.start()

    def update_lab_chunk(self, text):
        if self.lab_request_text.toPlainText() == "Escribiendo...":
            self.lab_request_text.clear()
        self.lab_request_text.moveCursor(QTextCursor.End)
        self.lab_request_text.insertPlainText(text)
        QApplication.processEvents()

    def lab_suggestion_finished(self, text):
        self.lab_request_text.setPlainText(text)
        self.ia_lab_suggestion_text.clear()

    def build_lab_prompt(self, current_lab_request, user_input):
                # Utilizar GPT para extraer información de la prescripción
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as config_file:
                config_data = json.load(config_file)
                model = apply_openai_key(config_data)

        prompt = f"""
        Current lab request:
        {current_lab_request}

        User input:
        {user_input}

        Lab record:
        {self.lab_record}

        Based on the current lab request, the user's input, and the lab record, suggest appropriate laboratory tests.
        
        lista sin enumerar
        sin espacio entre lista 
               si el usuario escribe rutina hombre cargar, y si es rutina mujer todo menos psa Y TESTOSTERONA CAMBIAR POR ESTROGENOS  “HEMOGRAMA”, “ORINA”, “COPROLOGICO”, “TIPIFICACION”, “VDRL”, “GLICEMIA”,
        “HEMOGLOBINA GLUCOSILADA”, “BUN”, “CREATINA”, “COLESTEROL”, “TRIGLICERIDO”, “HDLC”, “LDL”, 
        “PROTEINAS TOTALES”, “ALBUMINA”, “GLOBULINA”, “ACIDO URICO”, “FOSFATASA ALCALINA”, “ASO”,
        “TESTOSTERONA”, “hiv”,hepatitis b y c “FUNCION HEPATICA”, “PERFIL LIPIDICO”, “CALCIO”, “FERRITINA”, “VITAMINA D”, “PSA”, “TSH”, “T4 LIBRE”, “T3 TOTAL”
        
        Format the output as a clear and concise list of laboratory tests, one per line.
        Include all mandatory tests. If the user mentions routine tests and other routine tests, add them accordingly.
        """

        messages = [
            {"role": "system", "content": "You are a medical expert assistant specializing in suggesting appropriate laboratory tests based on lab requests and patient information."},
            {"role": "user", "content": prompt},
        ]
        return {"messages": messages, "model": model}
    def extraer_laboratorio(self, sugerencias):
                # Utilizar GPT para extraer información de la prescripción
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as config_file:
                config_data = json.load(config_file)
                model = apply_openai_key(config_data)

        prompt = f"""
        From the following text, extract the laboratory tests suggested. Provide a clear and concise list of laboratory tests no with their reasons and any relevant details.
        Format the output as a clear and concise list of laboratory tests,
        si apararecer Pruebas de función renal no colocar encabezado, destallarlas en lista uno por uno
        si apararecer Pruebas de electrolitos no colocar encabezado, destallarlas en lista uno por uno
        si apararecer Pruebas hepaticas no colocar encabezado, destallarlas en lista uno por uno
        sin espacio entre lista
        si el usuario escribe rutina hombre cargar, y si es rutina mujer todo menos psa Y TESTOSTERONA CAMBIAR POR ESTROGENOS  “HEMOGRAMA”, “ORINA”, “COPROLOGICO”, “TIPIFICACION”, “VDRL”, “GLICEMIA”,
        “HEMOGLOBINA GLUCOSILADA”, “BUN”, “CREATINA”, “COLESTEROL”, “TRIGLICERIDO”, “HDLC”, “LDL”, 
        “PROTEINAS TOTALES”, “ALBUMINA”, “GLOBULINA”, “ACIDO URICO”, “FOSFATASA ALCALINA”, “ASO”,
        “TESTOSTERONA”, “hiv”,hepatitis b y c “FUNCION HEPATICA”, “PERFIL LIPIDICO”, “CALCIO”, “FERRITINA”, “VITAMINA D”, “PSA”, “TSH”, “T4 LIBRE”, “T3 TOTAL”
        
        Format the output as a clear and concise list of laboratory tests, one per line.
        lista sin enumerar
        Text:
        {sugerencias}
        {self.lab_record}
        """

        try:
            response = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a medical expert assistant specialized in extracting suggested laboratory tests from text."},
                    {"role": "user", "content": prompt},
                ],
            )

            laboratorio_extraido = response['choices'][0]['message']['content'].strip()
            laboratorio_extraido = re.sub(r'#{1,6}\s', '', laboratorio_extraido)  # Eliminar encabezados
            laboratorio_extraido = re.sub(r'\*{1,3}(.+?)\*{1,3}', r'\1', laboratorio_extraido)  # Eliminar énfasis

            return laboratorio_extraido
        except Exception as e:
            QMessageBox.warning(self, "Error", f"No se pudo extraer la información del laboratorio: {str(e)}")
            return None

    def generate_lab_order(self):


        def sanitize_filename(filename, max_length=30):
            # Remove invalid characters and replace spaces
            filename = re.sub(r'[<>:"/\\|?*\s]', '_', filename)
            # Remove leading/trailing periods and spaces
            filename = filename.strip('. ')
            # Limit the patient name length
            return filename[:max_length]
        
        def get_unique_filename(base_path, filename):
            name, ext = os.path.splitext(filename)
            counter = 1
            while os.path.exists(os.path.join(base_path, filename)):
                filename = f"{name}_{counter}{ext}"
                counter += 1
            return filename
        
        suggested_tests = self.lab_request_text.toPlainText()
    
        document = self.create_lab_order_form(suggested_tests)
    
        # Preparar para guardar el documento
        sanitized_name = sanitize_filename(self.patient.nombre)
        date_str = datetime.now().strftime('%Y%m%d')
        default_filename = f"Lab_Order_{sanitized_name}_{date_str}.docx"
    
        home = os.path.expanduser("~")
        documents_path = os.path.join(home, "Documents")
    
        # Asegurar que el directorio de documentos existe
        try:
            os.makedirs(documents_path, exist_ok=True)
        except Exception as e:
            QMessageBox.warning(self, "Error", f"No se pudo crear el directorio de documentos: {str(e)}")
            return
    
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Guardar Orden de Laboratorio", 
            os.path.join(documents_path, default_filename),
            "Archivos Word (*.docx);;Todos los archivos (*)",
            options=options
        )
    
        if file_path:
            try:
                # Asegurar que el archivo tiene extensión .docx
                if not file_path.lower().endswith('.docx'):
                    file_path += '.docx'
    
                # Obtener nombre de archivo único
                dir_path, filename = os.path.split(file_path)
                unique_filename = get_unique_filename(dir_path, filename)
                final_path = os.path.join(dir_path, unique_filename)
    
                # Guardar el documento
                document.save(final_path)
                QMessageBox.information(self, "Éxito", f"Orden de laboratorio exportada exitosamente como:\n{unique_filename}")
                os.startfile(final_path)
            except PermissionError:
                QMessageBox.warning(self, "Error", "No se pudo guardar el archivo debido a un error de permiso.")
            except Exception as e:
                QMessageBox.warning(self, "Error", f"No se pudo guardar el archivo: {str(e)}")
        else:
            QMessageBox.warning(self, "Aviso", "No se seleccionó ninguna ubicación para guardar el archivo.")




    def create_lab_order_form(self, suggested_tests):
        doc = Document()
    
        # Agregar logo
        try:    
            logo = "logo.png"    
            logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), logo)    
            doc.add_picture(logo_path, width=Inches(0.8))  # Ajusta el tamaño según sea necesario
        except OSError:    
            # Si no se encuentra el logo, simplemente omitimos    
            pass
    
        # Agregar título    
        title = doc.add_paragraph("LABORATORIO CLINICO")    
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
        title.runs[0].bold = True    
        title.runs[0].font.size = Pt(14)    
    
        # Agregar subtítulo
        subtitle = doc.add_paragraph("Pedido de Laboratorio")    
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
        subtitle.runs[0].italic = True
    
        # Agregar la fecha actual
        fecha_actual = datetime.now().strftime("%d de %B de %Y")    
        fecha_paragraph = doc.add_paragraph(f"Fecha: {fecha_actual}", style='Normal')    
        fecha_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
        # Información del paciente
        patient_info = doc.add_paragraph()    
        patient_info.add_run(f"Paciente: {self.patient.nombre}").bold = True    
        patient_info.add_run(f"  Edad: {self.patient.edad}")  
    
        # Línea separadora
        doc.add_paragraph("_" * 50)
    
        # Pruebas de laboratorio sugeridas
        doc.add_heading("Pruebas de Laboratorio Sugeridas", level=2)
    
        # Crear una tabla con dos columnas para las pruebas
        table = doc.add_table(rows=0, cols=2)
    
        # Ajustar el ancho de las columnas
        for column in table.columns:
            column.width = Inches(3.0)
    
        # Añadir las pruebas a la tabla
        tests = suggested_tests.split('\n')
        for i in range(0, len(tests), 2):
            cells = table.add_row().cells
            test_1 = tests[i].strip()
            test_2 = tests[i + 1].strip() if i + 1 < len(tests) else ""
            
            cells[0].text = f"☑ {test_1}" if test_1 else ""
            cells[1].text = f"☑ {test_2}" if test_2 else ""
    
            # Ajustar el espaciado entre filas
            for cell in cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)  # Ajusta el espacio después de cada párrafo a 0 puntos
                    paragraph.paragraph_format.space_before = Pt(0)  # Ajusta el espacio antes de cada párrafo a 0 puntos
                    paragraph.paragraph_format.line_spacing = Pt(12)  # Ajusta el interlineado a 12 puntos (puedes ajustar este valor según sea necesario)
    
        # Agregar pie de página
                        # Obtener pie de página desde ConfigReceta
        config_receta = ConfigReceta()
        config_receta.load_config()
        footer_text = config_receta.get_footer_text()
        footer_details = config_receta.get_footer_details()
        footer_phone = config_receta.get_footer_phone()
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = f"{footer_text}\n{footer_details}\n{footer_phone}"
        footer_para.alignment = 1  # Centrar el pie de página
    
        return doc






class IAAnalysisWindow(QDialog):
    def __init__(self, parent=None, excel_data=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.chart-bar')
        self.setWindowTitle("Análisis de Historia Clínica con IA")
        self.setGeometry(150, 150, 800, 600)
        self.setWindowFlags(self.windowFlags() | Qt.WindowMaximizeButtonHint)
        self.excel_data = excel_data
        self.setup_ui()

    def setup_ui(self):
        self.setStyleSheet("""
            QDialog {
                background-color: #ffffff;
            }
            QPushButton {
                background-color: #007bff;
                color: #ecf0f1;
                border: none;
                padding: 10px;
                margin: 5px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QTextEdit {
                border: 1px solid #555;
                padding: 5px;
                border-radius: 3px;
            }
        """)
        self.query_input = QTextEdit()
        self.query_input.setPlaceholderText("Escriba su pregunta sobre la historia clínica...")
        self.response_output = QTextEdit()
        self.response_output.setReadOnly(True)
        self.send_button = QPushButton("Consultar a IA")
        self.save_report_button = QPushButton("Guardar Respuesta en Reporte")

        layout = QVBoxLayout()
        layout.addWidget(QLabel("Ingrese su pregunta:"))
        layout.addWidget(self.query_input)
        layout.addWidget(QLabel("Respuesta de la IA:"))
        layout.addWidget(self.response_output)
        layout.addWidget(self.send_button)
        layout.addWidget(self.save_report_button)
        self.setLayout(layout)

        self.send_button.clicked.connect(self.send_query_to_ia)
        self.save_report_button.clicked.connect(self.generate_word_report)

    def send_query_to_ia(self):
        query = self.query_input.toPlainText()
        if not query:
            self.response_output.setText("Por favor, ingrese una pregunta.")
            return
        try:
            patient_info = self.get_patient_info()
            
                    # Utilizar GPT para extraer información de la prescripción
            if os.path.exists(CONFIG_FILE):
                with open(CONFIG_FILE, "r") as config_file:
                    config_data = json.load(config_file)
                    model = apply_openai_key(config_data)
            prompt = f"""Eres un asistente médico y estadístico que analiza datos clínicos de pacientes.
            Basándote en la información proporcionada , proporciona una respuesta precisa, completa y basada en datos.
            
            Información de la data:
            {patient_info}
            Pregunta del usuario: {query}
            Respuesta:"""
            response = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "Eres analista de datos y revisa todos los datos del csv y da respuestas completas."},
                    {"role": "user", "content": prompt}
                ]
            )
            ai_response = response.choices[0].message['content']
            self.response_output.setText(ai_response)
        except Exception as e:
            self.response_output.setText(f"Error al procesar la consulta: {str(e)}")

    def get_patient_info(self):
        if self.excel_data is not None:
            # Convertir Excel a CSV y reducir el tamaño de los datos
            csv_data = self.convert_excel_to_csv(self.excel_data)
            
            # Organizar la información del paciente
            patient_info_dict = {}
            for _, row in csv_data.iterrows():
                patient_name = row['Patient']
                diagnosis = row['Diagnosis']
                service_connected = 'Con Cobertura' if row['ServiceConected'] else 'Sin Cobertura'
                
                if patient_name not in patient_info_dict:
                    patient_info_dict[patient_name] = []
                
                patient_info_dict[patient_name].append(f"{diagnosis} ({service_connected})")
            
            # Crear un resumen de la información del paciente
            patients_summary = ""
            for patient_name, diagnoses in patient_info_dict.items():
                patients_summary += f"Paciente: {patient_name}, Diagnósticos: {', '.join(diagnoses)}\n"
            
            return patients_summary
        return "No hay datos disponibles del paciente."

    def convert_excel_to_csv(self, excel_data):
        # Seleccionar solo las columnas necesarias
        columns_needed = ['Patient', 'Diagnosis', 'ServiceConected']
        filtered_data = excel_data[columns_needed]
        
        # Guardar como CSV temporal
        csv_path = "temp_data.csv"
        filtered_data.to_csv(csv_path, index=False)
        
        # Leer el CSV de nuevo para asegurarse de que está en el formato correcto
        csv_data = read_csv(csv_path)
       
        return csv_data

    def generate_word_report(self):
        response_text = self.response_output.toPlainText()
        if not response_text:
            self.response_output.setText("No hay respuesta para guardar en el reporte.")
            return
        
        try:
            doc = Document()
            doc.add_heading("Reporte de Análisis Clínico", 0)
            paragraph = doc.add_paragraph(response_text)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            run = paragraph.runs[0]
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            file_path, _ = QFileDialog.getSaveFileName(self, "Guardar Reporte", "", "Documentos de Word (*.docx)")
            if file_path:
                doc.save(file_path)
                self.response_output.setText(f"Reporte guardado en {file_path}.")
            else:
                self.response_output.setText("Guardado cancelado.")
        except Exception as e:
            self.response_output.setText(f"Error al generar el reporte: {str(e)}")


class IAInteractionDialog(QDialog):
    """Dialog for a short Q&A with the AI before generating the history."""

    ready = pyqtSignal(str)

    def __init__(self, history_text, context_text="", parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.comments')
        self.setWindowTitle("Interacción con IA")
        StyleHelper.style_dialog(self)
        self.history_text = history_text
        self.context_text = context_text
        self.resize(1200, 800)
        self.user_answers = []
        self.copy_btn = None
        first_message = context_text + ("\n" if context_text else "") + history_text
        self.messages = [
            {"role": "system", "content": IA_INTERACTION_PROMPT},
            {"role": "user", "content": first_message},
        ]

        layout = QVBoxLayout(self)

        # Scrollable chat area with message bubbles
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.chat_widget = QWidget()
        self.chat_layout = QVBoxLayout(self.chat_widget)
        self.chat_layout.addStretch()
        self.scroll.setWidget(self.chat_widget)
        layout.addWidget(self.scroll, 1)

        # Mensaje de bienvenida
        self.add_message(
            "\\u00a1Bienvenido! Estoy listo para ayudar a completar la historia cl\u00ednica.",
            False,
        )

        self.answer_edit = AutoAdjustingTextEdit()
        self.answer_edit.heightMin = 80
        # allow the field to grow without an arbitrary limit
        self.answer_edit.heightMax = 10000
        self.answer_edit.adjust_height_to_content()
        self.answer_edit.setPlaceholderText("Escribe tu respuesta aqu\u00ed...")
        self.answer_edit.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        StyleHelper.style_input(self.answer_edit)
        layout.addWidget(self.answer_edit)

        btn_row = QHBoxLayout()
        self.send_btn = QPushButton("Responder")
        width, _ = StyleHelper.get_screen_resolution()
        self.send_btn.setMinimumHeight(StyleHelper.adjust_size(50, width))
        StyleHelper.style_button(self.send_btn)
        btn_row.addWidget(self.send_btn)
        self.ok_btn = QPushButton("Generar Historia Cl\u00ednica")
        StyleHelper.style_button(self.ok_btn)
        self.ok_btn.hide()
        btn_row.addWidget(self.ok_btn)
        layout.addLayout(btn_row)

        self.send_btn.clicked.connect(self.send_answer)
        self.ok_btn.clicked.connect(self.finish_interaction)

        # Show initial history in chat bubbles
        if self.context_text:
            self.add_message(self.context_text, False)
        if self.history_text:
            self.add_message(self.history_text, True)

        self.ask_ai()

    def add_message(self, text, is_user):
        """Add a new bubble to the chat area."""
        label = QLabel(text)
        label.setWordWrap(True)
        label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        color = "#cce5ff" if is_user else "#e0e0e0"
        label.setStyleSheet(
            f"background-color:{color}; border-radius:10px; padding:8px;"
        )
        container = QWidget()
        row = QHBoxLayout(container)
        if is_user:
            row.addStretch()
            row.addWidget(label)
        else:
            row.addWidget(label)
            row.addStretch()
        self.chat_layout.insertWidget(self.chat_layout.count() - 1, container)
        self.scroll_to_bottom()
        label.container = container  # save container for later actions
        return label

    def scroll_to_bottom(self):
        bar = self.scroll.verticalScrollBar()
        bar.setValue(bar.maximum())

    def ask_ai(self):
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as f:
                cfg = json.load(f)
            model = apply_openai_key(cfg)
        else:
            openai.api_key = ""
            model = "gpt-3.5-turbo"

        if self.copy_btn:
            self.copy_btn.deleteLater()
            self.copy_btn = None

        if not openai.api_key:
            QMessageBox.warning(
                self,
                "Falta API Key",
                "Configura tu API Key en el men\u00fa de Configuraci\u00f3n para usar la IA.",
            )
            return

        self.current_ai_label = self.add_message("IA escribiendo...", False)
        self.send_btn.setEnabled(False)
        self.thread = AIStreamThread(self.messages, model)
        self.thread.chunk_received.connect(self.update_chunk)
        self.thread.finished.connect(self.ai_finished)
        self.thread.error.connect(lambda m: QMessageBox.warning(self, "Error", m))
        self.thread.start()

    def update_chunk(self, text):
        if self.current_ai_label.text().endswith("IA escribiendo..."):
            self.current_ai_label.setText("")
        self.current_ai_label.setText(self.current_ai_label.text() + text)
        self.scroll_to_bottom()

    def ai_finished(self, text):
        self.current_ai_label.setText(text)
        self.messages.append({"role": "assistant", "content": text})
        if "\u2705 Listo para generar historia cl\u00ednica" in text:
            self.ok_btn.show()
            self.ok_btn.click()
        else:
            self.send_btn.setEnabled(True)
            if "?" in text:
                self.show_copy_button()

    def send_answer(self):
        answer = self.answer_edit.toPlainText().strip()
        if not answer:
            return
        self.user_answers.append(answer)
        self.add_message(answer, True)
        self.answer_edit.clear()
        self.answer_edit.adjust_height_to_content()
        self.messages.append({"role": "user", "content": answer})
        self.ask_ai()

    def show_copy_button(self):
        """Display a temporary button to copy the AI question."""
        if self.copy_btn:
            self.copy_btn.deleteLater()
        self.copy_btn = QPushButton("\U0001F4CB Copiar preguntas")
        self.copy_btn.setMinimumSize(100, 36)
        self.copy_btn.setToolTip("Copiar preguntas de IA")
        self.copy_btn.setStyleSheet(
            """
            QPushButton {
                background-color: #0078D7;
                color: white;
                border-radius: 8px;
                padding: 6px 12px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #005fa3;
            }
            """
        )
        self.copy_btn.clicked.connect(self.copy_question)
        layout = self.current_ai_label.container.layout()
        layout.insertWidget(1, self.copy_btn)

    def copy_question(self):
        """Insert the AI questions into the user input field."""
        text = self.current_ai_label.text()
        questions = re.findall(r"[^\n?!.]*\?", text)
        if questions:
            self.answer_edit.setPlainText("\n".join(q.strip() for q in questions))
            self.answer_edit.adjust_height_to_content()
        if self.copy_btn:
            self.copy_btn.deleteLater()
            self.copy_btn = None

    def finish_interaction(self):
        conversation = "\n".join(self.user_answers)
        self.ready.emit(conversation)
        self.accept()



class VerificarCobertura(QMainWindow):
    def __init__(self):
        super().__init__()
        StyleHelper.set_window_icon(self, 'fa5s.hospital')
        self.setWindowTitle("Cobertura de Pacientes")
        self.setGeometry(100, 100, 800, 600)
        self.excel_data = None
        self.excel_path = ""
        self.setup_ui()
        self.load_last_excel()

    def setup_ui(self):
        # Configurar el estilo
        self.setStyleSheet("""
            QMainWindow {
                background-color: #2c3e50;
            }
            QPushButton {
                background-color: #4CAF50;
                color: #ecf0f1;
                border: none;
                padding: 10px;
                margin: 5px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QTextEdit, QComboBox {
                border: 1px solid #555;
                padding: 5px;
                border-radius: 3px;
            }
        """)

        # Crear widgets
        self.load_button = QPushButton("Cargar Excel")
        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        self.patient_combo.setInsertPolicy(QComboBox.NoInsert)
        self.query_input = QTextEdit()
        self.query_input.setPlaceholderText("Ingrese especialidad, medicamento o referimiento")
        self.query_input.setFixedHeight(100)
        self.response_output = QTextEdit()
        self.response_output.setReadOnly(True)
        self.diagnosis_output = QTextEdit()
        self.diagnosis_output.setReadOnly(True)
        self.send_button = QPushButton("Enviar a IA")
        self.generate_button = QPushButton("Generar Informe")
        self.analyze_button = QPushButton("Analizar con IA")

        # Crear layout
        layout = QVBoxLayout()
        layout.addWidget(self.load_button)
        layout.addWidget(QLabel("Seleccione o escriba el nombre del paciente:"))
        layout.addWidget(self.patient_combo)
        layout.addWidget(QLabel("Consulta:"))
        layout.addWidget(self.query_input)
        layout.addWidget(QLabel("Respuesta de IA:"))
        layout.addWidget(self.response_output)
        layout.addWidget(QLabel("Diagnósticos:"))
        layout.addWidget(self.diagnosis_output)

        button_layout = QHBoxLayout()
        button_layout.addWidget(self.send_button)
        button_layout.addWidget(self.generate_button)
        button_layout.addWidget(self.analyze_button)
        layout.addLayout(button_layout)

        # Crear widget central
        central_widget = QWidget()
        central_widget.setLayout(layout)
        self.setCentralWidget(central_widget)

        # Conectar señales
        self.load_button.clicked.connect(self.load_excel)
        self.send_button.clicked.connect(self.send_query)
        self.generate_button.clicked.connect(self.generate_report)
        self.patient_combo.currentTextChanged.connect(self.update_patient_info)
        self.analyze_button.clicked.connect(self.open_ia_analysis_window)

    def load_last_excel(self):
        if os.path.exists('config.json'):
            with open('config.json', 'r') as f:
                config = json.load(f)
                last_excel = config.get('last_excel', '')
                if os.path.exists(last_excel):
                    self.load_excel(last_excel)

    def save_excel_path(self, path):
        config = {'last_excel': path}
        with open('config.json', 'w') as f:
            json.dump(config, f)

    def load_excel(self, file_name=None):
        if not file_name:
            file_name, _ = QFileDialog.getOpenFileName(self, "Abrir archivo Excel", "", "Excel Files (*.xlsx *.xls)")
        if file_name:
            try:
                # Usar read_excel directamente sin la referencia 'pd'
                self.excel_data = read_excel(file_name)
                self.excel_path = file_name
                self.save_excel_path(file_name)
                self.response_output.setText("Archivo Excel cargado exitosamente.")
                self.update_patient_list()
            except Exception as e:
                self.response_output.setText(f"Error al cargar el archivo: {str(e)}")


    def update_patient_list(self):
        if self.excel_data is not None:
            patients = self.excel_data['Patient'].unique()
            self.patient_combo.clear()
            self.patient_combo.addItems(patients)
            completer = QCompleter(patients)
            completer.setCaseSensitivity(Qt.CaseInsensitive)
            self.patient_combo.setCompleter(completer)

    def update_patient_info(self, patient_name):
        if self.excel_data is not None and patient_name:
            patient_data = self.excel_data[self.excel_data['Patient'] == patient_name]
            if not patient_data.empty:
                diagnoses = patient_data[['Diagnosis', 'ServiceConected', 'Percent']].values.tolist()
                diagnosis_text = "Diagnósticos:\n"
                for diagnosis, service_connected, percent in diagnoses:
                    status = "Service Connected" if service_connected else "No Service Connected"
                    diagnosis_text += f"{diagnosis}: {status} ({percent}%)\n"
                self.diagnosis_output.setText(diagnosis_text)
            else:
                self.diagnosis_output.setText("No se encontró información para este paciente.")

    def send_query(self):
        if self.excel_data is None:
            self.response_output.setText("Por favor, cargue un archivo Excel primero.")
            return

        query = self.query_input.toPlainText()
        if not query:
            self.response_output.setText("Por favor, ingrese su consulta.")
            return
        
        try:
            selected_patient = self.patient_combo.currentText()
            patient_data = self.excel_data[self.excel_data['Patient'] == selected_patient]
            patient_info = self.get_patient_info(patient_data)
            
                    # Utilizar GPT para extraer información de la prescripción
            if os.path.exists(CONFIG_FILE):
                with open(CONFIG_FILE, "r") as config_file:
                    config_data = json.load(config_file)
                    model = apply_openai_key(config_data)
            prompt = f"""Eres un asistente médico especializado en la cobertura de pacientes.
            Basado en los datos proporcionados del paciente, responde la siguiente consulta:
            
            Información del paciente:
            {patient_info}
            Pregunta del usuario: {query}
            Respuesta:"""
            
            response = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "Eres un asistente médico."},
                    {"role": "user", "content": prompt}
                ]
            )
            ai_response = response.choices[0].message['content']
            self.response_output.setText(ai_response)
        except Exception as e:
            self.response_output.setText(f"Error al enviar la consulta a IA: {str(e)}")

    def get_patient_info(self, patient_data):
        if patient_data is not None and not patient_data.empty:
            patient_summary = ""
            for _, row in patient_data.iterrows():
                patient_summary += f"Paciente: {row['Patient']}, Diagnóstico: {row['Diagnosis']}, Service Connected: {row['ServiceConected']}, Porcentaje: {row['Percent']}%\n"
            return patient_summary
        return "No hay datos disponibles del paciente."

    def generate_report(self):
        if self.excel_data is None:
            self.response_output.setText("Por favor, cargue un archivo Excel primero.")
            return
        try:
            doc = Document()
            doc.add_heading("Reporte de Diagnósticos", 0)
            selected_patient = self.patient_combo.currentText()
            patient_data = self.excel_data[self.excel_data['Patient'] == selected_patient]
            if not patient_data.empty:
                for _, row in patient_data.iterrows():
                    doc.add_heading(row['Patient'], level=1)
                    doc.add_paragraph(f"Diagnóstico: {row['Diagnosis']}")
                    status = "Sí" if row['ServiceConected'] else "No"
                    doc.add_paragraph(f"Service Connected: {status}")
                    doc.add_paragraph(f"Porcentaje: {row['Percent']}%")
                file_path, _ = QFileDialog.getSaveFileName(self, "Guardar Informe", "", "Documentos de Word (*.docx)")
                if file_path:
                    doc.save(file_path)
                    self.response_output.setText(f"Informe guardado en {file_path}.")
            else:
                self.response_output.setText("No se encontró información para este paciente.")
        except Exception as e:
            self.response_output.setText(f"Error al generar el informe: {str(e)}")

    def open_ia_analysis_window(self):
        self.ia_analysis_window = IAAnalysisWindow(self, excel_data=self.excel_data)
        self.ia_analysis_window.exec_()




# ------------------------------------------------------------
#   DIÁLOGOS ADICIONALES: CITAS, FACTURAS, TELEMEDICINA Y ESTADÍSTICAS
# ------------------------------------------------------------

class AppointmentDialog(QDialog):
    def __init__(self, parent=None, patient=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.calendar-check')
        self.setWindowTitle("Programación de Citas")
        self._initial_patient = patient
        self.resize(700, 400)
        self.setWindowFlags(self.windowFlags() | Qt.WindowMinimizeButtonHint | Qt.WindowMaximizeButtonHint)
        self.setWindowState(Qt.WindowMaximized)

        layout = QVBoxLayout(self)
        form = QHBoxLayout()
        form.addWidget(QLabel("Paciente:"))
        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        StyleHelper.style_input(self.patient_combo)
        form.addWidget(self.patient_combo)
        self.completer = QCompleter()
        self.completer.setFilterMode(Qt.MatchContains)
        self.completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.patient_combo.setCompleter(self.completer)
        self.load_patients()
        if self._initial_patient:
            self.patient_combo.setCurrentText(self._initial_patient.nombre)
        form.addWidget(QLabel("Fecha:"))
        self.date_edit = QDateTimeEdit(datetime.now())
        self.date_edit.setCalendarPopup(True)
        StyleHelper.style_input(self.date_edit)
        form.addWidget(self.date_edit)
        form.addWidget(QLabel("Notas:"))
        self.notes_edit = QLineEdit()
        StyleHelper.style_input(self.notes_edit)
        form.addWidget(self.notes_edit)
        add_btn = QPushButton("Agregar")
        StyleHelper.style_button(add_btn)
        form.addWidget(add_btn)
        layout.addLayout(form)

        lang = parent.current_language if parent else 'es'
        t = TRANSLATIONS[lang]
        self.today_label = QLabel(t['today_appointments'])
        StyleHelper.style_label(self.today_label)
        layout.addWidget(self.today_label)
        self.today_table = QTableWidget()
        self.today_table.setColumnCount(3)
        self.today_table.setHorizontalHeaderLabels(['Hora', 'Paciente', 'Notas'])
        StyleHelper.style_table(self.today_table)
        self.today_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.today_table)

        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(["ID", "Paciente", "Fecha", "Notas"])
        StyleHelper.style_table(self.table)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.table)
        del_btn = QPushButton("Eliminar")
        StyleHelper.style_button(del_btn)
        layout.addWidget(del_btn)

        add_btn.clicked.connect(self.add_appointment)
        del_btn.clicked.connect(self.delete_appointment)
        self.load_appointments()
        self.load_today_appointments()

    def load_patients(self):
        session = Session()
        pacientes = session.query(Paciente).order_by(Paciente.nombre).all()
        names = [p.nombre for p in pacientes]
        self.patient_combo.clear()
        self.patient_combo.addItems(names)
        self.completer.setModel(QStringListModel(names))
        session.close()


    def load_appointments(self):
        session = Session()
        citas = session.query(Cita).order_by(Cita.fecha.desc()).all()
        self.table.setRowCount(len(citas))
        for row, c in enumerate(citas):
            name = c.paciente.nombre if c.paciente else str(c.paciente_id)
            self.table.setItem(row, 0, QTableWidgetItem(str(c.id)))
            self.table.setItem(row, 1, QTableWidgetItem(name))
            self.table.setItem(row, 2, QTableWidgetItem(c.fecha.strftime("%Y-%m-%d %H:%M")))
            self.table.setItem(row, 3, QTableWidgetItem(c.notas or ""))
        session.close()

    def load_today_appointments(self):
        session = Session()
        today = datetime.now().date()
        citas = session.query(Cita).filter(func.date(Cita.fecha) == today).order_by(Cita.fecha).all()
        self.today_table.setRowCount(len(citas))
        for row, c in enumerate(citas):
            name = c.paciente.nombre if c.paciente else str(c.paciente_id)
            self.today_table.setItem(row, 0, QTableWidgetItem(c.fecha.strftime('%H:%M')))
            self.today_table.setItem(row, 1, QTableWidgetItem(name))
            self.today_table.setItem(row, 2, QTableWidgetItem(c.notas or ''))
        session.close()

    def add_appointment(self):
        name = self.patient_combo.currentText()
        session = Session()
        pac = session.query(Paciente).filter_by(nombre=name).first()
        if not pac:
            session.close()
            QMessageBox.warning(self, "Error", "Paciente no encontrado")
            return
        pid = pac.id
        fecha = self.date_edit.dateTime().toPyDateTime()
        notas = self.notes_edit.text().strip()
        cita = Cita(paciente_id=pid, fecha=fecha, notas=notas)
        session.add(cita)
        session.commit()
        session.close()
        self.load_appointments()
        self.load_today_appointments()

    def delete_appointment(self):
        row = self.table.currentRow()
        if row < 0:
            return
        cita_id = int(self.table.item(row, 0).text())
        session = Session()
        session.query(Cita).filter_by(id=cita_id).delete()
        session.commit()
        session.close()
        self.load_appointments()
        self.load_today_appointments()


class BillingDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.file-invoice-dollar')
        self.setWindowTitle("Facturación")
        self.resize(700, 400)
        self.setWindowFlags(self.windowFlags() | Qt.WindowMinimizeButtonHint | Qt.WindowMaximizeButtonHint)
        self.setWindowState(Qt.WindowMaximized)

        lang = parent.current_language if parent else 'es'
        t = TRANSLATIONS[lang]

        layout = QVBoxLayout(self)
        form = QHBoxLayout()
        form.addWidget(QLabel("Paciente:"))
        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        StyleHelper.style_input(self.patient_combo)
        form.addWidget(self.patient_combo)
        self.completer = QCompleter()
        self.completer.setFilterMode(Qt.MatchContains)
        self.completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.patient_combo.setCompleter(self.completer)
        self.load_patients()
        layout.addLayout(form)

        service_form = QHBoxLayout()
        service_form.addWidget(QLabel("Servicio:"))
        self.service_combo = QComboBox()
        service_form.addWidget(self.service_combo)
        self.qty_spin = QSpinBox()
        self.qty_spin.setValue(1)
        service_form.addWidget(self.qty_spin)
        self.price_edit = QLineEdit()
        StyleHelper.style_input(self.price_edit)
        self.price_edit.setReadOnly(True)
        service_form.addWidget(self.price_edit)
        add_line_btn = QPushButton("Agregar")
        StyleHelper.style_button(add_line_btn)
        service_form.addWidget(add_line_btn)
        layout.addLayout(service_form)

        self.lines_table = QTableWidget()
        self.lines_table.setColumnCount(4)
        self.lines_table.setHorizontalHeaderLabels(["Servicio", "Cant.", "Precio", "Subtotal"])
        self.lines_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        StyleHelper.style_table(self.lines_table)
        layout.addWidget(self.lines_table)

        total_layout = QHBoxLayout()
        total_layout.addWidget(QLabel("Total:"))
        self.total_label = QLabel("0")
        total_layout.addWidget(self.total_label)
        save_btn = QPushButton("Guardar Factura")
        StyleHelper.style_button(save_btn)
        total_layout.addWidget(save_btn)
        layout.addLayout(total_layout)

        self.table = QTableWidget()
        self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels(["ID", "Paciente", "Fecha", "Total", "Descripción", "Pagado"])
        StyleHelper.style_table(self.table)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.table)
        btn_row = QHBoxLayout()
        del_btn = QPushButton(t.get('delete', 'Eliminar'))
        StyleHelper.style_button(del_btn)
        print_btn = QPushButton(t['print_invoice'])
        StyleHelper.style_button(print_btn)
        pay_btn = QPushButton("Registrar Pago")
        StyleHelper.style_button(pay_btn)
        btn_row.addWidget(print_btn)
        btn_row.addWidget(pay_btn)
        btn_row.addWidget(del_btn)
        layout.addLayout(btn_row)

        add_line_btn.clicked.connect(self.add_line)
        save_btn.clicked.connect(self.add_invoice)
        del_btn.clicked.connect(self.delete_invoice)
        print_btn.clicked.connect(self.print_invoice)
        pay_btn.clicked.connect(self.add_payment)
        self.load_services()
        self.service_combo.currentTextChanged.connect(self.update_price)
        self.patient_combo.currentTextChanged.connect(self.update_price)
        self.load_invoices()

    def load_patients(self):
        session = Session()
        pacientes = session.query(Paciente).order_by(Paciente.nombre).all()
        names = [p.nombre for p in pacientes]
        self.patient_combo.clear()
        self.patient_combo.addItems(names)
        self.completer.setModel(QStringListModel(names))
        session.close()

    def load_invoices(self):
        session = Session()
        facturas = session.query(Factura).order_by(Factura.fecha.desc()).all()
        self.table.setRowCount(len(facturas))
        for row, f in enumerate(facturas):
            name = f.paciente.nombre if f.paciente else str(f.paciente_id)
            pagado = "Sí" if f.pagado else "No"
            self.table.setItem(row, 0, QTableWidgetItem(str(f.id)))
            self.table.setItem(row, 1, QTableWidgetItem(name))
            self.table.setItem(row, 2, QTableWidgetItem(f.fecha.strftime("%Y-%m-%d")))
            self.table.setItem(row, 3, QTableWidgetItem(str(f.total)))
            self.table.setItem(row, 4, QTableWidgetItem(f.descripcion or ""))
            self.table.setItem(row, 5, QTableWidgetItem(pagado))
        session.close()

    def add_invoice(self):
        name = self.patient_combo.currentText()
        session = Session()
        pac = session.query(Paciente).filter_by(nombre=name).first()
        if not pac:
            session.close()
            QMessageBox.warning(self, "Error", "Paciente no encontrado")
            return
        pid = pac.id
        total = float(self.total_label.text()) if self.total_label.text() else 0
        descripcion = "; ".join(
            [self.lines_table.item(r, 0).text() for r in range(self.lines_table.rowCount())]
        )
        factura = Factura(paciente_id=pid, fecha=datetime.now(), total=total, monto=total, descripcion=descripcion)
        session.add(factura)
        session.commit()
        for r in range(self.lines_table.rowCount()):
            svc_name = self.lines_table.item(r, 0).text()
            qty = int(self.lines_table.item(r, 1).text())
            price = float(self.lines_table.item(r, 2).text())
            svc = session.query(Servicio).filter_by(nombre=svc_name).first()
            session.add(
                FacturaDetalle(
                    factura_id=factura.id,
                    servicio_id=svc.id if svc else None,
                    cantidad=qty,
                    precio_unitario=price,
                )
            )
        session.commit()
        session.close()
        self.lines_table.setRowCount(0)
        self.total_label.setText("0")
        self.load_invoices()

    def delete_invoice(self):
        row = self.table.currentRow()
        if row < 0:
            return
        fac_id = int(self.table.item(row, 0).text())
        session = Session()
        session.query(Factura).filter_by(id=fac_id).delete()
        session.commit()
        session.close()
        self.load_invoices()

    def add_payment(self):
        row = self.table.currentRow()
        if row < 0:
            return
        fac_id = int(self.table.item(row, 0).text())
        amount, ok = QInputDialog.getDouble(self, "Monto", "Monto del pago:", 0, 0, 1000000, 2)
        if not ok:
            return
        method, ok = QInputDialog.getText(self, "Método", "Método de pago:")
        if not ok:
            return
        session = Session()
        pago = Pago(factura_id=fac_id, fecha=datetime.now(), monto=amount, metodo=method)
        session.add(pago)
        total_pagado = session.query(func.coalesce(func.sum(Pago.monto), 0)).filter_by(factura_id=fac_id).scalar() + amount
        factura = session.get(Factura, fac_id)
        if factura and total_pagado >= factura.total:
            factura.pagado = True
        session.commit()
        session.close()
        QMessageBox.information(self, "Pago", "Pago registrado")
        self.load_invoices()

    def print_invoice(self):
        row = self.table.currentRow()
        if row < 0:
            return
        inv_id = int(self.table.item(row, 0).text())
        session = Session()
        inv = session.get(Factura, inv_id)
        if not inv:
            session.close()
            return
        patient_obj = inv.paciente
        patient = patient_obj.nombre if patient_obj else str(inv.paciente_id)
        doc_id = patient_obj.documento_id if patient_obj else ""
        direccion = patient_obj.direccion if patient_obj else ""
        cfg = ConfigReceta(); cfg.load_config()
        clinic = cfg.get_clinic_header() or ""
        footer_addr = cfg.get_footer_text()

        detalle_html = ""
        subtotal = 0
        for d in inv.detalles:
            name = d.servicio.nombre if d.servicio else ""
            sub = d.cantidad * d.precio_unitario
            detalle_html += (
                f"<tr><td align='right'>{d.cantidad:.2f}</td><td>UNIDAD</td>"
                f"<td>{name}</td><td align='right'>{d.precio_unitario:.2f}</td>"
                f"<td align='right'>0.00</td></tr>"
            )
            subtotal += sub
        igv = subtotal * 0.18
        total = subtotal + igv
        fecha = inv.fecha.strftime('%Y-%m-%d')
        letras = numero_a_letras(total)

        html = f"""
        <html><head><style>
        body {{font-family:'Arial'; font-size:11pt;}}
        table {{border-collapse: collapse; width: 100%;}}
        th, td {{border: 1px solid #000; padding:4px;}}
        .no-border td {{border:none;}}
        </style></head><body>
        <table class='no-border'>
        <tr>
            <td><b>{clinic}</b></td>
            <td align='right'><b>FACTURA ELECTRÓNICA</b><br>Factura N.º: E001-{inv.id}</td>
        </tr></table><br>
        <table>
        <tr><td>Fecha de emisión</td><td>{fecha}</td></tr>
        <tr><td>Fecha de vencimiento</td><td>{fecha}</td></tr>
        <tr><td>Nombre del cliente</td><td>{patient}</td></tr>
        <tr><td>Cédula/ID</td><td>{doc_id}</td></tr>
        <tr><td>Dirección del cliente</td><td>{direccion}</td></tr>
        <tr><td>Tipo de moneda</td><td>DOP</td></tr>
        </table><br>
        <table>
        <tr><th>Cantidad</th><th>Unidad Medida</th><th>Descripción</th><th>Valor Unitario</th><th>ICBPER</th></tr>
        {detalle_html}
        </table><br>
        <table style='width:40%; float:right;'>
        <tr><td>Sub Total Venta</td><td align='right'>{subtotal:.2f}</td></tr>
        <tr><td>IGV (18%)</td><td align='right'>{igv:.2f}</td></tr>
        <tr><td>Otros</td><td align='right'>0.00</td></tr>
        <tr><td><b>Importe Total</b></td><td align='right'><b>{total:.2f}</b></td></tr>
        </table><br style='clear:both;'>
        <p>SON: {letras} Y {int(round((total-int(total))*100)):02d}/100</p>
        <p style='text-align:center'>Esta factura electrónica tiene validez legal conforme a las normativas aplicables. Para verificar su autenticidad, comuníquese con el emisor.</p>
        <p style='text-align:center'>{footer_addr}</p>
        </body></html>
        """

        doc = QTextDocument()
        doc.setHtml(html)
        filename, _ = QFileDialog.getSaveFileName(
            self,
            "Guardar Factura",
            f"Factura_{patient.replace(' ', '')}_{fecha}.pdf",
            "PDF Files (*.pdf)"
        )
        if not filename:
            session.close()
            return
        printer = QPrinter(QPrinter.HighResolution)
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(filename)
        doc.print_(printer)
        session.close()
        QMessageBox.information(self, "PDF", "Factura guardada correctamente")
        QDesktopServices.openUrl(QUrl.fromLocalFile(filename))

    def load_services(self):
        session = Session()
        services = session.query(Servicio).order_by(Servicio.nombre).all()
        self.service_combo.clear()
        self.service_combo.addItems([s.nombre for s in services])
        self.service_map = {s.nombre: s for s in services}
        session.close()
        self.update_price()

    def update_price(self):
        name = self.service_combo.currentText()
        service = getattr(self, 'service_map', {}).get(name)
        if not service:
            self.price_edit.clear()
            return
        price = service.precio_base
        session = Session()
        pac = session.query(Paciente).filter_by(nombre=self.patient_combo.currentText()).first()
        if pac and pac.ars:
            tarifa = (
                session.query(TarifaARS)
                .join(InsuranceName)
                .filter(TarifaARS.servicio_id == service.id)
                .filter(InsuranceName.nombre == pac.ars)
                .first()
            )
            if tarifa:
                price = tarifa.precio_ars
        session.close()
        self.price_edit.setText(str(price))

    def add_line(self):
        try:
            price = float(self.price_edit.text())
        except ValueError:
            price = 0
        qty = self.qty_spin.value()
        subtotal = price * qty
        row = self.lines_table.rowCount()
        self.lines_table.insertRow(row)
        self.lines_table.setItem(row, 0, QTableWidgetItem(self.service_combo.currentText()))
        self.lines_table.setItem(row, 1, QTableWidgetItem(str(qty)))
        self.lines_table.setItem(row, 2, QTableWidgetItem(f"{price:.2f}"))
        self.lines_table.setItem(row, 3, QTableWidgetItem(f"{subtotal:.2f}"))
        self.update_total()

    def update_total(self):
        total = 0
        for r in range(self.lines_table.rowCount()):
            total += float(self.lines_table.item(r, 3).text())
        self.total_label.setText(f"{total:.2f}")


class TelemedicineDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.video')
        self.setWindowTitle("Telemedicina")
        self.resize(700, 300)
        self.setWindowFlags(self.windowFlags() | Qt.WindowMinimizeButtonHint | Qt.WindowMaximizeButtonHint)
        self.setWindowState(Qt.WindowMaximized)

        layout = QVBoxLayout(self)
        form = QHBoxLayout()
        form.addWidget(QLabel("Paciente:"))
        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        StyleHelper.style_input(self.patient_combo)
        form.addWidget(self.patient_combo)
        self.completer = QCompleter()
        self.completer.setFilterMode(Qt.MatchContains)
        self.completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.patient_combo.setCompleter(self.completer)
        self.load_patients()
        form.addWidget(QLabel("Enlace:"))
        self.link_edit = QLineEdit()
        StyleHelper.style_input(self.link_edit)
        form.addWidget(self.link_edit)
        add_btn = QPushButton("Guardar")
        StyleHelper.style_button(add_btn)
        form.addWidget(add_btn)
        layout.addLayout(form)

        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(["ID", "Paciente", "Fecha", "Enlace"])
        StyleHelper.style_table(self.table)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.table)

        add_btn.clicked.connect(self.add_session)
        self.table.itemDoubleClicked.connect(self.open_link)
        self.load_sessions()

    def load_patients(self):
        session = Session()
        pacientes = session.query(Paciente).order_by(Paciente.nombre).all()
        names = [p.nombre for p in pacientes]
        self.patient_combo.clear()
        self.patient_combo.addItems(names)
        self.completer.setModel(QStringListModel(names))
        session.close()

    def load_sessions(self):
        session = Session()
        sesiones = session.query(TelemedSession).order_by(TelemedSession.fecha.desc()).all()
        self.table.setRowCount(len(sesiones))
        for row, s in enumerate(sesiones):
            name = s.paciente.nombre if s.paciente else str(s.paciente_id)
            self.table.setItem(row, 0, QTableWidgetItem(str(s.id)))
            self.table.setItem(row, 1, QTableWidgetItem(name))
            self.table.setItem(row, 2, QTableWidgetItem(s.fecha.strftime("%Y-%m-%d %H:%M")))
            self.table.setItem(row, 3, QTableWidgetItem(s.enlace))
        session.close()

    def add_session(self):
        name = self.patient_combo.currentText()
        session = Session()
        pac = session.query(Paciente).filter_by(nombre=name).first()
        if not pac:
            session.close()
            QMessageBox.warning(self, "Error", "Paciente no encontrado")
            return
        pid = pac.id
        enlace = self.link_edit.text().strip()
        ses = TelemedSession(paciente_id=pid, fecha=datetime.now(), enlace=enlace)
        session.add(ses)
        session.commit()
        session.close()
        self.load_sessions()

    def open_link(self, item):
        row = item.row()
        enlace = self.table.item(row, 3).text()
        import webbrowser
        webbrowser.open(enlace)


class DocumentDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.file-alt')
        self.setWindowTitle("Documentos")
        self.resize(700, 400)
        self.setWindowFlags(self.windowFlags() | Qt.WindowMinimizeButtonHint | Qt.WindowMaximizeButtonHint)
        self.setWindowState(Qt.WindowMaximized)

        layout = QVBoxLayout(self)
        form = QHBoxLayout()
        form.addWidget(QLabel("Paciente:"))
        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        StyleHelper.style_input(self.patient_combo)
        form.addWidget(self.patient_combo)
        self.completer = QCompleter()
        self.completer.setFilterMode(Qt.MatchContains)
        self.completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.patient_combo.setCompleter(self.completer)
        self.load_patients()
        self.patient_combo.currentTextChanged.connect(self.load_histories)
        form.addWidget(QLabel("Historia:"))
        self.history_combo = QComboBox()
        StyleHelper.style_input(self.history_combo)
        form.addWidget(self.history_combo)
        form.addWidget(QLabel("Archivo:"))
        self.file_edit = QLineEdit()
        StyleHelper.style_input(self.file_edit)
        browse_btn = QPushButton("...")
        StyleHelper.style_button(browse_btn)
        browse_btn.clicked.connect(self.browse_file)
        form.addWidget(self.file_edit)
        form.addWidget(browse_btn)
        form.addWidget(QLabel("Descripción:"))
        self.desc_edit = QLineEdit()
        StyleHelper.style_input(self.desc_edit)
        form.addWidget(self.desc_edit)
        add_btn = QPushButton("Agregar")
        StyleHelper.style_button(add_btn)
        form.addWidget(add_btn)
        layout.addLayout(form)

        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["ID", "Paciente", "Historia", "Archivo", "Descripción"])
        StyleHelper.style_table(self.table)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.table)
        del_btn = QPushButton("Eliminar")
        StyleHelper.style_button(del_btn)
        open_btn = QPushButton("Abrir")
        StyleHelper.style_button(open_btn)
        btns = QHBoxLayout()
        btns.addWidget(open_btn)
        btns.addWidget(del_btn)
        layout.addLayout(btns)

        add_btn.clicked.connect(self.add_document)
        del_btn.clicked.connect(self.delete_document)
        open_btn.clicked.connect(self.open_document)
        self.load_documents()

    def load_patients(self):
        session = Session()
        pacientes = session.query(Paciente).order_by(Paciente.nombre).all()
        names = [p.nombre for p in pacientes]
        self.patient_combo.clear()
        self.patient_combo.addItems(names)
        self.completer.setModel(QStringListModel(names))
        session.close()
        if names:
            self.load_histories(names[0])

    def load_histories(self, name):
        session = Session()
        pac = session.query(Paciente).filter_by(nombre=name).first()
        self.history_combo.clear()
        if pac:
            historias = (
                session.query(HistoriaClinica)
                .filter_by(paciente_id=pac.id)
                .order_by(HistoriaClinica.fecha.desc())
                .all()
            )
            for h in historias:
                self.history_combo.addItem(h.fecha.strftime('%Y-%m-%d'), h.id)
        session.close()

    def browse_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Seleccionar Archivo", "", "Documentos (*.pdf *.png *.jpg *.jpeg *.docx)")
        if path:
            self.file_edit.setText(path)

    def load_documents(self):
        session = Session()
        docs = session.query(Documento).order_by(Documento.id.desc()).all()
        self.table.setRowCount(len(docs))
        for row, d in enumerate(docs):
            name = d.paciente.nombre if d.paciente else str(d.paciente_id)
            hist = d.historia.fecha.strftime('%Y-%m-%d') if d.historia else ''
            self.table.setItem(row, 0, QTableWidgetItem(str(d.id)))
            self.table.setItem(row, 1, QTableWidgetItem(name))
            self.table.setItem(row, 2, QTableWidgetItem(hist))
            self.table.setItem(row, 3, QTableWidgetItem(os.path.basename(d.ruta)))
            self.table.setItem(row, 4, QTableWidgetItem(d.descripcion or ""))
        session.close()

    def add_document(self):
        name = self.patient_combo.currentText()
        session = Session()
        pac = session.query(Paciente).filter_by(nombre=name).first()
        if not pac:
            session.close()
            QMessageBox.warning(self, "Error", "Paciente no encontrado")
            return
        pid = pac.id
        ruta = self.file_edit.text().strip()
        if not ruta:
            QMessageBox.warning(self, "Error", "Seleccione un archivo")
            return
        descripcion = self.desc_edit.text().strip()
        dest_dir = os.path.join(os.path.expanduser("~"), "Documents", "health_docs")
        os.makedirs(dest_dir, exist_ok=True)
        dest = os.path.join(dest_dir, os.path.basename(ruta))
        try:
            shutil.copy2(ruta, dest)
        except Exception:
            dest = ruta
        hid = self.history_combo.currentData()
        doc = Documento(paciente_id=pid, historia_id=hid, ruta=dest, descripcion=descripcion)
        session.add(doc)
        session.commit()
        session.close()
        self.load_documents()

    def delete_document(self):
        row = self.table.currentRow()
        if row < 0:
            return
        doc_id = int(self.table.item(row, 0).text())
        session = Session()
        session.query(Documento).filter_by(id=doc_id).delete()
        session.commit()
        session.close()
        self.load_documents()

    def open_document(self):
        row = self.table.currentRow()
        if row < 0:
            return
        path = self.table.item(row, 3).text()
        dest_dir = os.path.join(os.path.expanduser("~"), "Documents", "health_docs")
        full = os.path.join(dest_dir, path)
        if not os.path.exists(full):
            full = path
        try:
            os.startfile(full)
        except Exception:
            QMessageBox.warning(self, "Error", "No se pudo abrir el archivo")


class InsuranceManagerDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.id-card')
        self.setWindowTitle("Seguros Médicos")
        self.resize(600, 400)

        layout = QVBoxLayout(self)

        lang = parent.current_language if parent else 'es'
        t = TRANSLATIONS[lang]

        form = QHBoxLayout()
        self.name_edit = QLineEdit()
        StyleHelper.style_input(self.name_edit)
        add_btn = QPushButton("Agregar")
        StyleHelper.style_button(add_btn)
        form.addWidget(self.name_edit)
        form.addWidget(add_btn)
        layout.addLayout(form)

        self.table = QTableWidget()
        self.table.setColumnCount(2)
        self.table.setHorizontalHeaderLabels(["ID", "Nombre"])
        StyleHelper.style_table(self.table)
        layout.addWidget(self.table)

        del_btn = QPushButton("Eliminar")
        StyleHelper.style_button(del_btn)
        layout.addWidget(del_btn)

        stats_btn = QPushButton(t.get('insurance_stats', 'Estadísticas'))
        StyleHelper.style_button(stats_btn)
        layout.addWidget(stats_btn)

        add_btn.clicked.connect(self.add_name)
        del_btn.clicked.connect(self.delete_name)
        stats_btn.clicked.connect(self.show_stats)

        self.load_names()

    def load_names(self):
        session = Session()
        ars = session.query(InsuranceName).order_by(InsuranceName.nombre).all()
        self.table.setRowCount(len(ars))
        for row, a in enumerate(ars):
            self.table.setItem(row, 0, QTableWidgetItem(str(a.id)))
            self.table.setItem(row, 1, QTableWidgetItem(a.nombre))
        session.close()

    def add_name(self):
        name = self.name_edit.text().strip()
        if not name:
            return
        session = Session()
        if not session.query(InsuranceName).filter_by(nombre=name).first():
            session.add(InsuranceName(nombre=name))
            session.commit()
        session.close()
        self.name_edit.clear()
        self.load_names()

    def delete_name(self):
        row = self.table.currentRow()
        if row < 0:
            return
        id_ = int(self.table.item(row, 0).text())
        session = Session()
        session.query(InsuranceName).filter_by(id=id_).delete()
        session.commit()
        session.close()
        self.load_names()

    def show_stats(self):
        dialog = InsuranceStatsDialog(self)
        dialog.exec_()


class InsuranceStatsDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.chart-pie')
        lang = getattr(parent, 'current_language', 'es')
        t = TRANSLATIONS[lang]
        self.setWindowTitle(t.get('insurance_stats', 'Pacientes por ARS'))

        layout = QVBoxLayout(self)
        self.fig = plt.Figure(figsize=(6, 4))
        self.ax = self.fig.add_subplot(111)
        self.canvas = FigureCanvas(self.fig)
        layout.addWidget(self.canvas)

        close_btn = QPushButton('Cerrar')
        StyleHelper.style_button(close_btn)
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn, alignment=Qt.AlignRight)

        self.load_stats()

    def load_stats(self):
        session = Session()
        data = (
            session.query(Paciente.ars, func.count(Paciente.id))
            .group_by(Paciente.ars)
            .all()
        )
        session.close()
        names = [d[0] or 'Sin ARS' for d in data]
        values = [d[1] for d in data]
        self.ax.clear()
        bars = self.ax.bar(names, values, color='#2980b9')
        self.ax.bar_label(bars, fmt='%d')
        self.ax.set_ylabel('Pacientes')
        self.ax.set_title(self.windowTitle())
        self.canvas.draw()
        if signs:
            risk = self.parent().risk_level(signs[-1])
            self.risk_label.setText(f"Riesgo cardiovascular: {risk}")
        else:
            self.risk_label.setText("")

class ServiceManagerDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.toolbox')
        self.setWindowTitle("Servicios")
        self.resize(700, 400)

        layout = QVBoxLayout(self)
        form = QHBoxLayout()
        self.name_edit = QLineEdit()
        StyleHelper.style_input(self.name_edit)
        self.type_edit = QLineEdit()
        StyleHelper.style_input(self.type_edit)
        self.price_edit = QLineEdit()
        StyleHelper.style_input(self.price_edit)
        add_btn = QPushButton("Agregar")
        StyleHelper.style_button(add_btn)
        form.addWidget(self.name_edit)
        form.addWidget(self.type_edit)
        form.addWidget(self.price_edit)
        form.addWidget(add_btn)
        layout.addLayout(form)

        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(["ID", "Nombre", "Tipo", "Precio"])
        StyleHelper.style_table(self.table)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.table)

        del_btn = QPushButton("Eliminar")
        StyleHelper.style_button(del_btn)
        layout.addWidget(del_btn)

        add_btn.clicked.connect(self.add_service)
        del_btn.clicked.connect(self.delete_service)

        self.load_services()

    def load_services(self):
        session = Session()
        services = session.query(Servicio).order_by(Servicio.nombre).all()
        self.table.setRowCount(len(services))
        for row, s in enumerate(services):
            self.table.setItem(row, 0, QTableWidgetItem(str(s.id)))
            self.table.setItem(row, 1, QTableWidgetItem(s.nombre))
            self.table.setItem(row, 2, QTableWidgetItem(s.tipo or ""))
            self.table.setItem(row, 3, QTableWidgetItem(str(s.precio_base or 0)))
        session.close()

    def add_service(self):
        name = self.name_edit.text().strip()
        if not name:
            return
        try:
            price = float(self.price_edit.text())
        except ValueError:
            price = 0
        tipo = self.type_edit.text().strip()
        session = Session()
        if not session.query(Servicio).filter_by(nombre=name).first():
            session.add(Servicio(nombre=name, tipo=tipo, precio_base=price))
            session.commit()
        session.close()
        self.name_edit.clear()
        self.type_edit.clear()
        self.price_edit.clear()
        self.load_services()

    def delete_service(self):
        row = self.table.currentRow()
        if row < 0:
            return
        sid = int(self.table.item(row, 0).text())
        session = Session()
        session.query(Servicio).filter_by(id=sid).delete()
        session.commit()
        session.close()
        self.load_services()



class DataAnalysisDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.chart-bar')
        lang = parent.current_language if parent else 'es'
        t = TRANSLATIONS[lang]
        self.setWindowTitle(t['stats_dashboard'])
        layout = QVBoxLayout(self)

        filter_layout = QHBoxLayout()
        filter_layout.addWidget(QLabel(t['gender']))
        self.gender_combo = QComboBox()
        self.gender_combo.addItems([t['all'], 'M', 'F'])
        filter_layout.addWidget(self.gender_combo)
        filter_layout.addWidget(QLabel(t['diagnosis_filter']))
        self.diagnosis_edit = QLineEdit()
        filter_layout.addWidget(self.diagnosis_edit)
        self.apply_btn = QPushButton(t['apply_filters'])
        StyleHelper.style_button(self.apply_btn)
        filter_layout.addWidget(self.apply_btn)
        layout.addLayout(filter_layout)

        self.summary_label = QLabel()
        layout.addWidget(self.summary_label)

        self.fig = plt.Figure(figsize=(6,3))
        self.ax = self.fig.add_subplot(111)
        self.canvas = FigureCanvas(self.fig)
        layout.addWidget(self.canvas)

        self.gender_fig = plt.Figure(figsize=(6,3))
        self.gender_ax = self.gender_fig.add_subplot(111)
        self.gender_canvas = FigureCanvas(self.gender_fig)
        layout.addWidget(self.gender_canvas)

        self.age_fig = plt.Figure(figsize=(6,3))
        self.age_ax = self.age_fig.add_subplot(111)
        self.age_canvas = FigureCanvas(self.age_fig)
        layout.addWidget(self.age_canvas)

        close_btn = QPushButton('Cerrar')
        StyleHelper.style_button(close_btn)
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn, alignment=Qt.AlignRight)

        self.apply_btn.clicked.connect(self.update_stats)
        self.update_stats()

    def update_stats(self):
        lang = getattr(self.parent(), 'current_language', 'es')
        t = TRANSLATIONS[lang]
        session = Session()
        query = session.query(Paciente).join(HistoriaClinica, isouter=True).join(Diagnostico, isouter=True)
        gender = self.gender_combo.currentText()
        if gender in ('M', 'F'):
            query = query.filter(Paciente.sexo == gender)
        diag_filter = self.diagnosis_edit.text().strip()
        if diag_filter:
            like = f"%{diag_filter}%"
            query = query.filter(or_(Diagnostico.descripcion.ilike(like), Diagnostico.codigo_cie10.ilike(like)))
        patients = query.distinct().all()
        total = len(patients)

        male = sum(1 for p in patients if p.sexo == 'M')
        female = sum(1 for p in patients if p.sexo == 'F')
        age_groups = {'0-17':0, '18-30':0, '31-45':0, '46-60':0, '61+':0}
        for p in patients:
            age = p.edad or 0
            if age < 18:
                age_groups['0-17'] += 1
            elif age <= 30:
                age_groups['18-30'] += 1
            elif age <= 45:
                age_groups['31-45'] += 1
            elif age <= 60:
                age_groups['46-60'] += 1
            else:
                age_groups['61+'] += 1
        self.summary_label.setText(f"{t['total_patients']}: {total}")

        self.ax.clear()
        categories = [t['male'], t['female']]
        values = [male, female]
        bars = self.ax.bar(categories, values, color='#2980b9')
        self.ax.set_title(t['gender_distribution'])
        self.ax.bar_label(bars, fmt='%d')
        self.canvas.draw()

        self.gender_ax.clear()
        groups = list(age_groups.keys())
        counts = list(age_groups.values())
        age_bars = self.gender_ax.bar(groups, counts, color='#27ae60')
        self.gender_ax.set_title(t['age_distribution'])
        self.gender_ax.bar_label(age_bars, fmt='%d')
        self.gender_canvas.draw()

        # legacy global stats
        self.age_ax.clear()
        num_pacientes = len(patients)
        num_hist = session.query(func.count(HistoriaClinica.id)).scalar()
        num_citas = session.query(func.count(Cita.id)).scalar()
        pendientes = session.query(func.count(Factura.id)).filter(Factura.pagado==False).scalar()
        bars = self.age_ax.bar([t['stats_patients'], t['stats_histories'], t['stats_appointments'], t['stats_pending']],
                               [num_pacientes, num_hist, num_citas, pendientes], color='#8e44ad')
        self.age_ax.set_title(t['stats_dashboard'])
        self.age_ax.bar_label(bars, fmt='%d')
        self.age_canvas.draw()
        session.close()


class ValidationDialog(QDialog):
    """Dialog that runs the diagnosis validation workflow."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowFlags(
            self.windowFlags() | Qt.WindowMinimizeButtonHint | Qt.WindowMaximizeButtonHint
        )
        StyleHelper.set_window_icon(self, 'fa5s.flask')
        lang = parent.current_language if parent else 'es'
        t = TRANSLATIONS[lang]
        self.setWindowTitle(t['validate_ai_file'])
        self.validator = DiagnosticoValidator()
        self.thread = None
        self.autosave_path = ""

        layout = QVBoxLayout(self)

        file_box = QHBoxLayout()
        self.file_edit = QLineEdit()
        self.file_edit.setReadOnly(True)
        browse_btn = QPushButton(t['load_file'])
        StyleHelper.style_button(browse_btn)
        file_box.addWidget(self.file_edit)
        file_box.addWidget(browse_btn)
        layout.addLayout(file_box)

        self.start_btn = QPushButton(t['generate'])
        StyleHelper.style_button(self.start_btn)
        layout.addWidget(self.start_btn)

        self.progress = QProgressBar()
        self.progress.setStyleSheet(estilos.PROGRESS_BAR)
        self.progress.setValue(0)
        layout.addWidget(self.progress)

        # Tabla principal con proxy para filtrado
        self.table = QTableWidget()
        headers = [
            "Caso Clínico",
            "Diagnóstico Real",
            "Diagnóstico IA",
            "Evaluación",
            "Justificación Clínica",
        ]
        self.table.setColumnCount(len(headers))
        self.table.setHorizontalHeaderLabels(headers)

        self.proxy_model = QSortFilterProxyModel(self)
        self.proxy_model.setSourceModel(self.table.model())
        self.proxy_model.setFilterCaseSensitivity(Qt.CaseInsensitive)
        self.proxy_model.setFilterKeyColumn(3)

        self.table_view = QTableView()
        self.table_view.setModel(self.proxy_model)
        self.table_view.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        configurar_tabla(self.table_view)
        layout.addWidget(self.table_view)

        # Contadores y gráfico pastel 3D
        stats_widget = QFrame()
        stats_layout = QHBoxLayout(stats_widget)
        labels_layout = QHBoxLayout()
        self.lbl_total = QLabel("Total de casos: 0")
        self.lbl_correctos = QLabel("Correctos: 0")
        self.lbl_parciales = QLabel("Parciales: 0")
        self.lbl_incorrectos = QLabel("Incorrectos: 0")
        self.lbl_precision = QLabel("Precisión: 0.0%")
        for lbl in [self.lbl_total, self.lbl_correctos, self.lbl_parciales, self.lbl_incorrectos, self.lbl_precision]:
            lbl.setFont(QFont("Segoe UI", 10))
            lbl.setStyleSheet("color:#333333;")
            labels_layout.addWidget(lbl)
        labels_layout.addStretch()
        stats_layout.addLayout(labels_layout)

        self.grafico_pastel = grafico_pastel_3d(0, 0, 0, self)
        self.grafico_pastel.setFixedSize(250, 250)
        stats_layout.addWidget(self.grafico_pastel, alignment=Qt.AlignRight)
        layout.addWidget(stats_widget)

        # Filtro por evaluación
        filter_layout = QHBoxLayout()
        filter_layout.addWidget(QLabel("Evaluación:"))
        self.combo_filtro = QComboBox()
        self.combo_filtro.addItems(["Todos", "Correcto", "Parcial", "Incorrecto"])
        filter_layout.addWidget(self.combo_filtro)
        filter_layout.addStretch()
        layout.addLayout(filter_layout)

        self.combo_filtro.currentTextChanged.connect(self.on_filter_changed)
        self.on_filter_changed(self.combo_filtro.currentText())

        self.export_btn = QPushButton(t.get('export_results', 'Exportar Resultados'))
        StyleHelper.style_button(self.export_btn)
        layout.addWidget(self.export_btn)

        close_btn = QPushButton('Cerrar')
        StyleHelper.style_button(close_btn)
        layout.addWidget(close_btn, alignment=Qt.AlignRight)

        browse_btn.clicked.connect(self.select_file)
        self.start_btn.clicked.connect(self.run_validation)
        self.export_btn.clicked.connect(self.export_results)
        close_btn.clicked.connect(self.accept)

    def add_validation_result(self, res: ValidationResult, metrics: dict):
        insertar_resultado_en_tabla(self.table, res)

        total_casos = metrics.get("total_casos", 0)
        self.progress.setMaximum(total_casos)
        self.progress.setValue(metrics.get("casos_validados", 0))
        self.lbl_total.setText(
            f"Total de casos: {metrics.get('casos_validados', 0)}"
        )
        self.lbl_correctos.setText(
            f"Correctos: {metrics.get('CORRECTO', 0)}"
        )
        self.lbl_parciales.setText(
            f"Parciales: {metrics.get('PARCIAL', 0)}"
        )
        self.lbl_incorrectos.setText(
            f"Incorrectos: {metrics.get('INCORRECTO', 0)}"
        )
        self.lbl_precision.setText(
            f"Precisión: {metrics.get('precision', 0.0):.1f}%"
        )
        self.grafico_pastel.actualizar(
            {
                "CORRECTO": metrics.get("CORRECTO", 0),
                "PARCIAL": metrics.get("PARCIAL", 0),
                "INCORRECTO": metrics.get("INCORRECTO", 0),
            }
        )
        self.on_filter_changed(self.combo_filtro.currentText())

    def on_filter_changed(self, text: str) -> None:
        value = "" if text == "Todos" else text.upper()
        self.proxy_model.setFilterFixedString(value)

    def validation_finished(self, stats: dict):
        self.lbl_total.setText(f"Total de casos: {stats.get('total', 0)}")
        self.lbl_correctos.setText(f"Correctos: {stats.get('CORRECTO', 0)}")
        self.lbl_parciales.setText(f"Parciales: {stats.get('PARCIAL', 0)}")
        self.lbl_incorrectos.setText(f"Incorrectos: {stats.get('INCORRECTO', 0)}")
        self.lbl_precision.setText(
            f"Precisión: {stats.get('precision', 0.0):.1f}%"
        )
        self.grafico_pastel.actualizar(stats)
        self.start_btn.setEnabled(True)
        self.export_btn.setEnabled(True)

    def validation_error(self, msg: str):
        QMessageBox.critical(self, 'Error', msg)
        self.start_btn.setEnabled(True)
        self.export_btn.setEnabled(bool(self.validator.resultados))

    def cleanup_thread(self):
        if self.thread:
            self.thread.wait()
            self.thread = None

    def select_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, 'Archivo de Casos', '', 'Datos (*.csv *.xlsx)')
        if path:
            self.file_edit.setText(path)

    def run_validation(self):
        ruta = self.file_edit.text()
        if not ruta:
            QMessageBox.warning(self, 'Error', 'Selecciona un archivo primero.')
            return
        self.start_btn.setEnabled(False)
        self.export_btn.setEnabled(False)
        self.table.setRowCount(0)
        self.validator.reset()
        base = os.path.splitext(ruta)[0] + "_resultados.xlsx"
        self.autosave_path = base
        self.progress.setValue(0)
        self.grafico_pastel.actualizar({})
        self.lbl_total.setText("Total de casos: 0")
        self.lbl_correctos.setText("Correctos: 0")
        self.lbl_parciales.setText("Parciales: 0")
        self.lbl_incorrectos.setText("Incorrectos: 0")
        self.lbl_precision.setText("Precisión: 0.0%")
        self.combo_filtro.setCurrentIndex(0)
        self.on_filter_changed(self.combo_filtro.currentText())

        if self.thread and self.thread.isRunning():
            self.thread.quit()
            self.thread.wait()

        self.thread = ValidationThread(self.validator, ruta, base)
        self.thread.result_ready.connect(self.add_validation_result)
        self.thread.progress.connect(self.progress.setValue)
        self.thread.stats_ready.connect(self.validation_finished)
        self.thread.error.connect(self.validation_error)
        self.thread.finished.connect(self.cleanup_thread)
        self.thread.start()

    def closeEvent(self, event):
        if self.thread and self.thread.isRunning():
            self.thread.quit()
            self.thread.wait()
        super().closeEvent(event)

    def populate_results(self):
        results = self.validator.resultados
        headers = [
            "Caso Clínico",
            "Diagnóstico Real",
            "Diagnóstico IA",
            "Evaluación",
            "Justificación Clínica",
        ]
        self.table.setColumnCount(len(headers))
        self.table.setHorizontalHeaderLabels(headers)
        self.table.setRowCount(0)

        for r in results:
            insertar_resultado_en_tabla(self.table, r)
        self.on_filter_changed(self.combo_filtro.currentText())

    def export_results(self):
        if not self.validator.resultados:
            QMessageBox.warning(self, 'Error', 'No hay resultados para exportar.')
            return
        path, _ = QFileDialog.getSaveFileName(
            self,
            'Exportar Resultados',
            'resultados.xlsx',
            'Excel (*.xlsx)'
        )
        if path:
            try:
                clean = path.replace('.xlsx', '_clean.xlsx')
                self.validator.exportar_excel(path, clean)
                QMessageBox.information(self, 'Éxito', 'Resultados exportados.')
                QDesktopServices.openUrl(QUrl.fromLocalFile(path))
                if os.path.exists(clean):
                    QDesktopServices.openUrl(QUrl.fromLocalFile(clean))
            except Exception as e:
                QMessageBox.critical(self, 'Error', str(e))


class HousingStatsTab(QWidget):
    """Dashboard-style statistics tab for housing evaluations."""

    def __init__(self, parent=None):
        super().__init__(parent)
        grid = QGridLayout(self)

        # Summary of diagnosis distribution
        self.summary_box = QGroupBox("\U0001F4CA Distribuci\u00f3n por Diagn\u00f3stico")
        summ_layout = QVBoxLayout(self.summary_box)
        self.summary_label = QLabel()
        self.summary_label.setAlignment(Qt.AlignCenter)
        summ_layout.addWidget(self.summary_label)
        StyleHelper.style_groupbox(self.summary_box)
        grid.addWidget(self.summary_box, 0, 0, 1, 2)

        # Table listing counts per characteristic
        table_box = QGroupBox("\U0001F4CB Caracter\u00edsticas")
        tbl_layout = QVBoxLayout(table_box)
        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels([
            "Caracter\u00edstica", "Opci\u00f3n", "Cantidad", ""
        ])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        StyleHelper.style_table(self.table)
        tbl_layout.addWidget(self.table)
        StyleHelper.style_groupbox(table_box)
        grid.addWidget(table_box, 1, 0, 1, 2)

        # Quick stats section with checkbox and text area
        self.quick_box = QGroupBox("\u2705 Estad\u00edsticas r\u00e1pidas")
        q_layout = QVBoxLayout(self.quick_box)
        self.quick_check = QCheckBox("Habilitar")
        self.quick_edit = QTextEdit()
        self.quick_edit.setReadOnly(True)
        StyleHelper.style_input(self.quick_edit)
        self.quick_check.toggled.connect(self.on_quick_toggle)
        q_layout.addWidget(self.quick_check)
        q_layout.addWidget(self.quick_edit)
        q_layout.addStretch(1)
        StyleHelper.style_groupbox(self.quick_box)
        grid.addWidget(self.quick_box, 2, 0)

        # Action buttons
        btn_box = QWidget()
        btn_layout = QHBoxLayout(btn_box)
        btn_layout.addStretch(1)
        self.close_btn = QPushButton("Cerrar")
        StyleHelper.style_button(self.close_btn)
        self.close_btn.clicked.connect(self.parent().accept if parent else self.close)
        btn_layout.addWidget(self.close_btn)
        grid.addWidget(btn_box, 2, 1, alignment=Qt.AlignRight)

        self.update_stats()

    def update_stats(self):
        cleanup_orphan_viviendas()
        session = Session()
        viviendas = session.query(Vivienda).all()
        session.close()

        if not viviendas:
            self.summary_label.setText("Sin registros")
            self.total_records = 0
            self.quick_edit.clear()
        else:
            diag_counts = {"Buena": 0, "Regular": 0, "Mala": 0}
            for v in viviendas:
                diag = v.calificacion or "N/A"
                diag_counts[diag] = diag_counts.get(diag, 0) + 1

            lines = [f"{k}: {diag_counts.get(k, 0)}" for k in ["Buena", "Regular", "Mala"]]
            extra = [f"{k}: {v}" for k, v in diag_counts.items() if k not in ("Buena", "Regular", "Mala")]
            lines.extend(extra)
            self.summary_label.setText("\n".join(lines))
            total = sum(diag_counts.values())
            self.total_records = total
            self.quick_edit.clear()
            self.quick_edit.setPlainText(f"Total registros: {total}")

        rows = []
        self.row_map = []
        for attr, (title, mapping) in HOUSING_FIELDS.items():
            counts = {}
            for v in viviendas:
                val = getattr(v, attr)
                if val is None:
                    continue
                counts.setdefault(str(val), []).append(v)
            for val, vivs in counts.items():
                label = mapping.get(val)
                if label is None:
                    try:
                        label = mapping.get(int(val))
                    except (ValueError, TypeError):
                        label = None
                rows.append((title, label if label is not None else str(val), len(vivs), vivs))

        self.table.setRowCount(len(rows))
        for row, (title, opt, count, vivs) in enumerate(rows):
            self.table.setItem(row, 0, QTableWidgetItem(title))
            self.table.setItem(row, 1, QTableWidgetItem(opt))
            self.table.setItem(row, 2, QTableWidgetItem(str(count)))
            btn = QPushButton()
            if qta:
                try:
                    btn.setIcon(qta.icon('fa.search', color='white'))
                except Exception:
                    btn.setText('🔍')
            else:
                btn.setText('🔍')
            StyleHelper.style_button(btn)
            btn.clicked.connect(lambda _, v=vivs: self.show_detail(v))
            self.table.setCellWidget(row, 3, btn)

        self.rows_data = rows
        if self.quick_check.isChecked():
            self.populate_quick_stats(total)

    def populate_quick_stats(self, total):
        """Generate simple summary lines from table data."""
        lines = []
        total = max(total, 1)
        # rows_data: list of (title, option, count, vivs)
        by_title = {}
        for title, opt, count, _ in self.rows_data:
            by_title.setdefault(title, []).append((opt, count))
        for title, opts in by_title.items():
            opts.sort(key=lambda x: x[1], reverse=True)
            if not opts:
                continue
            opt, count = opts[0]
            percent = count * 100 / total
            lines.append(f"✔ La mayoría en {title.lower()} es {opt} ({percent:.0f}%)")
        self.quick_edit.setPlainText("\n".join(lines))

    def on_quick_toggle(self, checked):
        if checked:
            self.populate_quick_stats(getattr(self, 'total_records', 0))
        else:
            self.quick_edit.clear()

    def show_detail(self, vivs):
        # Clean up orphan records before showing details
        cleanup_orphan_viviendas()
        dialog = QDialog(self)
        dialog.setWindowTitle('Detalles')
        dlg_layout = QVBoxLayout(dialog)
        table = QTableWidget()
        table.setColumnCount(4)
        table.setHorizontalHeaderLabels(['Nombre', 'Dirección', 'Teléfono', 'Fecha'])
        table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        StyleHelper.style_table(table)
        dlg_layout.addWidget(table)
        session = Session()
        valid = [v for v in vivs if session.get(Paciente, v.paciente_id)]
        table.setRowCount(len(valid))
        for row, v in enumerate(valid):
            p = session.get(Paciente, v.paciente_id)
            table.setItem(row, 0, QTableWidgetItem(p.nombre if p else ''))
            table.setItem(row, 1, QTableWidgetItem(p.direccion if p else ''))
            phone = getattr(p, 'telefono', '') if p else ''
            table.setItem(row, 2, QTableWidgetItem(phone))
            table.setItem(row, 3, QTableWidgetItem(v.fecha.strftime('%Y-%m-%d')))
        session.close()
        close_btn = QPushButton('Cerrar')
        StyleHelper.style_button(close_btn)
        close_btn.clicked.connect(dialog.accept)
        dlg_layout.addWidget(close_btn, alignment=Qt.AlignRight)
        dialog.exec_()


class ExplanationDialog(QDialog):
    def __init__(self, messages, model, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.comment-medical')
        self.setWindowFlags(self.windowFlags() | Qt.WindowStaysOnTopHint | Qt.Tool)
        self.setWindowTitle('Explicación para el Paciente')
        self.resize(900, 600)
        layout = QVBoxLayout(self)
        self.text = QTextEdit()
        self.text.setReadOnly(True)
        self.text.setText('Escribiendo...')
        StyleHelper.style_input(self.text)
        self.text.setMinimumHeight(450)
        layout.addWidget(self.text)
        btn = QPushButton('Cerrar')
        StyleHelper.style_button(btn)
        btn.clicked.connect(self.accept)
        layout.addWidget(btn, alignment=Qt.AlignRight)

        self.thread = AIStreamThread(messages, model)
        self.thread.chunk_received.connect(self.update_chunk)
        self.thread.finished.connect(self.finish_text)
        self.thread.error.connect(lambda m: QMessageBox.warning(self, 'Error', m))
        self.thread.start()

    def update_chunk(self, text):
        if self.text.toPlainText() == 'Escribiendo...':
            self.text.clear()
        self.text.moveCursor(QTextCursor.End)
        self.text.insertPlainText(text)
        QApplication.processEvents()

    def finish_text(self, text):
        self.text.setPlainText(text)


class HistoryQueryDialog(QDialog):
    """Chat-like dialog for asking about patient history."""

    def __init__(self, patient, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.question')
        self.patient = patient
        self.setWindowTitle('Preguntar Historial')
        self.resize(1000, 1100)
        self.setWindowFlags(
            self.windowFlags()
            | Qt.Window
            | Qt.WindowMinMaxButtonsHint
            | Qt.WindowStaysOnTopHint
        )

        StyleHelper.style_dialog(self)
        # lighter background and smaller border radius
        self.setStyleSheet(self.styleSheet() + "\nQDialog{background-color:white;border-radius:12px;}")

        main_layout = QVBoxLayout(self)

        header = QHBoxLayout()
        title = QLabel(
            "🤖 Asistente IA – Preguntar sobre el Historial del Paciente"
        )
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("font-weight:bold;")
        header.addWidget(title)
        header.addStretch()
        self.max_btn = QPushButton()
        if qta:
            try:
                self.max_btn.setIcon(qta.icon('fa5s.expand', color='black'))
            except Exception:
                pass
        self.max_btn.setFlat(True)
        self.max_btn.clicked.connect(self.toggle_maximize)
        header.addWidget(self.max_btn)
        close_btn = QPushButton()
        if qta:
            try:
                close_btn.setIcon(qta.icon('fa5s.times', color='black'))
            except Exception:
                pass
        close_btn.setFlat(True)
        close_btn.clicked.connect(self.close)
        header.addWidget(close_btn)
        main_layout.addLayout(header)

        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.chat_widget = QWidget()
        self.chat_layout = QVBoxLayout(self.chat_widget)
        self.chat_layout.addStretch()
        self.scroll.setWidget(self.chat_widget)
        main_layout.addWidget(self.scroll, 1)
        # Mensaje de bienvenida
        self.add_message("\u00a1Bienvenido! Formule sus preguntas sobre el historial del paciente.", False)

        self.question_edit = AutoAdjustingTextEdit()
        self.question_edit.heightMin = 80
        self.question_edit.heightMax = 300
        self.question_edit.adjust_height_to_content()
        self.question_edit.setPlaceholderText("Escribe tu pregunta aqu\u00ed...")
        StyleHelper.style_input(self.question_edit)
        main_layout.addWidget(self.question_edit)

        ask_btn = QPushButton('Preguntar')
        StyleHelper.style_button(ask_btn)
        main_layout.addWidget(ask_btn, alignment=Qt.AlignCenter)

        self.context = self.build_context()
        self.messages = []
        self.current_ai_label = None

        ask_btn.clicked.connect(self.ask)


    def build_context(self):
        # Prefer the comprehensive summary from the parent window if available
        parent = self.parent()
        if parent and hasattr(parent, "assemble_patient_context"):
            context = parent.assemble_patient_context()
        else:
            context = ""

        session = Session()
        patient = session.merge(self.patient)
        notes = "\n".join(h.historia_enfermedad or "" for h in patient.historias)
        session.close()
        if notes:
            context += f"\nNotas previas:\n{notes}"
        return context

    def add_message(self, text, is_user):
        """Add a new bubble to the chat area."""
        label = QLabel(text)
        label.setWordWrap(True)
        label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        color = "#cce5ff" if is_user else "#e0e0e0"
        label.setStyleSheet(
            f"background-color:{color}; border-radius:10px; padding:8px;"
        )
        container = QWidget()
        row = QHBoxLayout(container)
        if is_user:
            row.addStretch()
            row.addWidget(label)
        else:
            row.addWidget(label)
            row.addStretch()
        self.chat_layout.insertWidget(self.chat_layout.count() - 1, container)
        self.scroll_to_bottom()
        return label

    def scroll_to_bottom(self):
        bar = self.scroll.verticalScrollBar()
        bar.setValue(bar.maximum())

    def ask(self):
        question = self.question_edit.toPlainText().strip()
        if not question:
            return
        self.last_question = question
        self.add_message(question, True)
        self.question_edit.clear()
        self.question_edit.adjust_height_to_content()
        session = Session()
        patient = session.merge(self.patient)
        session.close()
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r') as f:
                cfg = json.load(f)
            model = apply_openai_key(cfg)
        else:
            openai.api_key = os.getenv('OPENAI_API_KEY', '')
            model = ''
        if not openai.api_key:
            QMessageBox.warning(self, 'Error', 'No se ha configurado la clave API de OpenAI.')
            return
        prompt = self.context + f"\n\nPregunta: {question}"
        messages = [
            {
                "role": "system",
                "content": (
                    "Eres un asistente médico que responde a otro médico en español. "
                    "Analiza cuidadosamente los antecedentes personales y familiares, "
                    "incluidas cirugías previas y alergias, para que la respuesta sea "
                    "coherente con ellos. No indiques que consulte a un profesional y "
                    "responde de forma concisa."
                ),
            },
        ] + self.messages + [{"role": "user", "content": prompt}]
        self.current_ai_label = self.add_message('IA escribiendo...', False)
        self.thread = AIStreamThread(messages, model)
        self.thread.chunk_received.connect(self.update_chunk)
        self.thread.finished.connect(self.finish_text)
        self.thread.error.connect(lambda m: QMessageBox.warning(self, 'Error', m))
        self.thread.start()

    def update_chunk(self, text):
        if self.current_ai_label.text().endswith('IA escribiendo...'):
            self.current_ai_label.setText('')
        self.current_ai_label.setText(self.current_ai_label.text() + text)
        self.scroll_to_bottom()

    def finish_text(self, text):
        self.current_ai_label.setText(text)
        # update conversation history
        self.messages.append({"role": "user", "content": self.context + f"\n\nPregunta: {self.last_question}"})
        self.messages.append({"role": "assistant", "content": text})
        self.scroll_to_bottom()

    def toggle_maximize(self):
        if self.isMaximized():
            self.showNormal()
        else:
            self.showMaximized()



class ImageQueryDialog(QDialog):
    """Chat-style dialog for discussing an image with the AI."""

    def __init__(self, patient, img_data, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.image')
        self.patient = patient
        self.img_data = img_data
        self.setWindowTitle('Análisis de Imagen')
        self.resize(1000, 650)
        self.setWindowFlags(self.windowFlags() | Qt.WindowStaysOnTopHint | Qt.Tool)

        layout = QVBoxLayout(self)

        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.chat_widget = QWidget()
        self.chat_layout = QVBoxLayout(self.chat_widget)
        self.chat_layout.addStretch()
        self.scroll.setWidget(self.chat_widget)
        layout.addWidget(self.scroll, 1)

        self.question_edit = AutoAdjustingTextEdit()
        self.question_edit.heightMin = 80
        self.question_edit.heightMax = 300
        self.question_edit.adjust_height_to_content()
        self.question_edit.setPlaceholderText('Escribe tu respuesta aquí...')
        StyleHelper.style_input(self.question_edit)
        layout.addWidget(self.question_edit)

        send_btn = QPushButton('Responder')
        StyleHelper.style_button(send_btn)
        layout.addWidget(send_btn)

        self.messages = []
        self.current_ai_label = None
        self.context = parent.assemble_patient_context() if parent else ''

        send_btn.clicked.connect(self.send_answer)

        self.analyze_image()

    def add_message(self, text, is_user):
        label = QLabel(text)
        label.setWordWrap(True)
        label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        color = '#cce5ff' if is_user else '#e0e0e0'
        label.setStyleSheet(f'background-color:{color}; border-radius:10px; padding:8px;')
        container = QWidget()
        row = QHBoxLayout(container)
        if is_user:
            row.addStretch()
            row.addWidget(label)
        else:
            row.addWidget(label)
            row.addStretch()
        self.chat_layout.insertWidget(self.chat_layout.count() - 1, container)
        self.scroll_to_bottom()
        return label

    def scroll_to_bottom(self):
        bar = self.scroll.verticalScrollBar()
        bar.setValue(bar.maximum())

    def analyze_image(self):
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r') as f:
                cfg = json.load(f)
            model = apply_openai_key(cfg)
        else:
            openai.api_key = os.getenv('OPENAI_API_KEY', '')
            model = ''
        if not openai.api_key:
            QMessageBox.warning(self, 'Error', 'No se ha configurado la clave API de OpenAI.')
            # show message bubble indicating the missing key
            self.add_message('No hay API Key configurada.', False)
            return

        user_content = [
            {
                'type': 'text',
                'text': (
                    'Analiza esta imagen y ofrece un diagnóstico basado únicamente en lo '
                    'que observas.'
                ),
            },
            {
                'type': 'image_url',
                'image_url': {'url': f'data:image/jpeg;base64,{self.img_data}'},
            },
        ]
        self.messages = [
            {
                'role': 'system',
                'content': (
                    'Eres un asistente médico experto. Debes responder con un diagnóstico '
                    'breve basado solo en la imagen proporcionada. No hagas preguntas ni '
                    'solicites información adicional.'
                ),
            },
            {'role': 'user', 'content': user_content},
        ]
        # store the label returned by add_message so update_chunk works on the
        # QLabel itself instead of the container widget
        self.current_ai_label = self.add_message('IA escribiendo...', False)
        self.thread = AIStreamThread(self.messages, model)
        self.thread.chunk_received.connect(self.update_chunk)
        self.thread.finished.connect(self.analysis_finished)
        self.thread.error.connect(lambda m: QMessageBox.warning(self, 'Error', m))
        self.thread.start()

    def update_chunk(self, text):
        if self.current_ai_label.text().endswith('IA escribiendo...'):
            self.current_ai_label.setText('')
        self.current_ai_label.setText(self.current_ai_label.text() + text)
        self.scroll_to_bottom()

    def analysis_finished(self, text):
        self.current_ai_label.setText(text)
        self.messages.append({'role': 'assistant', 'content': text})

    def send_answer(self):
        """Handle the user's follow-up question about the image."""
        answer = self.question_edit.toPlainText().strip()
        if not answer:
            return
        self.add_message(answer, True)
        self.question_edit.clear()
        self.question_edit.adjust_height_to_content()
        self.messages.append({'role': 'user', 'content': answer})
        self.ask_ai()

    def ask_ai(self):
        """Send the current conversation to the AI model."""
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r') as f:
                cfg = json.load(f)
            model = apply_openai_key(cfg)
        else:
            openai.api_key = os.getenv('OPENAI_API_KEY', '')
            model = ''
        if not openai.api_key:
            QMessageBox.warning(self, 'Error', 'No se ha configurado la clave API de OpenAI.')
            return

        self.current_ai_label = self.add_message('IA escribiendo...', False)
        self.thread = AIStreamThread(self.messages, model)
        self.thread.chunk_received.connect(self.update_chunk)
        self.thread.finished.connect(self.analysis_finished)
        self.thread.error.connect(lambda m: QMessageBox.warning(self, 'Error', m))
        self.thread.start()

class ClinicalStudyDialog(QDialog):
    """Chat-like dialog that shows an AI-generated clinical study for a patient."""

    def __init__(self, patient, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.brain')
        self.patient = patient
        self.parent_window = parent
        self.setWindowTitle('Estudio Clínico IA')
        self.resize(1000, 650)
        self.setWindowFlags(self.windowFlags() | Qt.WindowStaysOnTopHint | Qt.Tool)

        layout = QVBoxLayout(self)

        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.chat_widget = QWidget()
        self.chat_layout = QVBoxLayout(self.chat_widget)
        self.chat_layout.addStretch()
        self.scroll.setWidget(self.chat_widget)
        layout.addWidget(self.scroll, 1)

        self.input_edit = AutoAdjustingTextEdit()
        self.input_edit.heightMin = 80
        self.input_edit.heightMax = 300
        self.input_edit.adjust_height_to_content()
        self.input_edit.setPlaceholderText('Escribe tu mensaje aquí...')
        StyleHelper.style_input(self.input_edit)
        layout.addWidget(self.input_edit)

        send_btn = QPushButton('Enviar')
        StyleHelper.style_button(send_btn)
        if qta:
            send_btn.setIcon(qta.icon('fa5s.paper-plane', color='white'))
        layout.addWidget(send_btn)

        tool_bar = QHBoxLayout()
        self.copy_btn = QPushButton('Copiar')
        StyleHelper.style_button(self.copy_btn)
        if qta:
            self.copy_btn.setIcon(qta.icon('fa5s.copy', color='white'))
        self.save_btn = QPushButton('Guardar Word')
        StyleHelper.style_button(self.save_btn)
        if qta:
            self.save_btn.setIcon(qta.icon('fa5s.file-word', color='white'))
        close_btn = QPushButton('Cerrar')
        StyleHelper.style_button(close_btn)
        if qta:
            close_btn.setIcon(qta.icon('fa5s.times', color='white'))
        tool_bar.addWidget(self.copy_btn)
        tool_bar.addWidget(self.save_btn)
        tool_bar.addStretch()
        tool_bar.addWidget(close_btn)
        layout.addLayout(tool_bar)

        send_btn.clicked.connect(self.send_message)
        self.copy_btn.clicked.connect(self.copy_text)
        self.save_btn.clicked.connect(self.save_word)
        close_btn.clicked.connect(self.accept)

        self.messages = []
        self.log = []
        self.current_ai_label = None

        self.generate_initial()

    def add_message(self, text, is_user):
        label = QLabel(text)
        label.setWordWrap(True)
        label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        color = '#cce5ff' if is_user else '#e0e0e0'
        label.setStyleSheet(f'background-color:{color}; border-radius:10px; padding:8px;')
        container = QWidget()
        row = QHBoxLayout(container)
        if is_user:
            row.addStretch(); row.addWidget(label)
        else:
            row.addWidget(label); row.addStretch()
        self.chat_layout.insertWidget(self.chat_layout.count() - 1, container)
        self.scroll_to_bottom()
        self.log.append(text)
        return label

    def scroll_to_bottom(self):
        bar = self.scroll.verticalScrollBar()
        bar.setValue(bar.maximum())

    def build_context(self):
        p = self.patient
        parent = self.parent_window
        months = months_between(p.fecha_nacimiento) if p.fecha_nacimiento else None
        age = months // 12 if months is not None else p.edad
        age_line = f"Edad: {age}"
        if months is not None:
            age_line += f" ({months} meses)"
        lines = [
            f"Nombre: {p.nombre}",
            age_line,
            f"Sexo: {p.sexo}",
        ]
        if p.documento_id:
            lines.append(f"Cédula: {p.documento_id}")
        if p.alergias:
            lines.append(f"Alergias: {p.alergias}")
        if p.diagnosticos_previos:
            lines.append(f"Diagnósticos Previos: {p.diagnosticos_previos}")
        if p.medicamentos_continuos:
            lines.append(f"Medicamentos Continuos: {p.medicamentos_continuos}")
        lines.append('Antecedentes Personales: ' + parent.get_antecedentes_personales())
        lines.append('Antecedentes Familiares: ' + parent.get_antecedentes_heredofamiliares())
        hist = parent.historia_enfermedad_input.toPlainText().strip()
        if hist:
            lines.append('Historia de la Enfermedad Actual:\n' + hist)
        labs = getattr(parent, 'lab_record', '') or self.fetch_recent_labs()
        if labs:
            lines.append('Resultados de Laboratorio:\n' + labs)
        appts = self.fetch_recent_appointments()
        if appts:
            lines.append('Citas Previas:\n' + appts)
        return '\n'.join(lines)

    def fetch_recent_labs(self, limit=3):
        session = Session()
        res = (
            session.query(LabResult)
            .filter_by(paciente_id=self.patient.id)
            .order_by(LabResult.fecha.desc())
            .limit(limit)
            .all()
        )
        session.close()
        lines = []
        for r in res:
            parts = [r.fecha.strftime('%Y-%m-%d')]
            if r.glicemia is not None:
                parts.append(f'glicemia {r.glicemia}')
            if r.colesterol is not None:
                parts.append(f'colesterol {r.colesterol}')
            if r.trigliceridos is not None:
                parts.append(f'trigliceridos {r.trigliceridos}')
            lines.append(', '.join(parts))
        return '\n'.join(lines)

    def fetch_recent_appointments(self, limit=5):
        session = Session()
        citas = (
            session.query(Cita)
            .filter_by(paciente_id=self.patient.id)
            .order_by(Cita.fecha.desc())
            .limit(limit)
            .all()
        )
        session.close()
        return '\n'.join(f"{c.fecha:%Y-%m-%d %H:%M} - {c.notas or ''}" for c in citas)

    def generate_initial(self):
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r') as f:
                cfg = json.load(f)
            model = apply_openai_key(cfg) or 'gpt-3.5-turbo'
        else:
            openai.api_key = os.getenv('OPENAI_API_KEY', '')
            model = 'gpt-3.5-turbo'
        if not openai.api_key:
            QMessageBox.warning(self, 'Error', 'No se ha configurado la clave API de OpenAI.')
            return

        self.messages = [
            {'role': 'system', 'content': CLINICAL_STUDY_PROMPT},
            {'role': 'user', 'content': self.build_context()},
        ]
        self.current_ai_label = self.add_message('IA escribiendo...', False)
        self.thread = AIStreamThread(self.messages, model)
        self.thread.chunk_received.connect(self.update_chunk)
        self.thread.finished.connect(self.finish_text)
        self.thread.error.connect(lambda m: QMessageBox.warning(self, 'Error', m))
        self.thread.start()

    def update_chunk(self, text):
        if self.current_ai_label.text().endswith('IA escribiendo...'):
            self.current_ai_label.setText('')
        self.current_ai_label.setText(self.current_ai_label.text() + text)
        self.scroll_to_bottom()

    def finish_text(self, text):
        self.current_ai_label.setText(text)
        self.messages.append({'role': 'assistant', 'content': text})

    def send_message(self):
        msg = self.input_edit.toPlainText().strip()
        if not msg:
            return
        self.add_message(msg, True)
        self.input_edit.clear()
        self.input_edit.adjust_height_to_content()
        self.messages.append({'role': 'user', 'content': msg})
        self.ask_ai()

    def ask_ai(self):
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r') as f:
                cfg = json.load(f)
            model = apply_openai_key(cfg) or 'gpt-3.5-turbo'
        else:
            openai.api_key = os.getenv('OPENAI_API_KEY', '')
            model = 'gpt-3.5-turbo'
        if not openai.api_key:
            QMessageBox.warning(self, 'Error', 'No se ha configurado la clave API de OpenAI.')
            return

        self.current_ai_label = self.add_message('IA escribiendo...', False)
        self.thread = AIStreamThread(self.messages, model)
        self.thread.chunk_received.connect(self.update_chunk)
        self.thread.finished.connect(self.finish_text)
        self.thread.error.connect(lambda m: QMessageBox.warning(self, 'Error', m))
        self.thread.start()

    def copy_text(self):
        QApplication.clipboard().setText('\n'.join(self.log))

    def save_word(self):
        text = '\n'.join(self.log)
        doc = Document()
        doc.add_paragraph(text)
        fname, _ = QFileDialog.getSaveFileName(self, 'Guardar Word', '', 'Documento Word (*.docx)')
        if fname:
            doc.save(fname)

class VitalSignsDialog(QDialog):
    def __init__(self, parent=None, patient=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.heartbeat')
        self.setWindowTitle('Signos Vitales')
        self.resize(600, 400)
        layout = QVBoxLayout(self)

        self.grafica_datos = []

        form = QHBoxLayout()
        form.addWidget(QLabel('Paciente:'))
        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        form.addWidget(self.patient_combo)
        self.patient_combo.currentTextChanged.connect(self.load_signs)

        for w in [self.patient_combo]:
            StyleHelper.style_input(w)

        form.addWidget(QLabel('Presión:'))
        self.bp_edit = QLineEdit()
        StyleHelper.style_input(self.bp_edit)
        form.addWidget(self.bp_edit)
        form.addWidget(QLabel('FC:'))
        self.hr_edit = QLineEdit()
        StyleHelper.style_input(self.hr_edit)
        form.addWidget(self.hr_edit)
        form.addWidget(QLabel('Peso:'))
        self.weight_edit = QLineEdit()
        StyleHelper.style_input(self.weight_edit)
        form.addWidget(self.weight_edit)
        form.addWidget(QLabel('Glicemia:'))
        self.gly_edit = QLineEdit()
        StyleHelper.style_input(self.gly_edit)
        form.addWidget(self.gly_edit)
        form.addWidget(QLabel('Colesterol:'))
        self.chol_edit = QLineEdit()
        StyleHelper.style_input(self.chol_edit)
        form.addWidget(self.chol_edit)
        form.addWidget(QLabel('Oximetría:'))
        self.oxi_edit = QLineEdit()
        StyleHelper.style_input(self.oxi_edit)
        form.addWidget(self.oxi_edit)
        form.addWidget(QLabel('Temp:'))
        self.temp_edit = QLineEdit()
        StyleHelper.style_input(self.temp_edit)
        form.addWidget(self.temp_edit)
        add_btn = QPushButton('Agregar')
        StyleHelper.style_button(add_btn)
        form.addWidget(add_btn)
        graph_btn = QPushButton('Ver Gráfica')
        StyleHelper.style_button(graph_btn)
        form.addWidget(graph_btn)
        graph_btn.clicked.connect(self.abrir_grafica)
        layout.addLayout(form)

        self.table = QTableWidget()
        self.table.setColumnCount(10)
        self.table.setHorizontalHeaderLabels([
            'Fecha', 'Presión', 'FC', 'Peso', 'IMC', 'Glicemia',
            'Colesterol', 'Oximetría', 'Temp', 'Eliminar'
        ])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        StyleHelper.style_table(self.table)
        layout.addWidget(self.table)



        self.risk_label = QLabel()
        StyleHelper.style_input(self.risk_label)
        layout.addWidget(self.risk_label)

        add_btn.clicked.connect(self.add_sign)
        self.load_patients(patient)

    def load_patients(self, patient=None):
        session = Session()
        pats = session.query(Paciente).order_by(Paciente.nombre).all()
        names = [p.nombre for p in pats]
        self.patient_combo.clear()
        self.patient_combo.addItems(names)
        if patient:
            idx = self.patient_combo.findText(patient.nombre)
            if idx >= 0:
                self.patient_combo.setCurrentIndex(idx)
            else:
                self.patient_combo.setEditText(patient.nombre)
        session.close()
        if names:
            self.load_signs(self.patient_combo.currentText())

    def load_signs(self, name=None):
        if name is None:
            name = self.patient_combo.currentText()
        session = Session()
        pat = session.query(Paciente).filter_by(nombre=name).first()
        if not pat:
            session.close()
            return
        signs = session.query(SignoVital).filter_by(paciente_id=pat.id).order_by(SignoVital.fecha).all()
        self.table.setRowCount(len(signs))
        dates = []
        weights = []
        bmis = []
        systolic = []
        diastolic = []
        oxi = []
        temps = []
        for row, s in enumerate(signs):
            self.table.setItem(row, 0, QTableWidgetItem(s.fecha.strftime('%Y-%m-%d')))
            self.table.setItem(row, 1, QTableWidgetItem(s.presion or ''))
            self.table.setItem(row, 2, QTableWidgetItem(str(s.frecuencia or '')))
            self.table.setItem(row, 3, QTableWidgetItem(str(s.peso or '')))
            self.table.setItem(row, 4, QTableWidgetItem(f"{s.imc:.1f}" if s.imc else ''))
            self.table.setItem(row, 5, QTableWidgetItem(str(s.glicemia or '')))
            self.table.setItem(row, 6, QTableWidgetItem(str(s.colesterol or '')))
            self.table.setItem(row, 7, QTableWidgetItem(str(s.oximetria or '')))
            self.table.setItem(row, 8, QTableWidgetItem(str(s.temperatura or '')))
            del_btn = QPushButton('Eliminar')
            StyleHelper.style_button(del_btn)
            del_btn.clicked.connect(lambda _, sid=s.id: self.eliminar_signo(sid))
            self.table.setCellWidget(row, 9, del_btn)
            dates.append(s.fecha)
            weights.append(s.peso or 0)
            bmis.append(s.imc or 0)
            sys_v, dia_v = self.parent().parse_bp(s.presion)
            systolic.append(sys_v)
            diastolic.append(dia_v)
            oxi.append(s.oximetria or 0)
            temps.append(s.temperatura or 0)
        session.close()

        self.grafica_datos = [
            [
                s.fecha.strftime('%Y-%m-%d'),
                s.presion or '',
                s.frecuencia,
                s.peso,
                None,
                None,
                None,
                s.oximetria,
            ]
            for s in signs
        ]



    def add_sign(self):
        name = self.patient_combo.currentText()
        session = Session()
        pat = session.query(Paciente).filter_by(nombre=name).first()
        if not pat:
            session.close()
            return
        weight = float(self.weight_edit.text() or 0)
        height = pat.altura or 0
        if height > 3:  # assume cm
            height = height / 100
        bmi = weight / height**2 if height else 0
        last = (
            session.query(SignoVital)
            .filter_by(paciente_id=pat.id)
            .order_by(SignoVital.fecha.desc())
            .first()
        )
        signo = SignoVital(
            paciente_id=pat.id,
            historia_id=None,
            fecha=datetime.now(),
            presion=self.bp_edit.text().strip(),
            frecuencia=int(self.hr_edit.text() or 0),
            peso=weight,
            imc=bmi,
            glicemia=float(self.gly_edit.text() or 0) if self.gly_edit.text() else None,
            colesterol=float(self.chol_edit.text() or 0) if self.chol_edit.text() else None,
            oximetria=float(self.oxi_edit.text() or 0) if self.oxi_edit.text() else None,
            temperatura=float(self.temp_edit.text() or 0) if self.temp_edit.text() else None,
        )
        session.add(signo)
        session.commit()
        change_msg = []
        if last:
            if last.peso and weight != last.peso:
                diff = weight - last.peso
                change_msg.append(f"Peso {'subió' if diff>0 else 'bajó'} {abs(diff):.1f} kg")
            if last.presion and self.bp_edit.text().strip():
                prev_s, prev_d = self.parent().parse_bp(last.presion)
                curr_s, curr_d = self.parent().parse_bp(self.bp_edit.text().strip())
                if prev_s is not None and curr_s is not None and (curr_s != prev_s or curr_d != prev_d):
                    direction = "subió" if (curr_s > prev_s or curr_d > prev_d) else "bajó"
                    msg = f"PA {direction} {abs(curr_s - prev_s)}/{abs(curr_d - prev_d)}"
                    if curr_s >= 140 or curr_d >= 90:
                        msg += " (ALTO)"
                    change_msg.append(msg)
            if last.frecuencia and self.hr_edit.text() and int(self.hr_edit.text()) != last.frecuencia:
                diff_hr = int(self.hr_edit.text()) - last.frecuencia
                change_msg.append(f"FC {'subió' if diff_hr>0 else 'bajó'} {abs(diff_hr)}")
            if last.imc and bmi and round(bmi,1) != round(last.imc,1):
                diff_imc = bmi - last.imc
                change_msg.append(f"IMC {'subió' if diff_imc>0 else 'bajó'} {abs(diff_imc):.1f}")
            if last.glicemia is not None and signo.glicemia is not None and signo.glicemia != last.glicemia:
                diff_g = signo.glicemia - last.glicemia
                change_msg.append(f"Glicemia {'subió' if diff_g>0 else 'bajó'} {abs(diff_g):.1f}")
            if last.colesterol is not None and signo.colesterol is not None and signo.colesterol != last.colesterol:
                diff_c = signo.colesterol - last.colesterol
                change_msg.append(f"Colesterol {'subió' if diff_c>0 else 'bajó'} {abs(diff_c):.1f}")
            if last.oximetria is not None and signo.oximetria is not None and signo.oximetria != last.oximetria:
                diff_o = signo.oximetria - last.oximetria
                change_msg.append(f"Oximetría {'subió' if diff_o>0 else 'bajó'} {abs(diff_o):.1f}")
            if last.temperatura is not None and signo.temperatura is not None and signo.temperatura != last.temperatura:
                diff_t = signo.temperatura - last.temperatura
                change_msg.append(f"Temp {'subió' if diff_t>0 else 'bajó'} {abs(diff_t):.1f}")
        if signo.glicemia is not None and signo.glicemia > 126:
            change_msg.append("Glicemia alta")
        if signo.colesterol is not None and signo.colesterol > 200:
            change_msg.append("Colesterol alto")

        # Store cardiovascular risk in patient record
        risk = self.parent().risk_level(signo)
        pat.riesgo_cardiovascular = risk
        signo.risk_processed = True
        session.commit()
        session.close()
        self.bp_edit.clear()
        self.hr_edit.clear()
        self.weight_edit.clear()
        self.gly_edit.clear()
        self.chol_edit.clear()
        self.oxi_edit.clear()
        self.temp_edit.clear()
        self.load_signs(name)
        if risk == "Alto":
            QMessageBox.warning(self, "Riesgo", "Riesgo cardiovascular alto")
        elif risk == "Moderado":
            QMessageBox.information(self, "Riesgo", "Riesgo cardiovascular moderado")
        if change_msg:
            parent = self.parent()
            if parent and hasattr(parent, 'historia_enfermedad_input'):
                parent.historia_enfermedad_input.append("\n" + "\n".join(change_msg))

    def eliminar_signo(self, signo_id):
        if QMessageBox.question(
            self,
            'Confirmar',
            '¿Eliminar este registro?',
            QMessageBox.Yes | QMessageBox.No,
        ) != QMessageBox.Yes:
            return
        session = Session()
        signo = session.get(SignoVital, signo_id)
        if signo:
            session.delete(signo)
            session.commit()
        session.close()
        self.load_signs(self.patient_combo.currentText())

    def abrir_grafica(self):
        if not self.grafica_datos:
            QMessageBox.information(self, 'Sin datos', 'No hay datos para mostrar.')
            return
        ventana = VentanaGraficaSignos(self.grafica_datos, self)
        ventana.setWindowState(ventana.windowState() | Qt.WindowMaximized)
        ventana.exec_()


class VentanaGraficaSignos(QDialog):
    def __init__(self, datos, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.chart-line')
        self.datos = datos
        self.setWindowTitle('Tendencia de Signos Vitales')
        self.resize(1000, 800)

        layout = QVBoxLayout(self)
        btn_layout = QHBoxLayout()
        self.export_btn = QPushButton('Exportar como Imagen')
        StyleHelper.style_button(self.export_btn)
        self.export_btn.clicked.connect(self.exportar_imagen)
        btn_layout.addWidget(self.export_btn)
        btn_layout.addStretch()
        close_btn = QPushButton('Cerrar')
        StyleHelper.style_button(close_btn)
        close_btn.clicked.connect(self.accept)
        btn_layout.addWidget(close_btn)
        layout.addLayout(btn_layout)

        self.figure = self.generar_figura()
        self.canvas = FigureCanvas(self.figure)
        layout.addWidget(self.canvas)

    def generar_figura(self):
        fig = Figure(figsize=(12, 10))
        axs = fig.subplots(4, 1, sharex=True)
        fig.suptitle('Tendencia de Signos Vitales', fontsize=16, fontweight='bold')

        fechas = []
        sistolica = []
        diastolica = []
        fc = []
        peso = []
        oxim = []
        for row in self.datos:
            try:
                fechas.append(datetime.strptime(row[0], '%Y-%m-%d'))
            except Exception:
                continue
            bp = row[1] or ''
            if bp:
                try:
                    s, d = bp.split('/')
                    sistolica.append(int(s))
                    diastolica.append(int(d))
                except Exception:
                    sistolica.append(None)
                    diastolica.append(None)
            else:
                sistolica.append(None)
                diastolica.append(None)
            fc.append(int(row[2]) if row[2] not in (None, '') else None)
            peso.append(float(row[3]) if row[3] not in (None, '') else None)
            oxim.append(float(row[7]) if len(row) > 7 and row[7] not in (None, '') else None)

        # Track valid x-values so we can safely set axis limits later
        used_dates = []

        def safe_plot(ax, x, y, *args, **kwargs):
            """Plot data if y values are present, skipping missing points."""
            pts = [(xi, yi) for xi, yi in zip(x, y) if xi and yi is not None]
            if not pts:
                return
            xs, ys = zip(*pts)
            used_dates.extend(xs)
            ax.plot(xs, ys, *args, **kwargs)

        safe_plot(
            axs[0],
            fechas,
            sistolica,
            marker="o",
            color="#3498db",
            label="Sistólica",
        )
        safe_plot(
            axs[0],
            fechas,
            diastolica,
            marker="o",
            color="#2ecc71",
            label="Diastólica",
        )
        axs[0].set_title("Presión Arterial (mmHg)")
        axs[0].legend()

        safe_plot(
            axs[1],
            fechas,
            fc,
            marker="s",
            color="#e67e22",
            label="Frecuencia Cardíaca",
        )
        axs[1].set_title("Frecuencia Cardíaca (lpm)")
        axs[1].legend()

        safe_plot(
            axs[2],
            fechas,
            peso,
            marker="^",
            color="#9b59b6",
            label="Peso (kg)",
        )
        axs[2].set_title("Peso (kg)")
        axs[2].legend()

        safe_plot(
            axs[3],
            fechas,
            oxim,
            marker="d",
            color="#e74c3c",
            label="Oximetría (%)",
        )
        axs[3].set_title("Oximetría (%)")
        axs[3].legend()

        axs[3].xaxis.set_major_formatter(mdates.DateFormatter('%d-%b'))

        for ax in axs:
            ax.grid(True, linestyle='--', alpha=0.5)
            ax.tick_params(axis='x', rotation=45)
        if used_dates:
            start = min(used_dates)
            end = max(used_dates)
            if start == end:
                start -= timedelta(days=1)
                end += timedelta(days=1)
            for ax in axs:
                if ax.lines:
                    ax.set_xlim(start, end)

        fig.tight_layout(rect=[0, 0, 1, 0.96])
        return fig

    def exportar_imagen(self):
        fname, _ = QFileDialog.getSaveFileName(
            self,
            'Guardar gráfica',
            '',
            'PNG (*.png);;PDF (*.pdf)'
        )
        if fname:
            if fname.lower().endswith('.pdf'):
                self.figure.savefig(fname, format='pdf')
            else:
                self.figure.savefig(fname, format='png')
            QMessageBox.information(self, 'Éxito', 'Gráfica guardada.')


class FamilySocialDialog(QDialog):
    def __init__(self, parent=None, patient=None):
        super().__init__(parent)
        self.patient = patient
        self.current_vivienda_id = None
        StyleHelper.set_window_icon(self, 'fa5s.home')
        self.setWindowTitle('\U0001F3E0 Información Familiar y del Entorno')
        self.resize(900, 700)

        layout = QVBoxLayout(self)
        self.family_id = ""
        if self.patient:
            session = Session()
            p = session.get(Paciente, self.patient.id)
            if p.id_familia:
                self.family_id = p.id_familia
            else:
                self.family_id = f"FFS-{p.id}"
                p.id_familia = self.family_id
                session.commit()
            session.close()
        self.id_label = QLabel(f"🧬 ID Familia: {self.family_id or 'N/A'}")
        StyleHelper.style_input(self.id_label)
        layout.addWidget(self.id_label, alignment=Qt.AlignLeft)
        self.tabs = QTabWidget()
        layout.addWidget(self.tabs)
        tab = self.create_housing_tab()
        self.tabs.addTab(tab, '\U0001F4D1 Ficha de la Vivienda')
        self.stats_tab = HousingStatsTab(self)
        self.tabs.addTab(self.stats_tab, '\U0001F4CA Estadísticas')
        if self.family_id:
            self.load_family(self.family_id)
        elif self.patient:
            self.load_latest_for_patient()

        self.save_btn = QPushButton('Guardar Ficha')
        StyleHelper.style_button(self.save_btn)
        self.save_btn.clicked.connect(self.save_housing)
        close_btn = QPushButton('Cerrar')
        StyleHelper.style_button(close_btn)
        close_btn.clicked.connect(self.accept)
        btn_row = QHBoxLayout()
        btn_row.addStretch(1)
        btn_row.addWidget(self.save_btn)
        btn_row.addWidget(close_btn)
        layout.addLayout(btn_row)

    def radio_group(self, title, options):
        """Return a group box with radio buttons laid out horizontally."""
        box = QGroupBox(title)
        box.base_title = title
        h = QHBoxLayout()
        h.setSpacing(15)
        btns = []
        for text, val in options:
            rb = QRadioButton(f"{text} ({val:02d})")
            StyleHelper.style_radio(rb)
            h.addWidget(rb)
            rb.toggled.connect(self.update_score)
            btns.append((rb, val))
        h.addStretch(1)
        box.setLayout(h)
        box.setProperty('nbi', False)
        StyleHelper.style_groupbox(box)
        return box, btns

    def create_housing_tab(self):
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)

        container = QWidget()
        container_layout = QVBoxLayout(container)
        container_layout.setContentsMargins(20, 20, 20, 20)
        container.setStyleSheet("background-color:#F7F9FB;")

        panel = QFrame()
        StyleHelper.add_drop_shadow(panel)
        panel.setStyleSheet("background-color:#FFFFFF; border-radius:12px;")
        panel_layout = QVBoxLayout(panel)
        panel_layout.setContentsMargins(20, 20, 20, 20)

        top_row = QHBoxLayout()
        top_row.addWidget(QLabel('Cargar ficha por ID:'))
        self.load_input = QLineEdit()
        StyleHelper.style_input(self.load_input)
        load_btn = QPushButton('Cargar')
        StyleHelper.style_button(load_btn)
        load_btn.setIcon(StyleHelper.window_icon('fa5s.folder-open'))
        load_btn.clicked.connect(self.load_existing)
        top_row.addWidget(self.load_input)
        top_row.addWidget(load_btn)
        top_row.addStretch(1)
        icon_lbl = QLabel()
        icon_lbl.setPixmap(StyleHelper.window_icon('fa5s.calendar').pixmap(24,24))
        top_row.addWidget(icon_lbl)
        self.date_edit = QDateEdit(QDate.currentDate())
        self.date_edit.setDisplayFormat('dd/MM/yy')
        self.date_edit.setCalendarPopup(True)
        StyleHelper.style_input(self.date_edit)
        top_row.addWidget(self.date_edit)
        panel_layout.addLayout(top_row)

        form_grid = QGridLayout()
        form_grid.setSpacing(10)
        panel_layout.addLayout(form_grid)

        self.fields = {}
        field_defs = [
            ('🏠 Tenencia de la vivienda', [('Propia', 5), ('Alquilada', 3), ('Cedida/Prestada', 2)]),
            ('🧱 Paredes de la vivienda', [('Cemento', 10), ('Madera', 9), ('Asbesto', 5), ('Zinc', 4), ('Cartón/Yagua/Desechos', 0)]),
            ('🏚️ Techo de la vivienda', [('Concreto', 10), ('Asbesto', 9), ('Cana/Yagua/Desechos', 0)]),
            ('🪵 Piso de la vivienda', [('Mosaico', 10), ('Madera', 9), ('Tierra', 0)]),
            ('🚽 Servicios sanitarios', [('Inodoro exclusivo', 10), ('Letrina exclusiva', 9), ('Letrina colectiva', 4), ('No tiene', 0)]),
            ('🚰 Agua instalación', [('Dentro y llega', 9), ('Fuera y llega', 8), ('No llega o no tiene', 0)]),
            ('💧 Abastecimiento de agua', [('Acueducto', 9), ('Manantial', 8), ('Río', 4), ('Pozo', 3), ('Lluvia', 2), ('No tiene', 0)]),
            ('🗑️ Eliminación de basura', [('Recoge ayuntamiento', 8), ('La entierran', 7), ('Queman', 6), ('Cañada/patio', 0)]),
            ('🔌 Electricidad', [('CDEEE/Inversor/Planta', 5), ('No tiene', 0)]),
            ('🛏️ Dormitorios', [('4 o más', 9), ('3', 8), ('2', 6), ('1', 0)]),
            ('🔥 Combustible de cocina', [('Gas', 5), ('Carbón', 4), ('Leña', 2)]),
            ('🐾 Animales domésticos', [('Sí (00)', 0), ('No (05)', 5)]),
            ('🦟 Vectores (criaderos)', [('Sí (00) (Mosquitos, ratas)', 0), ('No (05)', 5)]),
        ]

        nbi = {
            'Techo de la vivienda': [0],
            'Piso de la vivienda': [0],
            'Electricidad': [0],
            'Combustible de cocina': [2,4],
            'Agua instalación': [0],
            'Eliminación de basura': [0],
        }

        row = 0
        col = 0
        for title, opts in field_defs:
            box, btns = self.radio_group(title, opts)
            StyleHelper.add_drop_shadow(box, radius=12)
            form_grid.addWidget(box, row, col)
            key = title
            if key and not key[0].isalnum() and ' ' in key:
                key = key.split(' ', 1)[1]
            self.fields[key] = (box, btns)
            col += 1
            if col == 2:
                col = 0
                row += 1
        if col != 0:
            row += 1

        self.personas_spin = QSpinBox(); self.personas_spin.setMinimum(1); self.personas_spin.setMaximum(20)
        StyleHelper.style_input(self.personas_spin)
        form_grid.addWidget(QLabel('Personas en la vivienda'), row, 0)
        form_grid.addWidget(self.personas_spin, row, 1)
        row += 1

        self.total_label = QLabel('0')
        self.grade_label = QLabel('')
        self.hacinamiento_label = QLabel('')
        self.riesgo_label = QLabel('')
        self.vulnerabilidad_label = QLabel('')
        self.minvur_label = QLabel('')
        self.alerta_label = QLabel('')
        self.recom_label = QLabel('')
        StyleHelper.style_input(self.total_label)
        StyleHelper.style_input(self.grade_label)
        for lbl in [self.hacinamiento_label, self.riesgo_label, self.vulnerabilidad_label, self.minvur_label, self.alerta_label, self.recom_label]:
            StyleHelper.style_input(lbl)
        result_box = QGroupBox('📊 Calificación Final')
        rb_layout = QHBoxLayout(result_box)

        sec1 = QVBoxLayout()
        sec1.addWidget(self.total_label)
        sec1.addWidget(self.grade_label)
        rb_layout.addLayout(sec1)

        sec2 = QVBoxLayout()
        sec2.addWidget(self.hacinamiento_label)
        sec2.addWidget(self.riesgo_label)
        sec2.addWidget(self.vulnerabilidad_label)
        rb_layout.addLayout(sec2)

        sec3 = QVBoxLayout()
        sec3.addWidget(self.minvur_label)
        sec3.addWidget(self.alerta_label)
        sec3.addWidget(self.recom_label)
        rb_layout.addLayout(sec3)

        rb_layout.addStretch(1)

        StyleHelper.add_drop_shadow(result_box, radius=12)
        StyleHelper.style_groupbox(result_box)
        panel_layout.addWidget(result_box)

        container_layout.addWidget(panel)

        self.nbi_map = nbi
        scroll.setWidget(container)
        return scroll

    def update_score(self):
        total = 0
        for title, (box, btns) in self.fields.items():
            selected_val = None
            answered = False
            for rb, val in btns:
                if rb.isChecked():
                    selected_val = val
                    total += val
                    answered = True
                    break
            nbi_hit = title in self.nbi_map and selected_val in self.nbi_map[title]
            box.setProperty('nbi', nbi_hit)
            box.setProperty('answered', answered)
            new_title = f"✅ {box.base_title}" if answered else box.base_title
            if box.title() != new_title:
                box.setTitle(new_title)
            box.setStyleSheet(box.styleSheet())
        vect_val = self.get_value('Vectores (criaderos)')
        if vect_val is not None:
            total += vect_val
        self.total_label.setText(f'🧮 Puntuación Total: {total}')
        if total >= 75:
            grade = '⭐ Buena'
            color = 'green'
        elif total >= 50:
            grade = '⚠️ Regular'
            color = 'orange'
        else:
            grade = '❌ Mala'
            color = 'red'
        self.grade_label.setText(grade)
        self.grade_label.setStyleSheet(f'color:{color}; font-weight:bold')

        # Índice de Hacinamiento
        dorm_val = self.get_value('Dormitorios')
        dorm_map = {9: 4, 8: 3, 6: 2, 0: 1}
        dormitorios = dorm_map.get(dorm_val, 1)
        personas = self.personas_spin.value()
        hac = personas / dormitorios if dormitorios else 0
        if hac > 3:
            hac_cat = 'Grave'
        elif hac > 2:
            hac_cat = 'Moderado'
        elif hac >= 1:
            hac_cat = 'Leve'
        else:
            hac_cat = 'Sin hacinamiento'
        self.hacinamiento_label.setText(f'Índice de Hacinamiento: {hac:.1f} → {hac_cat}')

        # Índice de Riesgo Sanitario
        riesgo_score = 0
        for key in [
            'Piso de la vivienda',
            'Techo de la vivienda',
            'Paredes de la vivienda',
            'Servicios sanitarios',
            'Agua instalación',
            'Abastecimiento de agua',
            'Eliminación de basura',
            'Electricidad',
            'Dormitorios',
            'Combustible de cocina',
        ]:
            val = self.get_value(key)
            if val is not None:
                riesgo_score += val
        if riesgo_score >= 70:
            riesgo_nivel = 'Riesgo bajo'
        elif riesgo_score >= 45:
            riesgo_nivel = 'Riesgo moderado'
        else:
            riesgo_nivel = 'Riesgo alto'
        self.riesgo_label.setText(
            f'Índice de Riesgo Sanitario: {riesgo_score} → {riesgo_nivel}'
        )

        # Vulnerabilidad Social
        vul = 0
        if self.get_value('Tenencia de la vivienda') in (3, 2):
            vul += 1
        if self.get_value('Combustible de cocina') in (4, 2):
            vul += 1
        if vect_val == 0:
            vul += 1
        if self.get_value('Animales domésticos') == 0:
            vul += 1
        vul_cat = 'Alta' if vul >= 3 else 'Moderada' if vul == 2 else 'Baja'
        self.vulnerabilidad_label.setText(f'Vulnerabilidad Social: {vul} → {vul_cat}')

        # Clasificación MINVUR
        estructura = 'Adecuada'
        if self.get_value('Techo de la vivienda') == 0 or self.get_value('Piso de la vivienda') == 0:
            estructura = 'Inadecuada'
        servicios = 'Completos'
        if 0 in (
            self.get_value('Electricidad'),
            self.get_value('Servicios sanitarios'),
            self.get_value('Agua instalación'),
            self.get_value('Eliminación de basura'),
        ):
            servicios = 'Incompletos'
        saneamiento = 'Adecuado'
        if vect_val == 0 or self.get_value('Eliminación de basura') == 0:
            saneamiento = 'Deficiente'
        ten_map = {5: 'Propia', 3: 'Alquilada', 2: 'Cedida/Prestada'}
        ten_text = ten_map.get(self.get_value('Tenencia de la vivienda'), 'N/A')
        habitabilidad = f'Hacinamiento {hac_cat.lower()}'
        self.minvur_label.setText(
            f'🏠 Estructura: {estructura}  '
            f'🚿 Servicios básicos: {servicios}  '
            f'🧹 Saneamiento: {saneamiento}  '
            f'📜 Tenencia: {ten_text}  '
            f'👪 Habitabilidad: {habitabilidad}'
        )

        # Condiciones críticas
        crit = 0
        if self.get_value('Piso de la vivienda') == 0:
            crit += 1
        if self.get_value('Combustible de cocina') == 2:
            crit += 1
        if self.get_value('Servicios sanitarios') == 0:
            crit += 1
        if dormitorios and personas / dormitorios >= 5:
            crit += 1
        if crit >= 2:
            self.alerta_label.setText('⚠️ Vivienda en condiciones críticas')
        else:
            self.alerta_label.setText('')

        # Recomendaciones
        recs = []
        if self.get_value('Piso de la vivienda') == 0:
            recs.append('Mejorar piso')
        if self.get_value('Combustible de cocina') in (4, 2):
            recs.append('Usar combustible limpio')
        if self.get_value('Servicios sanitarios') == 0:
            recs.append('Instalar sanitarios')
        if vect_val == 0:
            recs.append('Eliminar criaderos')
        self.recom_label.setText('Sugerencias: ' + ', '.join(recs))

    def collect_values(self):
        values = {}
        for title, (box, btns) in self.fields.items():
            for rb, val in btns:
                if rb.isChecked():
                    values[title] = val
                    break
            else:
                values[title] = None
        values['Personas'] = self.personas_spin.value()
        return values

    def populate_from_vivienda(self, viv):
        mapping = {
            'Tenencia de la vivienda': viv.tenencia,
            'Paredes de la vivienda': viv.paredes,
            'Techo de la vivienda': viv.techo,
            'Piso de la vivienda': viv.piso,
            'Servicios sanitarios': viv.sanitarios,
            'Agua instalación': viv.agua_instalacion,
            'Abastecimiento de agua': viv.agua_abastecimiento,
            'Eliminación de basura': viv.basura,
            'Electricidad': viv.electricidad,
            'Dormitorios': viv.dormitorios,
            'Combustible de cocina': viv.combustible,
        }
        for title, val in mapping.items():
            if val is None:
                continue
            for rb, v in self.fields[title][1]:
                rb.setChecked(v == val)
        if viv.animales is not None:
            for rb, v in self.fields['Animales domésticos'][1]:
                rb.setChecked(str(v) == str(viv.animales))
        if viv.vectores is not None:
            for rb, v in self.fields['Vectores (criaderos)'][1]:
                rb.setChecked(str(v) == str(viv.vectores))
        if viv.num_personas:
            self.personas_spin.setValue(viv.num_personas)
        self.date_edit.setDate(QDate(viv.fecha))
        self.family_id = viv.familia_id
        self.id_label.setText(f"🧬 ID Familia: {self.family_id}")
        self.current_vivienda_id = viv.id
        self.update_score()

    def get_value(self, title):
        for rb, val in self.fields.get(title, (None, []))[1]:
            if rb.isChecked():
                return val
        return None

    def save_housing(self):
        vals = self.collect_values()
        total = 0
        for k, v in vals.items():
            if isinstance(v, int):
                total += v
        vect_val = vals.get('Vectores (criaderos)')
        if vect_val is not None:
            total += vect_val
        if total >= 75:
            grade = 'Buena'
        elif total >= 50:
            grade = 'Regular'
        else:
            grade = 'Mala'
        session = Session()
        if self.current_vivienda_id:
            vivienda = session.get(Vivienda, self.current_vivienda_id)
        else:
            vivienda = session.query(Vivienda).filter_by(familia_id=self.family_id).first()
            if not vivienda:
                vivienda = Vivienda()
                session.add(vivienda)
        vivienda.paciente_id = self.patient.id if self.patient else None
        vivienda.familia_id = self.family_id
        vivienda.fecha = self.date_edit.date().toPyDate()
        vivienda.tenencia = vals.get('Tenencia de la vivienda')
        vivienda.paredes = vals.get('Paredes de la vivienda')
        vivienda.techo = vals.get('Techo de la vivienda')
        vivienda.piso = vals.get('Piso de la vivienda')
        vivienda.sanitarios = vals.get('Servicios sanitarios')
        vivienda.agua_instalacion = vals.get('Agua instalación')
        vivienda.agua_abastecimiento = vals.get('Abastecimiento de agua')
        vivienda.basura = vals.get('Eliminación de basura')
        vivienda.electricidad = vals.get('Electricidad')
        vivienda.dormitorios = vals.get('Dormitorios')
        vivienda.combustible = vals.get('Combustible de cocina')
        vivienda.num_personas = vals.get('Personas')
        vivienda.animales = str(vals.get('Animales domésticos')) if vals.get('Animales domésticos') is not None else None
        vivienda.vectores = str(vals.get('Vectores (criaderos)')) if vals.get('Vectores (criaderos)') is not None else None
        vivienda.total = total
        vivienda.calificacion = grade
        if self.patient and not self.patient.id_familia:
            self.patient.id_familia = self.family_id
        session.commit()
        self.current_vivienda_id = vivienda.id
        session.close()
        QMessageBox.information(self, 'Éxito', 'Ficha guardada correctamente.')
        if hasattr(self, 'stats_tab'):
            self.stats_tab.update_stats()

    def load_existing(self):
        fam_id = self.load_input.text().strip()
        if not fam_id:
            return
        session = Session()
        viv = session.query(Vivienda).filter_by(familia_id=fam_id).order_by(Vivienda.fecha.desc()).first()
        session.close()
        if not viv:
            QMessageBox.warning(self, 'Aviso', 'No se encontró ficha para ese ID')
            return
        self.populate_from_vivienda(viv)

    def load_family(self, fam_id):
        session = Session()
        viv = (
            session.query(Vivienda)
            .filter_by(familia_id=fam_id)
            .order_by(Vivienda.fecha.desc())
            .first()
        )
        session.close()
        if viv:
            self.populate_from_vivienda(viv)

    def load_latest_for_patient(self):
        if not self.patient:
            return
        session = Session()
        viv = (
            session.query(Vivienda)
            .filter_by(paciente_id=self.patient.id)
            .order_by(Vivienda.fecha.desc())
            .first()
        )
        session.close()
        if viv:
            self.populate_from_vivienda(viv)



class LabResultsDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.vials')
        self.setWindowTitle('Resultados de Laboratorio')
        self.resize(700, 400)
        layout = QVBoxLayout(self)

        form = QHBoxLayout()
        form.addWidget(QLabel('Paciente:'))
        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        StyleHelper.style_input(self.patient_combo)
        form.addWidget(self.patient_combo)
        self.patient_combo.currentTextChanged.connect(self.load_results)

        self.date_edit = QDateEdit(QDate.currentDate())
        self.date_edit.setCalendarPopup(True)
        form.addWidget(self.date_edit)

        self.gly_edit = QLineEdit(); StyleHelper.style_input(self.gly_edit)
        form.addWidget(QLabel('Glicemia:'))
        form.addWidget(self.gly_edit)
        self.col_edit = QLineEdit(); StyleHelper.style_input(self.col_edit)
        form.addWidget(QLabel('Colesterol:'))
        form.addWidget(self.col_edit)
        add_btn = QPushButton('Agregar')
        StyleHelper.style_button(add_btn)
        form.addWidget(add_btn)
        layout.addLayout(form)

        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(['Fecha','Glicemia','Colesterol','Trigliceridos'])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        StyleHelper.style_table(self.table)
        layout.addWidget(self.table)

        add_btn.clicked.connect(self.add_result)
        self.load_patients()

    def load_patients(self):
        session = Session()
        names = [p.nombre for p in session.query(Paciente).order_by(Paciente.nombre).all()]
        self.patient_combo.clear(); self.patient_combo.addItems(names)
        session.close()
        if names:
            self.load_results(names[0])

    def load_results(self, name=None):
        if name is None:
            name = self.patient_combo.currentText()
        session = Session()
        pat = session.query(Paciente).filter_by(nombre=name).first()
        if not pat:
            session.close(); return
        res = session.query(LabResult).filter_by(paciente_id=pat.id).order_by(LabResult.fecha.desc()).all()
        self.table.setRowCount(len(res))
        for row, r in enumerate(res):
            self.table.setItem(row, 0, QTableWidgetItem(r.fecha.strftime('%Y-%m-%d')))
            self.table.setItem(row, 1, QTableWidgetItem(str(r.glicemia or '')))
            self.table.setItem(row, 2, QTableWidgetItem(str(r.colesterol or '')))
            self.table.setItem(row, 3, QTableWidgetItem(str(r.trigliceridos or '')))
        session.close()

    def add_result(self):
        name = self.patient_combo.currentText()
        session = Session()
        pat = session.query(Paciente).filter_by(nombre=name).first()
        if not pat:
            session.close(); return
        res = LabResult(
            paciente_id=pat.id,
            fecha=self.date_edit.date().toPyDate(),
            glicemia=float(self.gly_edit.text() or 0),
            colesterol=float(self.col_edit.text() or 0),
        )
        session.add(res)
        session.commit()
        session.close()
        schedule_lab_reminder(pat.id)
        self.gly_edit.clear(); self.col_edit.clear()
        self.load_results(name)


class TemplateManagerDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.file-alt')
        self.setWindowTitle('Plantillas')
        self.resize(600, 400)
        layout = QVBoxLayout(self)

        form = QHBoxLayout()
        form.addWidget(QLabel('Nombre:'))
        self.name_edit = QLineEdit()
        StyleHelper.style_input(self.name_edit)
        form.addWidget(self.name_edit)
        add_btn = QPushButton('Guardar')
        StyleHelper.style_button(add_btn)
        form.addWidget(add_btn)
        layout.addLayout(form)

        self.text_edit = QTextEdit()
        StyleHelper.style_input(self.text_edit)
        layout.addWidget(self.text_edit)

        self.table = QTableWidget()
        self.table.setColumnCount(2)
        self.table.setHorizontalHeaderLabels(['ID', 'Nombre'])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        StyleHelper.style_table(self.table)
        layout.addWidget(self.table)

        del_btn = QPushButton('Eliminar')
        StyleHelper.style_button(del_btn)
        layout.addWidget(del_btn)

        add_btn.clicked.connect(self.save_template)
        del_btn.clicked.connect(self.delete_template)
        self.table.itemSelectionChanged.connect(self.load_selected)
        self.load_templates()

    def load_templates(self):
        session = Session()
        temps = session.query(PlantillaNota).order_by(PlantillaNota.id.desc()).all()
        self.table.setRowCount(len(temps))
        for row, t in enumerate(temps):
            self.table.setItem(row, 0, QTableWidgetItem(str(t.id)))
            self.table.setItem(row, 1, QTableWidgetItem(t.nombre))
        session.close()

    def save_template(self):
        name = self.name_edit.text().strip()
        text = self.text_edit.toPlainText().strip()
        if not name or not text:
            return
        session = Session()
        if self.table.currentRow() >= 0:
            tid = int(self.table.item(self.table.currentRow(), 0).text())
            tpl = session.get(PlantillaNota, tid)
            if tpl:
                tpl.nombre = name
                tpl.contenido = text
        else:
            tpl = PlantillaNota(nombre=name, contenido=text)
            session.add(tpl)
        session.commit()
        session.close()
        self.name_edit.clear()
        self.text_edit.clear()
        self.load_templates()

    def load_selected(self):
        row = self.table.currentRow()
        if row < 0:
            return
        tid = int(self.table.item(row, 0).text())
        session = Session()
        tpl = session.get(PlantillaNota, tid)
        session.close()
        if tpl:
            self.name_edit.setText(tpl.nombre)
            self.text_edit.setPlainText(tpl.contenido)

    def delete_template(self):
        row = self.table.currentRow()
        if row < 0:
            return
        tid = int(self.table.item(row, 0).text())
        session = Session()
        session.query(PlantillaNota).filter_by(id=tid).delete()
        session.commit()
        session.close()
        self.load_templates()


class ReminderDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.bell')
        self.setWindowTitle('Recordatorios')
        self.resize(700, 500)

        layout = QVBoxLayout(self)

        form_box = QGroupBox()
        form_layout = QVBoxLayout(form_box)

        form_layout.addWidget(QLabel('Paciente'))
        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        StyleHelper.style_input(self.patient_combo)
        form_layout.addWidget(self.patient_combo)

        form_layout.addWidget(QLabel('Motivo'))
        self.desc_edit = QLineEdit()
        StyleHelper.style_input(self.desc_edit)
        form_layout.addWidget(self.desc_edit)

        form_layout.addWidget(QLabel('Tipo'))
        self.type_combo = QComboBox()
        self.type_combo.addItems(['Cita', 'Tratamiento', 'Laboratorio', 'Personalizado'])
        StyleHelper.style_input(self.type_combo)
        form_layout.addWidget(self.type_combo)

        form_layout.addWidget(QLabel('Fecha y hora'))
        self.date_edit = QDateTimeEdit(QDateTime.currentDateTime())
        self.date_edit.setCalendarPopup(True)
        StyleHelper.style_input(self.date_edit)
        form_layout.addWidget(self.date_edit)

        form_layout.addWidget(QLabel('Repetir'))
        self.repeat_combo = QComboBox()
        self.repeat_combo.addItems(['Ninguno', 'Diario', 'Semanal', 'Mensual'])
        StyleHelper.style_input(self.repeat_combo)
        form_layout.addWidget(self.repeat_combo)

        form_layout.addWidget(QLabel('Notificación'))
        self.notify_combo = QComboBox()
        self.notify_combo.addItems(['Visual', 'Correo', 'SMS'])
        StyleHelper.style_input(self.notify_combo)
        form_layout.addWidget(self.notify_combo)

        form_layout.addWidget(QLabel('Comentario'))
        self.comment_edit = QLineEdit()
        StyleHelper.style_input(self.comment_edit)
        form_layout.addWidget(self.comment_edit)

        add_btn = QPushButton('Agregar')
        StyleHelper.style_button(add_btn)
        form_layout.addWidget(add_btn)

        layout.addWidget(form_box)

        self.table = QTableWidget()
        self.table.setColumnCount(10)
        self.table.setHorizontalHeaderLabels([
            'ID', 'Paciente', 'Motivo', 'Fecha', 'Hora', 'Tipo', 'Repetir', 'Notificación', 'Comentario', 'Completado'
        ])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        StyleHelper.style_table(self.table)
        layout.addWidget(self.table)

        del_btn = QPushButton('Eliminar')
        StyleHelper.style_button(del_btn)
        layout.addWidget(del_btn)

        add_btn.clicked.connect(self.add_reminder)
        del_btn.clicked.connect(self.delete_reminder)
        self.table.itemChanged.connect(self.toggle_completed)
        self.load_patients()
        self.load_reminders()

    def show_today(self):
        self.only_today = True
        self.load_reminders()

    def load_patients(self):
        session = Session()
        pats = session.query(Paciente).order_by(Paciente.nombre).all()
        names = [p.nombre for p in pats]
        self.patient_combo.clear()
        self.patient_combo.addItems(names)
        session.close()

    def load_reminders(self):
        session = Session()
        query = session.query(Recordatorio, Paciente.nombre).outerjoin(
            Paciente, Recordatorio.paciente_id == Paciente.id)
        if getattr(self, 'only_today', False):
            today = date.today()
            query = query.filter(Recordatorio.fecha == today)
        rems = query.order_by(Recordatorio.fecha).all()

        self.table.blockSignals(True)
        self.table.setRowCount(len(rems))
        for row, (rec, name) in enumerate(rems):
            patient_name = name or str(rec.paciente_id)
            self.table.setItem(row, 0, QTableWidgetItem(str(rec.id)))
            self.table.setItem(row, 1, QTableWidgetItem(patient_name))
            self.table.setItem(row, 2, QTableWidgetItem(rec.descripcion or ''))
            self.table.setItem(row, 3, QTableWidgetItem(rec.fecha.strftime('%Y-%m-%d')))
            self.table.setItem(row, 4, QTableWidgetItem(rec.hora.strftime('%H:%M') if rec.hora else ''))
            self.table.setItem(row, 5, QTableWidgetItem(rec.tipo or ''))
            self.table.setItem(row, 6, QTableWidgetItem(rec.repetir or ''))
            self.table.setItem(row, 7, QTableWidgetItem(rec.notificacion or ''))
            self.table.setItem(row, 8, QTableWidgetItem(rec.comentario or ''))
            chk = QTableWidgetItem()
            chk.setCheckState(Qt.Checked if rec.completado else Qt.Unchecked)
            self.table.setItem(row, 9, chk)
        self.table.blockSignals(False)
        session.close()

    def add_reminder(self):
        name = self.patient_combo.currentText()
        session = Session()
        pat = session.query(Paciente).filter_by(nombre=name).first()
        if not pat:
            session.close()
            return
        rem = Recordatorio(
            paciente_id=pat.id,
            descripcion=self.desc_edit.text().strip(),
            fecha=self.date_edit.date().toPyDate(),
            hora=self.date_edit.time().toPyTime(),
            tipo=self.type_combo.currentText(),
            repetir=self.repeat_combo.currentText(),
            notificacion=self.notify_combo.currentText(),
            comentario=self.comment_edit.text().strip(),
        )
        session.add(rem)
        session.commit()
        session.close()
        self.desc_edit.clear()
        self.comment_edit.clear()
        self.load_reminders()

    def delete_reminder(self):
        row = self.table.currentRow()
        if row < 0:
            return
        rid = int(self.table.item(row, 0).text())
        session = Session()
        session.query(Recordatorio).filter_by(id=rid).delete()
        session.commit()
        session.close()
        self.load_reminders()

    def toggle_completed(self, item):
        if item.column() != 9:
            return
        rid = int(self.table.item(item.row(), 0).text())
        checked = item.checkState() == Qt.Checked
        session = Session()
        rec = session.get(Recordatorio, rid)
        if rec:
            rec.completado = checked
            rec.completado_fecha = datetime.now() if checked else None
        session.commit()
        session.close()


class ReminderListDialog(QDialog):
    def __init__(self, parent=None, only_today=False):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.bell')
        self.setWindowTitle('Recordatorios')
        self.resize(600, 400)

        layout = QVBoxLayout(self)
        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(['Paciente', 'Motivo', 'Fecha', 'Hora', 'Comentario'])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        StyleHelper.style_table(self.table)
        layout.addWidget(self.table)

        close_btn = QPushButton('Cerrar')
        StyleHelper.style_button(close_btn)
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)

        self.only_today = only_today
        self.load_reminders()

    def load_reminders(self):
        session = Session()
        query = session.query(Recordatorio, Paciente.nombre).outerjoin(
            Paciente, Recordatorio.paciente_id == Paciente.id)
        if self.only_today:
            today = date.today()
            query = query.filter(Recordatorio.fecha == today)
        rems = query.order_by(Recordatorio.fecha).all()
        self.table.setRowCount(len(rems))
        for row, (rec, name) in enumerate(rems):
            patient_name = name or str(rec.paciente_id)
            self.table.setItem(row, 0, QTableWidgetItem(patient_name))
            self.table.setItem(row, 1, QTableWidgetItem(rec.descripcion or ''))
            self.table.setItem(row, 2, QTableWidgetItem(rec.fecha.strftime('%Y-%m-%d')))
            self.table.setItem(row, 3, QTableWidgetItem(rec.hora.strftime('%H:%M') if rec.hora else ''))
            self.table.setItem(row, 4, QTableWidgetItem(rec.comentario or ''))
        session.close()


class AdvancedSearchDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.search')
        self.setWindowTitle('Búsqueda Avanzada')
        self.resize(700, 400)
        layout = QVBoxLayout(self)

        form = QHBoxLayout()
        form.addWidget(QLabel('Desde:'))
        self.start_edit = QDateEdit(QDate.currentDate().addMonths(-1))
        self.start_edit.setCalendarPopup(True)
        form.addWidget(self.start_edit)
        form.addWidget(QLabel('Hasta:'))
        self.end_edit = QDateEdit(QDate.currentDate())
        self.end_edit.setCalendarPopup(True)
        form.addWidget(self.end_edit)
        form.addWidget(QLabel('Diagnóstico:'))
        self.diagnosis_edit = QLineEdit()
        StyleHelper.style_input(self.diagnosis_edit)
        form.addWidget(self.diagnosis_edit)
        search_btn = QPushButton('Buscar')
        StyleHelper.style_button(search_btn)
        form.addWidget(search_btn)
        layout.addLayout(form)

        self.table = QTableWidget()
        self.table.setColumnCount(3)
        self.table.setHorizontalHeaderLabels(['Fecha', 'Paciente', 'Diagnóstico'])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        StyleHelper.style_table(self.table)
        layout.addWidget(self.table)

        search_btn.clicked.connect(self.search)

    def search(self):
        start = self.start_edit.date().toPyDate()
        end = self.end_edit.date().toPyDate()
        diag = self.diagnosis_edit.text().strip()
        session = Session()
        q = session.query(HistoriaClinica, Diagnostico).join(Diagnostico, Diagnostico.historia_id == HistoriaClinica.id)
        q = q.filter(HistoriaClinica.fecha.between(start, end))
        if diag:
            q = q.filter(Diagnostico.descripcion.ilike(f"%{diag}%"))
        rows = q.all()
        self.table.setRowCount(len(rows))
        for row, (h, d) in enumerate(rows):
            self.table.setItem(row, 0, QTableWidgetItem(h.fecha.strftime('%Y-%m-%d')))
            self.table.setItem(row, 1, QTableWidgetItem(h.paciente.nombre if h.paciente else ''))
            self.table.setItem(row, 2, QTableWidgetItem(d.descripcion))
        session.close()


class FormsDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5b.wpforms')
        lang = parent.current_language if parent else 'es'
        t = TRANSLATIONS[lang]
        self.setWindowTitle(t.get('forms', 'Formularios'))
        self.setWindowFlags(self.windowFlags() | Qt.WindowMinimizeButtonHint | Qt.WindowMaximizeButtonHint)
        self.setWindowState(Qt.WindowMaximized)

        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignTop)

        box = QGroupBox(t.get('forms', 'Formularios'))
        StyleHelper.style_groupbox(box)
        vbox = QVBoxLayout()

        self.cert_btn = QPushButton(t.get('medical_certificate', 'Certificado Médico'))
        StyleHelper.style_button(self.cert_btn)
        self.cert_btn.clicked.connect(self.open_certificate)
        vbox.addWidget(self.cert_btn)

        self.phq_btn = QPushButton(t.get('phq9', 'PHQ-9'))
        StyleHelper.style_button(self.phq_btn)
        self.phq_btn.clicked.connect(self.open_phq)
        vbox.addWidget(self.phq_btn)

        self.gad_btn = QPushButton(t.get('gad7', 'GAD-7'))
        StyleHelper.style_button(self.gad_btn)
        self.gad_btn.clicked.connect(self.open_gad)
        vbox.addWidget(self.gad_btn)

        self.mmse_btn = QPushButton(t.get('mmse', 'MMSE'))
        StyleHelper.style_button(self.mmse_btn)
        self.mmse_btn.clicked.connect(self.open_mmse)
        vbox.addWidget(self.mmse_btn)

        self.denver_btn = QPushButton(t.get('denver', 'Denver II'))
        StyleHelper.style_button(self.denver_btn)
        self.denver_btn.clicked.connect(self.open_denver)
        vbox.addWidget(self.denver_btn)

        box.setLayout(vbox)
        layout.addWidget(box)

    def open_certificate(self):
        dlg = MedicalCertificateDialog(self.parent())
        dlg.exec_()

    def open_phq(self):
        pat = getattr(self.parent(), 'current_patient', None)
        dlg = PHQ9Dialog(self.parent(), pat)
        dlg.setWindowState(dlg.windowState() | Qt.WindowMaximized)
        dlg.exec_()

    def open_gad(self):
        pat = getattr(self.parent(), 'current_patient', None)
        dlg = GAD7Dialog(self.parent(), pat)
        dlg.setWindowState(dlg.windowState() | Qt.WindowMaximized)
        dlg.exec_()

    def open_mmse(self):
        pat = getattr(self.parent(), 'current_patient', None)
        dlg = MMSEDialog(self.parent(), pat)
        dlg.setWindowState(dlg.windowState() | Qt.WindowMaximized)
        dlg.exec_()

    def open_denver(self):
        pat = getattr(self.parent(), 'current_patient', None)
        age = None
        if pat and pat.fecha_nacimiento:
            age = (date.today() - pat.fecha_nacimiento).days // 30
        dlg = DenverDialog(self.parent(), pat, age)
        dlg.setWindowState(dlg.windowState() | Qt.WindowMaximized)
        dlg.exec_()


class MedicalCertificateDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.file-medical')
        lang = parent.current_language if parent else 'es'
        t = TRANSLATIONS[lang]
        self.setWindowTitle(t.get('medical_certificate', 'Certificado Médico'))
        self.resize(700, 400)

        layout = QVBoxLayout(self)

        form = QFormLayout()
        form.addRow(QLabel(t.get('patient', 'Paciente') + ':'))
        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        StyleHelper.style_input(self.patient_combo)
        form.addRow(self.patient_combo)
        form.addRow(QLabel('Documento:'))
        self.id_edit = QLineEdit()
        StyleHelper.style_input(self.id_edit)
        form.addRow(self.id_edit)
        self.healthy_check = QCheckBox(t.get('healthy_patient', 'Paciente sano'))
        StyleHelper.style_checkbox(self.healthy_check)
        form.addRow(self.healthy_check)
        form.addRow(QLabel(t.get('diagnosis_filter', 'Diagnóstico') + ':'))
        self.diag_edit = QLineEdit()
        StyleHelper.style_input(self.diag_edit)
        form.addRow(self.diag_edit)
        load_btn = QPushButton(t.get('load_diag_history', 'Cargar diagnóstico'))
        StyleHelper.style_button(load_btn)
        form.addRow(load_btn)
        form.addRow(QLabel(t.get('days_off', 'Días de incapacidad') + ':'))
        self.days_spin = QSpinBox()
        self.days_spin.setRange(1, 30)
        StyleHelper.style_input(self.days_spin)
        form.addRow(self.days_spin)
        layout.addLayout(form)

        gen_btn = QPushButton(t.get('generate', 'Generar'))
        StyleHelper.style_button(gen_btn)
        layout.addWidget(gen_btn, alignment=Qt.AlignRight)

        gen_btn.clicked.connect(self.generate_certificate)
        load_btn.clicked.connect(self.load_last_diagnosis)
        self.load_patients()
        self.patient_combo.currentTextChanged.connect(self.fill_document)
        if parent and getattr(parent, 'current_patient', None):
            p = parent.current_patient
            self.patient_combo.setCurrentText(p.nombre)
            self.id_edit.setText(p.documento_id or '')

    def load_patients(self):
        session = Session()
        pats = session.query(Paciente).order_by(Paciente.nombre).all()
        names = [p.nombre for p in pats]
        self.patient_combo.addItems(names)
        self._pat_map = {p.nombre: p for p in pats}
        session.close()

    def fill_document(self, name):
        pat = self._pat_map.get(name)
        if pat:
            self.id_edit.setText(pat.documento_id or '')

    def load_last_diagnosis(self):
        name = self.patient_combo.currentText().strip()
        pat = self._pat_map.get(name)
        if not pat:
            return
        session = Session()
        rows = (
            session.query(Diagnostico.descripcion)
            .join(HistoriaClinica, Diagnostico.historia_id == HistoriaClinica.id)
            .filter(HistoriaClinica.paciente_id == pat.id)
            .order_by(Diagnostico.fecha.desc())
            .all()
        )
        session.close()
        if rows:
            text = "; ".join(r[0] for r in rows)
            self.diag_edit.setText(text)
            days = self.ask_ai_days(text)
            if days:
                self.days_spin.setValue(days)

    def ask_ai_days(self, diag):
        if not diag:
            return 1
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r') as f:
                cfg = json.load(f)
            model = apply_openai_key(cfg)
        else:
            return 1
        prompt = (
            "De acuerdo con el siguiente diagnóstico, recomienda un número de días de reposo. "
            "Responde solo con un número entero:\n" + diag
        )
        try:
            resp = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "Eres un asistente médico."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0
            )
            text = resp.choices[0].message.content.strip()
            m = re.search(r'(\d+)', text)
            if m:
                return int(m.group(1))
        except Exception:
            pass
        return 1

    def generate_certificate(self):
        name = self.patient_combo.currentText().strip()
        doc_id = self.id_edit.text().strip() or 'no tiene'
        sano = self.healthy_check.isChecked()
        diag = self.diag_edit.text().strip()
        cfg = ConfigReceta(); cfg.load_config()
        doc = Document()
        try:
            logo_path = os.path.join(os.path.dirname(__file__), cfg.logo_path)
            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell_logo = table.cell(0, 0)
            cell_head = table.cell(0, 1)
            cell_qr = table.cell(0, 2)
            cell_logo.paragraphs[0].add_run().add_picture(logo_path, width=Inches(0.8))
            header = cfg.get_clinic_header()
            if header:
                para = cell_head.paragraphs[0]
                para.text = header
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            qr_path = os.path.join(tempfile.gettempdir(), 'cert_qr.png')
            create_qr_code(f"Cert-{name}-{datetime.now().strftime('%Y%m%d%H%M')}", qr_path)
            cell_qr.paragraphs[0].add_run().add_picture(qr_path, width=Inches(1))
        except Exception:
            pass

        doc.add_heading('CERTIFICADO MEDICO', 0).alignment = 1
        exequatur = '725-09'
        pat = self._pat_map.get(name)
        genero = pat.sexo if pat else ''
        apto = 'APTO' if genero.startswith('Mas') else 'APTA'
        portador = 'portador' if genero.startswith('Mas') else 'portadora'
        dicho = 'dicho' if genero.startswith('Mas') else 'dicha'

        texto = (
            f"Yo DR. {cfg.get_doctor_name()} provisto del exequátur número {exequatur} "
            f"CERTIFICO haber examinado a {name} {portador} de la cédula de identidad y electoral número {doc_id} "
        )
        dias = self.days_spin.value()
        if sano:
            texto += f"Y he CONSTATADO QUE {dicho} paciente SE ENCUENTRA ACTUALMENTE {apto} PARA CUALQUIER ACTIVIDAD FISICA O MENTAL."
        else:
            texto += f"Y he CONSTATADO QUE {dicho} paciente CUENTA CON DIAGNOSTICO DE {diag}. SUGIERO {dias} DIAS DE REPOSO." 
        doc.add_paragraph(texto)
        fecha = datetime.now().strftime('%d de %b. de %y')
        doc.add_paragraph(
            f'Expido la presente certificación en Santiago de los Caballeros, a petición del interesado , hoy día {fecha}'
        )

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_sig, cell_seal = table.row_cells(0)
        sign = cfg.get_signature()
        seal = cfg.get_seal()
        if sign and os.path.exists(sign):
            cell_sig.paragraphs[0].add_run().add_picture(sign, width=Inches(1.5))
        cell_sig.add_paragraph(cfg.get_doctor_name()).alignment = 1
        cell_sig.add_paragraph(cfg.get_doctor_specialty()).alignment = 1
        if seal and os.path.exists(seal):
            cell_seal.paragraphs[0].add_run().add_picture(seal, width=Inches(2.5))

        footer = doc.sections[0].footer.paragraphs[0]
        footer.text = f"{cfg.get_footer_text()}\n{cfg.get_footer_details()}\n{cfg.get_footer_phone()}"
        footer.alignment = 1

        default = os.path.join(expanduser('~'), f'Certificado_{name}_{datetime.now().strftime("%Y%m%d")}.docx')
        path, _ = QFileDialog.getSaveFileName(self, 'Guardar Certificado', default, 'Archivos Word (*.docx)')
        if not path:
            return
        if not path.lower().endswith('.docx'):
            path += '.docx'
        try:
            doc.save(path)
            QMessageBox.information(self, 'Éxito', f'Certificado guardado en {path}')
            try:
                os.startfile(path)
            except Exception:
                pass
        except Exception as e:
            QMessageBox.critical(self, 'Error', str(e))
        self.accept()


class PHQ9Dialog(QDialog):
    """Formulario PHQ-9 para evaluar depresión."""
    QUESTIONS = [
        "Poco interés o placer en hacer cosas",
        "Se ha sentido decaído(a), deprimido(a) o sin esperanza",
        "Dificultad para dormir o dormir demasiado",
        "Se ha sentido cansado(a) o con poca energía",
        "Pérdida del apetito o comer en exceso",
        "Se siente mal consigo mismo o piensa que ha fallado",
        "Dificultad para concentrarse en actividades",
        "Se mueve o habla tan despacio que otras personas lo notan, o lo opuesto: se siente inquieto(a) o agitado(a)",
        "Pensamientos de que estaría mejor muerto(a) o de hacerse daño",
    ]

    def __init__(self, parent=None, patient=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.smile')
        self.patient = patient
        self.setWindowTitle("PHQ-9: Evaluación de Depresión")
        self.resize(700, 600)

        main_layout = QVBoxLayout(self)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        container = QWidget()
        layout = QVBoxLayout(container)
        form = QFormLayout()

        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        StyleHelper.style_input(self.patient_combo)
        form.addRow(QLabel("Paciente:"), self.patient_combo)







        self.option_texts = [
            "0: Nunca",
            "1: Varios días",
            "2: Más de la mitad de los días",
            "3: Casi todos los días",
        ]
        self.answer_boxes = []
        for i, qtext in enumerate(self.QUESTIONS, 1):
            combo = QComboBox()
            combo.addItems(self.option_texts)
            StyleHelper.style_input(combo)
            form.addRow(QLabel(f"{i}. {qtext}"), combo)
            combo.currentIndexChanged.connect(self.update_score_label)
            self.answer_boxes.append(combo)

        layout.addLayout(form)

        self.total_label = QLabel("Puntaje total: 0 - Mínima o ausente")
        layout.addWidget(self.total_label)

        self.report_edit = AutoAdjustingTextEdit()
        self.report_edit.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        StyleHelper.style_input(self.report_edit)
        layout.addWidget(self.report_edit)

        btn_layout = QHBoxLayout()
        self.gen_btn = QPushButton("Generar Informe")
        StyleHelper.style_button(self.gen_btn)
        self.pdf_btn = QPushButton("Exportar PDF")
        StyleHelper.style_button(self.pdf_btn)
        btn_layout.addWidget(self.gen_btn)
        btn_layout.addWidget(self.pdf_btn)
        layout.addLayout(btn_layout)

        scroll.setWidget(container)
        main_layout.addWidget(scroll)

        self.gen_btn.clicked.connect(self.generate_report)
        self.pdf_btn.clicked.connect(self.export_pdf)

        self.load_patients()
        if patient:
            self.patient_combo.setCurrentText(patient.nombre)

        self.update_score_label()

    def load_patients(self):
        session = Session()
        pats = session.query(Paciente).order_by(Paciente.nombre).all()
        self.patient_combo.addItems([p.nombre for p in pats])
        self.pat_map = {p.nombre: p for p in pats}
        session.close()

    def score(self):
        return sum(box.currentIndex() for box in self.answer_boxes)

    def update_score_label(self):
        total = self.score()
        interpretation = (
            "Mínima o ausente" if total <= 4 else
            "Leve" if total <= 9 else
            "Moderada" if total <= 14 else
            "Moderadamente severa" if total <= 19 else
            "Severa"
        )
        self.total_label.setText(f"Puntaje total: {total} - {interpretation}")

    def save_score(self, pat, total):
        session = Session()
        try:
            p = session.merge(pat)
            p.phq9_score = total
            session.commit()
        except Exception:
            session.rollback()
        finally:
            session.close()

    def generate_report(self):
        name = self.patient_combo.currentText().strip()
        pat = self.pat_map.get(name)
        total = self.score()
        if pat:
            self.save_score(pat, total)
        interpretation = (
            "Mínima o ausente" if total <= 4 else
            "Leve" if total <= 9 else
            "Moderada" if total <= 14 else
            "Moderadamente severa" if total <= 19 else
            "Severa"
        )

        summary = f"Puntaje total: {total} - {interpretation}"
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r') as f:
                    cfg = json.load(f)
                model = apply_openai_key(cfg)
                prompt = (
                    "Eres un profesional de salud mental. Resume brevemente las "
                    f"implicaciones de un puntaje PHQ-9 de {total} ({interpretation}) "
                    "y sugiere recomendaciones iniciales."
                )
                resp = openai.ChatCompletion.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "Asistente médico"},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.5,
                )
                ai_text = resp.choices[0].message.content.strip()
                ai_text = re.sub(r"[\*\u2022]+", "", ai_text)
                ai_text = re.sub(r"^[-\s]+", "", ai_text, flags=re.MULTILINE)
                summary += "\n" + ai_text
            except Exception as e:
                summary += f"\n[Error IA: {e}]"
        self.report_edit.setPlainText(summary)

    def export_pdf(self):
        name = self.patient_combo.currentText().strip()
        fname, _ = QFileDialog.getSaveFileName(
            self, 'Guardar PDF', f'PHQ9_{name}_{datetime.now().strftime("%Y%m%d")}.pdf', 'PDF (*.pdf)')
        if not fname:
            return
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import cm

        c = canvas.Canvas(fname, pagesize=letter)
        y = 27 * cm
        c.setFont('Helvetica-Bold', 14)
        c.drawCentredString(10.5 * cm, y, 'Informe de Evaluación PHQ-9')
        y -= 1 * cm
        c.setFont('Helvetica', 10)
        date_str = datetime.now().strftime('%d/%m/%Y')
        c.drawString(2 * cm, y, f'Paciente: {name}')
        c.drawRightString(19 * cm, y, f'Fecha: {date_str}')
        y -= 0.6 * cm
        total = self.score()
        interpretation = (
            'Mínima o ausente' if total <= 4 else
            'Leve' if total <= 9 else
            'Moderada' if total <= 14 else
            'Moderadamente severa' if total <= 19 else
            'Severa'
        )
        c.drawString(2 * cm, y, f'Puntaje total: {total} - {interpretation}')
        y -= 0.8 * cm
        for i, box in enumerate(self.answer_boxes, 1):
            txt = self.QUESTIONS[i-1]
            ans = self.option_texts[box.currentIndex()]
            c.drawString(2 * cm, y, f'{i}. {txt}')
            y -= 0.4 * cm
            c.drawString(2.5 * cm, y, f'Respuesta: {ans}')
            y -= 0.6 * cm
            if y < 3 * cm:
                c.showPage()
                y = 26 * cm
        y -= 0.4 * cm
        c.setFont('Helvetica', 10)
        for line in self.report_edit.toPlainText().splitlines():
            clean = re.sub(r'[\*\u2022]+', '', line).strip()
            if not clean:
                continue
            c.drawString(2 * cm, y, clean)
            y -= 0.4 * cm
            if y < 2 * cm:
                c.showPage()
                y = 26 * cm
        c.save()
        QMessageBox.information(self, 'Éxito', f'PDF guardado en {fname}')
        try:
            os.startfile(fname)
        except Exception:
            pass


class DenverDialog(QDialog):
    """Tamizaje de Desarrollo Denver II."""

    ITEMS = {
        2: {
            "personal_social": ["¿El niño sonríe socialmente?"],
            "motora_fina_adaptativa": ["¿El niño sigue objetos con la mirada?"],
            "lenguaje": ["¿Emite sonidos suaves?"],
            "motora_gruesa": ["¿Levanta la cabeza estando boca abajo (en prono)?"],
        },
        6: {
            "personal_social": ["¿Se lleva objetos a la boca?"],
            "motora_fina_adaptativa": ["¿Agarra objetos con ambas manos?"],
            "lenguaje": ["¿Balbucea consonantes (como 'ba', 'ma')?"],
            "motora_gruesa": ["¿Se sienta con apoyo?"],
        },
        12: {
            "personal_social": ["¿Aplaude por imitación?"],
            "motora_fina_adaptativa": ["¿Saca objetos de un recipiente?"],
            "lenguaje": ["¿Dice 'mamá' o 'papá' con sentido?"],
            "motora_gruesa": ["¿Camina con apoyo?"],
        },
        24: {
            "personal_social": ["¿Se quita alguna prenda de ropa?"],
            "motora_fina_adaptativa": ["¿Hace torres de al menos 6 bloques?"],
            "lenguaje": ["¿Combina dos palabras al hablar?"],
            "motora_gruesa": ["¿Corre sin caerse?"],
        },
        36: {
            "personal_social": ["¿Juega con otros niños?"],
            "motora_fina_adaptativa": ["¿Usa tijeras para cortar papel?"],
            "lenguaje": ["¿Habla con claridad suficiente para ser entendido por extraños?"],
            "motora_gruesa": ["¿Salta en un solo pie sin apoyo?"],
        },
    }

    OPTIONS = ["P", "F", "R", "NS"]

    def __init__(self, parent=None, patient=None, age_m=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.child')
        self.patient = patient
        self.age_m = age_m
        self.setWindowTitle("Denver II")
        self.resize(750, 650)

        main_layout = QVBoxLayout(self)
        form = QFormLayout()
        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        StyleHelper.style_input(self.patient_combo)
        form.addRow(QLabel("Paciente:"), self.patient_combo)
        self.age_label = QLabel("" if age_m is None else f"Edad: {age_m} meses")
        form.addRow(self.age_label)
        main_layout.addLayout(form)

        self.area_boxes = {}
        items = self.ITEMS.get(age_m, {})
        for area, questions in items.items():
            grp = QGroupBox(area.replace("_", " ").title())
            StyleHelper.style_groupbox(grp)
            gform = QFormLayout()
            boxes = []
            for q in questions:
                combo = QComboBox(); combo.addItems(self.OPTIONS)
                StyleHelper.style_input(combo)
                lbl = QLabel(q); lbl.setWordWrap(True)
                gform.addRow(lbl, combo)
                boxes.append((q, combo))
            grp.setLayout(gform)
            main_layout.addWidget(grp)
            self.area_boxes[area] = boxes

        main_layout.addWidget(QLabel("Observaciones:"))
        self.obs_edit = QTextEdit()
        StyleHelper.style_input(self.obs_edit)
        main_layout.addWidget(self.obs_edit)

        self.report_edit = AutoAdjustingTextEdit()
        self.report_edit.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        StyleHelper.style_input(self.report_edit)
        main_layout.addWidget(self.report_edit)

        btn_layout = QHBoxLayout()
        self.gen_btn = QPushButton("Generar Informe")
        StyleHelper.style_button(self.gen_btn)
        self.pdf_btn = QPushButton("Exportar PDF")
        StyleHelper.style_button(self.pdf_btn)
        btn_layout.addWidget(self.gen_btn)
        btn_layout.addWidget(self.pdf_btn)
        main_layout.addLayout(btn_layout)

        self.gen_btn.clicked.connect(self.generate_report)
        self.pdf_btn.clicked.connect(self.export_pdf)

        self.load_patients()
        if patient:
            self.patient_combo.setCurrentText(patient.nombre)

    def load_patients(self):
        session = Session()
        pats = session.query(Paciente).order_by(Paciente.nombre).all()
        self.patient_combo.addItems([p.nombre for p in pats])
        self.pat_map = {p.nombre: p for p in pats}
        session.close()

    def gather_responses(self):
        data = {}
        for area, items in self.area_boxes.items():
            data[area] = [(q, box.currentText()) for q, box in items]
        return data

    def generate_report(self):
        resp = self.gather_responses()
        result = interpretar_denver(resp)
        diagnosis = (
            "Desarrollo psicomotor adecuado" if result == "Normal" else
            "Retraso en el desarrollo psicomotor"
        )
        desc = (
            f"Se aplicó Denver II para {self.age_m} meses con resultado: {result}."
        )
        if result == "Anormal":
            desc += " Se recomienda referencia a neuropediatría, psicología infantil o intervención temprana."
        elif result == "Sospechoso":
            desc += " Se sugiere repetir el tamizaje en 2 semanas."
        if self.obs_edit.toPlainText().strip():
            desc += "\nObservaciones: " + self.obs_edit.toPlainText().strip()
        self.report_edit.setPlainText(desc)

        name = self.patient_combo.currentText().strip()
        pat = self.pat_map.get(name)
        if not pat:
            return
        session = Session()
        try:
            p = session.merge(pat)
            done = set(filter(None, (p.denver_done or "").split(';')))
            if self.age_m is not None:
                done.add(str(self.age_m))
            p.denver_done = ';'.join(sorted(done, key=int))
            hist = None
            if getattr(self.parent(), 'current_history_id', None):
                hist = session.get(HistoriaClinica, self.parent().current_history_id)
            if hist:
                hist.denver_resultado = result
                hist.denver_diagnostico = diagnosis
                hist.denver_descripcion = desc
                hist.historia_enfermedad = (hist.historia_enfermedad or '') + '\n' + desc
            session.commit()
        except Exception:
            session.rollback()
        finally:
            session.close()

        if getattr(self.parent(), 'historia_enfermedad_input', None):
            self.parent().historia_enfermedad_input.append(desc)

    def export_pdf(self):
        name = self.patient_combo.currentText().strip()
        fname, _ = QFileDialog.getSaveFileName(
            self, 'Guardar PDF', f'DenverII_{name}_{datetime.now().strftime("%Y%m%d")}.pdf', 'PDF (*.pdf)')
        if not fname:
            return
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import cm

        cfg = ConfigReceta(); cfg.load_config()
        c = canvas.Canvas(fname, pagesize=letter)
        y = 27 * cm
        c.setFont('Helvetica-Bold', 14)
        c.drawCentredString(10.5 * cm, y, 'Informe Denver II')
        y -= 1 * cm
        c.setFont('Helvetica', 10)
        date_str = datetime.now().strftime('%d/%m/%Y')
        c.drawString(2 * cm, y, f'Paciente: {name}')
        c.drawRightString(19 * cm, y, f'Fecha: {date_str}')
        y -= 0.6 * cm
        for area, items in self.area_boxes.items():
            c.setFont('Helvetica-Bold', 10)
            c.drawString(2 * cm, y, area.replace('_', ' ').title())
            y -= 0.5 * cm
            c.setFont('Helvetica', 10)
            for q, box in items:
                c.drawString(2 * cm, y, q)
                c.drawRightString(19 * cm, y, f'Respuesta: {box.currentText()}')
                y -= 0.4 * cm
                if y < 3 * cm:
                    c.showPage(); y = 26 * cm
        y -= 0.4 * cm
        for line in self.report_edit.toPlainText().splitlines():
            clean = re.sub(r'[\*\u2022]+', '', line).strip()
            if not clean:
                continue
            c.drawString(2 * cm, y, clean)
            y -= 0.4 * cm
            if y < 5 * cm:
                break

        sign = cfg.get_signature()
        seal = cfg.get_seal()
        if sign and os.path.exists(sign):
            c.drawImage(sign, 2 * cm, 3 * cm, width=4 * cm, preserveAspectRatio=True)
        if seal and os.path.exists(seal):
            c.drawImage(seal, 12 * cm, 2.5 * cm, width=4 * cm, preserveAspectRatio=True)
        c.drawCentredString(10.5 * cm, 2 * cm, cfg.get_doctor_name())
        c.drawCentredString(10.5 * cm, 1.6 * cm, cfg.get_doctor_specialty())

        c.save()
        QMessageBox.information(self, 'Éxito', f'PDF guardado en {fname}')
        try:
            os.startfile(fname)
        except Exception:
            pass


class GAD7Dialog(QDialog):
    """Formulario GAD-7 para evaluar ansiedad."""
    QUESTIONS = [
        "Sensación de nerviosismo, ansiedad o tensión",
        "Incapaz de parar o controlar las preocupaciones",
        "Se preocupa demasiado por diferentes cosas",
        "Dificultad para relajarse",
        "Estar tan inquieto(a) que es difícil quedarse quieto(a)",
        "Irritabilidad o molestarse fácilmente",
        "Sentir miedo como si algo terrible pudiera pasar",
    ]

    def __init__(self, parent=None, patient=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.diagnoses')
        self.patient = patient
        self.setWindowTitle("GAD-7: Evaluación de Ansiedad")
        self.resize(700, 600)

        main_layout = QVBoxLayout(self)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        container = QWidget()
        layout = QVBoxLayout(container)
        form = QFormLayout()

        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        StyleHelper.style_input(self.patient_combo)
        form.addRow(QLabel("Paciente:"), self.patient_combo)

        self.option_texts = [
            "0: Nunca",
            "1: Varios días",
            "2: Más de la mitad de los días",
            "3: Casi todos los días",
        ]
        self.answer_boxes = []
        for i, qtext in enumerate(self.QUESTIONS, 1):
            combo = QComboBox()
            combo.addItems(self.option_texts)
            StyleHelper.style_input(combo)
            form.addRow(QLabel(f"{i}. {qtext}"), combo)
            combo.currentIndexChanged.connect(self.update_score_label)
            self.answer_boxes.append(combo)

        layout.addLayout(form)

        self.total_label = QLabel("Puntaje total: 0 - Mínima o ausente")
        layout.addWidget(self.total_label)

        self.report_edit = AutoAdjustingTextEdit()
        self.report_edit.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        StyleHelper.style_input(self.report_edit)
        layout.addWidget(self.report_edit)

        btn_layout = QHBoxLayout()
        self.gen_btn = QPushButton("Generar Informe")
        StyleHelper.style_button(self.gen_btn)
        self.pdf_btn = QPushButton("Exportar PDF")
        StyleHelper.style_button(self.pdf_btn)
        btn_layout.addWidget(self.gen_btn)
        btn_layout.addWidget(self.pdf_btn)
        layout.addLayout(btn_layout)

        scroll.setWidget(container)
        main_layout.addWidget(scroll)

        self.gen_btn.clicked.connect(self.generate_report)
        self.pdf_btn.clicked.connect(self.export_pdf)

        self.load_patients()
        if patient:
            self.patient_combo.setCurrentText(patient.nombre)

        self.update_score_label()

    def load_patients(self):
        session = Session()
        pats = session.query(Paciente).order_by(Paciente.nombre).all()
        self.patient_combo.addItems([p.nombre for p in pats])
        self.pat_map = {p.nombre: p for p in pats}
        session.close()

    def score(self):
        return sum(box.currentIndex() for box in self.answer_boxes)

    def update_score_label(self):
        total = self.score()
        interpretation = (
            "Mínima o ausente" if total <= 4 else
            "Leve" if total <= 9 else
            "Moderada" if total <= 14 else
            "Severa"
        )
        self.total_label.setText(f"Puntaje total: {total} - {interpretation}")

    def save_score(self, pat, total):
        session = Session()
        try:
            p = session.merge(pat)
            p.gad7_score = total
            session.commit()
        except Exception:
            session.rollback()
        finally:
            session.close()

    def generate_report(self):
        name = self.patient_combo.currentText().strip()
        pat = self.pat_map.get(name)
        total = self.score()
        if pat:
            self.save_score(pat, total)
        interpretation = (
            "Mínima o ausente" if total <= 4 else
            "Leve" if total <= 9 else
            "Moderada" if total <= 14 else
            "Severa"
        )

        summary = f"Puntaje total: {total} - {interpretation}"
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r') as f:
                    cfg = json.load(f)
                model = apply_openai_key(cfg)
                prompt = (
                    "Eres un profesional de salud mental. Resume brevemente las "
                    f"implicaciones de un puntaje GAD-7 de {total} ({interpretation}) "
                    "y sugiere recomendaciones iniciales."
                )
                resp = openai.ChatCompletion.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "Asistente médico"},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.5,
                )
                ai_text = resp.choices[0].message.content.strip()
                ai_text = re.sub(r"[\*\u2022]+", "", ai_text)
                ai_text = re.sub(r"^[-\s]+", "", ai_text, flags=re.MULTILINE)
                summary += "\n" + ai_text
            except Exception as e:
                summary += f"\n[Error IA: {e}]"
        self.report_edit.setPlainText(summary)

    def export_pdf(self):
        name = self.patient_combo.currentText().strip()
        default = f'GAD7_{name}_{datetime.now().strftime("%Y%m%d")}.pdf'
        fname, _ = QFileDialog.getSaveFileName(self, 'Guardar PDF', default, 'PDF (*.pdf)')
        if not fname:
            return
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import cm

        cfg = ConfigReceta(); cfg.load_config()
        c = canvas.Canvas(fname, pagesize=letter)
        y = 27 * cm
        c.setFont('Helvetica-Bold', 14)
        c.drawCentredString(10.5 * cm, y, 'Informe de Evaluación GAD-7')
        y -= 1 * cm
        c.setFont('Helvetica', 10)
        date_str = datetime.now().strftime('%d/%m/%Y')
        c.drawString(2 * cm, y, f'Paciente: {name}')
        c.drawRightString(19 * cm, y, f'Fecha: {date_str}')
        y -= 0.6 * cm
        total = self.score()
        interpretation = (
            'Mínima o ausente' if total <= 4 else
            'Leve' if total <= 9 else
            'Moderada' if total <= 14 else
            'Severa'
        )
        c.drawString(2 * cm, y, f'Puntaje total: {total} - {interpretation}')
        y -= 0.8 * cm
        for i, box in enumerate(self.answer_boxes, 1):
            txt = self.QUESTIONS[i-1]
            ans = self.option_texts[box.currentIndex()]
            c.drawString(2 * cm, y, f'{i}. {txt}')
            y -= 0.4 * cm
            c.drawString(2.5 * cm, y, f'Respuesta: {ans}')
            y -= 0.6 * cm
            if y < 3 * cm:
                c.showPage()
                y = 26 * cm
        y -= 0.4 * cm
        c.setFont('Helvetica', 10)
        for line in self.report_edit.toPlainText().splitlines():
            clean = re.sub(r'[\*\u2022]+', '', line).strip()
            if not clean:
                continue
            c.drawString(2 * cm, y, clean)
            y -= 0.4 * cm
            if y < 5 * cm:
                break

        sign = cfg.get_signature()
        seal = cfg.get_seal()
        if sign and os.path.exists(sign):
            c.drawImage(sign, 2 * cm, 3 * cm, width=4 * cm, preserveAspectRatio=True)
        if seal and os.path.exists(seal):
            c.drawImage(seal, 12 * cm, 2.5 * cm, width=4 * cm, preserveAspectRatio=True)
        c.drawCentredString(10.5 * cm, 2 * cm, cfg.get_doctor_name())
        c.drawCentredString(10.5 * cm, 1.6 * cm, cfg.get_doctor_specialty())

        c.save()
        QMessageBox.information(self, 'Éxito', f'PDF guardado en {fname}')
        try:
            os.startfile(fname)
        except Exception:
            pass


class MMSEDialog(QDialog):
    """Versión simplificada del Mini-Mental State Examination."""

    GUIDE = (
        "Mini-Mental: Evaluación del Estado Mental\n\n"
        "Instrucciones al profesional:\n"
        "Este cuestionario se utiliza para evaluar la función cognitiva global del paciente. "
        "Se aplica de forma oral y algunas partes requieren observación directa. "
        "El puntaje máximo es 30 puntos.\n\n"
        "1. Orientación (10 puntos)\n"
        "¿Qué día de la semana es hoy? 1\n"
        "¿Qué día del mes es hoy? 1\n"
        "¿En qué mes estamos? 1\n"
        "¿En qué año estamos? 1\n"
        "¿En qué estación del año estamos? 1\n"
        "¿Dónde estamos ahora (país)? 1\n"
        "¿En qué ciudad? 1\n"
        "¿En qué lugar específico (clínica, casa)? 1\n"
        "¿En qué piso o planta está? 1\n"
        "¿En qué provincia o región estamos? 1\n\n"
        "2. Registro / Memoria inmediata (3 puntos)\n"
        "Diga al paciente 3 palabras simples (ej.: manzana, mesa, sol). Pídale que las repita inmediatamente.\n"
        "Repite correctamente 3 palabras /3\n\n"
        "3. Atención y cálculo (5 puntos)\n"
        "Restar de 7 en 7 a partir de 100 (máx. 5 restas): 100, 93, 86, 79, 72, 65\n"
        "Resta correcta por cada respuesta acertada 1 x 5\n\n"
        "4. Memoria diferida (3 puntos)\n"
        "Pida al paciente que repita las 3 palabras originales después de unos minutos.\n"
        "Reproduce correctamente las 3 palabras /3\n\n"
        "5. Lenguaje y comprensión (9 puntos)\n"
        "Nombrar dos objetos mostrados (por ejemplo, reloj y lápiz) 2\n"
        "Repetir una frase (“Ni sí, ni no, ni pero”) 1\n"
        "Dar una orden de 3 pasos (Ej: “Tome este papel con la mano derecha, dóblelo y colóquelo sobre la mesa.”) 3\n"
        "Leer y obedecer una orden escrita (“Cierre los ojos”) 1\n"
        "Escribir una frase con sujeto y verbo 1\n"
        "Copiar un dibujo de dos pentágonos que se superponen 1\n\n"
        "Interpretación del puntaje:\n"
        "25–30 Normal\n"
        "21–24 Deterioro cognitivo leve\n"
        "10–20 Deterioro cognitivo moderado\n"
        "<10 Deterioro severo"
    )

    QUESTIONS = [
        ("¿Qué día de la semana es hoy?", 1),
        ("¿Qué día del mes es hoy?", 1),
        ("¿En qué mes estamos?", 1),
        ("¿En qué año estamos?", 1),
        ("¿En qué estación del año estamos?", 1),
        ("¿Dónde estamos ahora (país)?", 1),
        ("¿En qué ciudad?", 1),
        ("¿En qué lugar específico (clínica, casa)?", 1),
        ("¿En qué piso o planta está?", 1),
        ("¿En qué provincia o región estamos?", 1),
        ("Repite correctamente 3 palabras", 3),
        ("Resta correcta por cada respuesta acertada (1 x 5)", 5),
        ("Reproduce correctamente las 3 palabras", 3),
        ("Nombrar dos objetos mostrados (por ejemplo, reloj y lápiz)", 2),
        ("Repetir una frase (“Ni sí, ni no, ni pero”)", 1),
        ("Dar una orden de 3 pasos (Ej: “Tome este papel con la mano derecha, dóblelo y colóquelo sobre la mesa.”)", 3),
        ("Leer y obedecer una orden escrita (“Cierre los ojos”)", 1),
        ("Escribir una frase con sujeto y verbo", 1),
        ("Copiar un dibujo de dos pentágonos que se superponen", 1),
    ]

    def __init__(self, parent=None, patient=None):
        super().__init__(parent)
        StyleHelper.set_window_icon(self, 'fa5s.brain')
        self.patient = patient
        self.setWindowTitle("Mini-Mental: Evaluación del Estado Mental")
        self.resize(700, 650)

        main_layout = QVBoxLayout(self)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        container = QWidget()
        layout = QVBoxLayout(container)

        header_form = QFormLayout()
        self.patient_combo = QComboBox()
        self.patient_combo.setEditable(True)
        StyleHelper.style_input(self.patient_combo)
        header_form.addRow(QLabel("Paciente:"), self.patient_combo)
        layout.addLayout(header_form)

        self.guide_btn = QPushButton("Ver Guía")
        StyleHelper.style_button(self.guide_btn)
        self.guide_btn.clicked.connect(self.show_guide)
        layout.addWidget(self.guide_btn, alignment=Qt.AlignRight)

        self.answer_boxes = []
        sections = [
            ("1. Orientación (10 puntos)", list(range(0, 10)), ""),
            ("2. Registro / Memoria inmediata (3 puntos)", [10],
             "Diga al paciente 3 palabras simples (ej.: manzana, mesa, sol). Pídale que las repita inmediatamente."),
            ("3. Atención y cálculo (5 puntos)", [11],
             "Restar de 7 en 7 a partir de 100 (máx. 5 restas): 100, 93, 86, 79, 72, 65"),
            ("4. Memoria diferida (3 puntos)", [12],
             "Pida al paciente que repita las 3 palabras originales después de unos minutos."),
            ("5. Lenguaje y comprensión (9 puntos)", list(range(13, 19)), ""),
        ]

        idx = 1
        for title, indexes, info in sections:
            grp = QGroupBox(title)
            StyleHelper.style_groupbox(grp)
            gf = QFormLayout()
            gf.setRowWrapPolicy(QFormLayout.WrapAllRows)
            if info:
                lbl = QLabel(info)
                lbl.setWordWrap(True)
                gf.addRow(lbl)
            for q_i in indexes:
                qtext, maxv = self.QUESTIONS[q_i]
                spin = QSpinBox()
                spin.setRange(0, maxv)
                StyleHelper.style_input(spin)
                lbl = QLabel(f"{idx}. {qtext}")
                lbl.setWordWrap(True)
                gf.addRow(lbl, spin)
                spin.valueChanged.connect(self.update_score_label)
                self.answer_boxes.append(spin)
                idx += 1
            grp.setLayout(gf)
            layout.addWidget(grp)

        self.total_label = QLabel("Puntaje total: 0 - Normal")
        layout.addWidget(self.total_label)

        self.report_edit = AutoAdjustingTextEdit()
        self.report_edit.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        StyleHelper.style_input(self.report_edit)
        layout.addWidget(self.report_edit)

        btn_layout = QHBoxLayout()
        self.gen_btn = QPushButton("Generar Informe")
        StyleHelper.style_button(self.gen_btn)
        self.pdf_btn = QPushButton("Exportar PDF")
        StyleHelper.style_button(self.pdf_btn)
        btn_layout.addWidget(self.gen_btn)
        btn_layout.addWidget(self.pdf_btn)
        layout.addLayout(btn_layout)

        scroll.setWidget(container)
        main_layout.addWidget(scroll)

        self.gen_btn.clicked.connect(self.generate_report)
        self.pdf_btn.clicked.connect(self.export_pdf)

        self.load_patients()
        if patient:
            self.patient_combo.setCurrentText(patient.nombre)

        self.update_score_label()

    def load_patients(self):
        session = Session()
        pats = session.query(Paciente).order_by(Paciente.nombre).all()
        self.patient_combo.addItems([p.nombre for p in pats])
        self.pat_map = {p.nombre: p for p in pats}
        session.close()

    def show_guide(self):
        dlg = QDialog(self)
        dlg.setWindowTitle("Guía MMSE")
        layout = QVBoxLayout(dlg)
        text = QTextEdit()
        text.setReadOnly(True)
        text.setPlainText(self.GUIDE)
        layout.addWidget(text)
        close_btn = QPushButton("Cerrar")
        StyleHelper.style_button(close_btn)
        close_btn.clicked.connect(dlg.accept)
        layout.addWidget(close_btn)
        dlg.resize(600, 700)
        dlg.exec_()

    def score(self):
        return sum(box.value() for box in self.answer_boxes)

    def update_score_label(self):
        total = self.score()
        interpretation = (
            "Normal" if total >= 25 else
            "Deterioro cognitivo leve" if total >= 21 else
            "Deterioro cognitivo moderado" if total >= 10 else
            "Deterioro severo"
        )
        self.total_label.setText(f"Puntaje total: {total} - {interpretation}")

    def save_score(self, pat, total):
        session = Session()
        try:
            p = session.merge(pat)
            p.mmse_score = total
            session.commit()
        except Exception:
            session.rollback()
        finally:
            session.close()

    def generate_report(self):
        name = self.patient_combo.currentText().strip()
        pat = self.pat_map.get(name)
        total = self.score()
        if pat:
            self.save_score(pat, total)
        interpretation = (
            "Normal" if total >= 25 else
            "Deterioro cognitivo leve" if total >= 21 else
            "Deterioro cognitivo moderado" if total >= 10 else
            "Deterioro severo"
        )

        summary = f"Puntaje total: {total} - {interpretation}"
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r') as f:
                    cfg = json.load(f)
                model = apply_openai_key(cfg)
                prompt = (
                    "Eres un médico neurólogo. Resume brevemente la "
                    f"interpretación de un MMSE de {total} puntos ({interpretation}) "
                    "y sugiere las acciones iniciales."
                )
                resp = openai.ChatCompletion.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "Asistente médico"},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.5,
                )
                ai_text = resp.choices[0].message.content.strip()
                ai_text = re.sub(r"[\*\u2022]+", "", ai_text)
                ai_text = re.sub(r"^[-\s]+", "", ai_text, flags=re.MULTILINE)
                summary += "\n" + ai_text
            except Exception as e:
                summary += f"\n[Error IA: {e}]"
        self.report_edit.setPlainText(summary)

    def export_pdf(self):
        name = self.patient_combo.currentText().strip()
        default = f'MMSE_{name}_{datetime.now().strftime("%Y%m%d")}.pdf'
        fname, _ = QFileDialog.getSaveFileName(self, 'Guardar PDF', default, 'PDF (*.pdf)')
        if not fname:
            return
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import cm

        cfg = ConfigReceta(); cfg.load_config()
        c = canvas.Canvas(fname, pagesize=letter)
        y = 27 * cm
        c.setFont('Helvetica-Bold', 14)
        c.drawCentredString(10.5 * cm, y, 'Informe Evaluación Mini-Mental')
        y -= 1 * cm
        c.setFont('Helvetica', 10)
        date_str = datetime.now().strftime('%d/%m/%Y')
        c.drawString(2 * cm, y, f'Paciente: {name}')
        c.drawRightString(19 * cm, y, f'Fecha: {date_str}')
        y -= 0.6 * cm
        total = self.score()
        interpretation = (
            'Normal' if total >= 25 else
            'Deterioro cognitivo leve' if total >= 21 else
            'Deterioro cognitivo moderado' if total >= 10 else
            'Deterioro severo'
        )
        c.drawString(2 * cm, y, f'Puntaje total: {total} - {interpretation}')
        y -= 0.8 * cm
        for i, (qtxt, _max) in enumerate(self.QUESTIONS, 1):
            ans = self.answer_boxes[i-1].value()
            c.drawString(2 * cm, y, f'{i}. {qtxt}')
            y -= 0.4 * cm
            c.drawString(2.5 * cm, y, f'Puntaje: {ans}')
            y -= 0.6 * cm
            if y < 3 * cm:
                c.showPage()
                y = 26 * cm
        y -= 0.4 * cm
        c.setFont('Helvetica', 10)
        for line in self.report_edit.toPlainText().splitlines():
            clean = re.sub(r'[\*\u2022]+', '', line).strip()
            if not clean:
                continue
            c.drawString(2 * cm, y, clean)
            y -= 0.4 * cm
            if y < 5 * cm:
                break

        sign = cfg.get_signature()
        seal = cfg.get_seal()
        if sign and os.path.exists(sign):
            c.drawImage(sign, 2 * cm, 3 * cm, width=4 * cm, preserveAspectRatio=True)
        if seal and os.path.exists(seal):
            c.drawImage(seal, 12 * cm, 2.5 * cm, width=4 * cm, preserveAspectRatio=True)
        c.drawCentredString(10.5 * cm, 2 * cm, cfg.get_doctor_name())
        c.drawCentredString(10.5 * cm, 1.6 * cm, cfg.get_doctor_specialty())

        c.save()
        QMessageBox.information(self, 'Éxito', f'PDF guardado en {fname}')
        try:
            os.startfile(fname)
        except Exception:
            pass

                     
def main():
    show_intro()
    openai.api_key = read_api_key()
    app = QApplication(sys.argv)
    theme = 'gpt'
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r') as f:
                cfg = json.load(f)
                theme = cfg.get('theme', 'gpt')
        except Exception:
            theme = 'gpt'

    if theme == 'light':
        StyleHelper.set_light_style(app)
    elif theme == 'dark':
        StyleHelper.set_dark_style(app)
    else:
        StyleHelper.set_gpt_style(app)


        # Simula un proceso de inicialización
    # Muestra la pantalla de carga
    loading_screen = LoadingScreen()
    loading_screen.show()
    # Initial database setup
    ensure_db_schema(engine)
    session = Session()
    if session.query(InsuranceName).count() == 0:
        for name in ["ARS Humano", "ARS Senasa", "ARS Universal", "Privada", "Ninguna"]:
            session.add(InsuranceName(nombre=name))
        session.commit()
    session.close()

    # Configura la ventana principal
    window = MainWindow()

   

    # Configura el hilo de inicialización
    init_thread = InitializationThread()
    init_thread.loading_complete.connect(loading_screen.stop_animation)
    loading_thread = LoadingThread(loading_screen)
    def create_ready_file():
        with open("app_ready.txt", "w") as f:
            f.write("App is ready")

    # Conecta la señal de carga completa con la creación del archivo indicador
    init_thread.loading_complete.connect(create_ready_file)
    loading_thread.start()   
    # Creamos un archivo indicador cuando la app esté lista

    # Inicia el hilo de inicialización
    init_thread.start()

    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
